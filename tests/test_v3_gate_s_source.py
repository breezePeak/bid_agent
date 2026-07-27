"""Gate S automated matrix: promoted Source authority, fidelity, fail-closed paths."""

from __future__ import annotations

import hashlib
import json
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from control_plane import ControlPlaneError, ControlStore, WorkspaceContext  # noqa: E402
from document_pipeline.contracts import InputRole, SourceIndex, TemplateStructureContract  # noqa: E402
from document_pipeline.input_manifest import InputManifestService, MANIFEST_PATH  # noqa: E402
from document_pipeline.source_artifacts import (  # noqa: E402
    load_promoted_input_manifest,
    load_promoted_source_index,
    load_promoted_template_structure,
    require_promoted_source_index,
)
from document_pipeline.source_normalizer import SOURCE_INDEX_PATH, SourceNormalizer  # noqa: E402
from document_pipeline.stage_runner import V3StageRunner  # noqa: E402
from document_pipeline.template_contract import TemplateContractCompiler  # noqa: E402
from utils import write_json  # noqa: E402


def _minimal_pdf(pages: list[str]) -> bytes:
    """Build a tiny PDF with optional text per page (empty string => blank page)."""
    objs: dict[int, str] = {1: "<< /Type /Catalog /Pages 2 0 R >>"}
    kids: list[str] = []
    next_id = 3
    content_ids: list[tuple[int, str]] = []
    page_ids: list[tuple[int, int]] = []
    for text in pages:
        content_id = next_id
        next_id += 1
        page_id = next_id
        next_id += 1
        content_ids.append((content_id, text))
        page_ids.append((page_id, content_id))
        kids.append(f"{page_id} 0 R")
    font_id = next_id
    objs[2] = f"<< /Type /Pages /Kids [{' '.join(kids)}] /Count {len(pages)} >>"
    for content_id, text in content_ids:
        stream = f"BT /F1 24 Tf 50 50 Td ({text}) Tj ET" if text else ""
        objs[content_id] = f"<< /Length {len(stream)} >>\nstream\n{stream}\nendstream"
    for page_id, content_id in page_ids:
        objs[page_id] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] "
            f"/Contents {content_id} 0 R /Resources << /Font << /F1 {font_id} 0 R >> >> >>"
        )
    objs[font_id] = "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
    out = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {0: 0}
    for i in sorted(objs):
        offsets[i] = len(out)
        out.extend(f"{i} 0 obj\n{objs[i]}\nendobj\n".encode("latin-1"))
    xref = len(out)
    out.extend(f"xref\n0 {max(objs) + 1}\n".encode("latin-1"))
    out.extend(b"0000000000 65535 f \n")
    for i in range(1, max(objs) + 1):
        out.extend(f"{offsets[i]:010d} 00000 n \n".encode("latin-1"))
    out.extend(
        f"trailer\n<< /Size {max(objs) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("latin-1")
    )
    return bytes(out)


class GateSSourceTests(unittest.TestCase):
    def _context(self, base: Path) -> WorkspaceContext:
        runs = base / "runs"
        (runs / "alpha").mkdir(parents=True)
        return WorkspaceContext.resolve(runs, "alpha")

    def test_input_manifest_and_source_index_are_promoted(self) -> None:
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            base = Path(tmp)
            tender = base / "tender.md"
            tender.write_text("项目目标。\n\n服务范围。", encoding="utf-8")
            context = self._context(base)
            InputManifestService(context).register_local_file(tender, InputRole.TENDER)
            SourceNormalizer(context).normalize_active_inputs()

            manifest = load_promoted_input_manifest(context)
            index = load_promoted_source_index(context)
            self.assertIsNotNone(manifest)
            self.assertIsNotNone(index)
            assert index is not None
            self.assertTrue(index.coverage.items)
            self.assertTrue(any(item.status == "normalized" for item in index.coverage.items))
            self.assertEqual(require_promoted_source_index(context).revision, index.revision)

    def test_disk_json_mutation_does_not_change_promoted_authority(self) -> None:
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            base = Path(tmp)
            tender = base / "tender.md"
            tender.write_text("唯一义务：按时交付。", encoding="utf-8")
            context = self._context(base)
            InputManifestService(context).register_local_file(tender, InputRole.TENDER)
            SourceNormalizer(context).normalize_active_inputs()
            before = load_promoted_source_index(context)
            assert before is not None
            block_ids = [b.block_id for b in before.blocks]

            # Corrupt both projections on disk.
            write_json(context.root / SOURCE_INDEX_PATH, {"schema_version": "v3", "revision": 99, "blocks": []})
            write_json(context.root / MANIFEST_PATH, {"schema_version": "v3", "revision": 99, "inputs": []})

            after_index = load_promoted_source_index(context)
            after_manifest = load_promoted_input_manifest(context)
            assert after_index is not None and after_manifest is not None
            self.assertEqual([b.block_id for b in after_index.blocks], block_ids)
            self.assertGreaterEqual(len(after_manifest.inputs), 1)
            self.assertNotEqual(after_index.revision, 99)

    def test_downstream_stage_rejects_unpromoted_source_json(self) -> None:
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            base = Path(tmp)
            context = self._context(base)
            (context.root / "workspace/v3").mkdir(parents=True)
            write_json(
                context.root / SOURCE_INDEX_PATH,
                {
                    "schema_version": "v3",
                    "revision": 1,
                    "source_hashes": {},
                    "blocks": [],
                    "coverage": {"items": []},
                    "input_manifest_revision": 1,
                },
            )
            with self.assertRaises(ControlPlaneError):
                V3StageRunner(context).run("build_requirement_ledger")

    def test_repeated_normalize_is_deterministic(self) -> None:
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            base = Path(tmp)
            path = base / "tender.docx"
            document = Document()
            document.add_heading("概述", level=1)
            document.add_paragraph("范围条款。")
            document.add_paragraph("1. 列表项甲")
            document.save(path)
            context = self._context(base)
            InputManifestService(context).register_local_file(path, InputRole.TENDER)
            first = SourceNormalizer(context).normalize_active_inputs()
            second = SourceNormalizer(context).normalize_active_inputs()
            self.assertEqual(
                [b["block_id"] for b in first["blocks"]],
                [b["block_id"] for b in second["blocks"]],
            )
            self.assertEqual(
                [b["ordinal"] for b in first["blocks"]],
                [b["ordinal"] for b in second["blocks"]],
            )
            # Active revision must remain single content-addressed stream.
            active = ControlStore(context).v3_active_artifact("SourceIndex")
            self.assertIsNotNone(active)

    def test_pdf_blank_page_is_structure_gap_not_whole_file_failure_when_other_pages_ok(self) -> None:
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            base = Path(tmp)
            pdf_path = base / "mixed.pdf"
            pdf_path.write_bytes(_minimal_pdf(["HelloSourcePage", ""]))
            context = self._context(base)
            InputManifestService(context).register_local_file(pdf_path, InputRole.TENDER)
            index = SourceNormalizer(context).normalize_active_inputs()
            kinds = [b["block_kind"] for b in index["blocks"]]
            self.assertIn("pdf_text", kinds)
            self.assertIn("ocr_gap", kinds)
            coverage_statuses = {item["status"] for item in index["coverage"]["items"]}
            self.assertIn("normalized", coverage_statuses)
            self.assertIn("structure_gap", coverage_statuses)
            # bbox present on extracted text when words are available
            text_blocks = [b for b in index["blocks"] if b["block_kind"] == "pdf_text"]
            self.assertTrue(text_blocks)
            self.assertTrue(any(b.get("bbox") for b in text_blocks) or text_blocks[0].get("page") == 1)

    def test_all_blank_pdf_is_blocked(self) -> None:
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            base = Path(tmp)
            pdf_path = base / "blank.pdf"
            pdf_path.write_bytes(_minimal_pdf(["", ""]))
            context = self._context(base)
            InputManifestService(context).register_local_file(pdf_path, InputRole.TENDER)
            with self.assertRaisesRegex(ValueError, "OCR_BLOCKED|SOURCE"):
                SourceNormalizer(context).normalize_active_inputs()

    def test_synthetic_deep_template_matches_frozen_structure(self) -> None:
        """A synthetic deep-template profile must match its committed structural snapshot."""
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            base = Path(tmp)
            path = base / "template_deep_structure.docx"
            document = Document()
            # Synthetic stress parameter only; it is not a business or release threshold.
            node_count = 0
            generated_heading_target = 198
            chapter = 0
            while node_count < generated_heading_target:
                chapter += 1
                document.add_heading(f"第{chapter}章 主题{chapter}", level=1)
                node_count += 1
                section = 0
                while node_count < generated_heading_target and section < 27:
                    section += 1
                    document.add_heading(f"{chapter}.{section} 小节{section}", level=2)
                    node_count += 1
                    if section == 1:
                        document.add_paragraph(f"{{{{slot-{chapter}}}}}")
            # One table under last H1 to check nearest-upstream binding.
            document.add_heading("附录表格章", level=1)
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "{{附录内容}}"
            document.save(path)

            context = self._context(base)
            item = InputManifestService(context).register_local_file(path, InputRole.TEMPLATE).item
            first = TemplateContractCompiler(context).compile_structure(item)
            # Second compile via new service instance; fingerprint and topology must match.
            second = TemplateContractCompiler(context).compile_structure(item)

            self.assertEqual(len(first.nodes), generated_heading_target + 1)
            self.assertEqual(len(first.nodes), len(second.nodes))
            self.assertEqual(
                [(n.order, n.level, n.title, n.parent_node_id, n.numbering) for n in first.nodes],
                [(n.order, n.level, n.title, n.parent_node_id, n.numbering) for n in second.nodes],
            )
            self.assertEqual(first.structural_fingerprint, second.structural_fingerprint)
            self.assertEqual(
                [(s.slot_id, s.node_id, s.kind, s.anchor) for s in first.slots],
                [(s.slot_id, s.node_id, s.kind, s.anchor) for s in second.slots],
            )
            promoted = load_promoted_template_structure(context)
            self.assertIsNotNone(promoted)
            assert promoted is not None
            self.assertEqual(promoted.structural_fingerprint, first.structural_fingerprint)

            actual_freeze = {
                "node_count": len(first.nodes),
                "slot_count": len(first.slots),
                "structural_fingerprint": first.structural_fingerprint,
                "nodes": [
                    {
                        "order": n.order,
                        "level": n.level,
                        "title": n.title,
                        "parent_node_id": n.parent_node_id,
                        "numbering": n.numbering,
                    }
                    for n in first.nodes
                ],
            }
            snapshot_path = ROOT / "tests" / "fixtures" / "v3_source" / "template_deep_structure_freeze.json"
            expected_freeze = json.loads(snapshot_path.read_text(encoding="utf-8"))
            self.assertEqual(expected_freeze["fixture_kind"], "synthetic_deep_template_structure_stress")
            self.assertEqual(expected_freeze["profile"]["generated_heading_count"], generated_heading_target)
            self.assertEqual(expected_freeze["profile"]["appended_table_heading_count"], 1)
            self.assertFalse(expected_freeze["profile"]["business_threshold"])
            self.assertEqual(
                {key: expected_freeze[key] for key in actual_freeze},
                actual_freeze,
            )

    def test_real_template_sample_is_promoted_and_stable(self) -> None:
        sample = ROOT / "sources" / "template" / "锡盟24年技术标投标框架.docx"
        if not sample.is_file():
            self.skipTest("real template sample missing")
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            base = Path(tmp)
            path = base / sample.name
            path.write_bytes(sample.read_bytes())
            context = self._context(base)
            item = InputManifestService(context).register_local_file(path, InputRole.TEMPLATE).item
            first = TemplateContractCompiler(context).compile_structure(item)
            second = TemplateContractCompiler(context).compile_structure(item)
            self.assertGreaterEqual(len(first.nodes), 1)
            self.assertEqual(first.structural_fingerprint, second.structural_fingerprint)
            self.assertEqual([n.title for n in first.nodes], [n.title for n in second.nodes])
            self.assertIsNotNone(ControlStore(context).v3_active_artifact("TemplateStructureContract"))

    def test_coverage_records_normalized_and_exempt_elements(self) -> None:
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            base = Path(tmp)
            path = base / "tender.md"
            path.write_text("标题行\n\n\n正文行\n", encoding="utf-8")
            context = self._context(base)
            InputManifestService(context).register_local_file(path, InputRole.TENDER)
            index = SourceNormalizer(context).normalize_active_inputs()
            statuses = {item["status"] for item in index["coverage"]["items"]}
            self.assertIn("normalized", statuses)
            self.assertIn("exempt", statuses)


if __name__ == "__main__":
    unittest.main()
