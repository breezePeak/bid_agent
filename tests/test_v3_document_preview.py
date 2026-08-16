from __future__ import annotations

import hashlib
import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from control_plane import ControlPlaneError, WorkspaceContext  # noqa: E402
from document_pipeline.document_preview import (  # noqa: E402
    FINAL_MARKDOWN_PATH,
    DocumentPreviewService,
    parse_markdown_preview,
)
from document_pipeline.quality import CONTENT_QUALITY_PATH  # noqa: E402
from document_pipeline.renderers.render_verifier import (  # noqa: E402
    RENDER_OUTPUT_PATH,
    RENDER_QUALITY_PATH,
    DeliveryVerifier,
)
from utils import write_json  # noqa: E402


class DocumentPreviewTests(unittest.TestCase):
    def test_markdown_preview_preserves_full_structure_and_duplicate_titles(self):
        toc, blocks = parse_markdown_preview(
            """# 投标文件

## 实施方案

第一章正文。

- 准备
- 执行

| 阶段 | 产物 |
| --- | --- |
| 实施 | 记录 |

## 实施方案

第二章正文。
"""
        )
        self.assertEqual([item["title"] for item in toc].count("实施方案"), 2)
        self.assertEqual(len({item["id"] for item in toc}), len(toc))
        self.assertTrue(any(item["type"] == "list" for item in blocks))
        table = next(item for item in blocks if item["type"] == "table")
        self.assertEqual(table["rows"][1], ["实施", "记录"])
        self.assertTrue(
            any(
                item.get("text") == "第二章正文。"
                for item in blocks
                if item["type"] == "paragraph"
            )
        )

    def test_ready_with_warnings_is_blocked_from_preview(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            runs = Path(tmp) / "runs"
            (runs / "alpha").mkdir(parents=True)
            context = WorkspaceContext.resolve(runs, "alpha")
            markdown_path = context.root / FINAL_MARKDOWN_PATH
            docx_path = context.root / RENDER_OUTPUT_PATH
            markdown_path.parent.mkdir(parents=True, exist_ok=True)
            markdown_path.write_text(
                "# 完整标书\n\n## 技术方案\n\n这里是完整正文。",
                encoding="utf-8",
            )
            docx_path.write_bytes(b"current-docx")
            digest = hashlib.sha256(docx_path.read_bytes()).hexdigest()
            write_json(
                context.root / RENDER_QUALITY_PATH,
                {
                    "status": "ready_with_warnings",
                    "artifact_sha256": digest,
                    "warnings": [
                        {
                            "code": "QUALITY_WARNING",
                            "message": "存在待复核覆盖项",
                        }
                    ],
                },
            )
            with self.assertRaises(ControlPlaneError) as captured:
                DocumentPreviewService(context).build()
            self.assertEqual(captured.exception.code, "DOCUMENT_PREVIEW_NOT_READY")

    def test_delivery_verifier_allows_quality_warning_and_hashes_both_formats(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            runs = Path(tmp) / "runs"
            (runs / "alpha").mkdir(parents=True)
            context = WorkspaceContext.resolve(runs, "alpha")
            docx_path = context.root / RENDER_OUTPUT_PATH
            markdown_path = context.root / FINAL_MARKDOWN_PATH
            docx_path.parent.mkdir(parents=True, exist_ok=True)
            document = Document()
            document.add_heading("技术标书", level=1)
            document.add_paragraph("完整正文")
            document.save(docx_path)
            markdown_path.write_text(
                "# 技术标书\n\n完整正文",
                encoding="utf-8",
            )
            write_json(
                context.root / CONTENT_QUALITY_PATH,
                {
                    "verdict": "warn",
                    "findings": [
                        {
                            "code": "COVERAGE_WARNING",
                            "message": "仍有一项覆盖关系待复核",
                        }
                    ],
                },
            )
            report = DeliveryVerifier(
                context,
                allow_quality_warnings=True,
            ).verify()
            self.assertEqual(report["status"], "ready_with_warnings")
            self.assertEqual(report["warning_count"], 1)
            self.assertEqual(
                report["artifact_sha256"],
                hashlib.sha256(docx_path.read_bytes()).hexdigest(),
            )
            self.assertEqual(
                report["markdown_sha256"],
                hashlib.sha256(markdown_path.read_bytes()).hexdigest(),
            )

    def test_latest_failed_render_attempt_rejects_old_preview(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            runs = Path(tmp) / "runs"
            (runs / "alpha").mkdir(parents=True)
            context = WorkspaceContext.resolve(runs, "alpha")
            markdown_path = context.root / FINAL_MARKDOWN_PATH
            docx_path = context.root / RENDER_OUTPUT_PATH
            markdown_path.parent.mkdir(parents=True, exist_ok=True)
            markdown_path.write_text("# 旧标书", encoding="utf-8")
            docx_path.write_bytes(b"old-docx")
            write_json(
                context.root / RENDER_QUALITY_PATH,
                {
                    "status": "ready",
                    "artifact_sha256": hashlib.sha256(
                        docx_path.read_bytes()
                    ).hexdigest(),
                },
            )
            service = DocumentPreviewService(context)
            stage_runs = [
                {
                    "stage_command": "render_document",
                    "attempt": 1,
                    "status": "succeeded",
                },
                {
                    "stage_command": "render_document",
                    "attempt": 2,
                    "status": "failed",
                },
                {
                    "stage_command": "verify_delivery",
                    "attempt": 1,
                    "status": "succeeded",
                },
            ]
            with (
                mock.patch.object(
                    service,
                    "_latest_generation_operation_id",
                    return_value="OP-current",
                ),
                mock.patch.object(
                    service.store,
                    "stage_runs",
                    return_value=stage_runs,
                ),
                self.assertRaises(ControlPlaneError) as raised,
            ):
                service.build()
            self.assertEqual(
                raised.exception.code,
                "DOCUMENT_PREVIEW_STALE",
            )


if __name__ == "__main__":
    unittest.main()
