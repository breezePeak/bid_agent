from __future__ import annotations

import asyncio,json,os,sys,tempfile,unittest
import io
from pathlib import Path
from types import SimpleNamespace
from unittest import mock

ROOT=Path(__file__).resolve().parents[1];sys.path.insert(0,str(ROOT/'src'))

from control_plane import CommandEnvelope,CommandGateway,ControlStore,WorkspaceContext
from document_pipeline.contracts import InputRole
from document_pipeline.execution_controller import V3ExecutionController
from document_pipeline.artifact_promotion import HumanGateService
from document_pipeline.input_manifest import InputManifestService
from document_pipeline.workspace_snapshot import V3WorkspaceSnapshotBuilder
import api.v3_app as v3_app
from api.settings_service import SettingsService
from fastapi import UploadFile
from fastapi.testclient import TestClient
from docx import Document


class _Request:
 def __init__(self, body, principal=None): self.body=body;self.state=SimpleNamespace(principal=principal)
 async def json(self): return self.body


class V3ExecutionControllerTests(unittest.TestCase):
 def test_public_fastapi_routes_expose_no_v1_or_v2_workspace_contract(self):
  paths=[getattr(route,'path','') for route in v3_app.app.routes]
  self.assertIn('/api/v3/workspaces',paths)
  self.assertFalse(any(path.startswith('/api/v2/') for path in paths))

 def test_login_shell_is_reachable_before_api_authentication(self):
  with TestClient(v3_app.app) as client:
   response=client.get('/login')
  self.assertEqual(response.status_code,200)
  self.assertIn('id=\"app\"',response.text)

  def test_legacy_api_namespaces_are_not_registered(self):
   with TestClient(v3_app.app) as client:
    for path in ('/api/v1/workspaces','/api/v2/workspaces/example/snapshot'):
     response=client.get(path)
     self.assertEqual(response.status_code,404)

 def test_command_gateway_is_the_only_execution_entry_point(self):
  with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as t:
   base=Path(t);runs=base/'runs';(runs/'alpha').mkdir(parents=True);context=WorkspaceContext.resolve(runs,'alpha')
   tender=base/'tender.docx';score=base/'score.md';doc=Document();doc.add_heading('项目技术需求',level=1);doc.add_paragraph('服务范围包括系统实施、培训、交付成果和验收条件，工期30日。');doc.add_heading('实施与验收',level=1);doc.add_paragraph('投标人应提交实施方案、服务响应和验收材料。');doc.add_heading('评标办法',level=1);doc.add_paragraph('项目实施方案完整性，满分10分。');doc.add_table(rows=1,cols=2).cell(0,0).text='交付成果';doc.save(tender);score.write_text('项目实施方案完整性，满分10分。',encoding='utf-8')
   inputs=InputManifestService(context);inputs.register_local_file(tender,InputRole.TENDER);inputs.register_local_file(score,InputRole.SCORE)
   controller=V3ExecutionController.for_deterministic_tests(context);gateway=CommandGateway(context,controller.handlers())
   envelope=CommandEnvelope.from_mapping({'kind':'document.prepare_outline','expected_revision':ControlStore(context).revision(),'idempotency_key':'v3-outline-run'},workspace_id='alpha')
   receipt=gateway.submit(envelope)
   self.assertEqual(receipt.status,'accepted')
   self.assertFalse((context.root/'outputs/v3/final.docx').exists())
   snapshot=V3WorkspaceSnapshotBuilder(context).build()
   # The legacy stage files are deliberately not runtime facts. Snapshot only
   # projects revisions that travelled through Proposal → Gate → Promotion.
   self.assertIsNone(snapshot['document']['mode'])
   self.assertIsNone(snapshot['document']['delivery'])
   self.assertEqual(snapshot['content_units'], [])
   stage_runs=ControlStore(context).stage_runs(str(receipt.operation_id))
   self.assertIsNotNone(ControlStore(context).v3_active_artifact('ProjectModel'))
   self.assertIsNone(ControlStore(context).v3_active_artifact('ResponseTopicGraph'))
   project_run=next(item for item in stage_runs if item['stage_command']=='plan_response')
   self.assertGreater(project_run['output']['input_chars'],0)
   self.assertGreater(project_run['output']['source_block_count'],0)
   self.assertEqual(
       {item['stage_command'] for item in stage_runs},
       {
        'ingest_inputs','normalize_sources','compile_template_structure',
        'build_requirement_ledger','analyze_scores','plan_response','compile_chapter_blueprint',
        'score_structure','score_semantic','confirm_planning',
       },
   )
   self.assertEqual(next(item for item in stage_runs if item['stage_command']=='confirm_planning')['status'],'blocked_human')
   store=ControlStore(context);store.grant_workspace_access('owner')
   planning_snapshot=HumanGateService(context).planning_snapshot()
   confirm=gateway.submit(CommandEnvelope.from_mapping({'kind':'document.confirm_planning','payload':{'decision':'confirm','planning_snapshot':planning_snapshot},'actor':{'type':'user','id':'owner'},'expected_revision':store.revision(),'idempotency_key':'v3-confirm'},workspace_id='alpha'))
   self.assertEqual(confirm.status,'accepted')
   self.assertEqual(confirm.status,'accepted')

 def test_prepare_outline_uses_embedded_tender_scores_and_stops_before_writing(self):
  with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as t:
   base=Path(t);runs=base/'runs';(runs/'alpha').mkdir(parents=True);context=WorkspaceContext.resolve(runs,'alpha')
   tender=base/'tender.md'
   tender.write_text(
    '# 项目需求\n投标人须提供实施方案并承诺30日内交付。\n\n'
    '# 评标办法\n## 技术评分\n项目实施方案完整性，满分10分。\n',
    encoding='utf-8',
   )
   InputManifestService(context).register_local_file(tender,InputRole.TENDER)
   controller=V3ExecutionController.for_deterministic_tests(context)
   gateway=CommandGateway(context,controller.handlers())
   store=ControlStore(context)
   receipt=gateway.submit(CommandEnvelope.from_mapping({
    'kind':'document.prepare_outline',
    'expected_revision':store.revision(),
    'idempotency_key':'v3-outline-only',
   },workspace_id='alpha'))
   self.assertEqual(receipt.status,'accepted',receipt.message)
   self.assertEqual(receipt.result['operation_status'],'blocked_human')
   self.assertTrue(receipt.result['planning_snapshot']['generation_trace'])
   snapshot=V3WorkspaceSnapshotBuilder(context).build()
   self.assertEqual(len(snapshot['analysis']['score_model']['points']),1)
   self.assertTrue(snapshot['analysis']['chapter_blueprint']['nodes'])
   self.assertEqual(snapshot['planning']['status'],'needs_human')
   self.assertIsNone(snapshot['document']['delivery'])
   self.assertFalse((context.root/'outputs/v3/final.docx').exists())
   stages={item['stage_command'] for item in store.stage_runs(str(receipt.operation_id))}
   self.assertEqual(stages,{
     'ingest_inputs','normalize_sources','compile_template_structure',
      'build_requirement_ledger','analyze_scores','score_structure','score_semantic','plan_response',
     'compile_chapter_blueprint','confirm_planning',
    })
   self.assertNotIn('execute_content_plan',stages)
   with mock.patch.object(
    controller.runner,
    'run',
    wraps=controller.runner.run,
   ) as resumed_run:
    second=gateway.submit(CommandEnvelope.from_mapping({
     'kind':'document.prepare_outline',
     'expected_revision':store.revision(),
     'idempotency_key':'v3-outline-only-resume',
    },workspace_id='alpha'))
   self.assertEqual(second.status,'accepted',second.message)
   resumed_commands=[
    item.args[0] for item in resumed_run.call_args_list
   ]
   self.assertNotIn('build_requirement_ledger',resumed_commands)
   self.assertNotIn('analyze_scores',resumed_commands)
   self.assertNotIn('plan_response',resumed_commands)
   self.assertNotIn('compile_chapter_blueprint',resumed_commands)
   second_stages={
    item['stage_command']:item['status']
    for item in store.stage_runs(str(second.operation_id))
   }
   self.assertEqual(second_stages['normalize_sources'],'reused')
   self.assertEqual(second_stages['build_requirement_ledger'],'reused')
   self.assertEqual(second_stages['analyze_scores'],'reused')
   self.assertEqual(second_stages['compile_chapter_blueprint'],'reused')

 def test_failed_outline_retry_marks_old_blueprint_and_h1_as_outdated(self):
  with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as t:
   base=Path(t);runs=base/'runs';(runs/'alpha').mkdir(parents=True);context=WorkspaceContext.resolve(runs,'alpha')
   tender=base/'tender.md'
   tender.write_text(
    '# 项目需求\n投标人须提供实施方案。\n\n'
    '# 评标办法\n## 技术评分\n项目实施方案完整性，满分10分。\n',
    encoding='utf-8',
   )
   InputManifestService(context).register_local_file(tender,InputRole.TENDER)
   controller=V3ExecutionController.for_deterministic_tests(context)
   store=ControlStore(context)
   first=CommandGateway(context,controller.handlers()).submit(CommandEnvelope.from_mapping({
    'kind':'document.prepare_outline',
    'expected_revision':store.revision(),
    'idempotency_key':'v3-outline-initial',
   },workspace_id='alpha'))
   self.assertEqual(first.status,'accepted',first.message)
   store.grant_workspace_access('owner')
   service=HumanGateService(context)
   service.confirm_planning(
    principal_id='owner',
    submitted_snapshot=service.planning_snapshot(),
    nonce='confirm-initial-outline',
   )
   self.assertEqual(V3WorkspaceSnapshotBuilder(context).build()['planning']['status'],'confirmed')

   def fail_retry(_context, _envelope, _operation_id):
    raise RuntimeError('ScoreModel 引用审计失败')

   handlers=controller.handlers()
   handlers['document.prepare_outline']=fail_retry
   failed=CommandGateway(context,handlers).submit(CommandEnvelope.from_mapping({
    'kind':'document.prepare_outline',
    'expected_revision':store.revision(),
    'idempotency_key':'v3-outline-failed-retry',
   },workspace_id='alpha'))
   self.assertEqual(failed.status,'rejected')

   snapshot=V3WorkspaceSnapshotBuilder(context).build()
   self.assertTrue(snapshot['analysis']['chapter_blueprint']['nodes'])
   self.assertTrue(snapshot['analysis']['latest_operation']['result_outdated'])
   self.assertTrue(snapshot['analysis']['stale'])
   self.assertEqual(snapshot['analysis']['status'],'failed')
   self.assertEqual(snapshot['planning']['status'],'outdated')
   self.assertNotIn('receipt_id',snapshot['planning'])

 def test_prepare_outline_records_failed_stage_and_exposes_pipeline_diagnostics(self):
   with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as t:
    runs=Path(t)/'runs';(runs/'alpha').mkdir(parents=True);context=WorkspaceContext.resolve(runs,'alpha')

    class _SemanticFailure(RuntimeError):
     code='score_semantic_candidate_invalid'
     attempts=2
     errors=(
      "JSONDecodeError: Expecting ',' delimiter: line 1 column 20 (char 19)",
      'SP-example/业绩: ValueError: 得分单元 u-performance 缺少满分原子条件',
     )

    class _Runner:
     def run(self, stage, *, operation_id=None):
      if stage=='analyze_scores':
       raise _SemanticFailure('评分语义推理失败')
      return object()

    store=ControlStore(context)
    controller=V3ExecutionController(context,runner=_Runner())
    receipt=CommandGateway(context,controller.handlers()).submit(CommandEnvelope.from_mapping({
     'kind':'document.prepare_outline',
     'expected_revision':store.revision(),
     'idempotency_key':'v3-outline-stage-failure',
    },workspace_id='alpha'))

    self.assertEqual(receipt.status,'rejected')
    stages=store.stage_runs(str(receipt.operation_id))
    failed=next(item for item in stages if item['stage_command']=='analyze_scores')
    self.assertEqual(failed['status'],'failed')
    self.assertEqual(failed['error']['code'],'score_semantic_candidate_invalid')
    self.assertEqual(failed['error']['details']['attempts'],2)
    snapshot=V3WorkspaceSnapshotBuilder(context).build()
    pipeline=snapshot['analysis']['pipeline']
    semantic=next(item for item in pipeline['stages'] if item['stage_id']=='score_semantic')
    self.assertEqual(semantic['status'],'failed')
    self.assertEqual(semantic['error']['details']['diagnostics'][1].split('/')[1].split(':')[0],'业绩')

 def test_controller_rejects_unregistered_stage(self):
   with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as t:
    runs=Path(t)/'runs';(runs/'alpha').mkdir(parents=True);context=WorkspaceContext.resolve(runs,'alpha');controller=V3ExecutionController.for_deterministic_tests(context);gateway=CommandGateway(context,controller.handlers())
    receipt=gateway.submit(CommandEnvelope.from_mapping({'kind':'document.run_stage','payload':{'stage':'legacy'},'expected_revision':0,'idempotency_key':'v3-bad-stage'},workspace_id='alpha'))
    self.assertEqual(receipt.status,'rejected')
    self.assertEqual(receipt.error['code'],'COMMAND_DISPATCH_FAILED')

 def test_company_input_does_not_invalidate_score_driven_outline(self):
   with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as t:
    base=Path(t);runs=base/'runs';(runs/'alpha').mkdir(parents=True);context=WorkspaceContext.resolve(runs,'alpha')
    tender=base/'tender.md';company=base/'company.md'
    tender.write_text(
     '# 项目需求\n投标人须提供实施方案。\n\n'
     '# 评标办法\n## 技术评分\n项目实施方案完整性，满分10分。\n',
     encoding='utf-8',
    )
    company.write_text('公司具备项目实施团队与质量管理制度。',encoding='utf-8')
    inputs=InputManifestService(context);inputs.register_local_file(tender,InputRole.TENDER)
    controller=V3ExecutionController.for_deterministic_tests(context);store=ControlStore(context)
    receipt=CommandGateway(context,controller.handlers()).submit(CommandEnvelope.from_mapping({
     'kind':'document.prepare_outline',
     'expected_revision':store.revision(),
     'idempotency_key':'v3-outline-before-company',
    },workspace_id='alpha'))
    self.assertEqual(receipt.status,'accepted',receipt.message)
    before=V3WorkspaceSnapshotBuilder(context).build()
    self.assertEqual(before['analysis']['status'],'current')
    inputs.register_local_file(company,InputRole.COMPANY)

    after=V3WorkspaceSnapshotBuilder(context).build()

    self.assertFalse(after['analysis']['stale'])
    self.assertEqual(after['analysis']['status'],'current')
    self.assertEqual(after['planning']['status'],'needs_human')
    self.assertEqual(len(after['analysis']['score_model']['points']),1)
    self.assertTrue(after['analysis']['chapter_blueprint']['nodes'])

 def test_new_score_input_invalidates_existing_outline(self):
   with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as t:
    base=Path(t);runs=base/'runs';(runs/'alpha').mkdir(parents=True);context=WorkspaceContext.resolve(runs,'alpha')
    tender=base/'tender.md';score=base/'score.md'
    tender.write_text(
     '# 项目需求\n投标人须提供实施方案。\n\n'
     '# 评标办法\n## 技术评分\n项目实施方案完整性，满分10分。\n',
     encoding='utf-8',
    )
    score.write_text('补充评分要求：服务保障，满分5分。',encoding='utf-8')
    inputs=InputManifestService(context);inputs.register_local_file(tender,InputRole.TENDER)
    controller=V3ExecutionController.for_deterministic_tests(context);store=ControlStore(context)
    receipt=CommandGateway(context,controller.handlers()).submit(CommandEnvelope.from_mapping({
     'kind':'document.prepare_outline',
     'expected_revision':store.revision(),
     'idempotency_key':'v3-outline-before-score',
    },workspace_id='alpha'))
    self.assertEqual(receipt.status,'accepted',receipt.message)
    inputs.register_local_file(score,InputRole.SCORE)

    after=V3WorkspaceSnapshotBuilder(context).build()

    self.assertTrue(after['analysis']['stale'])
    self.assertEqual(after['analysis']['status'],'stale')
    self.assertEqual(after['planning']['status'],'outdated')

 def test_v3_api_uses_the_v3_command_controller_and_snapshot(self):
  with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as t:
   base=Path(t);runs=base/'runs';(runs/'alpha').mkdir(parents=True);context=WorkspaceContext.resolve(runs,'alpha');tender=base/'tender.md';score=base/'score.md';tender.write_text('项目目标。\n\n服务范围；交付成果；验收条件；工期30日。\n\n# 评标办法\n项目实施方案完整性，满分10分。',encoding='utf-8');score.write_text('项目实施方案完整性，满分10分。',encoding='utf-8');inputs=InputManifestService(context);inputs.register_local_file(tender,InputRole.TENDER);inputs.register_local_file(score,InputRole.SCORE)
   with (
    mock.patch.object(v3_app,'RUNS_DIR',runs),
    mock.patch.object(
     v3_app,
     'V3ExecutionController',
     side_effect=V3ExecutionController.for_deterministic_tests,
    ),
   ):
    response=asyncio.run(v3_app.command('alpha',_Request({'kind':'document.prepare_outline','expected_revision':ControlStore(context).revision(),'idempotency_key':'v3-api-outline'},{'id':'owner','role':'admin'})))
    self.assertTrue(json.loads(response.body)['ok'])
    snapshot=v3_app.snapshot('alpha')
   payload=json.loads(snapshot.body)
   self.assertIsNone(payload['snapshot']['document']['delivery'])
   self.assertEqual([item['artifact_kind'] for item in payload['snapshot']['promoted_artifacts']], ['ChapterBlueprint', 'InputManifest', 'ProjectModel', 'RequirementLedger', 'ScoreModel', 'SourceIndex'])
   with mock.patch.object(v3_app,'RUNS_DIR',runs):
    self.assertTrue(json.loads(v3_app.latest_gate('alpha').body)['ok'])
    self.assertTrue(json.loads(v3_app.evidence('alpha').body)['ok'])
    self.assertTrue(json.loads(v3_app.events('alpha',0,200).body)['events'])

 def test_v3_upload_registers_an_explicit_input_role(self):
  with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as t:
   runs=Path(t)/'runs';(runs/'alpha').mkdir(parents=True);upload=UploadFile(filename='tender.md',file=io.BytesIO('项目目标'.encode('utf-8')))
   with mock.patch.object(v3_app,'RUNS_DIR',runs):
    response=asyncio.run(v3_app.upload('alpha','tender',upload,''))
   payload=json.loads(response.body)
   self.assertTrue(payload['ok'])
   self.assertEqual(payload['input']['role'],'tender')

 def test_v3_workspace_api_excludes_legacy_workspace_layouts(self):
  with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as t:
   runs=Path(t)/'runs';(runs/'legacy'/'workspace').mkdir(parents=True);principal={'id':'owner','role':'user'}
   with mock.patch.object(v3_app,'RUNS_DIR',runs):
    created=asyncio.run(v3_app.create_workspace(_Request({'name':'V3 项目'},principal)))
    body=json.loads(created.body);self.assertTrue(body['ok'])
    listed=json.loads(v3_app.list_workspaces(_Request({},principal)).body)
   self.assertEqual([item['id'] for item in listed['workspaces']],[body['workspace']['id']])

 def test_http_uploads_tender_and_company_then_returns_score_aware_outline(self):
  with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as t:
   root=Path(t);runs=root/'runs';settings=SettingsService(root)
   environment={
    'BID_AGENT_AUTH_USER':'ui-test',
    'BID_AGENT_AUTH_PASSWORD':'ui-password',
    'BID_AGENT_AUTH_SECURE_COOKIE':'0',
   }
   with (
    mock.patch.dict(os.environ,environment,clear=False),
    mock.patch.object(v3_app,'RUNS_DIR',runs),
    mock.patch.object(v3_app,'SETTINGS',settings),
    mock.patch.object(
     v3_app,
     'V3ExecutionController',
     side_effect=V3ExecutionController.for_deterministic_tests,
    ),
    TestClient(v3_app.app) as client,
   ):
    login=client.post('/api/auth/login',json={'username':'ui-test','password':'ui-password'})
    self.assertEqual(login.status_code,200)
    headers={'X-CSRF-Token':client.cookies.get('bid_agent_csrf')}
    created=client.post('/api/v3/workspaces',json={'name':'页面闭环验收'},headers=headers)
    workspace_id=created.json()['workspace']['id']
    tender=(
     '# 项目需求\n投标人须提供实施方案。\n\n'
     '# 评标办法\n## 技术评分\n项目实施方案完整性，满分10分。\n'
    ).encode('utf-8')
    company='公司具备项目实施团队与质量管理制度。'.encode('utf-8')
    tender_response=client.post(
     f'/api/v3/workspaces/{workspace_id}/uploads',
     data={'role':'tender'},
     files={'file':('tender.md',tender,'text/markdown')},
     headers=headers,
    )
    company_response=client.post(
     f'/api/v3/workspaces/{workspace_id}/uploads',
     data={'role':'company'},
     files={'file':('company.md',company,'text/markdown')},
     headers=headers,
    )
    self.assertEqual(tender_response.status_code,201)
    self.assertEqual(company_response.status_code,201)
    before=client.get(f'/api/v3/workspaces/{workspace_id}/snapshot').json()['snapshot']
    command=client.post(
     f'/api/v3/workspaces/{workspace_id}/commands',
     json={
      'command_id':'ui-outline-command',
      'kind':'document.prepare_outline',
      'payload':{},
      'expected_revision':before['workspace_revision'],
      'idempotency_key':'ui-outline-command',
     },
     headers=headers,
    )
    self.assertEqual(command.status_code,202)
    self.assertTrue(command.json()['ok'],command.json())
    after=client.get(f'/api/v3/workspaces/{workspace_id}/snapshot').json()['snapshot']
    self.assertEqual(len(after['inputs']['inputs']),2)
    self.assertEqual(
     [
      item['input_id']
      for item in after['inputs']['inputs']
      if item['active'] and item['role']=='company'
     ],
     [company_response.json()['input']['input_id']],
    )
    self.assertEqual(len(after['analysis']['score_model']['points']),1)
    self.assertTrue(after['analysis']['chapter_blueprint']['nodes'])
    self.assertEqual(after['planning']['status'],'needs_human')
   with v3_app._SESSION_LOCK:
    v3_app._SESSIONS.clear()

if __name__=='__main__':unittest.main()
