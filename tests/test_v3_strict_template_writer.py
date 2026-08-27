from __future__ import annotations

import sys
import tempfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from control_plane import ControlStore, WorkspaceContext  # noqa: E402
from document_pipeline.artifact_promotion import HumanGateService  # noqa: E402
from document_pipeline.contracts import (  # noqa: E402
    ContractNode,
    DocumentMode,
    InputRole,
    RequirementLedger,
    ScoreCondition,
    ScoreGroup,
    ScoreModel,
    ScorePoint,
    ScoreResponseUnit,
    SourceAnchor,
    TemplateContract,
    TemplateSlot,
    TemplateStructureContract,
)
from document_pipeline.input_manifest import InputManifestService, V3_ROOT  # noqa: E402
from document_pipeline.planning_agent import PlanningAgent  # noqa: E402
from document_pipeline.planning_inference import (  # noqa: E402
    ChapterOutlineCandidate,
    ChapterOutlineNodeCandidate,
)
from document_pipeline.scoring_outline_policy import (  # noqa: E402
    audit_chapter_blueprint,
)
from document_pipeline.chapter_writing_service import (  # noqa: E402
    ChapterWritingRequest,
    ChapterWritingService,
)
from document_pipeline.stage_runner import V3StageRunner  # noqa: E402


def _strict_planning_fixture():
    structure = TemplateStructureContract(
        revision=1,
        source_hashes={"template": "template-hash"},
        template_input_id="template",
        template_hash="template-hash",
        structural_fingerprint="template-fingerprint",
        nodes=[
            ContractNode(
                node_id="p-1",
                order=0,
                level=1,
                numbering="一、",
                writable_target="paragraph:1",
                title="一、项目整体响应",
            ),
            ContractNode(
                node_id="p-3",
                parent_node_id="p-1",
                order=1,
                level=2,
                numbering="1.1",
                writable_target="paragraph:3",
                title="1.1 实施计划",
            ),
        ],
        slots=[
            TemplateSlot(
                slot_id="text-p-2-1",
                node_id="p-1",
                kind="text_slot",
                anchor="paragraph:2:placeholder:{{总体响应}}",
            ),
            TemplateSlot(
                slot_id="text-p-4-1",
                node_id="p-3",
                kind="text_slot",
                anchor="paragraph:4:placeholder:{{实施计划}}",
            ),
        ],
    )
    ledger = RequirementLedger(
        revision=1,
        source_hashes={},
        requirements=[],
    )
    score_anchor = SourceAnchor(
        source_input_id="score-source",
        chunk_id="score-chunk",
        location="table[0]/row[1]",
    )
    condition = ScoreCondition(
        condition_id="SP-strict-C01",
        text="完整响应项目要求",
        normalized_condition="完整响应项目要求",
        condition_role="content",
        source_excerpt="完整响应项目要求",
        subject="项目响应",
        response_intent="完整编写项目整体响应",
        source_anchor=score_anchor,
    )
    scores = ScoreModel(
        revision=1,
        source_hashes={"score-source": "score-hash"},
        model_id="SM-strict",
        source_input_ids=["score-source"],
        total_points=1,
        groups=[ScoreGroup(group_id="technical", title="技术部分")],
        points=[
            ScorePoint(
                score_point_id="SP-strict",
                group_id="technical",
                title="项目整体响应",
                criterion="完整响应项目要求",
                max_points=1,
                score_conditions=[condition],
                response_units=[
                    ScoreResponseUnit(
                        unit_id="SP-strict-U01",
                        title="项目整体响应",
                        condition_ids=[condition.condition_id],
                        response_scope="section",
                        response_expectation="完整响应项目要求",
                    )
                ],
                response_expectation="完整响应项目要求",
                source_anchors=[score_anchor],
                confidence=1,
            )
        ],
    )
    candidate = ChapterOutlineCandidate(
        nodes=[
            ChapterOutlineNodeCandidate(
                local_id="root",
                order=0,
                title="一、项目整体响应",
                purpose="承载项目整体响应",
                primary_response_unit_ids=["SP-strict-U01"],
                score_condition_ids=[condition.condition_id],
                template_slot_ids=["text-p-2-1"],
                confidence=1,
            ),
            ChapterOutlineNodeCandidate(
                local_id="plan",
                parent_local_id="root",
                order=1,
                title="1.1 实施计划",
                purpose="保持模板中的实施计划结构",
                template_slot_ids=["text-p-4-1"],
                confidence=1,
            ),
        ]
    )
    blueprint = object.__new__(PlanningAgent).compile_outline_candidate(
        candidate,
        ledger,
        scores,
        revision=1,
        template_structure=structure,
    )
    return blueprint, ledger, scores, structure


def test_g2_independently_checks_every_frozen_template_mapping() -> None:
    blueprint, ledger, scores, structure = _strict_planning_fixture()

    assert audit_chapter_blueprint(
        blueprint,
        ledger,
        scores,
        structure,
    )["passed"]
    assert [
        (
            node.template_node_id,
            node.template_level,
            node.template_numbering,
            node.template_slot_ids,
        )
        for node in blueprint.nodes
    ] == [
        ("p-1", 1, "一、", ["text-p-2-1"]),
        ("p-3", 2, "1.1", ["text-p-4-1"]),
    ]

    mutations = {
        "title": (0, {"title": "未授权标题"}),
        "level": (1, {"template_level": 3}),
        "order": (1, {"order": 9}),
        "parent": (1, {"parent_chapter_id": None}),
        "numbering": (1, {"template_numbering": "9.9"}),
        "slots": (1, {"template_slot_ids": []}),
    }
    for expected_field, (node_index, update) in mutations.items():
        changed_nodes = list(blueprint.nodes)
        changed_nodes[node_index] = changed_nodes[node_index].model_copy(
            update=update
        )
        changed_blueprint = blueprint.model_copy(
            update={"nodes": changed_nodes}
        )

        audit = audit_chapter_blueprint(
            changed_blueprint,
            ledger,
            scores,
            structure,
        )

        assert not audit["passed"], expected_field
        template_messages = [
            item["message"]
            for item in audit["findings"]
            if item["code"] == "TEMPLATE_STRUCTURE_CHANGED"
        ]
        assert any(
            expected_field in message for message in template_messages
        ), (expected_field, template_messages)


def _assert_strict_template_g2_h1_document_contract_writer_path(
    tmp_path: Path,
) -> None:
    runs = tmp_path / "runs"
    (runs / "alpha").mkdir(parents=True)
    context = WorkspaceContext.resolve(runs, "alpha")

    tender = tmp_path / "tender.md"
    tender.write_text(
        "项目目标。\n\n服务范围；交付成果；验收条件；工期30日。",
        encoding="utf-8",
    )
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_heading("一、项目整体响应", level=1)
    template.add_paragraph("{{总体响应}}")
    template.save(template_path)

    inputs = InputManifestService(context)
    inputs.register_local_file(tender, InputRole.TENDER)
    inputs.register_local_file(template_path, InputRole.TEMPLATE)
    runner = V3StageRunner.for_deterministic_tests(context)
    for stage in (
        "normalize_sources",
        "compile_template_structure",
        "build_requirement_ledger",
        "analyze_scores",
        "plan_response",
        "compile_chapter_blueprint",
    ):
        runner.run(stage)

    store = ControlStore(context)
    blueprint_artifact = store.v3_active_artifact("ChapterBlueprint")
    assert blueprint_artifact is not None
    g2 = store.latest_v3_gate_receipt(
        str(blueprint_artifact["proposal_id"]),
        "G2_BLUEPRINT_INTEGRITY",
    )
    assert g2 is not None
    assert g2["verdict"] == "pass"
    assert blueprint_artifact["payload"]["mode"] == DocumentMode.TEMPLATE_STRICT

    store.grant_workspace_access("owner")
    human_gate = HumanGateService(context)
    h1 = human_gate.confirm_planning(
        principal_id="owner",
        submitted_snapshot=human_gate.planning_snapshot(),
        nonce="strict-template-writer-h1",
    )
    assert human_gate.require_current_confirmation().receipt_id == h1.receipt_id

    contract = runner.run("compile_document_contract")
    assert isinstance(contract, TemplateContract)
    assert contract.source_blueprint_hash == blueprint_artifact["artifact_hash"]
    assert contract.blocking_gaps == []
    slot_ids = {slot.slot_id for slot in contract.slots}
    assert slot_ids
    assert {
        node.writable_target for node in contract.nodes
    }.issubset(slot_ids)
    blueprint_nodes = {
        node["chapter_id"]: node
        for node in blueprint_artifact["payload"]["nodes"]
    }
    assert [
        (
            node.title,
            node.level,
            node.order,
            node.parent_node_id,
            node.numbering,
        )
        for node in contract.nodes
    ] == [
        (
            blueprint_nodes[node.node_id]["title"],
            blueprint_nodes[node.node_id]["template_level"],
            blueprint_nodes[node.node_id]["order"],
            blueprint_nodes[node.node_id]["parent_chapter_id"],
            blueprint_nodes[node.node_id]["template_numbering"],
        )
        for node in contract.nodes
    ]

    plan, units = runner.run("plan_document")
    assert plan.source_blueprint_hash == blueprint_artifact["artifact_hash"]
    assert {node.node_id for node in plan.nodes} == {
        node.node_id for node in contract.nodes
    }
    assert units

    document_contract_path = (
        context.root / V3_ROOT / "contracts" / "document_contract.json"
    )
    document_contract_path.unlink()
    assert not document_contract_path.exists()

    writing = ChapterWritingService(context, deterministic_test=True)
    blocks = [
        block
        for unit in units
        for block in writing.write(
            ChapterWritingRequest(
                unit_id=unit.unit_id,
                node_ids=tuple(unit.node_ids),
                run_research=False,
                commit_drafts=False,
            )
        ).blocks
    ]
    assert blocks
    assert {block.target_node_id for block in blocks}.issubset(slot_ids)


def test_strict_template_g2_h1_document_contract_writer_path() -> None:
    with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
        _assert_strict_template_g2_h1_document_contract_writer_path(
            Path(tmp)
        )
