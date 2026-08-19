"""Draft Word export must keep the on-page preview styles."""

from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from document_pipeline.renderers.word_styles import (  # noqa: E402
    BODY_EAST_ASIA_FONT,
    FIRST_LINE_INDENT_PT,
    HEADING_EAST_ASIA_FONT,
    PAGE_WIDTH_CM,
    write_composed_document,
)


def _east_asia_font(paragraph) -> str:
    run = paragraph.runs[0]
    rfonts = run._element.rPr.find(qn("w:rFonts"))
    return str(rfonts.get(qn("w:eastAsia")) or "")


def _style_east_asia(document: Document, style_name: str) -> str:
    rpr = document.styles[style_name].element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    return str(rfonts.get(qn("w:eastAsia")) or "")


class CurrentWordExportStyleTests(unittest.TestCase):
    def test_export_applies_preview_fonts_indent_tables_and_lists(self) -> None:
        composed = {
            "mode": "current_draft",
            "chapters": [
                {
                    "chapter_id": "ch-a",
                    "title": "技术方案",
                    "depth": 0,
                    "blocks": [
                        {
                            "type": "paragraph",
                            "content": "本方案采用分层架构，覆盖招标文件全部功能要求。",
                        },
                        {
                            "type": "heading",
                            "content": "## 总体架构",
                        },
                        {
                            "type": "list",
                            "content": "- 应用层负责业务编排\n- 数据层负责持久化",
                        },
                        {
                            "type": "table",
                            "content": "| 模块 | 职责 |\n| --- | --- |\n| 门户 | 统一入口 |",
                        },
                    ],
                },
                {
                    "chapter_id": "ch-b",
                    "title": "实施计划",
                    "depth": 1,
                    "blocks": [],
                },
            ],
        }
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "current.docx"
            write_composed_document(composed, output)
            document = Document(str(output))

            texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            self.assertEqual(texts[0], "技术方案")
            self.assertIn("本方案采用分层架构，覆盖招标文件全部功能要求。", texts)
            self.assertIn("总体架构", texts)
            self.assertIn("应用层负责业务编排", texts)
            self.assertIn("实施计划", texts)

            heading = document.paragraphs[0]
            self.assertEqual(heading.style.name, "Heading 1")
            self.assertEqual(heading.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(_east_asia_font(heading), HEADING_EAST_ASIA_FONT)
            self.assertEqual(int(heading.paragraph_format.first_line_indent.pt or 0), 0)

            body = next(
                paragraph
                for paragraph in document.paragraphs
                if "分层架构" in paragraph.text
            )
            self.assertEqual(body.style.name, "Normal")
            self.assertEqual(body.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assertEqual(_east_asia_font(body), BODY_EAST_ASIA_FONT)
            self.assertAlmostEqual(
                float(body.paragraph_format.first_line_indent.pt),
                float(FIRST_LINE_INDENT_PT),
                places=1,
            )

            subheading = next(
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.strip() == "总体架构"
            )
            self.assertEqual(subheading.style.name, "Heading 2")
            self.assertEqual(_east_asia_font(subheading), HEADING_EAST_ASIA_FONT)

            bullet = next(
                paragraph
                for paragraph in document.paragraphs
                if "应用层负责业务编排" in paragraph.text
            )
            self.assertEqual(bullet.style.name, "List Bullet")
            self.assertEqual(int(bullet.paragraph_format.first_line_indent.pt or 0), 0)

            self.assertEqual(len(document.tables), 1)
            table = document.tables[0]
            self.assertEqual(table.cell(0, 0).text.strip(), "模块")
            self.assertEqual(table.cell(1, 1).text.strip(), "统一入口")
            header_shd = table.cell(0, 0)._tc.tcPr.find(qn("w:shd"))
            self.assertIsNotNone(header_shd)
            self.assertEqual(header_shd.get(qn("w:fill")), "F8FAFC")
            borders = table.cell(0, 0)._tc.tcPr.find(qn("w:tcBorders"))
            self.assertIsNotNone(borders)
            self.assertEqual(
                _east_asia_font(table.cell(0, 0).paragraphs[0]),
                "微软雅黑",
            )

            section = document.sections[0]
            self.assertAlmostEqual(section.page_width.cm, PAGE_WIDTH_CM, places=1)
            self.assertEqual(_style_east_asia(document, "Normal"), BODY_EAST_ASIA_FONT)
            self.assertEqual(_style_east_asia(document, "Heading 1"), HEADING_EAST_ASIA_FONT)

    def test_markdown_inside_paragraph_block_becomes_structured_word(self) -> None:
        composed = {
            "chapters": [
                {
                    "chapter_id": "ch-md",
                    "title": "服务方案",
                    "depth": 0,
                    "blocks": [
                        {
                            "type": "paragraph",
                            "content": (
                                "### 响应时间\n\n"
                                "7×24 小时值守。\n\n"
                                "1. 一级故障 30 分钟响应\n"
                                "2. 二级故障 2 小时响应"
                            ),
                        }
                    ],
                }
            ]
        }
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "current.docx"
            write_composed_document(composed, output)
            document = Document(str(output))
            styles = {paragraph.text.strip(): paragraph.style.name for paragraph in document.paragraphs if paragraph.text.strip()}
            self.assertEqual(styles["响应时间"], "Heading 3")
            self.assertEqual(styles["7×24 小时值守。"], "Normal")
            self.assertEqual(styles["一级故障 30 分钟响应"], "List Number")


if __name__ == "__main__":
    unittest.main()
