from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from control_plane import WorkspaceContext  # noqa: E402
from document_pipeline.renderers.template_renderer import StrictTemplateRenderer  # noqa: E402


class V3TemplateRendererTests(unittest.TestCase):
    def test_renderer_copies_template_and_keeps_fixed_content(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp); runs = base / "runs"; workspace = runs / "alpha"; source_dir = workspace / "workspace" / "v3" / "sources" / "T1"; source_dir.mkdir(parents=True)
            doc = Document(); doc.add_heading("固定标题", 1); doc.add_paragraph("固定说明"); doc.add_paragraph("名称：{{name}}"); doc.save(source_dir / "template.docx")
            from document_pipeline.template_contract import TemplateContractCompiler
            context = WorkspaceContext.resolve(runs, "alpha")
            # Build a contract from the same immutable template source.
            (workspace / "workspace" / "v3" / "input_manifest.json").write_text(json.dumps({"schema_version":"v3","revision":1,"source_hashes":{"T1":"x"},"inputs":[{"input_id":"T1","role":"template","filename":"template.docx","mime_type":"application/docx","sha256":"x","version":1,"active":True,"replaces_input_id":None}]}), encoding="utf-8")
            # Minimal contract uses the compiler's structural fingerprint and known slot.
            fingerprint = TemplateContractCompiler._fingerprint(source_dir / "template.docx")
            contract = {"schema_version":"v3","revision":1,"source_hashes":{},"mode":"template_strict","template_hash":"x","structural_fingerprint":fingerprint,"nodes":[{"node_id":"p-1","parent_node_id":None,"order":0,"writable_target":"paragraph:1","title":"固定标题","requirement_ids":[]}],"slots":[{"slot_id":"text-p-3-1","node_id":"p-1","kind":"text_slot","anchor":"paragraph:3:placeholder:{{name}}"}],"warnings":[],"blocking_gaps":[]}
            (workspace / "workspace" / "v3" / "contracts").mkdir(parents=True); (workspace / "workspace" / "v3" / "contracts" / "document_contract.json").write_text(json.dumps(contract), encoding="utf-8")
            integrated = {"schema_version":"v3","revision":1,"source_hashes":{},"contract_revision":1,"plan_revision":1,"blocks":[{"block_id":"b1","target_node_id":"text-p-3-1","type":"paragraph","content":"测试项目","topic_ids":[],"requirement_ids":[],"score_point_ids":[],"evidence_ids":[],"fact_ids":[],"confidence":1,"human_locked":False,"critical_claims":[]}]}
            (workspace / "workspace" / "v3" / "integrated_document.json").write_text(json.dumps(integrated), encoding="utf-8")
            output = StrictTemplateRenderer(context).render()
            result = Document(str(output))
            self.assertEqual(result.paragraphs[0].text, "固定标题")
            self.assertEqual(result.paragraphs[1].text, "固定说明")
            self.assertIn("测试项目", result.paragraphs[2].text)


if __name__ == "__main__":
    unittest.main()
