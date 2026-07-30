from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from control_plane import WorkspaceContext  # noqa: E402
from document_pipeline.contracts import InputRole  # noqa: E402
from document_pipeline.input_manifest import InputManifestService  # noqa: E402
from document_pipeline.source_normalizer import SourceNormalizer  # noqa: E402
from document_pipeline.template_contract import TemplateContractCompiler  # noqa: E402


class V3SourceStructureTests(unittest.TestCase):
    def _context(self, base: Path) -> WorkspaceContext:
        runs = base / "runs"
        (runs / "alpha").mkdir(parents=True)
        return WorkspaceContext.resolve(runs, "alpha")

    def test_docx_source_index_preserves_heading_paragraph_and_table_order(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            path = base / "tender.docx"
            document = Document()
            document.add_heading("第一章 项目概述", level=1)
            document.add_paragraph("项目范围。")
            document.add_table(rows=1, cols=2).cell(0, 0).text = "交付成果"
            document.tables[0].cell(0, 1).text = "实施报告"
            document.add_heading("第二章 服务要求", level=1)
            document.save(path)
            context = self._context(base)
            InputManifestService(context).register_local_file(path, InputRole.TENDER)
            index = SourceNormalizer(context).normalize_active_inputs()
            blocks = index["blocks"]
            self.assertEqual(
                [block["block_kind"] for block in blocks],
                ["heading", "paragraph", "table", "table_cell", "table_cell", "heading"],
            )
            self.assertEqual(blocks[1]["heading_path"], ["第一章 项目概述"])
            self.assertEqual(blocks[3]["source_anchor"]["location"], "table:1:row:1:cell:1")
            self.assertTrue(all(block["content_hash"] for block in blocks))
            self.assertEqual(index.get("authority"), "promoted_artifact_projection")
            self.assertIn("coverage", index)
            # Mutating disk projection must not change promoted authority.
            from control_plane import ControlStore
            from document_pipeline.contracts import SourceIndex

            promoted = ControlStore(context).v3_active_artifact("SourceIndex")
            self.assertIsNotNone(promoted)
            index["blocks"] = []
            still = SourceIndex.model_validate(promoted["payload"])
            self.assertGreaterEqual(len(still.blocks), 5)

    def test_docx_merged_cell_aliases_are_deduplicated_per_logical_row(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            path = base / "scoring-table.docx"
            document = Document()
            table = document.add_table(rows=3, cols=4)
            table.cell(0, 0).merge(table.cell(0, 2)).text = "评分内容"
            table.cell(0, 3).text = "分值"
            table.cell(1, 0).merge(table.cell(2, 0)).text = "技术部分"
            table.cell(1, 1).text = "实施方案"
            table.cell(1, 3).text = "10分"
            table.cell(2, 1).text = "服务保障"
            table.cell(2, 3).text = "5分"
            document.save(path)

            context = self._context(base)
            InputManifestService(context).register_local_file(path, InputRole.TENDER)
            index = SourceNormalizer(context).normalize_active_inputs()

            cell_blocks = [
                block for block in index["blocks"] if block["block_kind"] == "table_cell"
            ]
            header_blocks = [block for block in cell_blocks if block["content"] == "评分内容"]
            self.assertEqual(len(header_blocks), 1)
            self.assertEqual(header_blocks[0]["column_index"], 0)

            # A vertical merge is intentionally represented once in every
            # logical row so row-aware score extraction keeps its category.
            vertical_context = [
                block for block in cell_blocks if block["content"] == "技术部分"
            ]
            self.assertEqual(
                [(block["row_index"], block["column_index"]) for block in vertical_context],
                [(1, 0), (2, 0)],
            )

            coverage = {
                item["locator"]: item for item in index["coverage"]["items"]
            }
            canonical_id = header_blocks[0]["block_id"]
            for alias_locator in (
                "table:1:row:1:cell:2",
                "table:1:row:1:cell:3",
            ):
                alias = coverage[alias_locator]
                self.assertEqual(alias["status"], "exempt")
                self.assertEqual(alias["block_id"], canonical_id)
                self.assertEqual(
                    alias["reason"],
                    "merged-cell alias of table:1:row:1:cell:1",
                )

    def test_amendment_keeps_issued_at_and_explicit_supersession(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            tender = base / "tender.md"
            amendment = base / "amendment.md"
            tender.write_text("原工期为30日。", encoding="utf-8")
            amendment.write_text("更正工期为20日。", encoding="utf-8")
            context = self._context(base)
            inputs = InputManifestService(context)
            original = inputs.register_local_file(tender, InputRole.TENDER)
            inputs.register_local_file(
                amendment,
                InputRole.AMENDMENT,
                issued_at="2026-07-27T10:00:00+08:00",
                supersedes_input_ids=[original.item.input_id],
            )
            index = SourceNormalizer(context).normalize_active_inputs()
            self.assertEqual(index["amendments"][0]["supersedes_input_ids"], [original.item.input_id])
            self.assertEqual(index["amendments"][0]["issued_at"], "2026-07-27T10:00:00+08:00")

    def test_template_structure_compiles_without_requirement_ledger(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            path = base / "template.docx"
            document = Document()
            document.add_heading("技术方案", level=1)
            document.add_paragraph("项目名称：{{项目名称}}")
            document.add_heading("实施计划", level=1)
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "{{实施内容}}"
            document.save(path)
            context = self._context(base)
            item = InputManifestService(context).register_local_file(path, InputRole.TEMPLATE).item
            structure = TemplateContractCompiler(context).compile_structure(item)
            self.assertEqual([node.title for node in structure.nodes], ["技术方案", "实施计划"])
            self.assertEqual(structure.nodes[0].level, 1)
            self.assertEqual(structure.slots[0].kind, "text_slot")
            self.assertEqual(structure.template_input_id, item.input_id)
            cell_slots = [slot for slot in structure.slots if slot.kind == "cell_slot"]
            self.assertEqual(len(cell_slots), 1)
            # Table under second heading must bind nearest upstream chapter, not the first heading.
            self.assertEqual(cell_slots[0].node_id, structure.nodes[1].node_id)
            from control_plane import ControlStore

            self.assertIsNotNone(ControlStore(context).v3_active_artifact("TemplateStructureContract"))

    def test_unreadable_company_input_is_reported_without_blocking_tender_sources(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            tender = base / "tender.md"
            company = base / "company.pdf"
            tender.write_text("# 评标办法\n技术方案满分10分。", encoding="utf-8")
            company.write_bytes(b"%PDF-scanned-placeholder")
            context = self._context(base)
            inputs = InputManifestService(context)
            inputs.register_local_file(tender, InputRole.TENDER)
            company_item = inputs.register_local_file(company, InputRole.COMPANY).item
            normalizer = SourceNormalizer(context)
            original = normalizer._blocks_for

            def parse_or_simulate_scan(item, source):
                if item.input_id == company_item.input_id:
                    raise ValueError("V3_SOURCE_OCR_BLOCKED: PDF 全部页面均无可提取文本（疑似扫描件）。")
                return original(item, source)

            with mock.patch.object(normalizer, "_blocks_for", side_effect=parse_or_simulate_scan):
                index = normalizer.normalize_active_inputs()

            status = {item["input_id"]: item for item in index["input_status"]}
            self.assertEqual(status[company_item.input_id]["status"], "blocked")
            self.assertTrue(any(block["input_role"] == "tender" for block in index["blocks"]))
            self.assertFalse(any(block["input_role"] == "company" for block in index["blocks"]))


if __name__ == "__main__":
    unittest.main()
