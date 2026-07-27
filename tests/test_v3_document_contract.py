from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from control_plane import WorkspaceContext  # noqa: E402
from document_pipeline.contracts import DocumentMode, InputRole, OutlineContract, TemplateContract  # noqa: E402
from document_pipeline.document_contract import DocumentContractCompiler  # noqa: E402
from document_pipeline.input_manifest import InputManifestService  # noqa: E402
from document_pipeline.project_model import ProjectModelBuilder  # noqa: E402
from document_pipeline.requirement_ledger import RequirementLedgerBuilder  # noqa: E402
from document_pipeline.source_normalizer import SourceNormalizer  # noqa: E402


class V3DocumentContractTests(unittest.TestCase):
    def _context(self, base: Path) -> WorkspaceContext:
        runs = base / "runs"
        (runs / "alpha").mkdir(parents=True)
        return WorkspaceContext.resolve(runs, "alpha")

    def _prepare_model(self, base: Path, *, template: bool) -> WorkspaceContext:
        tender = base / "tender.md"
        score = base / "score.md"
        tender.write_text("项目目标与实施要求。\n\n交付成果为实施报告；验收条件为采购人验收；工期为30日。", encoding="utf-8")
        score.write_text("评分要求：实施方案。", encoding="utf-8")
        context = self._context(base)
        inputs = InputManifestService(context)
        inputs.register_local_file(tender, InputRole.TENDER)
        inputs.register_local_file(score, InputRole.SCORE)
        if template:
            docx = base / "template.docx"
            document = Document()
            document.add_heading("项目目标与实施要求", level=1)
            document.add_paragraph("项目名称：{{项目名称}}")
            document.add_heading("实施方案", level=2)
            document.add_table(rows=1, cols=1).cell(0, 0).text = "【服务内容】"
            document.save(docx)
            inputs.register_local_file(docx, InputRole.TEMPLATE)
        SourceNormalizer(context).normalize_active_inputs()
        RequirementLedgerBuilder(context).build()
        ProjectModelBuilder(context).build()
        return context

    def test_template_mode_preserves_template_titles_and_discovers_slots(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            contract = DocumentContractCompiler(self._prepare_model(Path(tmp), template=True)).compile()
            self.assertIsInstance(contract, TemplateContract)
            self.assertEqual(contract.mode, DocumentMode.TEMPLATE_STRICT)
            self.assertEqual([node.title for node in contract.nodes], ["项目目标与实施要求", "实施方案"])
            self.assertEqual(contract.nodes[1].parent_node_id, contract.nodes[0].node_id)
            self.assertEqual({slot.kind for slot in contract.slots}, {"text_slot", "cell_slot"})
            self.assertTrue(any(gap.startswith("TEMPLATE_COVERAGE_GAP:") for gap in contract.blocking_gaps))

    def test_auto_outline_has_source_requirement_for_every_title(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            contract = DocumentContractCompiler(self._prepare_model(Path(tmp), template=False)).compile()
            self.assertIsInstance(contract, OutlineContract)
            self.assertEqual(contract.mode, DocumentMode.AUTO_OUTLINE)
            self.assertTrue(contract.nodes)
            self.assertTrue(all(node.requirement_ids for node in contract.nodes))
            self.assertFalse(any(node.title == "通用投标目录" for node in contract.nodes))


if __name__ == "__main__":
    unittest.main()
