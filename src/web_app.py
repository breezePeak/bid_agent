from __future__ import annotations

import asyncio
import hashlib
import hmac
import json
import os
import queue
import re
import shutil
import secrets
import subprocess
import sys
import threading
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from fastapi import FastAPI, File, Query, Request, UploadFile
from fastapi.responses import (
    FileResponse,
    HTMLResponse,
    JSONResponse,
    PlainTextResponse,
    StreamingResponse,
)
from fastapi.staticfiles import StaticFiles
from starlette.concurrency import run_in_threadpool

ROOT = Path(__file__).resolve().parent.parent
# Web 进程与流水线子进程统一：配置以项目根 .env / models.json 为准
os.environ.setdefault("BID_AGENT_CONFIG_ROOT", str(ROOT))
WEB_DIR = ROOT / "web"
VUE_DIST_DIR = ROOT / "frontend" / "dist"
RUNS_DIR = ROOT / "runs"
ACTIVE_RUN_FILE = RUNS_DIR / ".active_run"

sys.path.insert(0, str(ROOT / "src"))

from chat_store import clear_messages, close_chat_store, load_messages, save_message
from control_plane import (
    CommandEnvelope,
    CommandGateway,
    ControlPlaneError,
    ControlStore,
    WorkspaceContext,
)
from agent.repair_jobs import (
    RUNNING_REPAIR_STATUSES,
    TERMINAL_REPAIR_STATUSES,
    claim_repair_job,
    claim_repair_job_authorized,
    create_confirmation,
    create_authorized_repair_job,
    decline_repair_job,
    load_repair_job,
    load_v2_repair_job,
    reconcile_interrupted_repair,
    update_repair_job,
)
from session_orchestrator import plan as orchestrator_plan, resolve_execution as orchestrator_resolve
from graph.state_recorder import load_run_events, load_stage_metrics, save_run_state
from manual_review import apply_manual_review_update, manual_review_items, manual_review_summary
from pipeline_registry import (
    RunArtifact,
    auto_run_commands,
    artifact_exists,
    stage_command_map,
    stage_outputs_ready,
    stage_spec_by_command,
    workflow_stage_specs,
)
from pipeline_supervisor import PipelineSupervisor
from stage_validation import COLLECTION_STAGE_IDS, stage_collection_status
from project_profile_registry import load_project_profile, project_profile_choices, save_project_profile
from prompt_registry import AGENT_SPECS
from utils import read_json

app = FastAPI(title="标书 Agent 控制台", docs_url=None, redoc_url=None)

_AUTH_COOKIE = "bid_agent_session"
_CSRF_COOKIE = "bid_agent_csrf"
_AUTH_SESSION_SECONDS = 12 * 60 * 60
_AUTH_SESSIONS: dict[str, dict[str, Any]] = {}
_AUTH_LOCK = threading.Lock()


def _is_v1_compat_api(path: str) -> bool:
    if not path.startswith("/api/"):
        return False
    return not (
        path.startswith("/api/v2/")
        or path.startswith("/api/auth/")
        or path.startswith("/api/llm-settings")
    )


def _mark_v1_compat_response(path: str, response: Any) -> Any:
    if _is_v1_compat_api(path):
        response.headers["Deprecation"] = "true"
        response.headers["Warning"] = '299 bid-agent "V1 compatibility API; migrate to workspace-scoped V2"'
        response.headers["Link"] = '</api/v2/workspaces>; rel="successor-version"'
    return response


def _auth_credentials() -> tuple[str, str]:
    return (
        str(os.environ.get("BID_AGENT_AUTH_USER") or "admin"),
        str(os.environ.get("BID_AGENT_AUTH_PASSWORD") or ""),
    )


def _session_record(token: str) -> dict[str, Any] | None:
    value = str(token or "").strip()
    if not value:
        return None
    now = time.time()
    with _AUTH_LOCK:
        session = _AUTH_SESSIONS.get(value)
        if not session:
            return None
        if float(session.get("expires_at") or 0) <= now:
            _AUTH_SESSIONS.pop(value, None)
            return None
        return dict(session)


def _session_principal(token: str) -> dict[str, Any] | None:
    session = _session_record(token)
    return dict(session.get("principal") or {}) if session else None


def _ensure_workspace_acl(context: WorkspaceContext, principal: dict[str, Any], *, write: bool) -> None:
    principal_id = str(principal.get("id") or "").strip()
    store = ControlStore(context)
    acl = store.workspace_acl()
    if not acl:
        if str(principal.get("role") or "") != "admin":
            raise ControlPlaneError("WORKSPACE_FORBIDDEN", "工作区尚未分配所有者。", status_code=403)
        store.grant_workspace_access(principal_id, role="owner")
    store.require_workspace_access(principal_id, write=write)


@app.middleware("http")
async def api_auth_and_workspace_acl(request: Request, call_next):
    path = request.url.path
    if not path.startswith("/api/") or path == "/api/auth/login":
        return await call_next(request)
    session = _session_record(request.cookies.get(_AUTH_COOKIE, ""))
    principal = dict(session.get("principal") or {}) if session else None
    if not principal:
        return JSONResponse(
            {"ok": False, "error": {"code": "AUTH_REQUIRED", "message": "请先登录。"}, "message": "请先登录。"},
            status_code=401,
        )
    if request.method.upper() not in {"GET", "HEAD", "OPTIONS"}:
        expected_csrf = str(session.get("csrf_token") or "")
        header_csrf = str(request.headers.get("x-csrf-token") or "")
        cookie_csrf = str(request.cookies.get(_CSRF_COOKIE, "") or "")
        if not (
            expected_csrf
            and hmac.compare_digest(header_csrf, expected_csrf)
            and hmac.compare_digest(cookie_csrf, expected_csrf)
        ):
            return JSONResponse(
                {
                    "ok": False,
                    "error": {"code": "CSRF_REQUIRED", "message": "请求缺少有效的 CSRF 令牌。"},
                    "message": "请求缺少有效的 CSRF 令牌。",
                },
                status_code=403,
            )
    request.state.principal = principal
    try:
        workspace_id = ""
        context: WorkspaceContext | None = None
        prefix = "/api/v2/workspaces/"
        if path.startswith(prefix):
            workspace_id = path[len(prefix):].split("/", 1)[0]
        elif path not in {"/api/runs", "/api/start-run", "/api/select-run"}:
            # V1 read adapters may carry an explicit workspace query. Bind ACL
            # to that exact workspace before falling back to legacy ACTIVE_RUN.
            workspace_id = str(
                request.query_params.get("workspace_id")
                or request.query_params.get("run_id")
                or ACTIVE_RUN_ID
                or ""
            ).strip()
        if workspace_id:
            context = _workspace_context(workspace_id)
            _ensure_workspace_acl(
                context,
                principal,
                write=request.method.upper() not in {"GET", "HEAD", "OPTIONS"},
            )
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    response = await call_next(request)
    if _is_v1_compat_api(path) and context is not None:
        try:
            ControlStore(context).record_compatibility_usage(path, principal)
        except Exception:
            # Compatibility telemetry is advisory; it cannot make an otherwise
            # valid legacy request fail or become a second control authority.
            pass
    return _mark_v1_compat_response(path, response)


@app.post("/api/auth/login")
async def api_auth_login(request: Request) -> JSONResponse:
    try:
        body = await request.json()
    except Exception:
        body = {}
    username = str((body or {}).get("username") or "")
    password = str((body or {}).get("password") or "")
    expected_user, expected_password = _auth_credentials()
    if not expected_password:
        return JSONResponse(
            {"ok": False, "message": "服务端尚未配置 BID_AGENT_AUTH_PASSWORD。"},
            status_code=503,
        )
    if not (hmac.compare_digest(username, expected_user) and hmac.compare_digest(password, expected_password)):
        return JSONResponse({"ok": False, "message": "用户名或密码错误。"}, status_code=401)
    token = secrets.token_urlsafe(32)
    csrf_token = secrets.token_urlsafe(32)
    principal = {"type": "user", "id": username[:128], "role": "admin"}
    with _AUTH_LOCK:
        _AUTH_SESSIONS[token] = {
            "principal": principal,
            "csrf_token": csrf_token,
            "expires_at": time.time() + _AUTH_SESSION_SECONDS,
        }
    response = JSONResponse({"ok": True, "principal": principal, "csrf_token": csrf_token})
    secure_cookie = str(os.environ.get("BID_AGENT_AUTH_SECURE_COOKIE") or "0").lower() in {"1", "true", "yes"}
    response.set_cookie(
        _AUTH_COOKIE,
        token,
        max_age=_AUTH_SESSION_SECONDS,
        httponly=True,
        samesite="strict",
        secure=secure_cookie,
        path="/",
    )
    response.set_cookie(
        _CSRF_COOKIE,
        csrf_token,
        max_age=_AUTH_SESSION_SECONDS,
        httponly=False,
        samesite="strict",
        secure=secure_cookie,
        path="/",
    )
    return response


@app.post("/api/auth/logout")
def api_auth_logout(request: Request) -> JSONResponse:
    token = request.cookies.get(_AUTH_COOKIE, "")
    with _AUTH_LOCK:
        _AUTH_SESSIONS.pop(token, None)
    response = JSONResponse({"ok": True})
    response.delete_cookie(_AUTH_COOKIE, path="/")
    response.delete_cookie(_CSRF_COOKIE, path="/")
    return response


@app.get("/api/auth/me")
def api_auth_me(request: Request) -> JSONResponse:
    return JSONResponse({"ok": True, "principal": getattr(request.state, "principal", {})})

LOG_LINES: list[str] = []
LOG_MAX = 2000
_LOG_CONTEXT = threading.local()
_WORKSPACE_LOG_LOCK = threading.Lock()
RUNNING = False
CURRENT_TASK = ""
_REPAIR_START_LOCK = threading.RLock()
_CHAT_TURN_LOCK_GUARD = threading.Lock()
_CHAT_TURN_LOCKS: dict[str, asyncio.Lock] = {}
CURRENT_PROCESS: subprocess.Popen | None = None
CURRENT_RUN_ID = ""
CURRENT_RUN_ROOT: Path | None = None
PAUSE_REQUESTED = False
SUPERVISOR = PipelineSupervisor()
ACTIVE_RUN_ID = ""
ACTIVE_RUN_ROOT: Path | None = None
_PENDING_LINE_EDITS: dict[Path, dict[str, Any]] = {}

def _workflow_step_payload() -> list[dict[str, Any]]:
    steps: list[dict[str, Any]] = []
    for spec in workflow_stage_specs():
        steps.append(
            {
                "id": spec.id,
                "label": spec.label,
                "command": spec.command,
                "kind": spec.kind,
                "requires": [{"path": artifact.path, "kind": artifact.kind} for artifact in spec.requires],
                "produces": [{"path": artifact.path, "kind": artifact.kind} for artifact in spec.produces],
            }
        )
    return steps


WORKFLOW_STEPS: list[dict[str, Any]] = _workflow_step_payload()
WORKFLOW_COMMAND_LABELS = {step["command"]: step["label"] for step in WORKFLOW_STEPS}


def _append_log(line: str) -> None:
    global LOG_LINES
    LOG_LINES.append(line)
    if len(LOG_LINES) > LOG_MAX:
        LOG_LINES = LOG_LINES[-LOG_MAX:]
    run_root = getattr(_LOG_CONTEXT, "run_root", None)
    if isinstance(run_root, Path):
        try:
            path = run_root / "workspace" / "runtime_logs.jsonl"
            path.parent.mkdir(parents=True, exist_ok=True)
            record = json.dumps({"ts": datetime.now().isoformat(timespec="seconds"), "line": line}, ensure_ascii=False)
            with _WORKSPACE_LOG_LOCK, path.open("a", encoding="utf-8") as handle:
                handle.write(record + "\n")
        except Exception:
            pass


def _workspace_log_lines(root: Path, limit: int = LOG_MAX) -> list[str]:
    path = root / "workspace" / "runtime_logs.jsonl"
    if not path.exists():
        return []
    rows: list[str] = []
    try:
        raw_lines = path.read_text(encoding="utf-8", errors="replace").splitlines()[-max(1, limit):]
        for raw in raw_lines:
            try:
                record = json.loads(raw)
                rows.append(str(record.get("line") or "") if isinstance(record, dict) else str(record))
            except json.JSONDecodeError:
                rows.append(raw)
    except OSError:
        return []
    return rows


def _decode_log_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "gbk", "cp936"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _active_root() -> Path:
    _load_active_run_from_disk()
    return ACTIVE_RUN_ROOT or ROOT


def _load_active_run_from_disk() -> None:
    global ACTIVE_RUN_ID, ACTIVE_RUN_ROOT
    if ACTIVE_RUN_ROOT and ACTIVE_RUN_ROOT.exists():
        return
    run_id = ""
    if ACTIVE_RUN_FILE.exists():
        run_id = ACTIVE_RUN_FILE.read_text(encoding="utf-8").strip()
    if run_id:
        run_root = RUNS_DIR / run_id
        if run_root.exists():
            ACTIVE_RUN_ID = run_id
            ACTIVE_RUN_ROOT = run_root
            return
    if RUNS_DIR.exists():
        run_dirs = sorted(
            [path for path in RUNS_DIR.iterdir() if path.is_dir() and not path.name.startswith(".")],
            key=lambda path: path.stat().st_mtime,
        )
        if run_dirs:
            ACTIVE_RUN_ROOT = run_dirs[-1]
            ACTIVE_RUN_ID = ACTIVE_RUN_ROOT.name


def _active_run_payload() -> dict[str, Any]:
    root = _active_root()
    return _workspace_payload(ACTIVE_RUN_ID, root, isolated=ACTIVE_RUN_ROOT is not None)


def _workspace_payload(run_id: str, root: Path, *, isolated: bool = True) -> dict[str, Any]:
    return {
        "id": run_id,
        "root": str(root),
        "relative_root": str(root.relative_to(ROOT)) if root != ROOT and root.is_relative_to(ROOT) else str(root),
        "isolated": isolated,
    }


def _same_path(left: Path | None, right: Path | None) -> bool:
    if left is None or right is None:
        return False
    try:
        return left.resolve() == right.resolve()
    except Exception:
        return str(left) == str(right)


def _chat_turn_lock(root: Path, run_id: str) -> asyncio.Lock:
    key = f"{root.resolve()}::{run_id}"
    with _CHAT_TURN_LOCK_GUARD:
        lock = _CHAT_TURN_LOCKS.get(key)
        if lock is None:
            lock = asyncio.Lock()
            _CHAT_TURN_LOCKS[key] = lock
        return lock


def _read_run_state(root: Path) -> dict[str, Any]:
    run_state_path = root / "workspace" / "run_state.json"
    if not run_state_path.exists():
        return {}
    try:
        loaded = json.loads(run_state_path.read_text(encoding="utf-8"))
    except Exception:
        return {}
    return loaded if isinstance(loaded, dict) else {}


def _parse_timestamp(value: Any) -> datetime | None:
    if not isinstance(value, str) or not value.strip():
        return None
    try:
        return datetime.fromisoformat(value.strip())
    except ValueError:
        return None


def _format_duration(seconds: float | int | None) -> str:
    if seconds is None:
        return ""
    total = max(0, int(round(float(seconds))))
    minutes, second = divmod(total, 60)
    if minutes == 0:
        return f"{second}秒"
    hours, minute = divmod(minutes, 60)
    if hours == 0:
        return f"{minute}分{second:02d}秒"
    return f"{hours}小时{minute:02d}分{second:02d}秒"


def _load_run_history(root: Path) -> list[dict[str, Any]]:
    history_path = root / "workspace" / "run_state_history.jsonl"
    if not history_path.exists():
        return []
    records: list[dict[str, Any]] = []
    try:
        lines = history_path.read_text(encoding="utf-8").splitlines()
    except Exception:
        return []
    for line in lines:
        if not line.strip():
            continue
        try:
            record = json.loads(line)
        except Exception:
            continue
        if isinstance(record, dict):
            records.append(record)
    return records


def _copy_tree_contents(source: Path, destination: Path) -> None:
    if not source.exists():
        return
    destination.mkdir(parents=True, exist_ok=True)
    for item in source.iterdir():
        target = destination / item.name
        if item.is_dir():
            if target.exists():
                shutil.rmtree(str(target))
            shutil.copytree(str(item), str(target))
        elif item.is_file():
            target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(str(item), str(target))


def _normalize_run_name(name: str) -> str:
    cleaned = re.sub(r"\s+", "_", name.strip())
    cleaned = re.sub(r"[^0-9A-Za-z_\-\u4e00-\u9fff]+", "_", cleaned)
    cleaned = cleaned.strip("._-")
    return cleaned[:48]


def _create_run_workspace(name: str, project_type: str | None = None, expected_pages: int = 0) -> tuple[str, Path]:
    RUNS_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = _normalize_run_name(name)
    if not safe_name:
        raise ValueError("请先设置工作空间名称。")
    base_id = f"{safe_name}_{time.strftime('%Y%m%d_%H%M%S')}"
    run_id = base_id
    run_root = RUNS_DIR / run_id
    suffix = 1
    while run_root.exists():
        suffix += 1
        run_id = f"{base_id}_{suffix}"
        run_root = RUNS_DIR / run_id

    for relative_dir in [
        "sources/tender",
        "sources/company",
        "sources/template",
        "inputs",
        "workspace",
        "workspace/chunks",
        "workspace/imported",
        "workspace/jobs",
        "workspace/contexts",
        "workspace/chapters",
        "workspace/reviews",
        "workspace/rewrites",
        "workspace/summaries",
        "outputs",
    ]:
        (run_root / relative_dir).mkdir(parents=True, exist_ok=True)

    _copy_tree_contents(ROOT / "prompts", run_root / "prompts")
    save_project_profile(run_root, project_type, expected_pages=expected_pages)
    return run_id, run_root


# ---------------------------------------------------------------
#  Static files & templates
# ---------------------------------------------------------------

if (WEB_DIR / "static").exists():
    app.mount("/static", StaticFiles(directory=str(WEB_DIR / "static")), name="static")

USE_VUE = VUE_DIST_DIR.exists()
if USE_VUE:
    app.mount("/assets", StaticFiles(directory=str(VUE_DIST_DIR / "assets")), name="vue_assets")


@app.get("/", response_class=HTMLResponse)
def index(request: Request) -> str:
    if USE_VUE:
        return (VUE_DIST_DIR / "index.html").read_text(encoding="utf-8")
    index_path = WEB_DIR / "templates" / "index.html"
    if not index_path.exists():
        return "<h1>缺少 web/templates/index.html</h1>"
    return index_path.read_text(encoding="utf-8")


# ---------------------------------------------------------------
#  Status
# ---------------------------------------------------------------

def _exists(path: Path) -> bool:
    return path.exists() and path.stat().st_size > 0


def _count_glob(directory: Path, pattern: str) -> int:
    return len(list(directory.glob(pattern))) if directory.exists() else 0


def _list_source_files(category: str, root: Path | None = None) -> list[dict[str, Any]]:
    source_dir = (root or _active_root()) / "sources" / category
    if not source_dir.exists():
        return []
    files: list[dict[str, Any]] = []
    for path in sorted(source_dir.iterdir()):
        if path.is_file():
            files.append(
                {
                    "name": path.name,
                    "size": path.stat().st_size,
                    "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(path.stat().st_mtime)),
                }
            )
    return files


def _workspace_file_item(root: Path, path: Path) -> dict[str, Any]:
    st = path.stat()
    return {
        "name": path.name,
        "path": str(path.relative_to(root)).replace("\\", "/"),
        "size": st.st_size,
        "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(st.st_mtime)),
    }


def _expand_artifact_file_items(root: Path, artifact_path: str, artifact_kind: str = "file") -> list[dict[str, Any]]:
    if artifact_kind == "virtual":
        return []
    text = artifact_path.replace("\\", "/")
    if artifact_kind == "glob" or "*" in text:
        if "/" not in text:
            return []
        directory_text, pattern = text.rsplit("/", 1)
        directory = root / directory_text
        if not directory.exists() or not directory.is_dir():
            return []
        return [_workspace_file_item(root, path) for path in sorted(directory.glob(pattern)) if path.is_file()]
    target = root / text
    if target.is_file():
        return [_workspace_file_item(root, target)]
    if target.is_dir():
        return [_workspace_file_item(root, path) for path in sorted(target.iterdir()) if path.is_file()]
    return []


def _list_dir_file_items(root: Path, relative_dir: str) -> list[dict[str, Any]]:
    directory = root / relative_dir
    if not directory.exists() or not directory.is_dir():
        return []
    return [_workspace_file_item(root, path) for path in sorted(directory.iterdir()) if path.is_file()]


def build_workspace_file_tree(root: Path) -> dict[str, Any]:
    """List run inputs / stage intermediates / outputs for the frontend file explorer."""
    seen_paths: set[str] = set()
    sections: list[dict[str, Any]] = []

    def _push_section(
        key: str,
        label: str,
        items: list[dict[str, Any]],
        *,
        open_default: bool = False,
        stage_command: str = "",
        mark_seen: bool = True,
        skip_seen: bool = True,
    ) -> None:
        unique: list[dict[str, Any]] = []
        local_seen: set[str] = set()
        for item in items:
            path = str(item.get("path") or "")
            if not path or path in local_seen:
                continue
            if skip_seen and path in seen_paths:
                continue
            local_seen.add(path)
            if mark_seen:
                seen_paths.add(path)
            unique.append(item)
        if not unique and key not in {"tender", "company", "template", "outputs"}:
            return
        sections.append(
            {
                "key": key,
                "label": label,
                "open": open_default,
                "stage_command": stage_command,
                "items": unique,
            }
        )

    for category, label in (("tender", "招标文件"), ("company", "公司资料"), ("template", "标书模板")):
        _push_section(category, label, _list_dir_file_items(root, f"sources/{category}"), open_default=False)

    for stage in workflow_stage_specs(include_utility=False):
        items: list[dict[str, Any]] = []
        for artifact in stage.produces:
            items.extend(_expand_artifact_file_items(root, artifact.path, artifact.kind))
        _push_section(
            f"stage_{stage.id}",
            stage.label,
            items,
            open_default=False,
            stage_command=stage.command,
        )

    # Catch intermediate products not declared on a single stage produces list.
    extra_globs = [
        ("workspace/issues/*", "问题单"),
        ("workspace/manual_review/*", "人工复核"),
        ("workspace/source_traces/*", "来源追溯明细"),
        ("workspace/agent/*", "Agent 状态"),
        ("workspace/debug_*", "调试中间文件"),
    ]
    extra_items: list[dict[str, Any]] = []
    for pattern, _label in extra_globs:
        if pattern.startswith("workspace/debug_"):
            workspace_dir = root / "workspace"
            if workspace_dir.exists():
                for path in sorted(workspace_dir.glob("debug_*")):
                    if path.is_file():
                        extra_items.append(_workspace_file_item(root, path))
            continue
        extra_items.extend(_expand_artifact_file_items(root, pattern, "glob"))

    # Top-level workspace json not already covered (reports / state / etc.)
    workspace_dir = root / "workspace"
    if workspace_dir.exists():
        for path in sorted(workspace_dir.iterdir()):
            if not path.is_file():
                continue
            rel = str(path.relative_to(root)).replace("\\", "/")
            if rel in seen_paths:
                continue
            if path.suffix.lower() in {".json", ".jsonl", ".md", ".txt", ".log"}:
                extra_items.append(_workspace_file_item(root, path))

    _push_section("other_workspace", "其他中间产物", extra_items, open_default=False)
    _push_section(
        "outputs",
        "最终输出",
        _list_dir_file_items(root, "outputs"),
        open_default=False,
        skip_seen=False,
        mark_seen=True,
    )

    total = sum(len(section["items"]) for section in sections)
    return {"ok": True, "sections": sections, "total": total}


def _latest_mtime_in_dir(directory: Path) -> float:
    if not directory.exists():
        return 0.0
    mtimes = [directory.stat().st_mtime]
    mtimes.extend(item.stat().st_mtime for item in directory.iterdir() if item.is_file())
    return max(mtimes) if mtimes else 0.0


def _latest_mtime(paths: list[Path]) -> float:
    mtimes = [path.stat().st_mtime for path in paths if path.exists() and path.is_file()]
    return max(mtimes) if mtimes else 0.0


def _path_status(root: Path, path: str) -> bool:
    if path.endswith("/*"):
        target = root / path[:-2]
        return target.exists() and any(target.glob("*"))
    target = root / path
    return target.exists() and target.is_file() and target.stat().st_size > 0


def _read_json_file(path: Path) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None


def _write_json_file(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _safe_relative(root: Path, path: Path) -> str:
    try:
        return str(path.relative_to(root)).replace("\\", "/")
    except Exception:
        return str(path).replace("\\", "/")


def _load_agent_runs(
    root: Path,
    *,
    stage: str = "",
    chapter_id: str = "",
    agent_name: str = "",
    limit: int = 50,
) -> list[dict[str, Any]]:
    agent_runs_dir = root / "workspace" / "agent_runs"
    if not agent_runs_dir.exists():
        return []
    rows: list[dict[str, Any]] = []
    for path in sorted(agent_runs_dir.glob("*.json"), key=lambda item: item.stat().st_mtime, reverse=True):
        payload = _read_json_file(path)
        if not isinstance(payload, dict):
            continue
        if stage and str(payload.get("stage", "")).strip() != stage:
            continue
        if chapter_id and str(payload.get("chapter_id", "")).strip() != chapter_id:
            continue
        if agent_name and str(payload.get("agent_name", "")).strip() != agent_name:
            continue
        rows.append(
            {
                "path": _safe_relative(root, path),
                "stage": str(payload.get("stage", "")),
                "agent_name": str(payload.get("agent_name", "")),
                "chapter_id": str(payload.get("chapter_id", "")),
                "started_at": str(payload.get("started_at", "")),
                "finished_at": str(payload.get("finished_at", "")),
                "duration_ms": payload.get("duration_ms", 0),
                "llm_calls": payload.get("llm_calls", 0),
                "model": str(payload.get("model", "")),
                "temperature": payload.get("temperature", 0),
                "prompt_file": str(payload.get("prompt_file", "")),
                "prompt_version": str(payload.get("prompt_version", "")),
                "prompt_checksum": str(payload.get("prompt_checksum", "")),
                "project_type": str(payload.get("project_type", "")),
                "context_budget": payload.get("context_budget", {}),
                "input_summary": payload.get("input_summary", {}),
                "output_summary": payload.get("output_summary", {}),
                "input_tokens_est": payload.get("input_tokens_est", 0),
                "output_tokens_est": payload.get("output_tokens_est", 0),
            }
        )
        if len(rows) >= max(1, limit):
            break
    return rows


def _latest_agent_runs(root: Path, limit: int = 8) -> list[dict[str, Any]]:
    return _load_agent_runs(root, limit=limit)


def _budget_hits_for_command(root: Path, command: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    if command == "select-context-all":
        contexts_dir = root / "workspace" / "contexts"
        for path in sorted(contexts_dir.glob("*_context.json")) if contexts_dir.exists() else []:
            payload = _read_json_file(path)
            if not isinstance(payload, dict):
                continue
            meta = payload.get("selection_meta", {})
            if not isinstance(meta, dict):
                meta = {}
            rows.append(
                {
                    "chapter_id": str(payload.get("chapter_id", "")),
                    "metric": "context_selection",
                    "max_context_chars": meta.get("max_context_chars", 0),
                    "max_chunks_per_side": meta.get("max_chunks_per_side", 0),
                    "tender_candidates_total": meta.get("tender_candidates_total", 0),
                    "tender_candidates_in_prompt": meta.get("tender_candidates_in_prompt", 0),
                    "company_candidates_total": meta.get("company_candidates_total", 0),
                    "company_candidates_in_prompt": meta.get("company_candidates_in_prompt", 0),
                    "dropped_reason": str(meta.get("dropped_reason", "")),
                }
            )
    stage = _command_for_stage(command) if command not in STAGE_TO_COMMAND.values() else next(
        (stage_name for stage_name, stage_command in STAGE_TO_COMMAND.items() if stage_command == command),
        command,
    )
    for run in _load_agent_runs(root, stage=stage, limit=100):
        budget = run.get("context_budget", {})
        if not isinstance(budget, dict) or not budget:
            continue
        summary = run.get("input_summary", {})
        if not isinstance(summary, dict):
            summary = {}
        rows.append(
            {
                "chapter_id": run.get("chapter_id", ""),
                "agent_name": run.get("agent_name", ""),
                "metric": "agent_prompt_budget",
                "prompt_file": run.get("prompt_file", ""),
                "prompt_checksum": run.get("prompt_checksum", ""),
                "context_budget": budget,
                "input_summary": summary,
            }
        )
    return rows[:24]


def _file_payload(root: Path, relative_path: str) -> dict[str, Any]:
    path = root / relative_path
    exists = path.exists()
    payload: dict[str, Any] = {
        "path": relative_path,
        "exists": exists,
        "size": path.stat().st_size if exists and path.is_file() else 0,
        "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(path.stat().st_mtime)) if exists else "",
    }
    if path.is_dir():
        payload["type"] = "dir"
        payload["count"] = len(list(path.iterdir()))
    else:
        payload["type"] = "file"
    return payload


def _artifact_path(artifact: Any) -> str:
    if isinstance(artifact, dict):
        return str(artifact.get("path", ""))
    return str(artifact)


def _artifact_kind(artifact: Any) -> str:
    if isinstance(artifact, dict):
        return str(artifact.get("kind", "file"))
    text = str(artifact)
    if text in {"基础目录", "默认提示词"}:
        return "virtual"
    if "*" in text:
        return "glob"
    return "file"


def _artifact_payloads(root: Path, artifacts: list[Any]) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []
    for artifact in artifacts:
        artifact_path = _artifact_path(artifact)
        artifact_kind = _artifact_kind(artifact)
        if artifact_kind == "glob":
            directory_text, pattern = artifact_path.rsplit("/", 1)
            directory = root / directory_text
            files = sorted(directory.glob(pattern)) if directory.exists() else []
            latest = max((path.stat().st_mtime for path in files), default=0)
            output.append(
                {
                    "path": artifact_path,
                    "type": "glob",
                    "exists": bool(files),
                    "count": len(files),
                    "size": sum(path.stat().st_size for path in files if path.is_file()),
                    "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(latest)) if latest else "",
                    "samples": [path.name for path in files[:8]],
                    "previewable": True,
                }
            )
            continue
        if artifact_kind == "virtual":
            output.append({"path": artifact_path, "type": "virtual", "exists": True, "size": 0, "modified": "", "previewable": False})
            continue
        payload = _file_payload(root, artifact_path)
        payload["previewable"] = payload.get("exists", False) and payload.get("type") == "file"
        output.append(payload)
    return output


def _json_count(value: Any) -> int:
    if isinstance(value, list):
        return len(value)
    if isinstance(value, dict):
        for key in ("items", "data", "chapters", "score_points", "score_requirements"):
            nested = value.get(key)
            if isinstance(nested, list):
                return len(nested)
        return len(value)
    return 0


def _summarize_review_step(root: Path) -> dict[str, Any]:
    reviews_dir = root / "workspace" / "reviews"
    rewrites_dir = root / "workspace" / "rewrites"
    review_files = sorted(reviews_dir.glob("*_review.json")) if reviews_dir.exists() else []
    rewrite_files = sorted(rewrites_dir.glob("*_rewrite_log.json")) if rewrites_dir.exists() else []
    need_rewrite: list[str] = []
    need_evidence: list[str] = []
    stuck: list[str] = []
    problem_count = 0
    priority_fix_count = 0
    severity = {"blocker": 0, "major": 0, "minor": 0}
    coverage = {"high": 0, "medium": 0, "low": 0, "none": 0}
    problem_samples: list[str] = []
    for path in review_files:
        data = _read_json_file(path)
        if not isinstance(data, dict):
            continue
        chapter_id = str(data.get("chapter_id") or path.name.replace("_review.json", ""))
        status = str(data.get("rewrite_status") or "")
        if status == "stuck" or data.get("stuck"):
            stuck.append(chapter_id)
        elif status == "need_evidence" or (
            data.get("need_evidence") and not data.get("has_writing_fixes", True)
        ):
            need_evidence.append(chapter_id)
        elif data.get("need_rewrite"):
            need_rewrite.append(chapter_id)
        if data.get("need_evidence") and chapter_id not in need_evidence:
            need_evidence.append(chapter_id)
        problems = data.get("problems") if isinstance(data.get("problems"), list) else []
        problem_count += len(problems)
        fixes = data.get("priority_fixes") if isinstance(data.get("priority_fixes"), list) else []
        priority_fix_count += len(fixes)
        for problem in problems[:2]:
            if isinstance(problem, dict):
                sev = str(problem.get("severity") or "").lower()
                if sev in severity:
                    severity[sev] += 1
                problem_samples.append(f"{chapter_id}: {problem.get('description') or problem.get('type') or '需修订'}")
        for item in data.get("score_coverage") or []:
            if isinstance(item, dict):
                level = str(item.get("coverage_level") or "").lower()
                if level in coverage:
                    coverage[level] += 1
    return {
        "审核文件": len(review_files),
        "打回重写日志": len(rewrite_files),
        "当前仍需重写": len(need_rewrite),
        "需补证据": len(need_evidence),
        "卡住": len(stuck),
        "问题数": problem_count,
        "优先修复项": priority_fix_count,
        "阻断": severity["blocker"],
        "重要": severity["major"],
        "次要": severity["minor"],
        "高覆盖": coverage["high"],
        "中覆盖": coverage["medium"],
        "低覆盖": coverage["low"],
        "未覆盖": coverage["none"],
        "仍需重写章节": need_rewrite[:12],
        "需补证据章节": need_evidence[:12],
        "卡住章节": stuck[:12],
        "问题示例": problem_samples[:8],
    }


def _review_detail_rows(root: Path) -> list[dict[str, Any]]:
    reviews_dir = root / "workspace" / "reviews"
    rewrites_dir = root / "workspace" / "rewrites"
    review_files = sorted(reviews_dir.glob("*_review.json")) if reviews_dir.exists() else []
    rewrite_ids = {
        path.name.replace("_rewrite_log.json", "")
        for path in rewrites_dir.glob("*_rewrite_log.json")
    } if rewrites_dir.exists() else set()
    rows: list[dict[str, Any]] = []
    for path in review_files:
        data = _read_json_file(path)
        if not isinstance(data, dict):
            continue
        chapter_id = str(data.get("chapter_id") or path.name.replace("_review.json", ""))
        problems = [
            str(item.get("description") or item.get("type") or "")
            for item in (data.get("problems") or [])
            if isinstance(item, dict)
        ]
        weak_coverage: list[str] = []
        for item in data.get("score_coverage") or []:
            if not isinstance(item, dict):
                continue
            level = str(item.get("coverage_level") or "").lower()
            if level in {"low", "none"} or not item.get("covered", True):
                score_id = str(item.get("score_point_id") or "")
                weak_coverage.append(f"{score_id}({level or '未覆盖'})")
        priority_fixes = [
            str(item.get("target") or item.get("action") or item.get("id") or "")
            for item in (data.get("priority_fixes") or [])
            if isinstance(item, dict)
        ]
        rows.append(
            {
                "chapter_id": chapter_id,
                "chapter_title": str(data.get("chapter_title") or ""),
                "need_rewrite": bool(data.get("need_rewrite")),
                "need_evidence": bool(data.get("need_evidence")),
                "stuck": bool(data.get("stuck")) or str(data.get("rewrite_status") or "") == "stuck",
                "rewrite_status": str(data.get("rewrite_status") or ""),
                "max_severity": str(data.get("max_severity") or ""),
                "rewritten": chapter_id in rewrite_ids,
                "problem_count": len(problems),
                "priority_fix_count": len(priority_fixes),
                "problems": problems[:3],
                "priority_fixes": priority_fixes[:3],
                "weak_coverage": weak_coverage[:6],
                "review_path": str(path.relative_to(root)).replace("\\", "/"),
                "rewrite_path": f"workspace/rewrites/{chapter_id}_rewrite_log.json" if chapter_id in rewrite_ids else "",
            }
        )
    rows.sort(key=lambda item: (not item["need_rewrite"], item["chapter_id"]))
    return rows


def _score_point_rows(root: Path) -> list[dict[str, Any]]:
    points = _read_json_file(root / "workspace" / "score_points.json")
    if not isinstance(points, list):
        return []
    rows: list[dict[str, Any]] = []
    for item in points[:80]:
        if not isinstance(item, dict):
            continue
        rows.append(
            {
                "id": str(item.get("id", "")),
                "title": str(item.get("title", "")),
                "category": str(item.get("category", "")),
                "score": item.get("score"),
                "requirement": str(item.get("requirement", "")),
                "response_strategy": str(item.get("response_strategy", "")),
                "keywords": item.get("keywords", []) if isinstance(item.get("keywords"), list) else [],
            }
        )
    return rows


def _score_requirement_rows(root: Path) -> list[dict[str, Any]]:
    requirements = _read_json_file(root / "workspace" / "score_requirements.json")
    if not isinstance(requirements, list):
        return []
    rows: list[dict[str, Any]] = []
    for item in requirements[:80]:
        if not isinstance(item, dict):
            continue
        rows.append(
            {
                "category": str(item.get("category", "")),
                "title": str(item.get("title", "")),
                "score": item.get("score"),
                "requirement": str(item.get("requirement", "")),
                "scoring_criteria": str(item.get("scoring_criteria", "")),
                "keywords": item.get("keywords", []) if isinstance(item.get("keywords"), list) else [],
            }
        )
    return rows


def _stage_prompt_summary(root: Path, command: str) -> list[dict[str, Any]]:
    from pipeline_registry import STAGE_SPECS

    stage = next((st for st in STAGE_SPECS if st.command == command), None)
    if stage is None or not stage.prompt_agents:
        return []
    results: list[dict[str, Any]] = []
    for agent_name in stage.prompt_agents:
        spec = AGENT_SPECS.get(agent_name)
        if spec is None:
            results.append({"agent_name": agent_name, "prompt_file": "", "version": "", "file_exists": False, "missing": True})
            continue
        prompt_path = root / "prompts" / spec.prompt_file
        results.append({
            "agent_name": spec.name,
            "prompt_file": spec.prompt_file,
            "version": spec.version,
            "file_exists": prompt_path.exists() and prompt_path.is_file() and prompt_path.stat().st_size > 0,
            "context_budget": spec.context_budget if spec.context_budget else None,
        })
    return results


def _step_detail_summary(root: Path, command: str) -> dict[str, Any]:
    workspace = root / "workspace"
    if command == "prepare-inputs":
        report = _read_json_file(workspace / "imported" / "tender_classification_report.json")
        schema = _read_json_file(workspace / "template_schema.json")
        return {
            "分类块总数": report.get("total_blocks", 0) if isinstance(report, dict) else 0,
            "评分块": report.get("score_blocks", 0) if isinstance(report, dict) else 0,
            "需求块": report.get("tender_blocks", 0) if isinstance(report, dict) else 0,
            "其他块": report.get("other_blocks", 0) if isinstance(report, dict) else 0,
            "分类警告": report.get("warnings", []) if isinstance(report, dict) else [],
            "模板标题": len(schema.get("headings", [])) if isinstance(schema, dict) and isinstance(schema.get("headings"), list) else 0,
            "模板表格": len(schema.get("tables", [])) if isinstance(schema, dict) and isinstance(schema.get("tables"), list) else 0,
        }
    if command == "split-docs":
        tender = _read_json_file(workspace / "chunks" / "tender_chunks.json")
        company = _read_json_file(workspace / "chunks" / "company_chunks.json")
        return {"招标片段": _json_count(tender), "公司片段": _json_count(company)}
    if command == "parse-score":
        points = _read_json_file(workspace / "score_points.json")
        requirements = _read_json_file(workspace / "score_requirements.json")
        return {"评分点": _json_count(points), "原始评分要求": _json_count(requirements)}
    if command == "extract-facts":
        facts = _read_json_file(workspace / "global_facts.json")
        if isinstance(facts, dict):
            return {"事实字段": len(facts), "项目名称": facts.get("project_name", ""), "投标人": facts.get("bidder_name", "")}
        return {}
    if command == "build-template-evidence":
        quality = _read_json_file(workspace / "template_quality_report.json")
        if isinstance(quality, dict):
            return {
                "失败项": quality.get("fail_count", 0),
                "警告项": quality.get("warn_count", 0),
                "通过项": quality.get("ok_count", 0),
            }
        return {}
    if command == "generate-outline":
        outline = _read_json_file(workspace / "outline.json")
        return {"章节数": len(outline.get("chapters", [])) if isinstance(outline, dict) else 0}
    if command == "plan-jobs":
        return {"任务文件": _count_glob(workspace / "jobs", "*.json")}
    if command == "select-context-all":
        return {
            "上下文文件": _count_glob(workspace / "contexts", "*_context.json"),
            "排序片段文件": _count_glob(workspace / "contexts", "*_ranked_chunks.json"),
        }
    if command == "write-all":
        chapters = list((workspace / "chapters").glob("*.md")) if (workspace / "chapters").exists() else []
        return {"章节文件": len(chapters), "总字数": sum(len(path.read_text(encoding="utf-8", errors="ignore")) for path in chapters)}
    if command == "review-fix-all":
        return _summarize_review_step(root)
    if command == "build-source-trace":
        index = _read_json_file(workspace / "source_trace_index.json")
        summary = index.get("summary") if isinstance(index, dict) and isinstance(index.get("summary"), dict) else {}
        return {
            "来源索引项": _json_count(index),
            "章节来源文件": _count_glob(workspace / "source_traces", "*_sources.json"),
            "claim数": summary.get("claim_count", 0),
            "claim已对齐": summary.get("claim_aligned_count", 0),
        }
    if command == "build-score-coverage":
        matrix = _read_json_file(workspace / "score_coverage_matrix.json")
        if not isinstance(matrix, dict):
            return {}
        summary = matrix.get("summary", {}) if isinstance(matrix.get("summary"), dict) else {}
        return {
            **summary,
            "硬指标未覆盖": len(matrix.get("hard_uncovered_score_points") or []),
            "硬指标偏弱": len(matrix.get("hard_weak_score_points") or []),
            "硬指标较强": len(matrix.get("hard_strong_score_points") or []),
        }
    if command == "estimate-score":
        estimate = _read_json_file(workspace / "final_score_estimate.json")
        if not isinstance(estimate, dict):
            return {}
        summary = estimate.get("summary") if isinstance(estimate.get("summary"), dict) else {}
        return {
            "满分合计": summary.get("full_score_total"),
            "预估得分": summary.get("estimated_score_total"),
            "预估得分率": summary.get("estimated_percent"),
            "等级": summary.get("grade"),
            "保守分": summary.get("conservative_score_total"),
            "乐观分": summary.get("optimistic_score_total"),
            "预计失分": summary.get("lost_points"),
            "无分值项": summary.get("unscored_point_count"),
            "主要失分": [
                f"{item.get('score_point_id')}:{item.get('lost_points')}"
                for item in (estimate.get("top_score_losses") or [])[:5]
                if isinstance(item, dict)
            ],
        }
    if command == "summarize-all":
        return {"章节摘要": _count_glob(workspace / "summaries", "*_summary.json")}
    if command == "global-review":
        review = _read_json_file(workspace / "global_review.json")
        if isinstance(review, dict):
            return {
                "需人工复核": bool(review.get("need_manual_review")),
                "未覆盖评分点": len(review.get("uncovered_score_points") or []),
                "章节冲突": len(review.get("chapter_conflicts") or []),
                "编造风险": len(review.get("fabrication_risks") or []),
                "建议": review.get("suggestions", [])[:6] if isinstance(review.get("suggestions"), list) else [],
                "风险示例": review.get("fabrication_risks", [])[:6] if isinstance(review.get("fabrication_risks"), list) else [],
            }
        return {}
    if command == "compliance-check":
        report = _read_json_file(workspace / "compliance_report.json")
        if isinstance(report, dict):
            summary = report.get("summary") if isinstance(report.get("summary"), dict) else {}
            counts = summary.get("counts") if isinstance(summary.get("counts"), dict) else {}
            items = report.get("items") if isinstance(report.get("items"), list) else []
            failed = [
                f"{item.get('check_id')}:{item.get('check_name')}"
                for item in items
                if isinstance(item, dict) and item.get("status") in {"fail", "warn"}
            ][:8]
            price_report = _read_json_file(workspace / "price_table_report.json")
            dev_report = _read_json_file(workspace / "deviation_table_report.json")
            claim_report = _read_json_file(workspace / "claim_validation_report.json")
            return {
                "检查通过": bool(report.get("ok")),
                "阻断": bool(report.get("blocking") or summary.get("blocking")),
                "需人工复核": bool(report.get("need_manual_review") or summary.get("need_manual_review")),
                "最高风险": report.get("max_severity") or summary.get("max_severity"),
                "通过项": counts.get("pass", 0),
                "警告项": counts.get("warn", 0),
                "失败项": counts.get("fail", 0),
                "跳过项": counts.get("skip", 0),
                "报价表问题": (price_report or {}).get("issue_count", 0) if isinstance(price_report, dict) else 0,
                "偏离表问题行": (dev_report or {}).get("fail_row_count", 0) if isinstance(dev_report, dict) else 0,
                "claim阻断": ((claim_report or {}).get("summary") or {}).get("blocker_count", 0)
                if isinstance(claim_report, dict)
                else 0,
                "结果": failed,
            }
        return {}
    if command == "build-md":
        final_md = root / "outputs" / "final.md"
        text = final_md.read_text(encoding="utf-8", errors="ignore") if final_md.exists() else ""
        return {"字数": len(text), "标题数": len(re.findall(r"(?m)^#{1,6}\\s+", text))}
    if command == "build-docx":
        report = _read_json_file(workspace / "template_fill_report.json")
        return {"模板填充项": _json_count(report), "填充报告存在": bool(report)}
    if command == "check-format":
        report = _read_json_file(workspace / "format_check_report.json")
        if isinstance(report, dict):
            return {
                "检查通过": bool(report.get("ok")),
                "通过项": report.get("ok_count", 0),
                "警告项": report.get("warn_count", 0),
                "失败项": report.get("fail_count", 0),
                "结果": [
                    item.get("message", "")
                    for item in report.get("results", [])
                    if isinstance(item, dict) and item.get("level") != "ok"
                ][:8],
            }
        return {}
    return {}


def _step_history(root: Path, command: str) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for record in _load_run_history(root):
        if _command_for_stage(str(record.get("stage", ""))) == command:
            records.append(
                {
                    "status": record.get("status", ""),
                    "message": record.get("message", ""),
                    "updated_at": record.get("updated_at", ""),
                }
            )
    return records[-8:]


def _step_extra_details(root: Path, command: str) -> dict[str, Any]:
    workspace = root / "workspace"
    if command == "parse-score":
        return {
            "score_point_rows": _score_point_rows(root),
            "score_requirement_rows": _score_requirement_rows(root),
        }
    if command == "review-fix-all":
        return {"review_rows": _review_detail_rows(root)}
    if command == "global-review":
        data = _read_json_file(workspace / "global_review.json")
        return {"global_review": data if isinstance(data, dict) else {}}
    if command == "compliance-check":
        data = _read_json_file(workspace / "compliance_report.json")
        if not isinstance(data, dict):
            return {"compliance_report": {}}
        items = [i for i in (data.get("items") or []) if isinstance(i, dict)]
        return {
            "compliance_report": {
                "blocking": data.get("blocking"),
                "need_manual_review": data.get("need_manual_review"),
                "max_severity": data.get("max_severity"),
                "summary": data.get("summary") if isinstance(data.get("summary"), dict) else {},
                "failed_items": [
                    i for i in items if i.get("status") in {"fail", "warn"}
                ][:80],
            }
        }
    if command == "build-score-coverage":
        data = _read_json_file(workspace / "score_coverage_matrix.json")
        if not isinstance(data, dict):
            return {"score_coverage": {}}
        matrix = data.get("matrix") if isinstance(data.get("matrix"), list) else []
        return {
            "score_coverage": {
                "summary": data.get("summary") if isinstance(data.get("summary"), dict) else {},
                "uncovered_score_points": data.get("uncovered_score_points") or [],
                "weak_score_points": data.get("weak_score_points") or [],
                "fully_covered_score_points": data.get("fully_covered_score_points") or [],
                "matrix_preview": matrix[:40],
            }
        }
    if command == "estimate-score":
        data = _read_json_file(workspace / "final_score_estimate.json")
        return {"score_estimate": data if isinstance(data, dict) else {}}
    if command == "generate-outline":
        data = _read_json_file(workspace / "outline.json")
        chapters = data.get("chapters") if isinstance(data, dict) and isinstance(data.get("chapters"), list) else []
        return {
            "outline_chapters": [
                {
                    "id": c.get("id"),
                    "title": c.get("title"),
                    "score_point_ids": c.get("score_point_ids") if isinstance(c.get("score_point_ids"), list) else [],
                    "description": str(c.get("description") or "")[:200],
                }
                for c in chapters[:60]
                if isinstance(c, dict)
            ]
        }
    if command == "extract-facts":
        data = _read_json_file(workspace / "global_facts.json")
        return {"global_facts": data if isinstance(data, dict) else {}}
    if command == "write-all":
        chapters_dir = workspace / "chapters"
        files = sorted(chapters_dir.glob("*.md")) if chapters_dir.exists() else []
        return {
            "chapter_files": [
                {"chapter_id": f.stem, "path": f"workspace/chapters/{f.name}", "size": f.stat().st_size}
                for f in files[:80]
            ]
        }
    if command == "summarize-all":
        summaries_dir = workspace / "summaries"
        files = sorted(summaries_dir.glob("*.json")) if summaries_dir.exists() else []
        rows = []
        for f in files[:40]:
            data = _read_json_file(f)
            if isinstance(data, dict):
                rows.append({
                    "chapter_id": data.get("chapter_id") or f.stem.replace("_summary", ""),
                    "title": data.get("title") or data.get("chapter_title") or "",
                    "summary": str(data.get("summary") or data.get("brief") or "")[:240],
                })
            else:
                rows.append({"chapter_id": f.stem, "title": "", "summary": ""})
        return {"chapter_summaries": rows}
    if command == "check-format":
        data = _read_json_file(workspace / "format_check_report.json")
        return {"format_check": data if isinstance(data, dict) else {}}
    return {}


STAGE_TO_COMMAND: dict[str, str] = stage_command_map()


def _command_for_stage(stage: str) -> str:
    return STAGE_TO_COMMAND.get(stage, stage)


def _recovery_has_live_progress(
    command: str,
    recovery: dict[str, Any],
    events: list[dict[str, Any]],
) -> bool:
    recovered_after = _parse_timestamp(recovery.get("updated_at"))
    if recovered_after is None:
        return False
    stage_id = next(
        (str(step.get("id", "")) for step in WORKFLOW_STEPS if step.get("command") == command),
        command.replace("-", "_"),
    )
    for event in reversed(events):
        event_time = _parse_timestamp(event.get("ts"))
        if event_time is None:
            continue
        if event_time <= recovered_after:
            break
        if str(event.get("stage", "")) != stage_id:
            continue
        if str(event.get("event_type", "")) == "success":
            return True
        metrics = event.get("metrics", {}) if isinstance(event.get("metrics"), dict) else {}
        if str(event.get("event_type", "")) == "agent_artifact" and int(metrics.get("llm_calls", 0) or 0) > 0:
            return True
    return False


def _current_command_from_status(status: dict[str, Any]) -> str:
    if not status.get("running"):
        return ""
    task = str(status.get("current_task", "")).strip()
    if task in {"run", "graph-run"}:
        run_state = status.get("run_state", {}) if isinstance(status.get("run_state"), dict) else {}
        return _command_for_stage(str(run_state.get("stage", "")))
    return task


def _workflow_timings(
    root: Path,
    *,
    running: bool,
    current_task: str,
    run_state: dict[str, Any],
) -> dict[str, dict[str, Any]]:
    starts: dict[str, datetime] = {}
    timings: dict[str, dict[str, Any]] = {}
    workflow_commands = {step["command"] for step in WORKFLOW_STEPS}

    for record in _load_run_history(root):
        command = _command_for_stage(str(record.get("stage", "")))
        if command not in workflow_commands:
            continue
        updated_at = _parse_timestamp(record.get("updated_at"))
        if updated_at is None:
            continue

        state = str(record.get("status", "")).strip()
        timing = timings.setdefault(command, {})
        if state == "running":
            starts[command] = updated_at
            timing["started_at"] = updated_at.isoformat(timespec="seconds")
            timing["status"] = "running"
            timing["message"] = record.get("message", "")
            continue

        if state in {"ok", "error", "warn"}:
            started_at = starts.get(command) or _parse_timestamp(timing.get("started_at"))
            if started_at is not None:
                duration_seconds = max(0, (updated_at - started_at).total_seconds())
                timing["duration_seconds"] = int(round(duration_seconds))
                timing["duration_label"] = _format_duration(duration_seconds)
            timing["finished_at"] = updated_at.isoformat(timespec="seconds")
            timing["status"] = state
            timing["message"] = record.get("message", "")

    active_command = current_task
    if current_task in {"run", "graph-run"}:
        active_command = _command_for_stage(str(run_state.get("stage", "")))
    if running and active_command in workflow_commands:
        timing = timings.setdefault(active_command, {})
        started_at = _parse_timestamp(timing.get("started_at"))
        if started_at is None and str(run_state.get("status", "")) == "running":
            started_at = _parse_timestamp(run_state.get("updated_at"))
            if started_at is not None:
                timing["started_at"] = started_at.isoformat(timespec="seconds")
        if started_at is not None:
            duration_seconds = max(0, (datetime.now() - started_at).total_seconds())
            timing["duration_seconds"] = int(round(duration_seconds))
            timing["duration_label"] = _format_duration(duration_seconds)
        timing["status"] = "running"

    return timings


def _step_status(root: Path, step: dict[str, Any], status: dict[str, Any]) -> dict[str, Any]:
    key_map = {
        "inputs/tender.md": status["inputs"]["tender_md"],
        "inputs/company.md": status["inputs"]["company_md"],
        "inputs/score.md": status["inputs"]["score_md"],
        "inputs/template.docx": status["inputs"]["template_docx"],
        "workspace/imported/*": all(status["imported"].values()),
        "workspace/imported/tender_raw.md": status["imported"]["tender_raw"],
        "workspace/imported/tender_blocks.json": status["imported"]["tender_blocks"],
        "workspace/imported/tender_classified_blocks.json": status["imported"]["tender_classified_blocks"],
        "workspace/imported/tender_classification_report.json": status["imported"]["tender_classification_report"],
        "workspace/imported/tender_other.md": status["imported"]["tender_other"],
        "workspace/chunks/tender_chunks.json": status["chunks"]["tender_chunks"],
        "workspace/chunks/company_chunks.json": status["chunks"]["company_chunks"],
        "workspace/tender_requirements.json": status["workspace"]["tender_requirements"],
        "workspace/company_facts.json": status["workspace"]["company_facts"],
        "workspace/score_requirements.json": status["workspace"]["score_requirements"],
        "workspace/score_points.json": status["workspace"]["score_points"],
        "workspace/global_facts.json": status["workspace"]["global_facts"],
        "workspace/materials_checklist.json": status["workspace"]["materials_checklist"],
        "workspace/template_evidence_map.json": status["workspace"]["template_evidence_map"],
        "workspace/template_quality_report.json": status["workspace"]["template_quality_report"],
        "workspace/outline.json": status["workspace"]["outline"],
        "workspace/jobs/*.json": status["workspace"]["jobs_count"] > 0,
        "workspace/contexts/*_context.json": status["workspace"]["contexts_count"] > 0,
        "workspace/contexts/*_ranked_chunks.json": _count_glob(root / "workspace" / "contexts", "*_ranked_chunks.json") > 0,
        "workspace/chapters/*.md": status["workspace"]["chapters_count"] > 0,
        "workspace/reviews/*_review.json": status["workspace"]["reviews_count"] > 0,
        "workspace/rewrites/*_rewrite_log.json": status["workspace"]["rewrites_count"] > 0,
        "workspace/source_trace_index.json": status["workspace"]["source_trace_index"],
        "workspace/score_coverage_matrix.json": status["workspace"]["score_coverage_matrix"],
        "workspace/final_score_estimate.json": status["workspace"]["final_score_estimate"],
        "workspace/summaries/*_summary.json": status["workspace"]["summaries_count"] > 0,
        "workspace/global_review.json": status["workspace"]["global_review"],
        "workspace/compliance_report.json": status["workspace"]["compliance_report"],
        "outputs/final.md": status["outputs"]["final_md"],
        "outputs/final.docx": status["outputs"]["final_docx"],
        "outputs/score_estimate.md": status["outputs"]["score_estimate_md"],
        "workspace/format_check_report.json": status["workspace"]["format_check_report"],
        "workspace/template_schema.json": status["workspace"]["template_schema"],
        "workspace/template_fill_report.json": status["workspace"]["template_fill_report"],
        "sources/tender": _path_status(root, "sources/tender/*"),
        "sources/company": _path_status(root, "sources/company/*"),
        "sources/template": _path_status(root, "sources/template/*"),
    }

    requirement_paths = [_artifact_path(req) for req in step.get("requires", [])]
    produce_paths = [_artifact_path(prod) for prod in step.get("produces", [])]

    def _artifact_ok(artifact: Any) -> bool:
        if _artifact_kind(artifact) == "virtual":
            return True
        return bool(key_map.get(_artifact_path(artifact), False))

    requirements = [_artifact_ok(req) for req in step.get("requires", [])]
    ready = all(requirements)
    done = stage_outputs_ready(root, str(step.get("id", "")))
    step_index = next((idx for idx, item in enumerate(WORKFLOW_STEPS) if item.get("id") == step.get("id")), -1)
    if step_index > 0 and step.get("kind") != "utility":
        previous = WORKFLOW_STEPS[step_index - 1]
        ready = stage_outputs_ready(root, str(previous.get("id", "")))

    missing_requires = [req for req in requirement_paths if not _artifact_ok(req)]
    source_stale = status["sync"]["source_stale"]
    run_state = status.get("run_state", {}) if isinstance(status.get("run_state"), dict) else {}
    failed_command = _command_for_stage(str(run_state.get("stage", "")))
    active_command = _current_command_from_status(status)
    timing = {}
    if isinstance(status.get("timings"), dict):
        timing = status["timings"].get(step["command"], {}) or {}
    duration_label = str(timing.get("duration_label") or "")
    if step["command"] == "prepare-inputs" and source_stale:
        done = False
        state = "ready"
        message = "sources/ 有新文件，需重新导入"
    elif source_stale and step["command"] != "prepare-inputs" and step["command"] != "init":
        done = False
        state = "blocked"
        message = "请先重新执行导入资料"
    elif status.get("running") and step["command"] == active_command:
        done = False
        recovery_status = str(run_state.get("status", ""))
        if recovery_status in {"recovering", "retrying"}:
            state = recovery_status
            message = str(run_state.get("message", "")).strip() or ("正在尝试自主修复" if recovery_status == "recovering" else "正在自动重试")
        else:
            state = "running"
            message = f"已运行 {duration_label}" if duration_label else "运行中"
    elif done:
        state = "done"
        message = f"用时 {duration_label}" if duration_label else "用时 --"
    elif ready:
        state = "ready"
        message = "可执行"
    else:
        state = "blocked"
        message = "等待前置步骤"

    if not done and run_state.get("status") in {"error", "paused", "recovery_failed"} and step["command"] == failed_command:
        done = False
        state = "ready" if ready else "blocked"
        detail = str(run_state.get("message", "")).strip()
        message = f"上次执行失败，流程已暂停{f'：{detail}' if detail else ''}"

    collection = stage_collection_status(root, str(step.get("id"))) if step.get("id") in COLLECTION_STAGE_IDS else None
    return {
        **step,
        "ready": ready,
        "done": done,
        "state": state,
        "message": message,
        "missing_requires": missing_requires,
        "collection": collection,
    }


def _step_label(command: str) -> str:
    return WORKFLOW_COMMAND_LABELS.get(command, command)


def _chat_step_payload(status: dict[str, Any], command: str) -> dict[str, Any] | None:
    workflow = status.get("workflow", []) if isinstance(status.get("workflow"), list) else []
    return next((item for item in workflow if isinstance(item, dict) and item.get("command") == command), None)


def _chat_has_any(text: str, normalized: str, keywords: tuple[str, ...]) -> bool:
    return any(keyword in text or keyword.lower() in normalized for keyword in keywords)


def _diagnose_error_payload(status: dict[str, Any]) -> tuple[str, list[dict[str, str]]]:
    run_state = status.get("run_state", {}) if isinstance(status.get("run_state"), dict) else {}
    error = status.get("failed_stage_error", {}) if isinstance(status.get("failed_stage_error"), dict) else {}
    recovery = status.get("recovery", {}) if isinstance(status.get("recovery"), dict) else {}
    command = str(error.get("command") or _command_for_stage(str(run_state.get("stage", ""))))
    lines = [str(line) for line in error.get("lines", [])[-8:]] if isinstance(error.get("lines"), list) else []
    tail = "\n".join(lines).strip()
    status_text = str(run_state.get("status", "未知"))
    if status_text in {"recovering", "retrying"}:
        reply = (
            f"流程没有直接停下，我正在尝试自主修复“{_step_label(command)}”。"
            f"原因：{recovery.get('reason', '正在分析失败原因')}；"
            f"动作：{recovery.get('action', '自动重试')}；"
            f"次数：{recovery.get('attempt', 0)}/{recovery.get('max_attempts', AUTO_RECOVERY_MAX_ATTEMPTS)}。"
        )
        return reply, [{"type": "chat_prompt", "label": "查看诊断"}]
    if tail:
        reply = f"“{_step_label(command)}”上次失败。最后日志显示：\n{tail}"
    else:
        reply = f"“{_step_label(command)}”上次失败，但没有捕获到详细错误日志。"
    actions: list[dict[str, str]] = []
    if command:
        actions.extend(
            [
                {"type": "retry_stage", "command": command, "label": "手动重试"},
                {"type": "skip_stage", "command": command, "label": "跳过此阶段"},
            ]
        )
    return reply, actions


def _chat_reply(root: Path, message: str, selected_command: str = "") -> dict[str, Any]:
    status = api_status()
    text = message.strip()
    normalized = text.lower()
    next_step = status.get("next_step")
    blocked_step = status.get("blocked_step")
    manual_summary = status.get("manual_review_summary", {})
    sources = status.get("sources", {}) if isinstance(status.get("sources"), dict) else {}
    inputs = status.get("inputs", {}) if isinstance(status.get("inputs"), dict) else {}
    outputs = status.get("outputs", {}) if isinstance(status.get("outputs"), dict) else {}
    workspace = status.get("workspace", {}) if isinstance(status.get("workspace"), dict) else {}

    if not text:
        return {
            "reply": "可以直接问我当前状态、下一步、某个节点详情，或者让我继续执行流程。",
            "actions": [],
        }

    if _chat_has_any(text, normalized, ("继续", "下一步", "执行下一步", "继续执行", "next")):
        if isinstance(next_step, dict):
            return {
                "reply": f"当前下一步是“{next_step.get('label', '')}”，命令 `{next_step.get('command', '')}`，可以直接执行。",
                "actions": [
                    {"type": "run_command", "command": str(next_step.get("command", "")), "label": f"执行 {next_step.get('label', '')}"},
                    {"type": "show_step", "command": str(next_step.get("command", "")), "label": "查看节点详情"},
                ],
            }
        return {"reply": "当前没有可继续执行的下一步，可能已经完成或还未创建工作空间。", "actions": []}

    if _chat_has_any(text, normalized, ("人工复核", "未覆盖评分点", "弱证据", "模板缺口")):
        reply = (
            f"当前人工复核待处理总数 {manual_summary.get('total_pending', 0)} 项，"
            f"其中弱证据/缺口 {manual_summary.get('template_evidence_pending', 0)}，"
            f"评分点覆盖 {manual_summary.get('score_coverage_pending', 0)}，"
            f"章节问题 {manual_summary.get('chapter_review_pending', 0)}，"
            f"全文风险 {manual_summary.get('global_review_pending', 0)}。"
        )
        return {
            "reply": reply,
            "actions": [
                {"type": "show_manual_review", "category": "score_coverage", "label": "看未覆盖评分点"},
                {"type": "show_manual_review", "category": "template_evidence", "label": "看弱证据项"},
            ],
        }

    if _chat_has_any(text, normalized, ("评分", "分数", "得分", "score", "覆盖")):
        score_done = bool(workspace.get("score_points"))
        matrix_done = bool(workspace.get("score_coverage_matrix"))
        pending = manual_summary.get("score_coverage_pending", 0)
        reply = (
            f"你关心的是解析评分和评分覆盖。评分点解析：{'已完成' if score_done else '未完成'}；"
            f"评分覆盖矩阵：{'已生成' if matrix_done else '未生成'}；"
            f"未覆盖评分点待处理 {pending} 项。"
        )
        actions: list[dict[str, str]] = [
            {"type": "show_step", "command": "parse-score", "label": "查看评分解析"},
            {"type": "show_step", "command": "build-score-coverage", "label": "查看覆盖矩阵"},
            {"type": "show_manual_review", "category": "score_coverage", "label": "看未覆盖项"},
        ]
        if isinstance(next_step, dict):
            actions.append({"type": "run_command", "command": str(next_step.get("command", "")), "label": "执行下一步"})
        return {"reply": reply, "actions": actions}

    if _chat_has_any(text, normalized, ("风险", "质量", "问题", "审核", "复核", "合规", "废标", "risk", "review", "compliance")):
        compliance_done = bool(workspace.get("compliance_report"))
        reply = (
            f"你关心的是质量风险。当前人工复核待处理总数 {manual_summary.get('total_pending', 0)} 项，"
            f"章节问题 {manual_summary.get('chapter_review_pending', 0)}，"
            f"全文风险 {manual_summary.get('global_review_pending', 0)}，"
            f"弱证据/模板缺口 {manual_summary.get('template_evidence_pending', 0)}；"
            f"专项合规检查：{'已完成' if compliance_done else '未完成'}。"
        )
        return {
            "reply": reply,
            "actions": [
                {"type": "show_manual_review", "category": "global_review", "label": "看全文风险"},
                {"type": "show_manual_review", "category": "chapter_review", "label": "看章节问题"},
                {"type": "show_step", "command": "global-review", "label": "查看全文审核"},
                {"type": "show_step", "command": "compliance-check", "label": "查看专项合规"},
            ],
        }

    if _chat_has_any(text, normalized, ("输出", "结果", "word", "docx", "final", "下载", "文件")):
        reply = (
            f"你关心的是最终输出。Markdown：{'已生成' if outputs.get('final_md') else '未生成'}；"
            f"Word：{'已生成' if outputs.get('final_docx') else '未生成'}。"
        )
        actions = [
            {"type": "show_step", "command": "build-md", "label": "查看 Markdown 节点"},
            {"type": "show_step", "command": "build-docx", "label": "查看 Word 节点"},
        ]
        if isinstance(next_step, dict):
            actions.append({"type": "run_command", "command": str(next_step.get("command", "")), "label": "继续生成"})
        return {"reply": reply, "actions": actions}

    if _chat_has_any(text, normalized, ("资料", "上传", "齐全", "招标", "公司", "模板", "source", "input")):
        tender_count = len(sources.get("tender", [])) if isinstance(sources.get("tender"), list) else 0
        company_count = len(sources.get("company", [])) if isinstance(sources.get("company"), list) else 0
        template_count = len(sources.get("template", [])) if isinstance(sources.get("template"), list) else 0
        reply = (
            f"你关心的是输入资料。招标文件 {tender_count} 个，公司资料 {company_count} 个，Word 模板 {template_count} 个。"
            f" 已导入状态：招标 {'是' if inputs.get('tender_md') else '否'}，"
            f"公司 {'是' if inputs.get('company_md') else '否'}，模板 {'是' if inputs.get('template_docx') else '否'}。"
        )
        return {
            "reply": reply,
            "actions": [
                {"type": "show_step", "command": "prepare-inputs", "label": "查看导入资料"},
                {"type": "run_command", "command": "prepare-inputs", "label": "重新导入资料"},
            ],
        }

    if _chat_has_any(text, normalized, ("诊断", "失败", "错误", "修复", "error", "failed")) and status.get("failed_stage_error"):
        reply, actions = _diagnose_error_payload(status)
        return {"reply": reply, "actions": actions}

    if _chat_has_any(text, normalized, ("状态", "进度", "卡", "为什么", "失败", "暂停", "blocked", "error")):
        run_state = status.get("run_state", {})
        active = status.get("active_run", {})
        reply = (
            f"当前工作空间是 `{active.get('relative_root', active.get('id', ''))}`。"
            f" 运行状态：{run_state.get('status', '未知')}。"
        )
        if run_state.get("status") in {"recovering", "retrying"} and status.get("recovery"):
            recovery = status.get("recovery", {})
            reply += (
                f" 正在尝试自主修复：{recovery.get('reason', '')}；"
                f"{recovery.get('action', '')}（{recovery.get('attempt', 0)}/{recovery.get('max_attempts', AUTO_RECOVERY_MAX_ATTEMPTS)}）。"
            )
        if isinstance(next_step, dict):
            reply += f" 下一步是“{next_step.get('label', '')}”。"
        if isinstance(blocked_step, dict):
            reply += f" 当前阻塞点是“{blocked_step.get('label', '')}”。"
        if run_state.get("message"):
            reply += f" 说明：{run_state.get('message')}。"
        actions: list[dict[str, str]] = []
        if isinstance(next_step, dict):
            actions.append({"type": "show_step", "command": str(next_step.get("command", "")), "label": "打开下一步详情"})
        return {"reply": reply, "actions": actions}

    for step in WORKFLOW_STEPS:
        command = str(step.get("command", ""))
        label = str(step.get("label", ""))
        if command and command in normalized or label and label in text:
            step_payload = _chat_step_payload(status, command) or {}
            summary = _step_detail_summary(root, command)
            summary_text = "；".join(f"{key}: {value}" for key, value in list(summary.items())[:6]) or "暂无摘要。"
            reply = f"“{label}”当前状态是 {step_payload.get('state', '未知')}。{summary_text}"
            actions = [{"type": "show_step", "command": command, "label": f"打开 {label} 详情"}]
            if step_payload.get("ready") and not step_payload.get("done"):
                actions.append({"type": "run_command", "command": command, "label": f"执行 {label}"})
            return {"reply": reply, "actions": actions}

    if selected_command:
        label = _step_label(selected_command)
        summary = _step_detail_summary(root, selected_command)
        summary_text = "；".join(f"{key}: {value}" for key, value in list(summary.items())[:6]) or "暂无摘要。"
        return {
            "reply": f"当前你正在查看“{label}”。它的关键结果是：{summary_text}",
            "actions": [{"type": "show_step", "command": selected_command, "label": "重新打开节点详情"}],
        }

    return {
        "reply": "我现在更适合回答流程状态、节点详情、人工复核和下一步执行。你可以试试：当前状态怎么样 / 打开解析评分详情 / 继续执行下一步。",
        "actions": [],
    }


def _status_payload(
    root: Path,
    run_id: str,
    *,
    persist_manual_review_summary: bool = True,
    v2_read_only: bool = False,
) -> dict[str, Any]:
    run_state = _read_run_state(root)
    run_events = load_run_events(root)
    project_profile = load_project_profile(root)
    review_summary = manual_review_summary(root) if persist_manual_review_summary else {}
    pipeline_control = SUPERVISOR.load(root)
    repair_job = load_v2_repair_job(root) if v2_read_only else load_repair_job(root)
    pending_confirmation = None
    if str(repair_job.get("status") or "") == "awaiting_confirmation":
        pending_confirmation = {
            "type": "minimal_repair",
            "confirmation_id": str(repair_job.get("confirmation_id") or ""),
            "job_id": str(repair_job.get("job_id") or ""),
            "count": int(repair_job.get("total_count") or 0),
        }
    pipeline_running = SUPERVISOR.is_running(root) or str(pipeline_control.get("status", "")) in {
        "running",
        "recovering",
        "retrying",
        "pausing",
    }
    selected_running = (RUNNING and _same_path(CURRENT_RUN_ROOT, root)) or pipeline_running
    effective_task = CURRENT_TASK if RUNNING and _same_path(CURRENT_RUN_ROOT, root) else str(pipeline_control.get("current_stage", ""))
    run_state_stage = str(run_state.get("stage", ""))
    run_state_command = _command_for_stage(run_state_stage)
    run_state_status = str(run_state.get("status", ""))
    run_state_message = str(run_state.get("message", ""))
    recovery_file = root / "workspace" / "recovery_state.json"
    recovery_payload = _read_json_file(recovery_file) if recovery_file.exists() else {}
    if not isinstance(recovery_payload, dict):
        recovery_payload = {}
    recovery_resolved = (
        selected_running
        and run_state_status in {"recovering", "retrying"}
        and _recovery_has_live_progress(run_state_command, recovery_payload, run_events)
    )
    if recovery_resolved:
        run_state_status = "running"
        run_state_message = "LLM 已恢复，正在继续执行当前阶段"
    stale_run_error = run_state_status in {"error", "paused", "recovery_failed"} and _step_outputs_present(root, run_state_command)
    if stale_run_error:
        run_state_status = "progress"
        run_state_message = ""

    source_latest = max(
        _latest_mtime_in_dir(root / "sources" / "tender"),
        _latest_mtime_in_dir(root / "sources" / "company"),
        _latest_mtime_in_dir(root / "sources" / "template"),
    )
    imported_latest = _latest_mtime(
        [
            root / "inputs" / "tender.md",
            root / "inputs" / "score.md",
            root / "inputs" / "company.md",
            root / "inputs" / "template.docx",
            root / "workspace" / "imported" / "tender_raw.md",
            root / "workspace" / "imported" / "tender_blocks.json",
            root / "workspace" / "imported" / "tender_classified_blocks.json",
            root / "workspace" / "imported" / "tender_classification_report.json",
            root / "workspace" / "imported" / "tender_other.md",
        ]
    )
    source_stale = bool(source_latest and imported_latest and source_latest > imported_latest)

    status = {
        "inputs": {
            "tender_md": _exists(root / "inputs" / "tender.md"),
            "company_md": _exists(root / "inputs" / "company.md"),
            "score_md": _exists(root / "inputs" / "score.md"),
            "template_docx": _exists(root / "inputs" / "template.docx"),
        },
        "sources": {
            "tender": _list_source_files("tender", root),
            "company": _list_source_files("company", root),
            "template": _list_source_files("template", root),
        },
        "imported": {
            "tender_raw": _exists(root / "workspace" / "imported" / "tender_raw.md"),
            "tender_blocks": _exists(root / "workspace" / "imported" / "tender_blocks.json"),
            "tender_classified_blocks": _exists(root / "workspace" / "imported" / "tender_classified_blocks.json"),
            "tender_classification_report": _exists(root / "workspace" / "imported" / "tender_classification_report.json"),
            "tender_other": _exists(root / "workspace" / "imported" / "tender_other.md"),
        },
        "chunks": {
            "tender_chunks": _exists(root / "workspace" / "chunks" / "tender_chunks.json"),
            "company_chunks": _exists(root / "workspace" / "chunks" / "company_chunks.json"),
        },
        "workspace": {
            "tender_requirements": _exists(root / "workspace" / "tender_requirements.json"),
            "company_facts": _exists(root / "workspace" / "company_facts.json"),
            "score_requirements": _exists(root / "workspace" / "score_requirements.json"),
            "score_points": _exists(root / "workspace" / "score_points.json"),
            "global_facts": _exists(root / "workspace" / "global_facts.json"),
            "materials_checklist": _exists(root / "workspace" / "materials_checklist.json"),
            "template_evidence_map": _exists(root / "workspace" / "template_evidence_map.json"),
            "template_quality_report": _exists(root / "workspace" / "template_quality_report.json"),
            "outline": _exists(root / "workspace" / "outline.json"),
            "jobs_count": _count_glob(root / "workspace" / "jobs", "*.json"),
            "contexts_count": _count_glob(root / "workspace" / "contexts", "*_context.json"),
            "chapters_count": _count_glob(root / "workspace" / "chapters", "*.md"),
            "reviews_count": _count_glob(root / "workspace" / "reviews", "*_review.json"),
            "source_trace_index": _exists(root / "workspace" / "source_trace_index.json"),
            "score_coverage_matrix": _exists(root / "workspace" / "score_coverage_matrix.json"),
            "final_score_estimate": _exists(root / "workspace" / "final_score_estimate.json"),
            "format_check_report": _exists(root / "workspace" / "format_check_report.json"),
            "template_schema": _exists(root / "workspace" / "template_schema.json"),
            "template_fill_report": _exists(root / "workspace" / "template_fill_report.json"),
            "summaries_count": _count_glob(root / "workspace" / "summaries", "*_summary.json"),
            "global_review": _exists(root / "workspace" / "global_review.json"),
            "compliance_report": _exists(root / "workspace" / "compliance_report.json"),
            "rewrites_count": _count_glob(root / "workspace" / "rewrites", "*_rewrite_log.json"),
        },
        "outputs": {
            "final_md": _exists(root / "outputs" / "final.md"),
            "final_docx": _exists(root / "outputs" / "final.docx"),
            "score_estimate_md": _exists(root / "outputs" / "score_estimate.md"),
        },
        "sync": {
            "source_stale": source_stale,
            "source_latest": source_latest,
            "imported_latest": imported_latest,
        },
        "running": selected_running,
        "current_task": effective_task if selected_running else "",
        "global_running": RUNNING or SUPERVISOR.is_running(),
        "pipeline": pipeline_control,
        "recovery_resolved": recovery_resolved,
        "running_run": {
            "id": CURRENT_RUN_ID,
            "root": str(CURRENT_RUN_ROOT) if CURRENT_RUN_ROOT else "",
            "relative_root": str(CURRENT_RUN_ROOT.relative_to(ROOT)) if CURRENT_RUN_ROOT and CURRENT_RUN_ROOT.is_relative_to(ROOT) else (str(CURRENT_RUN_ROOT) if CURRENT_RUN_ROOT else ""),
            "command": effective_task,
        },
        "active_run": {
            "id": run_id,
            "root": str(root),
            "relative_root": str(root.relative_to(ROOT)) if root != ROOT and root.is_relative_to(ROOT) else str(root),
            "isolated": root != ROOT,
        },
        "run_state": {
            "stage": run_state.get("stage", ""),
            "status": run_state_status,
            "message": run_state_message,
            "updated_at": run_state.get("updated_at", ""),
            "summary": run_state.get("summary", {}),
        },
        "project_profile": project_profile,
        "project_profile_choices": project_profile_choices(),
        "llm_config": _active_llm_summary(),
        "agent_activity": (
            ControlStore(WorkspaceContext.resolve(root.parent, root.name)).agent_activity_state()
            if v2_read_only
            else _safe_agent_activity(root)
        ),
        "issues_summary": _v2_issues_summary(root) if v2_read_only else _safe_issues_summary(root),
        "pending_confirmation": pending_confirmation,
        "repair_job": repair_job or None,
        "manual_review_summary": review_summary,
        "latest_agent_runs": _latest_agent_runs(root),
        "run_metrics": load_stage_metrics(root),
        "run_events_tail": run_events[-20:],
    }
    # 失败/恢复阶段错误日志，供编排器诊断和前端恢复
    error_file = root / "workspace" / "run_error.json"
    if error_file.exists() and run_state_status in {"error", "paused", "progress", "recovering", "retrying", "recovery_failed"}:
        try:
            status["failed_stage_error"] = json.loads(error_file.read_text(encoding="utf-8"))
        except Exception:
            pass
    if recovery_payload and run_state_status in {"recovering", "retrying", "recovery_failed"}:
        status["recovery"] = recovery_payload
    status["timings"] = _workflow_timings(
        root,
        running=selected_running,
        current_task=effective_task if selected_running else "",
        run_state=status["run_state"],
    )

    workflow = [_step_status(root, step, status) for step in WORKFLOW_STEPS]
    core_workflow = [step for step in workflow if step["kind"] != "utility"]
    next_step = next(
        (step for step in core_workflow if not step["done"] and step["ready"]),
        None,
    )
    if next_step is None and not any(step["done"] for step in core_workflow):
        next_step = next((step for step in workflow if not step["done"] and step["ready"]), None)
    blocked_step = next((step for step in core_workflow if not step["done"] and not step["ready"]), None)

    compliance_summary: dict[str, Any] = {}
    compliance_path = root / "workspace" / "compliance_report.json"
    if compliance_path.exists():
        try:
            compliance = json.loads(compliance_path.read_text(encoding="utf-8"))
            if isinstance(compliance, dict):
                summary = compliance.get("summary") if isinstance(compliance.get("summary"), dict) else {}
                counts = summary.get("counts") if isinstance(summary.get("counts"), dict) else {}
                failed_items = [
                    {
                        "check_id": item.get("check_id"),
                        "check_name": item.get("check_name"),
                        "check_type": item.get("check_type"),
                        "status": item.get("status"),
                        "severity": item.get("severity"),
                        "requirement": item.get("requirement"),
                        "suggestion": item.get("suggestion"),
                        "auto_fixable": item.get("auto_fixable"),
                        "need_manual_review": item.get("need_manual_review"),
                    }
                    for item in (compliance.get("items") or [])
                    if isinstance(item, dict) and item.get("status") in {"fail", "warn"}
                ]
                # sort fatal/critical first
                _sev_order = {"fatal": 0, "critical": 1, "major": 2, "minor": 3, "info": 4}
                failed_items.sort(
                    key=lambda x: (
                        0 if x.get("status") == "fail" else 1,
                        _sev_order.get(str(x.get("severity") or ""), 9),
                        str(x.get("check_id") or ""),
                    )
                )
                compliance_summary = {
                    "exists": True,
                    "ok": bool(compliance.get("ok")),
                    "blocking": bool(compliance.get("blocking") or summary.get("blocking")),
                    "need_manual_review": bool(compliance.get("need_manual_review") or summary.get("need_manual_review")),
                    "max_severity": compliance.get("max_severity") or summary.get("max_severity"),
                    "phase": compliance.get("phase") or "",
                    "counts": counts,
                    "failed_items": failed_items,
                }
        except Exception:
            compliance_summary = {"exists": True, "ok": False, "blocking": False, "error": "read_failed"}

    materials_summary: dict[str, Any] = {
        "exists": False,
        "total": 0,
        "deferred": 0,
        "ready": 0,
        "waived": 0,
    }
    try:
        from materials_checklist import load_materials_checklist

        checklist = load_materials_checklist(root)
        summary = checklist.get("summary") if isinstance(checklist.get("summary"), dict) else {}
        materials_path = root / "workspace" / "materials_checklist.json"
        materials_summary = {
            "exists": materials_path.exists(),
            "total": int(summary.get("total") or 0),
            "deferred": int(summary.get("deferred") or 0),
            "ready": int(summary.get("ready") or 0),
            "waived": int(summary.get("waived") or 0),
        }
    except Exception:
        pass

    # Unified runtime view (single aggregator for goal/activity/repair/pipeline)
    runtime: dict[str, Any] = {}
    if not v2_read_only:
        try:
            from agent.runtime_status import build_runtime_status

            runtime = build_runtime_status(root, reevaluate_goal=False)
        except Exception as exc:
            runtime = {"ok": False, "message": str(exc), "warnings": [], "consistent": True}

    goal_view: dict[str, Any] | None = None
    goal_full: dict[str, Any] | None = None
    try:
        if v2_read_only:
            g = ControlStore(WorkspaceContext.resolve(root.parent, root.name)).goal_state()
            goal_summary = lambda value: str(value.get("summary") or "")  # noqa: E731
        else:
            from agent.goal import load_goal, goal_summary

            g = load_goal(root)
        if g:
            goal_full = g
            goal_view = {
                "goal_id": g.get("goal_id"),
                "status": g.get("status"),
                "all_criteria_ok": g.get("all_criteria_ok"),
                "blocked_reason": g.get("blocked_reason") or "",
                "raw_user_goal": str(g.get("raw_user_goal") or "")[:200],
                "summary": goal_summary(g),
                "plan": g.get("plan") or [],
                "criteria_results": g.get("criteria_results") or [],
                "progress": g.get("progress") or {},
            }
    except Exception:
        pass

    return {
        **status,
        "workflow": workflow,
        "next_step": next_step,
        "blocked_step": blocked_step,
        "compliance_summary": compliance_summary,
        "materials_summary": materials_summary,
        "goal": goal_view,
        "goal_full": goal_full,
        "runtime": runtime,
        "consistency_warnings": list(runtime.get("warnings") or []),
        "product_mode": runtime.get("product_mode") or "",
        "product_mode_label": runtime.get("product_mode_label") or "",
        "consistent": bool(runtime.get("consistent", True)),
    }


@app.get("/api/status")
def api_status() -> dict[str, Any]:
    root = _active_root()
    return _status_payload(root, ACTIVE_RUN_ID or root.name)


@app.get("/api/v2/workspaces/{workspace_id}/workflow-step-detail")
@app.get("/api/workflow-step-detail")
def api_workflow_step_detail(command: str = Query(..., min_length=1), workspace_id: str = "") -> JSONResponse:
    context = _workspace_context(workspace_id) if workspace_id else None
    root = context.root if context is not None else _active_root()
    step = next((item for item in WORKFLOW_STEPS if item.get("command") == command), None)
    if step is None:
        return JSONResponse({"ok": False, "message": f"未知流程节点: {command}"}, status_code=404)

    status = _status_payload(
        root,
        workspace_id or ACTIVE_RUN_ID or root.name,
        persist_manual_review_summary=context is None,
    )
    workflow = status.get("workflow", []) if isinstance(status, dict) else []
    step_status = next((item for item in workflow if isinstance(item, dict) and item.get("command") == command), {})
    timings = status.get("timings", {}) if isinstance(status.get("timings"), dict) else {}
    return JSONResponse(
        {
            "ok": True,
            "step": step_status or step,
            "summary": _step_detail_summary(root, command),
            "details": _step_extra_details(root, command),
            "requires": _artifact_payloads(root, list(step.get("requires", []))),
            "produces": _artifact_payloads(root, list(step.get("produces", []))),
            "timing": timings.get(command, {}),
            "history": _step_history(root, command),
            "stage_metrics": load_stage_metrics(root).get(next((stage_name for stage_name, stage_command in STAGE_TO_COMMAND.items() if stage_command == command), command), {}),
            "agent_runs": _load_agent_runs(
                root,
                stage=next((stage_name for stage_name, stage_command in STAGE_TO_COMMAND.items() if stage_command == command), command),
                limit=12,
            ),
            "budget_hits": _budget_hits_for_command(root, command),
            "prompt_summary": _stage_prompt_summary(root, command),
            "manual_review_summary": (
                _v2_manual_review_summary(context)
                if context is not None
                else manual_review_summary(root)
            ),
            "project_profile": load_project_profile(root),
            "run_root": str(root),
        }
    )


def _minimal_repair_candidates(
    root: Path,
    *,
    context: WorkspaceContext | None = None,
) -> list[dict[str, Any]]:
    """Return actionable blocking issues for a chat-based minimal-repair prompt."""
    if context is not None:
        issues = _ensure_v2_issue_import(context).issue_states()
        return [
            issue for issue in issues
            if str(issue.get("status")) in {"open", "in_progress"}
            and str(issue.get("severity")) == "block"
        ]
    try:
        from agent.issues import load_open_issues
        from agent.root_cause import sync_issues_from_compliance, sync_issues_from_global_review

        sync_issues_from_global_review(root)
        sync_issues_from_compliance(root)
        return [
            issue for issue in load_open_issues(root)
            if str(issue.get("status")) in {"open", "in_progress"}
            and str(issue.get("severity")) == "block"
        ]
    except Exception:
        return []


def _minimal_repair_resume_command(root: Path) -> str:
    commands = auto_run_commands()
    pipeline_command = str(SUPERVISOR.load(root).get("current_stage") or "")
    if pipeline_command in commands:
        return pipeline_command
    state_command = _command_for_stage(str(_read_run_state(root).get("stage") or ""))
    if state_command in commands:
        return state_command
    for command in commands:
        try:
            if not stage_outputs_ready(root, stage_spec_by_command(command).id):
                return command
        except Exception:
            continue
    return commands[-1] if commands else ""


def _issue_repair_fingerprint(issue: dict[str, Any]) -> str:
    try:
        from agent.repair import issue_fingerprint

        return str(issue_fingerprint(issue))
    except Exception:
        target = issue.get("target") if isinstance(issue.get("target"), dict) else {}
        target_key = json.dumps(target, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
        return f"{issue.get('stage_id', '')}|{issue.get('code', '')}|{target_key}"


def _issue_has_auto_repair(issue: dict[str, Any]) -> bool:
    auto_types = {"fix_compliance", "fix_coverage", "rewrite_chapters", "rerun_stage", "regenerate_artifact"}
    evidence = issue.get("evidence") if isinstance(issue.get("evidence"), dict) else {}
    if bool(evidence.get("need_manual_review")) and evidence.get("auto_fixable") is not True:
        return False
    actions = issue.get("suggested_actions") if isinstance(issue.get("suggested_actions"), list) else []
    return any(isinstance(action, dict) and str(action.get("type") or "") in auto_types for action in actions)


def _ensure_minimal_repair_confirmation(
    root: Path,
    *,
    context: WorkspaceContext | None = None,
) -> dict[str, Any]:
    issues = _minimal_repair_candidates(root, context=context)
    if not issues:
        current = load_repair_job(root)
        if str(current.get("status") or "") == "awaiting_confirmation":
            update_repair_job(
                root,
                str(current.get("job_id") or ""),
                status="completed",
                phase="complete",
                resolved_count=int(current.get("total_count") or 0),
                remaining_count=0,
                message="阻断问题已不存在，无需执行最小修复",
                result={"no_longer_blocked": True},
            )
        return {}
    auto_count = sum(1 for issue in issues if _issue_has_auto_repair(issue))
    return create_confirmation(
        root,
        issue_fingerprints=[_issue_repair_fingerprint(issue) for issue in issues],
        total_count=len(issues),
        auto_count=auto_count,
        manual_count=max(0, len(issues) - auto_count),
        resume_command=_minimal_repair_resume_command(root),
    )


def _persistent_minimal_repair_prompt(root: Path) -> tuple[str, list[dict[str, Any]]]:
    current = load_repair_job(root)
    if str(current.get("status") or "") in {"completed", "partial", "failed"}:
        return "", []
    job = _ensure_minimal_repair_confirmation(root)
    if not job or str(job.get("status") or "") != "awaiting_confirmation":
        return "", []
    count = int(job.get("total_count") or 0)
    auto_count = int(job.get("auto_count") or 0)
    manual_count = int(job.get("manual_count") or 0)
    return (
        f"发现 {count} 个阻断问题，其中预计 {auto_count} 项可自动处理、{manual_count} 项需人工处理。"
        "是否执行最小修复？确认后会自动修改、强制重验，并尝试从原阻塞步骤继续一次。",
        [
            {
                "type": "confirm_minimal_repair",
                "label": "是，自动最小修复",
                "confirmation_id": str(job.get("confirmation_id") or ""),
                "job_id": str(job.get("job_id") or ""),
            },
            {
                "type": "decline_minimal_repair",
                "label": "否，暂不修复",
                "confirmation_id": str(job.get("confirmation_id") or ""),
                "job_id": str(job.get("job_id") or ""),
            },
        ],
    )


def _minimal_repair_intent(message: str, *, has_pending: bool) -> str:
    normalized = re.sub(r"[\s，。！？!、,.？]", "", message.strip().lower())
    if not normalized:
        return ""
    negative_exact = {"否", "不用", "暂不", "取消", "不需要", "先不", "不要"}
    negative_repair = re.search(
        r"(?:不|别|勿|无需|不用|暂不|取消)(?:要|用|需要|执行|进行|自动|最小|再|先){0,3}(?:修复|处理)",
        normalized,
    )
    if normalized in negative_exact or negative_repair:
        return "decline" if has_pending else "query"
    if re.search(r"(?:怎么修|如何修|为什么失败|失败原因|查看(?:问题|详情|报告|修复计划)|修复会改什么|哪些问题)", normalized):
        return "query"
    if has_pending and (
        normalized in {"是", "好的", "好", "确认", "同意", "需要", "可以", "执行"}
        or normalized.startswith(("是的", "确认执行", "同意执行"))
    ):
        return "confirm"
    if normalized in {"修复", "修复吧"}:
        return "confirm" if has_pending else "start"
    # Restart / resume synonyms after interrupt or failed job (must not fall through to Supervisor)
    if re.search(
        r"(?:继续修复|重新修复|重新发起|再修一次|重试修复|再自动修复|重启修复|"
        r"重新发起最小修复|重新自动修复|再次修复|继续自动修复)",
        normalized,
    ):
        return "start"
    if re.search(
        r"(?:自动修复|最小修复|修复啊|修一下|开始修复|立即修复|赶紧修复|一键修复|"
        r"帮我.{0,8}修复|执行.{0,6}修复|修复这些|处理.{0,8}阻断)",
        normalized,
    ):
        return "confirm" if has_pending else "start"
    return ""


def _repair_result_count(result: dict[str, Any], key: str, count_key: str) -> int:
    value = result.get(key)
    if isinstance(value, list):
        return len(value)
    try:
        return int(result.get(count_key) or 0)
    except (TypeError, ValueError):
        return 0


def _trigger_repair_job(
    root: Path,
    confirmation_id: str,
    *,
    allow_remint: bool = True,
    control_operation_id: str = "",
    control_fencing_token: int = 0,
    resume_pipeline: bool = True,
) -> dict[str, Any]:
    """Claim exactly one repair slot and run the persisted job in the background.

    V2 calls use the confirmed control Operation as authorization. V1 calls use
    the compatibility confirmation token and may remint it after interruption.
    """
    global RUNNING, CURRENT_TASK, CURRENT_RUN_ID, CURRENT_RUN_ROOT, PAUSE_REQUESTED
    with _REPAIR_START_LOCK:
        current = load_v2_repair_job(root) if control_operation_id else load_repair_job(root)
        if str(current.get("status") or "") in RUNNING_REPAIR_STATUSES:
            return {"ok": True, "duplicate": True, "job": current, "message": current.get("message", "修复任务正在执行")}
        pipeline_status = str(SUPERVISOR.load(root).get("status") or "")
        pipeline_busy = pipeline_status in {"running", "recovering", "retrying", "pausing"}
        if RUNNING or SUPERVISOR.is_running() or pipeline_busy:
            return {"ok": False, "busy": True, "job": current, "message": "当前已有任务正在运行，修复确认已保留，请稍后重试"}
        claimed = (
            claim_repair_job_authorized(root, control_operation_id)
            if control_operation_id
            else claim_repair_job(root, confirmation_id)
        )
        if not claimed.get("ok"):
            # Stale terminal job or invalid confirmation → remint and claim once
            if not control_operation_id and allow_remint and (
                claimed.get("stale")
                or str(current.get("status") or "") in TERMINAL_REPAIR_STATUSES
                or "失效" in str(claimed.get("message") or "")
            ):
                fresh = _ensure_minimal_repair_confirmation(root)
                fresh_id = str(fresh.get("confirmation_id") or "")
                if not fresh_id:
                    return {
                        "ok": False,
                        "job": fresh or current,
                        "message": "当前没有可自动修复的阻断问题",
                    }
                claimed = claim_repair_job(root, fresh_id)
            if not claimed.get("ok"):
                return {
                    "ok": False,
                    "job": claimed.get("job") or current,
                    "message": claimed.get("message", "无法确认修复任务"),
                }
        job = claimed.get("job") if isinstance(claimed.get("job"), dict) else {}
        if claimed.get("duplicate"):
            return {"ok": True, "duplicate": True, "job": job, "message": job.get("message", "已处理该修复确认")}
        job_id = str(job.get("job_id") or "")
        job_run_id = ACTIVE_RUN_ID or root.name
        RUNNING = True
        CURRENT_TASK = "minimal-repair"
        CURRENT_RUN_ID = job_run_id
        CURRENT_RUN_ROOT = root
        PAUSE_REQUESTED = False

    def _progress(phase: str, progress: dict[str, Any] | None = None) -> None:
        details = progress if isinstance(progress, dict) else {}
        status = "revalidating" if phase in {"revalidate", "revalidating"} else "running"
        message = str(details.get("message") or "")
        if not message:
            message = {
                "analysis": "正在分析阻断问题并合并根因动作",
                "edit": "正在执行根因修复",
                "revalidate": "正在强制重验质量门禁",
                "complete": "正在汇总修复结果",
            }.get(phase, "正在执行最小修复")
        completed = int(details.get("completed") or 0)
        total = int(details.get("total") or 0)
        ratio = (completed / total) if total > 0 else 0.0
        if phase == "analysis":
            percent = 10
        elif phase == "edit":
            percent = 10 + round(60 * ratio)
        elif phase in {"revalidate", "revalidating"}:
            percent = 70 + round(25 * ratio)
        elif phase == "complete":
            percent = 98
        else:
            percent = 5
        changes: dict[str, Any] = {
            "status": status,
            "phase": phase,
            "message": message,
            "phase_completed": completed,
            "phase_total": total,
            "progress_percent": max(0, min(98, percent)),
        }
        for key in ("total_count", "auto_count", "manual_count", "resolved_count", "remaining_count", "failed_count"):
            if key in details:
                changes[key] = details[key]
        update_repair_job(root, job_id, **changes)
        if control_operation_id:
            try:
                context = _workspace_context(root.name)
                ControlStore(context).sync_operation(
                    control_operation_id,
                    "running",
                    message=message,
                    fencing_token=control_fencing_token,
                )
            except Exception as exc:
                raise RuntimeError(f"repair control state sync failed: {exc}") from exc

    def _run() -> None:
        global RUNNING, CURRENT_TASK, CURRENT_RUN_ID, CURRENT_RUN_ROOT
        result: dict[str, Any]
        failure: Exception | None = None
        try:
            from agent.repair import execute_repair_batch

            repair_context = _workspace_context(root.name) if control_operation_id else None
            issues = _minimal_repair_candidates(root, context=repair_context)
            issue_ids = [str(issue.get("id") or "") for issue in issues if issue.get("id")]
            result = execute_repair_batch(
                root,
                issue_ids,
                confirm=True,
                dry_run=False,
                issue_snapshot=(repair_context and ControlStore(repair_context).issue_states()),
                progress_callback=_progress,
            )
            if control_operation_id:
                from artifact_manifest import record_external_chapter_mutation

                record_external_chapter_mutation(
                    _workspace_context(root.name),
                    disposition="issue_repair",
                )
        except Exception as exc:
            failure = exc
            result = {"ok": False, "failed": [str(exc)], "message": f"最小修复失败：{exc}"}

        resolved_count = _repair_result_count(result, "resolved", "resolved_count")
        remaining_count = _repair_result_count(result, "still_open", "remaining_count")
        manual_count = _repair_result_count(result, "manual", "manual_count")
        failed_count = _repair_result_count(result, "failed", "failed_count")
        if failure and failed_count == 0:
            failed_count = 1
        if failure:
            terminal_status = "failed"
        elif remaining_count or manual_count or failed_count:
            terminal_status = "partial"
        else:
            terminal_status = "completed"

        with _REPAIR_START_LOCK:
            RUNNING = False
            CURRENT_TASK = ""
            CURRENT_RUN_ID = ""
            CURRENT_RUN_ROOT = None

        resume_command = str(job.get("resume_command") or "")
        resume_started = False
        # Only auto-resume when fully clear — partial/failed jobs would hit the same gate again
        should_resume = (
            resume_pipeline
            and bool(resume_command)
            and terminal_status == "completed"
            and remaining_count == 0
            and failed_count == 0
        )
        if should_resume:
            try:
                # Prefer gate-aware API path when available so 409 messages surface in logs
                resume_started = SUPERVISOR.start(
                    job_run_id,
                    root,
                    _run_sync,
                    start_command=resume_command,
                )
                if not resume_started:
                    result["resume_error"] = "流水线已在运行或未能启动"
            except Exception as exc:
                result["resume_error"] = str(exc)
        result["resume_command"] = resume_command
        result["resume_started"] = resume_started
        result_message = str(result.get("message") or "").strip()
        summary = (
            f"最小修复结束：已解决 {resolved_count} 项，仍存在 {remaining_count} 项，"
            f"需人工 {manual_count} 项，失败 {failed_count} 项。"
        )
        if should_resume:
            if resume_started:
                summary += f" 已从 {resume_command} 尝试恢复流水线一次。"
            else:
                summary += f" 未能自动从 {resume_command} 恢复流水线，请点击「继续流水线」。"
                if result.get("resume_error"):
                    summary += f"（{result['resume_error']}）"
        elif resume_command and terminal_status != "completed":
            summary += " 仍有未关闭问题，未自动恢复流水线。"
        if result_message and result_message not in summary:
            summary = f"{summary}\n{result_message}"
        try:
            save_message(root, job_run_id, "assistant", summary, actions=[], kind="repair_result")
        except Exception:
            pass
        final_job = update_repair_job(
            root,
            job_id,
            status=terminal_status,
            phase="complete" if terminal_status == "completed" else terminal_status,
            resolved_count=resolved_count,
            remaining_count=remaining_count,
            manual_count=manual_count,
            failed_count=failed_count,
            progress_percent=100,
            resume_attempted=bool(should_resume),
            message=summary,
            result=result,
        )
        _append_log(f"[最小修复] {summary}")
        if not final_job:
            _append_log("[警告] 最小修复结果未能写入任务状态")
        if control_operation_id:
            operation_status = {
                "completed": "succeeded",
                "partial": "blocked",
                "failed": "failed",
            }.get(terminal_status, "failed")
            try:
                context = _workspace_context(root.name)
                ControlStore(context).sync_operation(
                    control_operation_id,
                    operation_status,
                    message=summary,
                    error=result if operation_status == "failed" else None,
                    fencing_token=control_fencing_token,
                )
            except Exception as exc:
                _append_log(f"[警告] 修复 Operation 状态回写失败: {exc}")

    worker_thread = threading.Thread(target=_run, daemon=True, name=f"repair-{job_id}")
    worker_thread.start()
    return {
        "ok": True,
        "duplicate": False,
        "job": job,
        "message": "已开始最小修复，将按根因合并处理并统一重验",
        "_worker_thread": worker_thread,
    }


@app.post("/api/chat")
async def api_chat(request: Request) -> JSONResponse:
    root = _active_root()
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    message = str(body.get("message", "")).strip()
    selected_command = str(body.get("selected_command", "")).strip()
    payload = _chat_reply(root, message, selected_command=selected_command)
    return JSONResponse({"ok": True, **payload})


def _load_review_context(root: Path, limit: int = 50) -> list[dict[str, Any]]:
    """加载各章 review 摘要，供编排器综合改稿目标。"""
    reviews_dir = root / "workspace" / "reviews"
    if not reviews_dir.exists():
        return []
    items: list[dict[str, Any]] = []
    for review_path in sorted(reviews_dir.glob("*_review.json"))[:limit]:
        try:
            data = read_json(review_path)
        except Exception:
            continue
        if not isinstance(data, dict):
            continue
        problems = data.get("problems", []) if isinstance(data.get("problems"), list) else []
        top = [
            {"type": str(p.get("type", "")), "description": str(p.get("description", ""))[:120]}
            for p in problems[:3]
            if isinstance(p, dict)
        ]
        items.append(
            {
                "chapter_id": str(data.get("chapter_id") or review_path.stem.replace("_review", "")),
                "need_rewrite": bool(data.get("need_rewrite", False)),
                "need_evidence": bool(data.get("need_evidence", False)),
                "stuck": bool(data.get("stuck")) or str(data.get("rewrite_status") or "") == "stuck",
                "rewrite_status": str(data.get("rewrite_status") or ""),
                "problem_count": len(problems),
                "top_problems": top,
            }
        )
    return items


def _trigger_rewrite_targets_inline(
    targets: list[dict[str, Any]],
    *,
    root: Path | None = None,
    run_id: str = "",
    control_operation_id: str = "",
    control_fencing_token: int = 0,
) -> dict[str, Any]:
    global RUNNING
    if not targets:
        return {"ok": False, "message": "没有定向改稿目标。"}
    if not str(control_operation_id or "").strip() or int(control_fencing_token or 0) <= 0:
        return {"ok": False, "message": "定向改稿缺少权威 Operation/fencing token，已拒绝执行。"}
    if RUNNING:
        return {"ok": False, "message": "当前已有任务正在运行，请等待完成。"}
    if root is None and ACTIVE_RUN_ROOT is None:
        return {"ok": False, "message": "请先创建本次运行工作空间。"}
    chapter_ids = [str(t.get("chapter_id")) for t in targets if t.get("chapter_id")]
    if not chapter_ids:
        return {"ok": False, "message": "改稿目标缺少 chapter_id。"}
    run_root = root or _active_root()
    resolved_run_id = run_id or ACTIVE_RUN_ID or run_root.name

    def _sync_control(status: str, message: str, error: Any = None) -> None:
        if not control_operation_id:
            return
        context = _workspace_context(run_root.name)
        ControlStore(context).sync_operation(
            control_operation_id,
            status,
            message=message,
            error=error,
            fencing_token=control_fencing_token,
        )

    def _run_rewrite_sync(chapters: list[str], worker_root: Path) -> None:
        global RUNNING, CURRENT_TASK, CURRENT_RUN_ID, CURRENT_RUN_ROOT, PAUSE_REQUESTED
        RUNNING = True
        CURRENT_TASK = "dispatch-rewrite"
        CURRENT_RUN_ID = resolved_run_id
        CURRENT_RUN_ROOT = worker_root
        PAUSE_REQUESTED = False
        operation_status = "failed"
        operation_error: dict[str, Any] | None = None
        state_status = "error"
        state_message = "定向改稿未启动。"
        try:
            save_run_state(
                worker_root,
                {"root_dir": str(worker_root), "current_command": "dispatch-rewrite"},
                stage="review-fix-all",
                status="running",
                message=f"定向改稿: {chapters}",
            )
            # The authoritative Operation must be writable before any artifact
            # mutation starts; otherwise the worker fails closed.
            _sync_control("running", f"正在定向改稿: {chapters}")
            _append_log(f"--- [{time.strftime('%H:%M:%S')}] 定向改稿: {chapters} ---")
            from subagent_runner import run_rewrite_all

            from concurrency import workers_default

            result = run_rewrite_all(worker_root, workers=workers_default(), chapter_ids=chapters)
            if control_operation_id:
                from artifact_manifest import record_external_chapter_mutation

                record_external_chapter_mutation(
                    _workspace_context(worker_root.name),
                    disposition="chapter_rewrite",
                )
            failed = result.get("failed", [])
            state_status = "ok" if not failed else "error"
            state_message = f"定向改稿完成: 成功 {len(result.get('completed', []))}, 失败 {len(failed)}"
            operation_status = "succeeded" if not failed else "blocked"
            if failed:
                operation_error = {"failed_chapters": failed}
        except Exception as exc:
            state_status = "error"
            state_message = f"定向改稿失败: {exc}"
            operation_error = {"message": str(exc)}
            _append_log(f"[错误] 定向改稿异常: {exc}")
        try:
            save_run_state(
                worker_root,
                {"root_dir": str(worker_root), "current_command": "dispatch-rewrite"},
                stage="review-fix-all",
                status=state_status,
                message=state_message,
            )
        except Exception as exc:
            operation_status = "failed"
            operation_error = {"message": f"run state write failed: {exc}"}
            state_message = f"定向改稿终态保存失败: {exc}"
            _append_log(f"[错误] {state_message}")
        finally:
            # Publish the terminal Operation only after the in-process worker
            # state is terminal too, so observers never see succeeded + RUNNING.
            RUNNING = False
            CURRENT_TASK = ""
            CURRENT_RUN_ID = ""
            CURRENT_RUN_ROOT = None
        try:
            _sync_control(operation_status, state_message, operation_error)
        except Exception as exc:
            _append_log(f"[警告] 改稿 Operation 终态回写失败: {exc}")

    threading.Thread(target=_run_rewrite_sync, args=(chapter_ids, run_root), daemon=True).start()
    return {"ok": True, "message": f"定向改稿已启动: {chapter_ids}"}


def _chat_response(
    root: Path,
    run_id: str,
    reply: str,
    *,
    actions: list[dict[str, Any]] | None = None,
    thinking: str = "",
    kind: str = "message",
    **extra: Any,
) -> JSONResponse:
    action_items = actions or []
    assistant = save_message(root, run_id, "assistant", reply, thinking, action_items, kind)
    payload: dict[str, Any] = {
        "ok": True,
        "reply": reply,
        "actions": action_items,
        "assistant": assistant,
        **extra,
    }
    if thinking:
        payload["thinking"] = thinking
    return JSONResponse(payload)


def _repair_retry_actions(job: dict[str, Any]) -> list[dict[str, Any]]:
    if str(job.get("status") or "") != "awaiting_confirmation":
        return []
    return [
        {
            "type": "confirm_minimal_repair",
            "label": "是，自动最小修复",
            "confirmation_id": str(job.get("confirmation_id") or ""),
            "job_id": str(job.get("job_id") or ""),
        },
        {
            "type": "decline_minimal_repair",
            "label": "否，暂不修复",
            "confirmation_id": str(job.get("confirmation_id") or ""),
            "job_id": str(job.get("job_id") or ""),
        },
    ]


@app.post("/api/chat/orchestrate")
async def api_chat_orchestrate(request: Request) -> JSONResponse:
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON"}, status_code=400)
    if not isinstance(body, dict):
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON 对象"}, status_code=400)
    sent_run_id = str(body.get("run_id") or "").strip()
    is_v2_workspace_chat = bool(getattr(request, "v2_workspace_id", ""))
    if sent_run_id:
        try:
            context = _workspace_context(sent_run_id)
        except ControlPlaneError as exc:
            return _command_error_response(exc)
        root = context.root
        run_id = context.workspace_id
    else:
        # One-version compatibility for callers not yet sending run_id.
        root = _active_root()
        run_id = ACTIVE_RUN_ID or root.name
    action = body.get("action") if isinstance(body.get("action"), dict) else {}
    message = str(body.get("message") or "").strip()
    action_type = str(action.get("type") or "").strip()
    if not message and action_type == "confirm_minimal_repair":
        message = "是，执行最小修复"
    elif not message and action_type == "decline_minimal_repair":
        message = "否，暂不修复"

    async with _chat_turn_lock(root, run_id):
        history = load_messages(root, run_id, limit=20)
        if message:
            save_message(root, run_id, "user", message, actions=[], kind="message")
        if not message:
            return _chat_response(
                root,
                run_id,
                "可以直接告诉我你想做什么，例如查看状态、继续流程或自动修复阻断问题。",
                actions=[],
            )

        # V2 has no compatibility repair-confirmation token.  Fresh V2 repair
        # proposals use ActionProposal/repair.start; never let a stale legacy
        # chat button mutate repair_job.json through the V2 chat route.
        if is_v2_workspace_chat and action_type in {"confirm_minimal_repair", "decline_minimal_repair"}:
            return _chat_response(
                root,
                run_id,
                "该旧修复确认已不适用于 V2。请重新发送“自动修复”，系统会创建新的 V2 确认操作。",
                actions=[],
                intent="legacy_repair_confirmation_rejected",
            )

        control_text = re.sub(r"\s+", "", message)
        pause_intent = control_text in {"暂停", "暂停流程", "暂停任务", "先暂停", "停一下"}
        cancel_intent = control_text in {"取消", "取消流程", "取消任务", "终止流程", "终止任务"}
        if pause_intent or cancel_intent:
            try:
                context = _workspace_context(run_id)
                gateway = _command_gateway(context)
                store = gateway.store
                envelope = CommandEnvelope.from_mapping(
                    {
                        "kind": "pipeline.pause" if pause_intent else "pipeline.cancel",
                        "payload": {},
                        "expected_revision": store.revision(),
                        "idempotency_key": str(body.get("idempotency_key") or f"chat-control:{uuid.uuid4()}"),
                        "actor": _request_actor(request, source="chat"),
                    },
                    workspace_id=run_id,
                )
                if cancel_intent:
                    action_proposal = gateway.propose(envelope, label="确认取消当前任务", risk="high")
                    decline_proposal = {
                        **action_proposal,
                        "type": "decline_v2_command",
                        "label": "保留任务，不取消",
                    }
                    return _chat_response(
                        root,
                        run_id,
                        "取消会终止当前 Operation，且不能从正在写入的中间位置继续。请确认是否取消。",
                        actions=[action_proposal, decline_proposal],
                        intent="pipeline_cancel_confirmation",
                    )
                receipt = gateway.submit(envelope)
                return _chat_response(
                    root,
                    run_id,
                    receipt.message or "已发送暂停指令。",
                    actions=[],
                    intent="pipeline_control",
                    command_receipts=[receipt.as_dict()],
                    triggered_auto_run=False,
                )
            except ControlPlaneError as exc:
                return _chat_response(
                    root,
                    run_id,
                    exc.message,
                    actions=[],
                    intent="pipeline_control_error",
                    command_error=exc.as_dict(),
                )

        current_job = load_v2_repair_job(root) if is_v2_workspace_chat else load_repair_job(root)
        job_status = str(current_job.get("status") or "")
        has_pending = job_status == "awaiting_confirmation"
        terminal_repair = job_status in {"completed", "partial", "failed"}
        interrupted_repair = (
            terminal_repair
            and (
                str(current_job.get("phase") or "") == "interrupted"
                or "服务重启中断" in str(current_job.get("message") or "")
            )
        )
        intent = _minimal_repair_intent(message, has_pending=has_pending)
        if action_type == "confirm_minimal_repair":
            intent = "confirm"
        elif action_type == "decline_minimal_repair":
            intent = "decline"
        elif action_type == "restart_minimal_repair":
            intent = "start"
        # Failed/interrupted jobs: "confirm" with old confirmation_id cannot resume — force restart
        if intent == "confirm" and (terminal_repair or interrupted_repair):
            intent = "start"

        if intent == "decline" and has_pending:
            confirmation_id = str(action.get("confirmation_id") or current_job.get("confirmation_id") or "")
            job = decline_repair_job(root, confirmation_id)
            return _chat_response(
                root,
                run_id,
                "好的，暂不执行最小修复。阻断问题会保留，之后仍可直接说“自动修复”。",
                actions=[],
                intent="decline_minimal_repair",
                repair_job=job or None,
            )

        if intent in {"confirm", "start"}:
            try:
                candidates = _minimal_repair_candidates(
                    root,
                    context=_workspace_context(run_id) if is_v2_workspace_chat else None,
                )
            except ControlPlaneError as exc:
                return _chat_response(
                    root,
                    run_id,
                    exc.message,
                    actions=[],
                    intent="minimal_repair_error",
                    triggered_repair=False,
                    command_error=exc.as_dict(),
                )
            if not candidates:
                return _chat_response(
                    root,
                    run_id,
                    "当前没有可自动修复的阻断问题（质量问题单为空或均已关闭）。"
                    "若仍有材料缺口，请先到「材料」页补料；其它阻断请在「问题」页查看。",
                    actions=[],
                    intent="minimal_repair",
                    triggered_repair=False,
                    repair_job=None,
                )
            try:
                context = _workspace_context(run_id)
                gateway = _command_gateway(context)
                envelope = CommandEnvelope.from_mapping(
                    {
                        "kind": "repair.start",
                        "payload": {
                            "issue_ids": [str(item.get("id") or "") for item in candidates if item.get("id")],
                        },
                        "expected_revision": gateway.store.revision(),
                        "idempotency_key": str(body.get("idempotency_key") or f"chat-repair:{uuid.uuid4()}"),
                        "actor": _request_actor(request, source="chat"),
                    },
                    workspace_id=run_id,
                )
                proposal = gateway.propose(envelope, label="确认执行最小修复", risk="high")
                return _chat_response(
                    root,
                    run_id,
                    f"将对 {len(candidates)} 个阻断问题执行最小修复并重新验证门禁，请确认。",
                    actions=[
                        proposal,
                        {**proposal, "type": "decline_v2_command", "label": "暂不修复"},
                    ],
                    intent="minimal_repair_confirmation",
                    triggered_repair=False,
                )
            except ControlPlaneError as exc:
                return _chat_response(
                    root,
                    run_id,
                    exc.message,
                    actions=[],
                    intent="minimal_repair_error",
                    triggered_repair=False,
                    command_error=exc.as_dict(),
                )

        status = _status_payload(
            root,
            run_id,
            # Chat queries are read-only in V2.  The compatibility summary
            # writer remains available only to the legacy chat adapter.
            persist_manual_review_summary=not is_v2_workspace_chat,
            v2_read_only=is_v2_workspace_chat,
        )
        review_context = _load_review_context(root)
        # Frontend confirm buttons: tool_scope only when tool is present (PR-1)
        tool_name = str(action.get("tool") or "").strip()
        confirmed_tools: list[str] = []
        if tool_name:
            confirmed_tools.append(tool_name)
        # Also accept tool name from args when action.command is a stage command
        stage_from_action = str(
            action.get("command")
            or (action.get("args") or {}).get("command")
            or (action.get("args") or {}).get("start_command")
            or ""
        ).strip()
        user_confirmed = bool(action.get("user_confirmed")) or action_type in {
            "confirm_tool",
            "confirm_execute",
        }
        # Without explicit tool, do not grant all_mutations via bare user_confirmed
        if user_confirmed and not tool_name and action_type in {"confirm_tool", "confirm_execute"}:
            # only allow all_mutations when label explicitly says so
            label = str(action.get("label") or message or "")
            if not any(k in label for k in ("全部剩余", "全部操作", "所有剩余", "确认执行全部")):
                user_confirmed = False

        # Fast path: confirm run_stage / continue whole pipeline → start backend pipeline
        # instead of re-entering supervisor confirm loop (product control, not chat).
        resume_phrases = (
            "继续",
            "继续啊",
            "继续吧",
            "继续执行",
            "继续进行",
            "继续整个",
            "继续整个流程",
            "整体推进",
            "一键跑完",
            "跑完剩余",
            "接着跑",
            "继续流程",
            "启动流水线",
        )
        msg_compact = re.sub(r"\s+", "", message or "")
        phrase_continue = (
            msg_compact in resume_phrases
            or (
                len(msg_compact) <= 20
                and any(msg_compact == p or msg_compact.startswith(p) for p in resume_phrases)
            )
        )
        confirm_run_stage = (
            action_type == "confirm_tool"
            and tool_name in {"run_stage", "run_pipeline_remaining", ""}
        ) or (
            action_type == "confirm_tool"
            and stage_from_action in set(auto_run_commands())
        )
        take_pipeline_fast_path = (
            confirm_run_stage
            or action_type in {"auto_run", "start_pipeline"}
            or (phrase_continue and action_type in {"", "confirm_tool", "auto_run"})
            or (user_confirmed and tool_name in {"run_stage", "run_pipeline_remaining"})
        )
        if take_pipeline_fast_path:
            auto_cmds = set(auto_run_commands())
            start_cmd = stage_from_action if stage_from_action in auto_cmds else ""
            if not start_cmd:
                try:
                    start_cmd = _minimal_repair_resume_command(root)
                    if start_cmd not in auto_cmds:
                        start_cmd = ""
                except Exception:
                    start_cmd = ""
            try:
                context = _workspace_context(run_id)
                gateway = _command_gateway(context)
                snapshot = gateway.store.snapshot()
                active_operation = snapshot.get("operation") if isinstance(snapshot.get("operation"), dict) else {}
                active_status = str(active_operation.get("status") or "")
                if active_status in {"running", "queued", "pausing", "cancelling"}:
                    return _chat_response(
                        root,
                        run_id,
                        "流水线已在运行，请稍候查看进度。",
                        actions=[],
                        intent="pipeline_control",
                        triggered_auto_run=False,
                        pipeline_already_running=True,
                    )
                command_kind = "pipeline.resume" if active_status in {"paused", "blocked"} else "pipeline.start"
                payload: dict[str, Any] = {"start_command": start_cmd or ""}
                if command_kind == "pipeline.resume" and active_operation.get("operation_id"):
                    payload["operation_id"] = active_operation["operation_id"]
                envelope = CommandEnvelope.from_mapping(
                    {
                        "kind": command_kind,
                        "payload": payload,
                        "expected_revision": int(snapshot.get("revision") or 0),
                        "idempotency_key": str(body.get("idempotency_key") or f"chat-continue:{uuid.uuid4()}"),
                        "actor": _request_actor(request, source="chat"),
                    },
                    workspace_id=run_id,
                )
                receipt = gateway.submit(envelope)
                if receipt.status != "rejected":
                    stage_label = WORKFLOW_COMMAND_LABELS.get(start_cmd, start_cmd) if start_cmd else "当前进度"
                    return _chat_response(
                        root,
                        run_id,
                        f"已通过统一命令入口从「{stage_label}」继续流水线。进度见下方执行计划。",
                        actions=[],
                        intent="pipeline_control",
                        triggered_auto_run=True,
                        triggered_command=start_cmd or "",
                        command_receipts=[receipt.as_dict()],
                    )
                reply = receipt.message or "流水线未能启动。"
                repair_prompt, repair_actions = (
                    ("", []) if is_v2_workspace_chat else _persistent_minimal_repair_prompt(root)
                )
                if repair_prompt:
                    reply = f"{reply}\n\n{repair_prompt}"
                return _chat_response(
                    root,
                    run_id,
                    reply,
                    actions=repair_actions,
                    intent="pipeline_blocked",
                    command_receipts=[receipt.as_dict()],
                )
            except ControlPlaneError as exc:
                return _chat_response(
                    root,
                    run_id,
                    exc.message,
                    actions=[],
                    intent="pipeline_control_error",
                    command_error=exc.as_dict(),
                )

        try:
            plan_result = await run_in_threadpool(
                orchestrator_plan,
                message,
                history,
                status,
                review_context=review_context,
                root=root,
                user_confirmed=user_confirmed,
                confirmed_tools=confirmed_tools or None,
            )
        except Exception as exc:
            plan_result = {
                "reply": f"对话编排暂时失败：{exc}",
                "actions": [],
                "error": str(exc),
            }
        resolved = orchestrator_resolve(plan_result, status)

        trigger_command = str(resolved.get("trigger_command") or "").strip()
        trigger_auto_run = bool(resolved.get("trigger_auto_run", False))
        trigger_rewrite_targets = (
            resolved.get("trigger_rewrite_targets", [])
            if isinstance(resolved.get("trigger_rewrite_targets"), list)
            else []
        )
        actions = resolved.get("actions", []) if isinstance(resolved.get("actions"), list) else []
        reply = str(resolved.get("reply") or "").strip()

        execution_notes: list[str] = []
        if trigger_command:
            if trigger_command in set(auto_run_commands()):
                try:
                    context = _workspace_context(run_id)
                    gateway = _command_gateway(context)
                    envelope = CommandEnvelope.from_mapping(
                        {
                            "kind": "pipeline.run_stage",
                            "payload": {"start_command": trigger_command},
                            "expected_revision": gateway.store.revision(),
                            "idempotency_key": str(body.get("idempotency_key") or f"chat-stage:{uuid.uuid4()}"),
                            "actor": _request_actor(request, source="chat"),
                        },
                        workspace_id=run_id,
                    )
                    receipt = gateway.submit(envelope)
                    label = WORKFLOW_COMMAND_LABELS.get(trigger_command, trigger_command)
                    note = (
                        f"已通过统一命令入口启动「{label}」。"
                        if receipt.status != "rejected"
                        else f"启动失败：{receipt.message}"
                    )
                    execution_notes.append(note)
                    if note not in reply:
                        reply = f"{reply}\n\n{note}".strip() if reply else note
                    if receipt.status == "rejected":
                        actions = [item for item in actions if item.get("command") != trigger_command]
                except ControlPlaneError as exc:
                    note = f"启动失败：{exc.message}"
                    execution_notes.append(note)
                    if note not in reply:
                        reply = f"{reply}\n\n{note}".strip() if reply else note
            else:
                note = f"「{trigger_command}」尚未迁入统一命令入口，本轮未执行。"
                execution_notes.append(note)
                if note not in reply:
                    reply = f"{reply}\n\n{note}".strip() if reply else note

        rewrite_proposed = False
        if trigger_rewrite_targets:
            try:
                context = _workspace_context(run_id)
                gateway = _command_gateway(context)
                envelope = CommandEnvelope.from_mapping(
                    {
                        "kind": "rewrite.chapters",
                        "payload": {"targets": trigger_rewrite_targets},
                        "expected_revision": gateway.store.revision(),
                        "idempotency_key": str(body.get("idempotency_key") or f"chat-rewrite:{uuid.uuid4()}"),
                        "actor": _request_actor(request, source="chat"),
                    },
                    workspace_id=run_id,
                )
                proposal = gateway.propose(envelope, label="确认执行定向改稿", risk="high")
                actions = [
                    proposal,
                    {**proposal, "type": "decline_v2_command", "label": "暂不改稿"},
                    *actions,
                ]
                ids = ", ".join(
                    str(item.get("chapter_id") or "")
                    for item in trigger_rewrite_targets
                    if item.get("chapter_id")
                )
                note = f"定向改稿将影响：{ids}。请确认后执行。"
                rewrite_proposed = True
            except ControlPlaneError as exc:
                note = f"定向改稿提案失败：{exc.message}"
            execution_notes.append(note)
            if note not in reply:
                reply = f"{reply}\n\n{note}".strip() if reply else note

        if trigger_auto_run:
            actions = [item for item in actions if item.get("type") != "auto_run"]
            actions = [{"type": "auto_run", "label": "一键跑完剩余步骤"}, *actions]

        repair_prompt, repair_actions = (
            ("", []) if is_v2_workspace_chat else _persistent_minimal_repair_prompt(root)
        )
        if repair_prompt and not has_pending:
            if repair_prompt not in reply:
                reply = f"{reply}\n\n{repair_prompt}".strip() if reply else repair_prompt
            actions = [*actions, *repair_actions]

        thinking = str(
            plan_result.get("thinking")
            or plan_result.get("reasoning")
            or resolved.get("thinking")
            or ""
        ).strip()
        response_repair_job = load_v2_repair_job(root) if is_v2_workspace_chat else load_repair_job(root)
        extra: dict[str, Any] = {
            "action": resolved.get("action", "chat"),
            "intent": resolved.get("intent", ""),
            "auto_execute": bool(resolved.get("auto_execute", False)),
            "triggered_command": trigger_command,
            "triggered_auto_run": trigger_auto_run,
            "triggered_rewrite": False,
            "rewrite_proposed": rewrite_proposed,
            "triggered_repair": False,
            "job_id": str(response_repair_job.get("job_id") or ""),
            "repair_job": response_repair_job or None,
        }
        if plan_result.get("error"):
            # 仅作决策过程备注，不再单独弹「编排器」系统消息
            extra["execution_note"] = str(plan_result.get("error"))
            extra["orchestrator_note"] = plan_result.get("error")
        elif execution_notes:
            extra["execution_note"] = "；".join(execution_notes)
        if plan_result.get("supervisor") or plan_result.get("supervisor_steps"):
            extra["supervisor"] = True
            extra["supervisor_steps"] = plan_result.get("supervisor_steps") or []
            if plan_result.get("goal_id"):
                extra["goal_id"] = plan_result.get("goal_id")
            if isinstance(plan_result.get("goal"), dict):
                extra["goal"] = plan_result.get("goal")
        return _chat_response(
            root,
            run_id,
            reply,
            actions=actions,
            thinking=thinking,
            **extra,
        )


@app.post("/api/v2/workspaces/{workspace_id}/chat/turn")
async def api_v2_chat_turn(workspace_id: str, request: Request) -> JSONResponse:
    try:
        _workspace_context(workspace_id)
        body = await request.json()
        if not isinstance(body, dict):
            raise ControlPlaneError("COMMAND_INVALID", "请求体必须是 JSON 对象。", status_code=400)
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON 对象。"}, status_code=400)

    class _WorkspaceChatRequest:
        def __init__(self) -> None:
            # Preserve the authenticated server context so downstream Command
            # proposals bind the same actor as buttons and CLI/API requests.
            self.state = getattr(request, "state", None)
            self.v2_workspace_id = workspace_id

        async def json(self) -> dict[str, Any]:
            return {**body, "run_id": workspace_id}

    return await api_chat_orchestrate(_WorkspaceChatRequest())  # type: ignore[arg-type]


@app.get("/api/repair-jobs/current")
def api_current_repair_job() -> JSONResponse:
    return JSONResponse({"ok": True, "repair_job": load_repair_job(_active_root()) or None})


@app.get("/api/v2/workspaces/{workspace_id}/export-preflight")
@app.get("/api/export-preflight")
def api_export_preflight(workspace_id: str = "") -> JSONResponse:
    try:
        if workspace_id:
            return JSONResponse(_v2_export_preflight(_workspace_context(workspace_id)))
        from agent.issues import export_preflight

        return JSONResponse(export_preflight(_active_root()))
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.get("/api/issues/metrics")
def api_issue_metrics() -> JSONResponse:
    root = _active_root()
    try:
        from agent.issues import load_issue_metrics, issues_summary

        return JSONResponse(
            {
                "ok": True,
                "summary": issues_summary(root),
                "metrics": load_issue_metrics(root),
            }
        )
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.get("/api/concurrency")
def api_concurrency_metrics() -> JSONResponse:
    """Process-wide worker / LLM concurrency telemetry (PR-A0)."""
    try:
        from concurrency import concurrency_snapshot

        return JSONResponse({"ok": True, "concurrency": concurrency_snapshot()})
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.get("/api/agent/flags")
def api_agent_flags() -> JSONResponse:
    """Expose runtime Agent mode for UI badge (PR-A3)."""
    try:
        from api.agent import agent_mode_payload

        return JSONResponse(agent_mode_payload())
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.post("/api/materials-checklist/verify")
async def api_materials_verify(request: Request) -> JSONResponse:
    """Compatibility adapter: route verification through the V2 gateway."""
    try:
        body = await request.json()
    except Exception:
        body = {}
    try:
        context = _workspace_context(ACTIVE_RUN_ID)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "materials.verify",
                "payload": body if isinstance(body, dict) else {},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-material-verify:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        receipt = gateway.submit(envelope)
        return JSONResponse(
            {"ok": receipt.status != "rejected", "receipt": receipt.as_dict(), "message": receipt.message},
            status_code=202 if receipt.status != "rejected" else 409,
            headers={"Deprecation": "true", "Link": f'</api/v2/workspaces/{context.workspace_id}/commands>; rel="successor-version"'},
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/materials-checklist/confirm-verify")
async def api_materials_confirm_verify(request: Request) -> JSONResponse:
    """Compatibility adapter: create a persisted V2 confirmation proposal."""
    try:
        body = await request.json()
    except Exception:
        body = {}
    try:
        context = _workspace_context(ACTIVE_RUN_ID)
        gateway = _command_gateway(context)
        payload = dict(body) if isinstance(body, dict) else {}
        payload.pop("operator", None)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "materials.confirm_verification",
                "payload": payload,
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-material-confirm:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        action = gateway.propose(envelope, label="确认材料人工核验结论", risk="high")
        return JSONResponse(
            {"ok": True, "status": "requires_confirmation", "action": action},
            status_code=202,
            headers={"Deprecation": "true", "Link": f'</api/v2/workspaces/{context.workspace_id}/commands>; rel="successor-version"'},
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.get("/api/v2/workspaces/{workspace_id}/issues")
@app.get("/api/issues")
def api_list_issues(status: str = "open", workspace_id: str = "") -> JSONResponse:
    """List quality issues (open snapshot)."""
    context = _workspace_context(workspace_id) if workspace_id else None
    root = context.root if context else _active_root()
    try:
        if context:
            from agent.issues import quality_gate_mode

            store = ControlStore(context)
            issue_import_pending = (
                store.issue_v1_import_pending()
                and (context.root / "workspace" / "issues" / "open.json").exists()
            )
            all_issues = store.issue_states()
            open_issues = [i for i in all_issues if str(i.get("status")) in {"open", "in_progress"}]
            blocks = [i for i in open_issues if str(i.get("severity")) == "block"]
            warns = [i for i in open_issues if str(i.get("severity")) == "warn"]
            mode = quality_gate_mode()
            summary = {
                "open_count": len(open_issues),
                "block_count": len(blocks),
                "warn_count": len(warns),
                "can_proceed": mode == "soft" or not blocks,
                "mode": mode,
                "top_blocks": [
                    {"id": i.get("id"), "code": i.get("code"), "title": i.get("title"), "stage_id": i.get("stage_id")}
                    for i in blocks[:8]
                ],
                "source": "migration_required" if issue_import_pending else "control.db",
            }
        else:
            from agent.issues import issues_summary, load_open_issues
            from agent.root_cause import sync_issues_from_compliance, sync_issues_from_global_review

            try:
                sync_issues_from_global_review(root)
            except Exception:
                pass
            try:
                sync_issues_from_compliance(root)
            except Exception:
                pass
            all_issues = load_open_issues(root)
            summary = issues_summary(root)
        issues = all_issues
        if status and status != "all":
            if status == "open":
                issues = [i for i in issues if str(i.get("status")) in {"open", "in_progress"}]
            elif status == "block":
                issues = [
                    i
                    for i in issues
                    if str(i.get("severity")) == "block" and str(i.get("status")) in {"open", "in_progress"}
                ]
            else:
                issues = [i for i in issues if str(i.get("status")) == status]
        return JSONResponse(
            {
                "ok": True,
                "summary": summary,
                "issues": issues,
                "count": len(issues),
            }
        )
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)



@app.get("/api/issues/{issue_id}")
def api_get_issue(issue_id: str) -> JSONResponse:
    root = _active_root()
    try:
        from agent.issues import load_open_issues

        issue = next((i for i in load_open_issues(root) if str(i.get("id")) == issue_id), None)
        if not issue:
            return JSONResponse({"ok": False, "message": "未找到问题"}, status_code=404)
        return JSONResponse({"ok": True, "issue": issue})
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.post("/api/v2/workspaces/{workspace_id}/issues/{issue_id}/actions/preview")
@app.post("/api/issues/{issue_id}/actions/preview")
def api_preview_repair(issue_id: str, workspace_id: str = "") -> JSONResponse:
    context = _workspace_context(workspace_id) if workspace_id else None
    root = context.root if context else _active_root()
    try:
        from agent.repair import build_repair_plan

        if context:
            try:
                store = _ensure_v2_issue_import(context)
            except ControlPlaneError as exc:
                return JSONResponse(
                    {"ok": False, "code": exc.code, "message": exc.message},
                    status_code=exc.status_code,
                )
            issue = next((item for item in store.issue_states() if str(item.get("id")) == issue_id), None)
            plan = build_repair_plan(root, issue_id, issue=issue) if issue else {"ok": False, "message": f"未找到问题: {issue_id}"}
        else:
            plan = build_repair_plan(root, issue_id)
        status = 200 if plan.get("ok") else 404
        return JSONResponse(plan, status_code=status)
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.post("/api/issues/{issue_id}/actions/execute")
async def api_execute_repair(issue_id: str, request: Request) -> JSONResponse:
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}
    dry_run = bool(body.get("dry_run", False))
    if dry_run:
        try:
            from agent.repair import execute_repair_plan

            return JSONResponse(execute_repair_plan(_active_root(), issue_id, confirm=False, dry_run=True))
        except Exception as exc:
            return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)
    try:
        context = _workspace_context(ACTIVE_RUN_ID)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "repair.issues",
                "payload": {"issue_ids": [issue_id]},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-repair-issue:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        action = gateway.propose(envelope, label="确认执行问题最小修复", risk="high")
        return JSONResponse({"ok": True, "status": "requires_confirmation", "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)



@app.post("/api/issues/{issue_id}/actions/accept")
async def api_accept_issue_risk(issue_id: str, request: Request) -> JSONResponse:
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}
    reason = str(body.get("reason") or "").strip()
    try:
        context = _workspace_context(ACTIVE_RUN_ID)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "issues.accept_risk",
                "payload": {"issue_id": issue_id, "reason": reason},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-accept-risk:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        action = gateway.propose(envelope, label="确认接受问题风险", risk="high")
        return JSONResponse({"ok": True, "status": "requires_confirmation", "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/v2/workspaces/{workspace_id}/issues/{issue_id}/actions/explain")
@app.post("/api/issues/{issue_id}/actions/explain")
async def api_explain_issue_cause(issue_id: str, request: Request, workspace_id: str = "") -> JSONResponse:
    """Rule + optional LLM whitelist root-cause refinement."""
    context = _workspace_context(workspace_id) if workspace_id else None
    root = context.root if context else _active_root()
    try:
        from agent.issues import load_open_issues
        from agent.root_cause import refine_issue_cause_with_llm

        if context:
            try:
                store = _ensure_v2_issue_import(context)
            except ControlPlaneError as exc:
                return JSONResponse(
                    {"ok": False, "code": exc.code, "message": exc.message},
                    status_code=exc.status_code,
                )
            source = store.issue_states()
        else:
            source = load_open_issues(root)
        issue = next((i for i in source if str(i.get("id")) == issue_id), None)
        if not issue:
            return JSONResponse({"ok": False, "message": "未找到问题"}, status_code=404)
        result = refine_issue_cause_with_llm(root, issue)
        return JSONResponse(result)
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.post("/api/v2/workspaces/{workspace_id}/issues/actions/batch-preview")
@app.post("/api/issues/actions/batch-preview")
async def api_batch_preview_repair(request: Request, workspace_id: str = "") -> JSONResponse:
    context = _workspace_context(workspace_id) if workspace_id else None
    root = context.root if context else _active_root()
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON"}, status_code=400)
    ids = body.get("issue_ids") if isinstance(body, dict) else None
    if not isinstance(ids, list):
        return JSONResponse({"ok": False, "message": "issue_ids 必须是数组"}, status_code=400)
    try:
        from agent.repair import execute_repair_batch

        if context:
            try:
                issue_snapshot = _ensure_v2_issue_import(context).issue_states()
            except ControlPlaneError as exc:
                return JSONResponse(
                    {"ok": False, "code": exc.code, "message": exc.message},
                    status_code=exc.status_code,
                )
        else:
            issue_snapshot = None
        result = execute_repair_batch(
            root,
            [str(x) for x in ids],
            confirm=False,
            dry_run=True,
            issue_snapshot=issue_snapshot,
        )
        return JSONResponse(result)
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.post("/api/issues/actions/batch-execute")
async def api_batch_execute_repair(request: Request) -> JSONResponse:
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON"}, status_code=400)
    if not isinstance(body, dict):
        return JSONResponse({"ok": False, "message": "请求体必须是对象"}, status_code=400)
    ids = body.get("issue_ids")
    if not isinstance(ids, list) or not ids:
        return JSONResponse({"ok": False, "message": "issue_ids 必须是非空数组"}, status_code=400)
    try:
        context = _workspace_context(ACTIVE_RUN_ID)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "repair.issues",
                "payload": {"issue_ids": [str(x) for x in ids]},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-repair-batch:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        action = gateway.propose(envelope, label="确认执行批量最小修复", risk="high")
        return JSONResponse({"ok": True, "status": "requires_confirmation", "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/gates/revalidate")
async def api_revalidate_gate(request: Request) -> JSONResponse:
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON"}, status_code=400)
    command = str(body.get("command") or "").strip()
    if not command:
        return JSONResponse({"ok": False, "message": "缺少 command"}, status_code=400)
    try:
        context = _workspace_context(ACTIVE_RUN_ID)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "quality.revalidate",
                "payload": {"command": command},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-quality-revalidate:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        receipt = gateway.submit(envelope)
        return JSONResponse(
            {"ok": receipt.status != "rejected", "receipt": receipt.as_dict(), "message": receipt.message},
            status_code=202 if receipt.status != "rejected" else 409,
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.get("/api/v2/workspaces/{workspace_id}/compliance-report")
@app.get("/api/compliance-report")
def api_compliance_report(workspace_id: str = "") -> JSONResponse:
    """Full compliance report for right-side issues panel."""
    root = _workspace_context(workspace_id).root if workspace_id else _active_root()
    path = root / "workspace" / "compliance_report.json"
    if not path.exists():
        return JSONResponse(
            {
                "ok": True,
                "exists": False,
                "blocking": False,
                "items": [],
                "summary": {},
                "message": "尚未生成合规报告，请先执行 compliance-check。",
            }
        )
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        return JSONResponse({"ok": False, "message": f"读取失败: {exc}"}, status_code=500)
    if not isinstance(data, dict):
        return JSONResponse({"ok": False, "message": "报告格式无效"}, status_code=500)
    items = [i for i in (data.get("items") or []) if isinstance(i, dict)]
    sev_order = {"fatal": 0, "critical": 1, "major": 2, "minor": 3, "info": 4}
    items_sorted = sorted(
        items,
        key=lambda x: (
            0 if x.get("status") == "fail" else 1 if x.get("status") == "warn" else 2,
            sev_order.get(str(x.get("severity") or ""), 9),
            str(x.get("check_id") or ""),
        ),
    )
    summary = data.get("summary") if isinstance(data.get("summary"), dict) else {}
    return JSONResponse(
        {
            "ok": True,
            "exists": True,
            "blocking": bool(data.get("blocking") or summary.get("blocking")),
            "need_manual_review": bool(data.get("need_manual_review") or summary.get("need_manual_review")),
            "max_severity": data.get("max_severity") or summary.get("max_severity"),
            "summary": summary,
            "counts": summary.get("counts") if isinstance(summary.get("counts"), dict) else {},
            "items": items_sorted,
            "total": len(items_sorted),
        }
    )


@app.get("/api/agent/activity")
def api_agent_activity() -> JSONResponse:
    """Current sub-agent workbench snapshot for UI cards."""
    root = _active_root()
    try:
        from agent.activity import activity_for_api

        data = activity_for_api(root)
        return JSONResponse({"ok": True, "activity": data})
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc), "activity": {"status": "idle", "agents": []}})


@app.get("/api/agent/goal")
def api_agent_goal(reevaluate: bool = False) -> JSONResponse:
    """Current GoalState for active workspace.

    Default is read-only (no reevaluate on poll) to avoid status flip-flops.
    Pass ?reevaluate=1 after tools/material upload if a fresh check is needed.
    """
    root = _active_root()
    try:
        from agent.goal import goal_summary, load_goal, next_plan_step, reevaluate_goal
        from agent.runtime_status import build_runtime_status

        goal = load_goal(root)
        if goal and reevaluate:
            try:
                goal = reevaluate_goal(root, goal)
            except Exception:
                pass
        nxt = next_plan_step(root, goal) if goal else None
        runtime = build_runtime_status(root, reevaluate_goal=False)
        return JSONResponse(
            {
                "ok": True,
                "goal": goal,
                "summary": goal_summary(goal) if goal else "无活动目标",
                "next_plan_step": nxt,
                "product_mode": runtime.get("product_mode"),
                "product_mode_label": runtime.get("product_mode_label"),
                "consistent": runtime.get("consistent"),
                "consistency_warnings": runtime.get("warnings") or [],
            }
        )
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.get("/api/runtime")
def api_runtime_status(heal: bool = False) -> JSONResponse:
    """Unified live runtime status (goal + activity + repair + pipeline + materials)."""
    root = _active_root()
    try:
        from agent.runtime_status import build_runtime_status, soft_heal_inconsistencies

        if heal:
            payload = soft_heal_inconsistencies(root)
        else:
            payload = build_runtime_status(root, reevaluate_goal=False)
        return JSONResponse(payload)
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.get("/api/agent/snapshot")
def api_agent_snapshot() -> JSONResponse:
    root = _active_root()
    try:
        from agent.snapshot import build_snapshot

        snap = build_snapshot(root, for_llm=False)
        return JSONResponse({"ok": True, "snapshot": snap})
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.post("/api/agent/goal/resume")
async def api_agent_goal_resume(request: Request) -> JSONResponse:
    try:
        body = await request.json()
    except Exception:
        body = {}
    note = str((body or {}).get("note") or "web_resume")
    try:
        context = _workspace_context(ACTIVE_RUN_ID)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "goal.resume",
                "payload": {"note": note},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-goal-resume:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        receipt = gateway.submit(envelope)
        return JSONResponse(
            {"ok": receipt.status != "rejected", "receipt": receipt.as_dict(), "message": receipt.message},
            status_code=202 if receipt.status != "rejected" else 409,
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/agent/goal/confirm")
async def api_agent_goal_confirm(request: Request) -> JSONResponse:
    return _command_error_response(
        ControlPlaneError(
            "POLICY_DENIED",
            "V2 不再支持为 Goal 批量授权 mutation；请逐个确认持久化 Action。",
            status_code=410,
        )
    )


@app.get("/api/v2/workspaces/{workspace_id}/agent/decisions")
@app.get("/api/agent/decisions")
def api_agent_decisions(tail: int = 20, workspace_id: str = "") -> JSONResponse:
    root = _workspace_context(workspace_id).root if workspace_id else _active_root()
    try:
        from agent.trace import load_decisions

        items = load_decisions(root, tail=max(1, min(int(tail or 20), 100)))
        return JSONResponse({"ok": True, "decisions": items, "count": len(items)})
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.get("/api/agent/tools")
def api_agent_tools() -> JSONResponse:
    try:
        from agent.tool_registry import tool_manifest

        return JSONResponse({"ok": True, "tools": tool_manifest()})
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.post("/api/agent/tools/invoke")
async def api_agent_tools_invoke(request: Request) -> JSONResponse:
    """Advanced/debug invoke. Mutations still go through tool_runtime policy at call site."""
    root = _active_root()
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    name = str(body.get("name") or body.get("tool") or "").strip()
    args = body.get("args") if isinstance(body.get("args"), dict) else {}
    dry_run = bool(body.get("dry_run", False))
    if not name:
        return JSONResponse({"ok": False, "message": "缺少 tool name"}, status_code=400)
    try:
        from agent.tool_registry import get_tool
        from agent.tool_runtime import invoke as tool_invoke

        spec = get_tool(name)
        if spec is None:
            raise ControlPlaneError("COMMAND_INVALID", f"未知工具: {name}", status_code=404)
        if not dry_run and str(spec.kind) != "analysis":
            raise ControlPlaneError(
                "POLICY_DENIED",
                "调试 Tool API 只允许只读 analysis 工具；mutation/export 必须通过 V2 CommandGateway。",
                status_code=409,
            )
        result = tool_invoke(name, args, root=root, dry_run=dry_run, actor="api")
        return JSONResponse({"ok": result.ok, "result": result.to_dict()})
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.get("/api/v2/workspaces/{workspace_id}/chat/messages")
@app.get("/api/chat/messages")
def api_chat_messages_get(workspace_id: str = "") -> JSONResponse:
    context = _workspace_context(workspace_id) if workspace_id else None
    root = context.root if context else _active_root()
    run_id = context.workspace_id if context else ACTIVE_RUN_ID or root.name
    messages = load_messages(root, run_id)
    return JSONResponse({"ok": True, "run_id": run_id, "messages": messages})


@app.post("/api/v2/workspaces/{workspace_id}/chat/messages")
@app.post("/api/chat/messages")
async def api_chat_messages_post(request: Request, workspace_id: str = "") -> JSONResponse:
    context = _workspace_context(workspace_id) if workspace_id else None
    root = context.root if context else _active_root()
    run_id = context.workspace_id if context else ACTIVE_RUN_ID or root.name
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    role = str(body.get("role", "system")).strip()
    content = str(body.get("content", ""))
    thinking = str(body.get("thinking", ""))
    actions = body.get("actions") if isinstance(body.get("actions"), list) else []
    kind = str(body.get("kind", "message")).strip()
    if role not in {"user", "assistant", "system"}:
        role = "system"
    saved = save_message(root, run_id, role, content, thinking, actions, kind)
    return JSONResponse({"ok": True, "message": saved})


@app.delete("/api/v2/workspaces/{workspace_id}/chat/messages")
@app.delete("/api/chat/messages")
def api_chat_messages_delete(workspace_id: str = "") -> JSONResponse:
    context = _workspace_context(workspace_id) if workspace_id else None
    root = context.root if context else _active_root()
    run_id = context.workspace_id if context else ACTIVE_RUN_ID or root.name
    removed = clear_messages(root, run_id)
    return JSONResponse({"ok": True, "removed": removed})


@app.get("/api/v2/workspaces/{workspace_id}/materials-checklist")
@app.get("/api/materials-checklist")
def api_materials_checklist(workspace_id: str = "") -> JSONResponse:
    context = _workspace_context(workspace_id) if workspace_id else _workspace_context(ACTIVE_RUN_ID or _active_root().name)
    root = context.root
    try:
        from materials_checklist import (
            chapters_ready_for_refill,
            chapters_with_material_gaps,
            load_materials_checklist,
        )

        data = load_materials_checklist(root)
        store = ControlStore(context)
        material_import_pending = (
            store.v1_import_pending("materials")
            and (root / "workspace" / "materials_checklist.json").exists()
        )
        # This V2 read endpoint must never promote the legacy projection into
        # SQLite.  Only the administrator-confirmed migration.scan Command may
        # do that; while it is pending, return an explicit empty V2 snapshot.
        authoritative_items = [] if material_import_pending else store.material_states()
        audit_summary = store.material_audit_summary()
        enriched_items: list[dict[str, Any]] = []
        for item in authoritative_items:
            item_id = str(item.get("item_id") or "")
            audit = audit_summary.get(item_id) or {}
            enriched_items.append(
                {
                    **item,
                    "submission_count": int(audit.get("submission_count") or 0),
                    "latest_submission": audit.get("latest_submission"),
                    "latest_verification": audit.get("latest_verification"),
                }
            )
        authoritative_items = enriched_items
        summary = {
            "total": len(authoritative_items),
            "ready": sum(1 for item in authoritative_items if item.get("response_status") == "ready"),
            "deferred": sum(1 for item in authoritative_items if item.get("response_status") == "deferred"),
            "waived": sum(1 for item in authoritative_items if item.get("response_status") == "waived"),
        }
        data = {**data, "items": authoritative_items, "summary": summary}
        exists = (root / "workspace" / "materials_checklist.json").exists()
        return JSONResponse(
            {
                "ok": True,
                "exists": exists,
                "checklist": data,
                "summary": summary,
                "items": authoritative_items,
                "gap_chapters": {} if material_import_pending else chapters_with_material_gaps(root),
                "refill_plans": [] if material_import_pending else chapters_ready_for_refill(root),
                "source": "migration_required" if material_import_pending else "control.db",
            }
        )
    except Exception as exc:
        return JSONResponse({"ok": False, "message": str(exc)}, status_code=500)


@app.post("/api/materials-checklist/update")
async def api_materials_checklist_update(request: Request) -> JSONResponse:
    root = _active_root()
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    try:
        context = _workspace_context(ACTIVE_RUN_ID or root.name)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "materials.update",
                "payload": body,
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-material-update:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        action = gateway.propose(envelope, label="确认更新材料状态", risk="high")
        return JSONResponse(
            {"ok": True, "status": "requires_confirmation", "action": action},
            status_code=202,
            headers={"Deprecation": "true", "Link": f'</api/v2/workspaces/{context.workspace_id}/commands>; rel="successor-version"'},
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/materials-checklist/rebuild")
def api_materials_checklist_rebuild(request: Request) -> JSONResponse:
    try:
        context = _workspace_context(ACTIVE_RUN_ID)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "materials.rebuild",
                "payload": {},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-material-rebuild:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        receipt = gateway.submit(envelope)
        return JSONResponse(
            {"ok": receipt.status != "rejected", "receipt": receipt.as_dict(), "message": receipt.message},
            status_code=202 if receipt.status != "rejected" else 409,
            headers={"Deprecation": "true", "Link": f'</api/v2/workspaces/{context.workspace_id}/commands>; rel="successor-version"'},
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/materials-checklist/refill")
async def api_materials_checklist_refill(request: Request) -> JSONResponse:
    root = _active_root()
    try:
        body = await request.json()
    except Exception:
        body = {}
    try:
        context = _workspace_context(ACTIVE_RUN_ID or root.name)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "materials.refill",
                "payload": body if isinstance(body, dict) else {},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-material-refill:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        action = gateway.propose(envelope, label="确认将已验证材料回填正文", risk="high")
        return JSONResponse(
            {"ok": True, "status": "requires_confirmation", "action": action},
            status_code=202,
            headers={"Deprecation": "true", "Link": f'</api/v2/workspaces/{context.workspace_id}/commands>; rel="successor-version"'},
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/materials-checklist/upload")
async def api_materials_checklist_upload(request: Request) -> JSONResponse:
    """Compatibility adapter: create a persisted V2 upload proposal."""
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    try:
        if not str(body.get("upload_token") or "").strip():
            raise ControlPlaneError(
                "UPLOAD_TOKEN_REQUIRED",
                "旧材料登记接口不再接受服务器路径，请先调用 V2 暂存上传接口取得 upload_token。",
                status_code=400,
            )
        context = _workspace_context(ACTIVE_RUN_ID)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "materials.upload",
                "payload": body,
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-material-upload:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        action = gateway.propose(envelope, label="确认登记并验证上传材料", risk="high")
        return JSONResponse(
            {"ok": True, "status": "requires_confirmation", "action": action},
            status_code=202,
            headers={"Deprecation": "true", "Link": f'</api/v2/workspaces/{context.workspace_id}/commands>; rel="successor-version"'},
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)


_MANUAL_REVIEW_CLOSED = {"accepted", "resolved", "dismissed", "confirmed"}


def _v2_manual_review_items(
    context: WorkspaceContext,
    category: str,
    *,
    include_closed: bool = False,
) -> list[dict[str, Any]]:
    rows = manual_review_items(context.root, category, include_closed=True)
    prefix = f"manual-review:{category}:"
    decisions: dict[str, dict[str, Any]] = {}
    for decision in ControlStore(context).policy_decisions():
        issue_id = str(decision.get("issue_id") or "")
        if decision.get("decision_type") == "manual_review" and issue_id.startswith(prefix):
            decisions[issue_id[len(prefix) :]] = decision
    result: list[dict[str, Any]] = []
    for raw in rows:
        row = dict(raw) if isinstance(raw, dict) else {}
        item_id = str(row.get("item_id") or "")
        decision = decisions.get(item_id)
        if decision:
            value = decision.get("decision") if isinstance(decision.get("decision"), dict) else {}
            effective = value.get("payload") if isinstance(value.get("payload"), dict) else value
            row["override"] = dict(effective)
            row["control_source"] = "control.db"
            row["policy_decision_id"] = decision.get("decision_id")
        else:
            effective = row.get("override") if isinstance(row.get("override"), dict) else {}
            row["control_source"] = "v1_projection"
        if not include_closed and str(effective.get("status") or "").lower() in _MANUAL_REVIEW_CLOSED:
            continue
        result.append(row)
    return result


def _v2_manual_review_summary(context: WorkspaceContext) -> dict[str, Any]:
    path = context.root / "workspace" / "manual_review" / "summary.json"
    base = _read_json_file(path) if path.exists() else {}
    summary = dict(base) if isinstance(base, dict) else {}
    counts = {
        category: len(_v2_manual_review_items(context, category))
        for category in ("template_evidence", "score_coverage", "chapter_review", "global_review")
    }
    summary.update(
        {
            "project_type": load_project_profile(context.root).get("project_type", "general"),
            "template_evidence_pending": counts["template_evidence"],
            "score_coverage_pending": counts["score_coverage"],
            "chapter_review_pending": counts["chapter_review"],
            "global_review_pending": counts["global_review"],
            "total_pending": sum(counts.values()) + int(summary.get("compliance_pending") or 0),
            "source": "control.db",
        }
    )
    return summary


@app.get("/api/v2/workspaces/{workspace_id}/manual-review/summary")
@app.get("/api/manual-review/summary")
def api_manual_review_summary(workspace_id: str = "") -> JSONResponse:
    if workspace_id:
        context = _workspace_context(workspace_id)
        return JSONResponse({"ok": True, "summary": _v2_manual_review_summary(context)})
    return JSONResponse({"ok": True, "summary": manual_review_summary(_active_root())})


@app.get("/api/v2/workspaces/{workspace_id}/manual-review/items")
@app.get("/api/manual-review/items")
def api_manual_review_items(category: str = Query(..., min_length=1), workspace_id: str = "") -> JSONResponse:
    if workspace_id:
        context = _workspace_context(workspace_id)
        return JSONResponse({"ok": True, "category": category, "items": _v2_manual_review_items(context, category)})
    return JSONResponse({"ok": True, "category": category, "items": manual_review_items(_active_root(), category)})


@app.post("/api/manual-review/update")
async def api_manual_review_update(request: Request) -> JSONResponse:
    root = _active_root()
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    category = str(body.get("category", "")).strip()
    payload = body.get("payload")
    if not category or not isinstance(payload, dict):
        return JSONResponse({"ok": False, "message": "缺少 category 或 payload。"}, status_code=400)
    try:
        context = _workspace_context(ACTIVE_RUN_ID or root.name)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "review.update",
                "payload": {"category": category, "payload": payload},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-review-update:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        action = gateway.propose(envelope, label="确认更新人工复核结论", risk="high")
        return JSONResponse({"ok": True, "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"控制状态不可用，已拒绝更新: {exc}", status_code=503)
        )


@app.get("/api/agent-runs")
def api_agent_runs(
    stage: str = Query("", min_length=0),
    chapter_id: str = Query("", min_length=0),
    agent_name: str = Query("", min_length=0),
) -> JSONResponse:
    root = _active_root()
    return JSONResponse(
        {
            "ok": True,
            "items": _load_agent_runs(root, stage=stage, chapter_id=chapter_id, agent_name=agent_name, limit=100),
        }
    )


@app.get("/api/project-profile")
def api_project_profile() -> JSONResponse:
    root = _active_root()
    return JSONResponse({"ok": True, "profile": load_project_profile(root), "choices": project_profile_choices()})


@app.get("/api/v2/project-profiles")
def api_v2_project_profiles() -> JSONResponse:
    return JSONResponse({"ok": True, "choices": project_profile_choices()})


@app.post("/api/project-profile")
async def api_set_project_profile(request: Request) -> JSONResponse:
    root = _active_root()
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    project_type = str(body.get("project_type", "")).strip()
    try:
        context = _workspace_context(ACTIVE_RUN_ID or root.name)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "workspace.set_profile",
                "payload": {"project_type": project_type},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-set-profile:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        action = gateway.propose(envelope, label=f"确认切换项目类型为 {project_type}", risk="high")
        return JSONResponse({"ok": True, "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)


# ---------------------------------------------------------------
#  LLM settings (multi-model presets + .env sync)
# ---------------------------------------------------------------

LLM_ENV_KEYS: list[tuple[str, str]] = [
    ("OPENAI_BASE_URL", "base_url"),
    ("OPENAI_API_KEY", "api_key"),
    ("OPENAI_MODEL", "model"),
    ("OPENAI_TIMEOUT", "timeout"),
    ("OPENAI_MAX_RETRIES", "max_retries"),
    ("OPENAI_RETRY_INITIAL_DELAY", "retry_initial_delay"),
    ("OPENAI_RETRY_MAX_DELAY", "retry_max_delay"),
    ("OPENAI_STREAM", "stream"),
    ("OPENAI_VERIFY_SSL", "verify_ssl"),
    ("OPENAI_PROVIDER", "provider"),
]
_LLM_ALIAS_TO_KEY = {alias: key for key, alias in LLM_ENV_KEYS}
LLM_MODELS_FILE = ROOT / "models.json"


def _llm_env_path() -> Path:
    return ROOT / ".env"


def _to_int(value: Any, default: int) -> int:
    try:
        if value is None or value == "":
            return default
        return int(value)
    except Exception:
        return default


def _to_float(value: Any, default: float) -> float:
    try:
        if value is None or value == "":
            return default
        return float(value)
    except Exception:
        return default


def _parse_bool_str(value: Any, default: bool) -> bool:
    if value is None or value == "":
        return default
    return str(value).strip().lower() not in {"0", "false", "no", "off"}


def _read_llm_env_values() -> dict[str, str]:
    from config import _parse_env_file

    values = _parse_env_file(_llm_env_path())
    return {alias: values.get(key, "") for key, alias in LLM_ENV_KEYS}


def _write_llm_env(settings: dict[str, Any]) -> None:
    env_path = _llm_env_path()
    existing_lines: list[str] = []
    if env_path.exists():
        existing_lines = env_path.read_text(encoding="utf-8").splitlines()
    known_keys = {key for key, _ in LLM_ENV_KEYS}
    updated_keys: set[str] = set()
    new_lines: list[str] = []
    for raw_line in existing_lines:
        stripped = raw_line.strip()
        if not stripped or stripped.startswith("#") or "=" not in stripped:
            new_lines.append(raw_line)
            continue
        key = stripped.split("=", 1)[0].strip()
        if key in known_keys:
            alias = _LLM_ALIAS_TO_KEY.get(key, key)
            if alias in settings and settings[alias] is not None:
                new_lines.append(f"{key}={settings[alias]}")
                updated_keys.add(key)
                continue
        new_lines.append(raw_line)
    for alias, value in settings.items():
        key = _LLM_ALIAS_TO_KEY.get(alias)
        if key and key not in updated_keys and value is not None:
            new_lines.append(f"{key}={value}")
            updated_keys.add(key)
    temp_path = env_path.with_name(env_path.name + ".tmp")
    temp_path.write_text("\n".join(new_lines) + "\n", encoding="utf-8")
    temp_path.replace(env_path)
    # Web 进程自身也可能直接调用 LLM；同步进程环境可让它无需重启立即使用新配置。
    for alias, value in settings.items():
        key = _LLM_ALIAS_TO_KEY.get(alias)
        if key and value is not None:
            os.environ[key] = str(value)


def _llm_config_revision() -> str:
    env_path = _llm_env_path()
    if not env_path.exists():
        return ""
    stat = env_path.stat()
    return f"{stat.st_mtime_ns:x}-{stat.st_size:x}"


def _normalize_model(raw: dict[str, Any]) -> dict[str, Any]:
    provider = str(raw.get("provider") or "openai").strip().lower()
    if provider in {"claude"}:
        provider = "anthropic"
    if provider not in {"openai", "anthropic"}:
        provider = "openai"
    return {
        "id": str(raw.get("id", "")).strip(),
        "name": str(raw.get("name", "")).strip(),
        "provider": provider,
        "base_url": str(raw.get("base_url", "")).strip(),
        "api_key": str(raw.get("api_key", "")).strip(),
        "model": str(raw.get("model", "")).strip(),
        "timeout": _to_int(raw.get("timeout"), 300),
        "max_retries": _to_int(raw.get("max_retries"), 3),
        "retry_initial_delay": _to_float(raw.get("retry_initial_delay"), 2),
        "retry_max_delay": _to_float(raw.get("retry_max_delay"), 30),
        "stream": _parse_bool_str(raw.get("stream"), False),
        "verify_ssl": _parse_bool_str(raw.get("verify_ssl"), True),
    }


def _sync_model_to_env(model: dict[str, Any]) -> None:
    _write_llm_env(
        {
            "base_url": model.get("base_url", ""),
            "api_key": model.get("api_key", ""),
            "model": model.get("model", ""),
            "provider": model.get("provider", "openai") or "openai",
            "timeout": model.get("timeout", 300),
            "max_retries": model.get("max_retries", 3),
            "retry_initial_delay": model.get("retry_initial_delay", 2),
            "retry_max_delay": model.get("retry_max_delay", 30),
            "stream": "true" if model.get("stream") else "false",
            "verify_ssl": "true" if model.get("verify_ssl") else "false",
        }
    )


def _read_models_store() -> dict[str, Any]:
    if LLM_MODELS_FILE.exists():
        try:
            data = json.loads(LLM_MODELS_FILE.read_text(encoding="utf-8"))
            if isinstance(data, dict) and isinstance(data.get("models"), list):
                return data
        except Exception:
            pass
    env_values = _read_llm_env_values()
    default_model = _normalize_model(
        {
            "id": "default",
            "name": "默认模型",
            "base_url": env_values.get("base_url", ""),
            "api_key": env_values.get("api_key", ""),
            "model": env_values.get("model", ""),
            "provider": env_values.get("provider", "openai") or "openai",
            "timeout": env_values.get("timeout", 300),
            "max_retries": env_values.get("max_retries", 3),
            "retry_initial_delay": env_values.get("retry_initial_delay", 2),
            "retry_max_delay": env_values.get("retry_max_delay", 30),
            "stream": env_values.get("stream", "false"),
            "verify_ssl": env_values.get("verify_ssl", "true"),
        }
    )
    store: dict[str, Any] = {
        "models": [default_model],
        "active_id": default_model["id"] if default_model["base_url"] else "",
    }
    _write_models_store(store)
    return store


def _gate_can_proceed(next_command: str = "") -> dict:
    try:
        from agent.issues import can_proceed
        from agent.root_cause import sync_issues_from_compliance, sync_issues_from_global_review

        root = _active_root()
        # refresh issues from latest reports (non-destructive)
        try:
            sync_issues_from_global_review(root)
        except Exception:
            pass
        try:
            sync_issues_from_compliance(root)
        except Exception:
            pass
        return can_proceed(root, next_command=next_command)
    except Exception as exc:
        return {"ok": False, "can_proceed": True, "message": f"门禁检查异常(放行): {exc}", "blocks": []}


def _safe_issues_summary(root: Path | None = None) -> dict:
    try:
        from agent.issues import issues_summary
        return issues_summary(root or _active_root())
    except Exception:
        return {"open_count": 0, "block_count": 0, "can_proceed": True}


def _v2_issues_summary(root: Path) -> dict[str, Any]:
    """Read Issue authority for V2 presentation without invoking the V1 loader."""
    try:
        from agent.issues import quality_gate_mode

        issues = ControlStore(WorkspaceContext.resolve(root.parent, root.name)).issue_states()
        open_issues = [
            item for item in issues
            if str(item.get("status") or "") in {"open", "in_progress"}
        ]
        blocks = [item for item in open_issues if str(item.get("severity") or "") == "block"]
        warns = [item for item in open_issues if str(item.get("severity") or "") == "warn"]
        mode = quality_gate_mode()
        return {
            "open_count": len(open_issues),
            "block_count": len(blocks),
            "warn_count": len(warns),
            "can_proceed": mode == "soft" or not blocks,
            "mode": mode,
            "source": "control.db",
        }
    except Exception:
        # A read-side control failure must not let Chat infer that a blocked
        # workspace is clear; keep the result explicitly non-proceedable.
        return {"open_count": 0, "block_count": 0, "can_proceed": False, "source": "unavailable"}


def _safe_agent_activity(root: Path | None = None) -> dict:
    try:
        from agent.activity import activity_for_api
        return activity_for_api(root or _active_root())
    except Exception:
        return {"status": "idle", "agents": [], "summary": {}}


def _active_llm_summary() -> dict[str, Any]:
    store = _read_models_store()
    active_id = str(store.get("active_id", ""))
    models = store.get("models", []) if isinstance(store.get("models"), list) else []
    active = next((item for item in models if isinstance(item, dict) and str(item.get("id", "")) == active_id), {})
    return {
        "active_id": active_id,
        "name": str(active.get("name", "")),
        "model": str(active.get("model", "")),
        "config_revision": _llm_config_revision(),
    }


def _write_models_store(store: dict[str, Any]) -> None:
    LLM_MODELS_FILE.write_text(
        json.dumps(store, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


@app.get("/api/llm-settings")
def api_get_llm_settings() -> JSONResponse:
    store = _read_models_store()
    return JSONResponse(
        {
            "ok": True,
            "models": store.get("models", []),
            "active_id": store.get("active_id", ""),
            "config_revision": _llm_config_revision(),
        }
    )


@app.post("/api/llm-settings")
async def api_set_llm_settings(request: Request) -> JSONResponse:
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    if not isinstance(body, dict):
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON 对象。"}, status_code=400)
    raw_model = body.get("model")
    if not isinstance(raw_model, dict):
        return JSONResponse({"ok": False, "message": "缺少 model 字段。"}, status_code=400)
    model = _normalize_model(raw_model)
    if not model["name"]:
        return JSONResponse({"ok": False, "message": "模型别名（name）不能为空。"}, status_code=400)
    if not model["base_url"] or not model["api_key"] or not model["model"]:
        return JSONResponse({"ok": False, "message": "Base URL、API Key、模型均为必填项。"}, status_code=400)

    store = _read_models_store()
    models: list[dict[str, Any]] = list(store.get("models", []))
    model_id = model["id"]
    if model_id:
        index = next((i for i, m in enumerate(models) if str(m.get("id", "")) == model_id), -1)
        if index < 0:
            return JSONResponse({"ok": False, "message": "未找到要更新的模型。"}, status_code=404)
        models[index] = {**models[index], **model}
    else:
        model_id = uuid.uuid4().hex[:12]
        model["id"] = model_id
        models.append(model)

    store["models"] = models
    set_active = bool(body.get("set_active"))
    active_id = str(store.get("active_id", ""))
    applied_live = set_active or not active_id or active_id == model_id
    if set_active or not active_id:
        store["active_id"] = model_id
    if applied_live:
        _sync_model_to_env(next(m for m in models if m["id"] == model_id))

    _write_models_store(store)
    return JSONResponse(
        {
            "ok": True,
            "models": models,
            "active_id": store.get("active_id", ""),
            "saved_id": model_id,
            "applied_live": applied_live,
            "config_revision": _llm_config_revision(),
        }
    )


@app.post("/api/llm-settings/activate")
async def api_activate_llm_model(request: Request) -> JSONResponse:
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    model_id = str(body.get("id", "")).strip()
    store = _read_models_store()
    models: list[dict[str, Any]] = list(store.get("models", []))
    target = next((m for m in models if m["id"] == model_id), None)
    if target is None:
        return JSONResponse({"ok": False, "message": "未找到该模型。"}, status_code=404)
    store["active_id"] = model_id
    _write_models_store(store)
    _sync_model_to_env(target)
    return JSONResponse(
        {
            "ok": True,
            "models": models,
            "active_id": model_id,
            "applied_live": True,
            "config_revision": _llm_config_revision(),
        }
    )


@app.post("/api/llm-settings/delete")
async def api_delete_llm_model(request: Request) -> JSONResponse:
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    model_id = str(body.get("id", "")).strip()
    store = _read_models_store()
    models: list[dict[str, Any]] = list(store.get("models", []))
    new_models = [m for m in models if m["id"] != model_id]
    if len(new_models) == len(models):
        return JSONResponse({"ok": False, "message": "未找到该模型。"}, status_code=404)
    store["models"] = new_models
    if store.get("active_id") == model_id:
        store["active_id"] = new_models[0]["id"] if new_models else ""
        if new_models:
            _sync_model_to_env(new_models[0])
    _write_models_store(store)
    return JSONResponse(
        {
            "ok": True,
            "models": new_models,
            "active_id": store.get("active_id", ""),
            "applied_live": bool(store.get("active_id")),
            "config_revision": _llm_config_revision(),
        }
    )



@app.post("/api/llm-settings/test")
async def api_test_llm_settings(request: Request) -> JSONResponse:
    """Probe LLM with a tiny hello request using form or active model config."""
    import json as _json
    import time as _time
    import urllib.error
    import urllib.request
    import ssl
    import certifi

    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}

    raw_model = body.get("model") if isinstance(body.get("model"), dict) else None
    if raw_model:
        model = _normalize_model(raw_model)
    else:
        store = _read_models_store()
        active_id = str(store.get("active_id") or "")
        models = store.get("models") if isinstance(store.get("models"), list) else []
        active = next((m for m in models if isinstance(m, dict) and str(m.get("id")) == active_id), None)
        if not active and models and isinstance(models[0], dict):
            active = models[0]
        if not active:
            return JSONResponse({"ok": False, "message": "没有可测试的模型，请先填写配置。"}, status_code=400)
        model = _normalize_model(active)

    base_url = str(model.get("base_url") or "").strip().rstrip("/")
    api_key = str(model.get("api_key") or "").strip()
    model_id = str(model.get("model") or "").strip()
    provider = str(model.get("provider") or "openai").strip().lower()
    if provider not in {"openai", "anthropic"}:
        provider = "openai"
    if not base_url or not api_key or not model_id:
        return JSONResponse(
            {"ok": False, "message": "Base URL、API Key、模型 ID 均不能为空。"},
            status_code=400,
        )

    timeout = max(5, min(int(model.get("timeout") or 60), 90))
    verify_ssl = bool(model.get("verify_ssl", True))
    if verify_ssl:
        ctx = ssl.create_default_context(cafile=certifi.where())
    else:
        ctx = ssl._create_unverified_context()

    if provider == "anthropic":
        endpoint = base_url
        if endpoint.endswith("/messages"):
            pass
        elif endpoint.endswith("/v1"):
            endpoint = f"{endpoint}/messages"
        else:
            endpoint = f"{endpoint.rstrip('/')}/v1/messages"
        payload = {
            "model": model_id,
            "max_tokens": 64,
            "temperature": 0,
            "messages": [{"role": "user", "content": "hello"}],
            "system": "You are a connectivity probe. Reply briefly.",
        }
        headers = {
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "Content-Type": "application/json",
            "Accept": "application/json",
            "Connection": "close",
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/137.0.0.0 Safari/537.36"
            ),
        }
        if not api_key.startswith("sk-ant"):
            headers["Authorization"] = f"Bearer {api_key}"
    else:
        endpoint = base_url if base_url.endswith("/chat/completions") else f"{base_url}/chat/completions"
        payload = {
            "model": model_id,
            "temperature": 0,
            "messages": [
                {"role": "system", "content": "You are a connectivity probe. Reply briefly."},
                {"role": "user", "content": "hello"},
            ],
        }
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
            "Accept": "application/json",
            "Connection": "close",
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/137.0.0.0 Safari/537.36"
            ),
        }

    data = _json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(endpoint, data=data, method="POST", headers=headers)

    try:
        t0 = _time.time()
        with urllib.request.urlopen(req, timeout=timeout, context=ctx) as resp:
            raw = resp.read().decode("utf-8", errors="replace")
            status_code = getattr(resp, "status", 200)
        elapsed_ms = int((_time.time() - t0) * 1000)
        try:
            parsed = _json.loads(raw)
        except Exception:
            return JSONResponse(
                {
                    "ok": False,
                    "message": f"HTTP {status_code} 但响应不是 JSON: {raw[:200]}",
                    "model": model_id,
                    "provider": provider,
                    "base_url": base_url,
                    "elapsed_ms": elapsed_ms,
                }
            )

        text = ""
        if provider == "anthropic":
            content = parsed.get("content")
            if isinstance(content, list):
                text = "".join(
                    str(item.get("text") or "") if isinstance(item, dict) else str(item) for item in content
                ).strip()
            elif isinstance(content, str):
                text = content.strip()
        else:
            choices = parsed.get("choices") or []
            if choices:
                message = (choices[0] or {}).get("message") or {}
                content = message.get("content") or ""
                if isinstance(content, list):
                    content = "".join(
                        str(item.get("text") or item.get("content") or "") if isinstance(item, dict) else str(item)
                        for item in content
                    )
                text = str(content).strip()

        if not text:
            return JSONResponse(
                {
                    "ok": False,
                    "message": f"HTTP 成功但模型返回空内容（status={status_code}）。",
                    "model": model_id,
                    "provider": provider,
                    "base_url": base_url,
                    "elapsed_ms": elapsed_ms,
                }
            )
        preview = text if len(text) <= 300 else text[:300] + "…"
        return JSONResponse(
            {
                "ok": True,
                "message": "连接成功",
                "reply": preview,
                "model": model_id,
                "provider": provider,
                "name": model.get("name") or "",
                "base_url": base_url,
                "elapsed_ms": elapsed_ms,
            }
        )
    except urllib.error.HTTPError as exc:
        err_body = ""
        try:
            err_body = exc.read().decode("utf-8", errors="replace")[:400]
        except Exception:
            pass
        return JSONResponse(
            {
                "ok": False,
                "message": f"连接失败: HTTP {exc.code} {exc.reason}" + (f" | {err_body}" if err_body else ""),
                "model": model_id,
                "provider": provider,
                "base_url": base_url,
            }
        )
    except Exception as exc:
        return JSONResponse(
            {
                "ok": False,
                "message": f"连接失败: {type(exc).__name__}: {exc}",
                "model": model_id,
                "provider": provider,
                "base_url": base_url,
            }
        )



@app.get("/api/v2/workspaces/{workspace_id}/files")
@app.get("/api/workspace-files")
def api_workspace_files(workspace_id: str = "") -> JSONResponse:
    root = (_workspace_context(workspace_id).root if workspace_id else _active_root()).resolve()
    return JSONResponse(build_workspace_file_tree(root))


@app.get("/api/v2/workspaces/{workspace_id}/files/preview")
@app.get("/api/file-preview")
def api_file_preview(path: str = Query(..., min_length=1), workspace_id: str = "") -> JSONResponse:
    root = (_workspace_context(workspace_id).root if workspace_id else _active_root()).resolve()
    relative = path.strip().replace("\\", "/")
    if not relative or relative.startswith("/") or ".." in Path(relative).parts:
        return JSONResponse({"ok": False, "message": "无效文件路径。"}, status_code=400)

    if "*" in relative:
        directory_text, pattern = relative.rsplit("/", 1)
        directory = (root / directory_text).resolve()
        if not directory.is_relative_to(root) or not directory.exists() or not directory.is_dir():
            return JSONResponse({"ok": False, "message": f"目录不存在: {directory_text}"}, status_code=404)
        files = sorted(path for path in directory.glob(pattern) if path.is_file())
        return JSONResponse(
            {
                "ok": True,
                "path": relative,
                "kind": "list",
                "items": [
                    {
                        "name": item.name,
                        "path": str(item.relative_to(root)).replace("\\", "/"),
                        "size": item.stat().st_size,
                        "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(item.stat().st_mtime)),
                    }
                    for item in files[:500]
                ],
                "total": len(files),
            }
        )

    target = (root / relative).resolve()
    if not target.is_relative_to(root) or not target.exists():
        return JSONResponse({"ok": False, "message": f"文件不存在: {relative}"}, status_code=404)

    if target.is_dir():
        files = sorted(path for path in target.iterdir() if path.is_file())
        return JSONResponse(
            {
                "ok": True,
                "path": relative,
                "kind": "list",
                "items": [
                    {
                        "name": item.name,
                        "path": str(item.relative_to(root)).replace("\\", "/"),
                        "size": item.stat().st_size,
                        "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(item.stat().st_mtime)),
                    }
                    for item in files[:500]
                ],
                "total": len(files),
            }
        )

    if not target.is_file():
        return JSONResponse({"ok": False, "message": f"文件不存在: {relative}"}, status_code=404)

    suffix = target.suffix.lower()
    metadata = {
        "size": target.stat().st_size,
        "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(target.stat().st_mtime)),
    }
    if suffix in {".json", ".md", ".txt", ".log", ".jsonl", ".csv"}:
        text = target.read_text(encoding="utf-8", errors="replace")
        kind = "json" if suffix == ".json" else "text"
        if suffix == ".json":
            loaded = _read_json_file(target)
            if loaded is not None:
                text = json.dumps(loaded, ensure_ascii=False, indent=2)
        truncated = len(text) > 30000
        return JSONResponse(
            {
                "ok": True,
                "path": relative,
                "kind": kind,
                "metadata": metadata,
                "content": text[:30000],
                "truncated": truncated,
            }
        )

    if suffix == ".docx":
        try:
            from docx import Document

            document = Document(str(target))
            blocks: list[dict[str, Any]] = []
            for paragraph in document.paragraphs:
                content = paragraph.text.strip()
                if content:
                    blocks.append({"type": "paragraph", "text": content})
            for table_index, table in enumerate(document.tables, start=1):
                rows: list[list[str]] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        rows.append(cells)
                if rows:
                    blocks.append({"type": "table", "index": table_index, "rows": rows})
            return JSONResponse(
                {
                    "ok": True,
                    "path": relative,
                    "kind": "docx",
                    "metadata": metadata,
                    "blocks": blocks[:300],
                    "truncated": len(blocks) > 300,
                }
            )
        except Exception as exc:
            return JSONResponse(
                {
                    "ok": False,
                    "message": f"Word 预览失败: {exc}",
                    "path": relative,
                    "metadata": metadata,
                },
                status_code=500,
            )

    return JSONResponse(
        {
            "ok": True,
            "path": relative,
            "kind": "binary",
            "metadata": metadata,
            "message": "该文件类型不支持内嵌文本预览，请使用下载或本地 Word/PDF 工具查看。",
        }
    )


# ---------------------------------------------------------------
#  Command execution
# ---------------------------------------------------------------

COMMANDS: dict[str, list[str]] = {
    "init": [],
    "init-demo": [],
    "prepare-inputs": [],
    "analyze-template": [],
    "split-docs": [],
    "parse-score": [],
    "extract-facts": [],
    "build-template-evidence": [],
    "generate-outline": [],
    "plan-jobs": [],
    "select-context-all": [],
    "write-all": [],
    "review-fix-all": [],
    "build-source-trace": [],
    "build-score-coverage": [],
    "estimate-score": [],
    "summarize-all": [],
    "global-review": [],
    "compliance-check": [],
    "check-price-tables": [],
    "check-deviation-tables": [],
    "validate-claims": [],
    "build-md": [],
    "build-docx": [],
    "check-format": [],
    "validate": [],
    "run": [],
    "graph-run": [],
}

AUTO_RECOVERY_MAX_ATTEMPTS = 2

_NON_RECOVERABLE_ERROR_PATTERNS = (
    "api key",
    "401",
    "403",
    "unauthorized",
    "未授权",
    "无效或未授权",
    "permission denied",
    "access denied",
    "用户暂停",
    "已暂停",
)

_TRANSIENT_ERROR_PATTERNS = (
    "timeout",
    "timed out",
    "remote end closed",
    "remotedisconnected",
    "temporarily unavailable",
    "rate limit",
    "llm 请求失败",
    "流式响应为空",
    "connection reset",
    "无进展超时",
)

_PARSE_ERROR_PATTERNS = (
    "jsondecodeerror",
    "expecting value",
    "解析失败",
    "未返回合法 json",
    "invalid json",
)


def _run_process_once(command: str, run_root: Path) -> int:
    global CURRENT_PROCESS
    args = ["src/main.py", command, *COMMANDS.get(command, [])]
    env = os.environ.copy()
    env["PYTHONUTF8"] = "1"
    env["PYTHONIOENCODING"] = "utf-8"
    env["PYTHONUNBUFFERED"] = "1"
    env["BID_AGENT_ROOT"] = str(run_root)
    env["BID_AGENT_CONFIG_ROOT"] = str(ROOT)
    env["BID_AGENT_EXECUTION_WORKER"] = "1"
    env["BID_AGENT_RUNS_ROOT"] = str(RUNS_DIR)
    process = subprocess.Popen(
        [sys.executable, *args],
        cwd=str(ROOT),
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=False,
        bufsize=0,
    )
    CURRENT_PROCESS = process
    assert process.stdout is not None
    output_queue: queue.Queue[bytes | None] = queue.Queue()

    def _read_output() -> None:
        assert process.stdout is not None
        while True:
            raw_line = process.stdout.readline()
            if not raw_line:
                break
            output_queue.put(raw_line)
        output_queue.put(None)

    threading.Thread(target=_read_output, daemon=True, name=f"output-{process.pid}").start()
    try:
        stall_timeout = max(60, int(os.environ.get("BID_AGENT_STAGE_STALL_TIMEOUT", "900")))
    except ValueError:
        stall_timeout = 900
    last_progress = time.monotonic()
    last_heartbeat = 0.0
    output_closed = False
    timed_out = False
    while True:
        try:
            raw_line = output_queue.get(timeout=1)
            if raw_line is None:
                output_closed = True
            else:
                line = _decode_log_bytes(raw_line).rstrip("\n").rstrip("\r")
                if line:
                    _append_log(line)
                    last_progress = time.monotonic()
        except queue.Empty:
            pass

        now = time.monotonic()
        if now - last_heartbeat >= 5:
            SUPERVISOR.heartbeat(
                run_root,
                command=command,
                worker_pid=process.pid,
                progress_at=datetime.fromtimestamp(time.time() - (now - last_progress)).isoformat(timespec="seconds"),
            )
            last_heartbeat = now
        if process.poll() is None and now - last_progress > stall_timeout:
            _append_log(f"[错误] {command} 连续 {stall_timeout} 秒无进展超时，正在终止并自动恢复。")
            _terminate_process_tree(process)
            timed_out = True
            break
        if process.poll() is not None and output_closed:
            break
    try:
        process.wait(timeout=10 if timed_out else None)
    except subprocess.TimeoutExpired:
        try:
            process.kill()
        finally:
            process.wait(timeout=5)
    CURRENT_PROCESS = None
    return 124 if timed_out else int(process.returncode or 0)


def _terminate_process_tree(process: subprocess.Popen | None) -> None:
    if not process or process.poll() is not None:
        return
    _terminate_pid_tree(process.pid)


def _terminate_pid_tree(pid: int) -> None:
    if pid <= 0:
        return
    try:
        if os.name == "nt":
            subprocess.run(
                ["taskkill", "/PID", str(pid), "/T", "/F"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                check=False,
            )
        else:
            os.kill(pid, 15)
    except Exception:
        pass


def _error_text(lines: list[str]) -> str:
    return "\n".join(str(line) for line in lines).lower()


def _is_non_recoverable_error(lines: list[str]) -> bool:
    text = _error_text(lines)
    return any(pattern in text for pattern in _NON_RECOVERABLE_ERROR_PATTERNS)


def _is_transient_error(lines: list[str]) -> bool:
    text = _error_text(lines)
    return any(pattern in text for pattern in _TRANSIENT_ERROR_PATTERNS)


def _is_parse_error(lines: list[str]) -> bool:
    text = _error_text(lines)
    return any(pattern in text for pattern in _PARSE_ERROR_PATTERNS)


def _stage_index(command: str) -> int:
    commands = [str(step["command"]) for step in WORKFLOW_STEPS]
    try:
        return commands.index(command)
    except ValueError:
        return -1


def _missing_required_artifacts(command: str, run_root: Path) -> list[str]:
    try:
        spec = stage_spec_by_command(command)
    except KeyError:
        return []
    missing: list[str] = []
    for artifact in spec.requires:
        if not artifact_exists(run_root, artifact):
            missing.append(artifact.path)
    return missing


def _producer_command_for_artifact(artifact_path: str, before_command: str) -> str:
    before_idx = _stage_index(before_command)
    if before_idx < 0:
        return ""
    for step in reversed(WORKFLOW_STEPS[:before_idx]):
        for produced in step.get("produces", []):
            produced_path = _artifact_path(produced)
            if produced_path == artifact_path:
                return str(step.get("command", ""))
            if "*" in produced_path and Path(produced_path).parent == Path(artifact_path).parent:
                return str(step.get("command", ""))
    return ""


def _dependency_recovery_commands(command: str, run_root: Path) -> list[str]:
    commands: list[str] = []
    try:
        spec = stage_spec_by_command(command)
    except KeyError:
        spec = None
    if spec is not None:
        for required in spec.requires:
            producer = _producer_command_for_artifact(required.path, command)
            producer_step = _step_by_command(producer) if producer else None
            producer_spec = stage_spec_by_command(producer) if producer else None
            if (
                required.kind == "glob"
                and producer_step
                and producer_spec is not None
                and producer_spec.validator == "collection"
                and not stage_outputs_ready(run_root, str(producer_step.get("id", "")))
            ):
                commands.append(producer)
    for missing in _missing_required_artifacts(command, run_root):
        producer = _producer_command_for_artifact(missing, command)
        if producer and producer not in commands:
            commands.append(producer)
    return commands


def _save_recovery_state(run_root: Path, payload: dict[str, Any]) -> None:
    _write_json_file(run_root / "workspace" / "recovery_state.json", payload)


def _attempt_auto_recovery(command: str, run_root: Path, error_lines: list[str]) -> int | None:
    if PAUSE_REQUESTED or _is_non_recoverable_error(error_lines):
        return None

    dependency_commands = _dependency_recovery_commands(command, run_root)
    recoverable = bool(dependency_commands) or _is_transient_error(error_lines) or _is_parse_error(error_lines)
    if not recoverable:
        return None

    for attempt in range(1, AUTO_RECOVERY_MAX_ATTEMPTS + 1):
        reason = "缺少前置产物" if dependency_commands else ("LLM/网络临时异常" if _is_transient_error(error_lines) else "模型输出解析异常")
        action = "回退补跑前置阶段后重试" if dependency_commands else "等待后自动重试当前阶段"
        payload = {
            "command": command,
            "attempt": attempt,
            "max_attempts": AUTO_RECOVERY_MAX_ATTEMPTS,
            "reason": reason,
            "action": action,
            "dependency_commands": dependency_commands,
            "updated_at": datetime.now().isoformat(timespec="seconds"),
        }
        _save_recovery_state(run_root, payload)
        save_run_state(
            run_root,
            {"root_dir": str(run_root), "current_command": command, "recovery": payload},
            stage=command,
            status="recovering",
            message=f"正在尝试自主修复：{reason}；{action}（{attempt}/{AUTO_RECOVERY_MAX_ATTEMPTS}）",
        )
        _append_log(f"[自动恢复] {command}: {reason}；{action}（{attempt}/{AUTO_RECOVERY_MAX_ATTEMPTS}）")

        for dependency in dependency_commands:
            if PAUSE_REQUESTED:
                return None
            _append_log(f"[自动恢复] 补跑前置阶段: {dependency}")
            dep_exit = _run_process_once(dependency, run_root)
            if dep_exit != 0:
                _append_log(f"[自动恢复] 前置阶段 {dependency} 补跑失败，停止自动恢复。")
                return None

        if not dependency_commands:
            time.sleep(min(3 * attempt, 6))

        save_run_state(
            run_root,
            {"root_dir": str(run_root), "current_command": command, "recovery": payload},
            stage=command,
            status="retrying",
            message=f"自动修复后重试 {command}（{attempt}/{AUTO_RECOVERY_MAX_ATTEMPTS}）",
        )
        _append_log(f"[自动恢复] 重试阶段: {command}")
        retry_log_start = len(LOG_LINES)
        retry_exit = _run_process_once(command, run_root)
        if retry_exit == 0:
            _append_log(f"[自动恢复] {command} 已恢复成功。")
            return 0
        retry_error_lines = LOG_LINES[retry_log_start:][-40:]
        error_lines = retry_error_lines or error_lines
        dependency_commands = _dependency_recovery_commands(command, run_root)
        if PAUSE_REQUESTED or _is_non_recoverable_error(error_lines):
            return None

    _append_log(f"[自动恢复] {command} 已达到最大重试次数，停止自动恢复。")
    return None


def _run_sync(command: str, run_id: str, run_root: Path) -> int:
    _LOG_CONTEXT.run_root = run_root
    try:
        return _run_sync_impl(command, run_id, run_root)
    finally:
        if hasattr(_LOG_CONTEXT, "run_root"):
            del _LOG_CONTEXT.run_root


def _run_sync_impl(command: str, run_id: str, run_root: Path) -> int:
    global RUNNING, CURRENT_TASK, CURRENT_PROCESS, CURRENT_RUN_ID, CURRENT_RUN_ROOT, PAUSE_REQUESTED
    RUNNING = True
    CURRENT_TASK = command
    CURRENT_PROCESS = None
    CURRENT_RUN_ID = run_id
    CURRENT_RUN_ROOT = run_root
    PAUSE_REQUESTED = False

    log_start = len(LOG_LINES)
    args = ["src/main.py", command, *COMMANDS.get(command, [])]
    _append_log(f"--- [{time.strftime('%H:%M:%S')}] 开始: python {' '.join(args)} ---")
    _append_log(f"[运行目录] {run_root}")
    record_state = command not in {"validate", "init-demo"}
    if record_state:
        save_run_state(
            run_root,
            {"root_dir": str(run_root), "current_command": command},
            stage=command,
            status="running",
            message=f"开始执行: {COMMANDS.get(command, [])}",
        )

    try:
        exit_code = _run_process_once(command, run_root)
    except Exception as exc:
        _append_log(f"[错误] 命令执行异常: {exc}")
        exit_code = 1

    was_paused = PAUSE_REQUESTED
    if exit_code != 0 and was_paused:
        _append_log(f"[暂停] {command} 已暂停，可从断点继续。")
    elif exit_code != 0:
        _append_log(f"[错误] 流程已停止: {command} 执行失败，请查看上方报错。")
    if record_state:
        state_status = "paused" if exit_code != 0 and was_paused else ("ok" if exit_code == 0 else "error")
        state_message = (
            f"{command} 已暂停，可从断点继续"
            if state_status == "paused"
            else (f"{command} 执行完成" if exit_code == 0 else f"{command} 执行失败，exit_code={exit_code}")
        )
        # 失败时保存最后 40 行日志，供编排器/前端诊断和恢复
        if state_status == "error":
            error_lines = LOG_LINES[log_start:][-40:]
            error_path = run_root / "workspace" / "run_error.json"
            try:
                _write_json_file(error_path, {"command": command, "exit_code": exit_code, "lines": error_lines})
            except Exception:
                pass
            recovered_exit = _attempt_auto_recovery(command, run_root, error_lines)
            if recovered_exit == 0:
                exit_code = 0
                state_status = "ok"
                state_message = f"{command} 自动恢复后执行完成"
            elif recovered_exit is None and not was_paused:
                state_status = "recovery_failed"
                state_message = f"{command} 自动恢复未成功，exit_code={exit_code}"
        save_run_state(
            run_root,
            {"root_dir": str(run_root), "current_command": command},
            stage=command,
            status=state_status,
            message=state_message,
        )
    _append_log(f"--- [{time.strftime('%H:%M:%S')}] 完成: exit_code={exit_code} ---")
    RUNNING = False
    CURRENT_TASK = ""
    CURRENT_PROCESS = None
    CURRENT_RUN_ID = ""
    CURRENT_RUN_ROOT = None
    PAUSE_REQUESTED = False
    return exit_code


def _workspace_context(workspace_id: str) -> WorkspaceContext:
    return WorkspaceContext.resolve(RUNS_DIR, workspace_id)


def _ensure_v2_issue_import(context: WorkspaceContext) -> ControlStore:
    """Return migrated V2 Issue state; V1 imports only happen via migration.scan."""
    store = ControlStore(context)
    legacy_path = context.root / "workspace" / "issues" / "open.json"
    if store.issue_v1_import_pending() and legacy_path.exists():
        raise ControlPlaneError(
            "MIGRATION_SCAN_REQUIRED",
            "Issue 权威状态尚未迁移，请先执行管理员 migration.scan。",
            status_code=409,
        )
    return store


def _ensure_v2_material_import(context: WorkspaceContext) -> ControlStore:
    """Return material authority only after explicit migration when legacy data exists."""
    store = ControlStore(context)
    legacy_path = context.root / "workspace" / "materials_checklist.json"
    if store.v1_import_pending("materials") and legacy_path.exists():
        raise ControlPlaneError(
            "MIGRATION_SCAN_REQUIRED",
            "材料权威状态尚未迁移，请先执行管理员 migration.scan。",
            status_code=409,
        )
    return store


def _ensure_v2_repair_import(context: WorkspaceContext) -> ControlStore:
    """Reject V2 repair mutation until an existing V1 job is explicitly migrated."""
    store = ControlStore(context)
    legacy_path = context.root / "workspace" / "repair_job.json"
    if store.v1_import_pending("repair_job") and legacy_path.exists():
        raise ControlPlaneError(
            "MIGRATION_SCAN_REQUIRED",
            "RepairJob 权威状态尚未迁移，请先执行管理员 migration.scan。",
            status_code=409,
        )
    return store


def _v1_migration_dry_run(context: WorkspaceContext) -> dict[str, Any]:
    """Inventory legacy control files without mutating SQLite or compatibility files."""
    store = ControlStore(context)
    domains: dict[str, Any] = {}
    unrecognized: list[dict[str, Any]] = []
    manifest: list[dict[str, Any]] = []
    sources = {
        "goal": context.root / "workspace" / "agent" / "goal_state.json",
        "materials": context.root / "workspace" / "materials_checklist.json",
        "issues": context.root / "workspace" / "issues" / "open.json",
        "repair_job": context.root / "workspace" / "repair_job.json",
        "agent_activity": context.root / "workspace" / "agent" / "activity.json",
    }
    for domain, path in sources.items():
        if not path.exists():
            continue
        content = path.read_bytes()
        manifest.append(
            {
                "path": path.relative_to(context.root).as_posix(),
                "sha256": hashlib.sha256(content).hexdigest(),
                "size_bytes": len(content),
                "domain": domain,
                "import_pending": store.v1_import_pending(domain),
            }
        )
        if not store.v1_import_pending(domain):
            continue
        try:
            payload = json.loads(content.decode("utf-8"))
            if domain == "materials":
                payload = payload.get("items") if isinstance(payload, dict) else payload
            elif domain == "issues":
                payload = payload.get("issues") if isinstance(payload, dict) else payload
            valid = isinstance(payload, dict) if domain in {"goal", "repair_job", "agent_activity"} else isinstance(payload, list)
            if not valid:
                raise ValueError("unexpected JSON shape")
            domains[domain] = payload
        except Exception as exc:
            unrecognized.append(
                {"path": path.relative_to(context.root).as_posix(), "reason": str(exc)}
            )
    orphans: list[dict[str, Any]] = []
    for name in ("goal_state.json", "decision_trace.json", "decision_trace.jsonl"):
        path = context.root / name
        if path.exists() and path.is_file():
            content = path.read_bytes()
            manifest.append(
                {
                    "path": path.relative_to(context.root).as_posix(),
                    "sha256": hashlib.sha256(content).hexdigest(),
                    "size_bytes": len(content),
                    "domain": "orphan",
                    "import_pending": True,
                }
            )
            orphans.append(
                {
                    "path": path.relative_to(context.root).as_posix(),
                    "kind": "root_legacy_control_state",
                    "reason": "根目录旧状态未绑定到工作区 Agent，不自动导入。",
                }
            )
    for relative, kind, reason in (
        ("workspace/pipeline_control.json", "legacy_pipeline_checkpoint", "旧 Pipeline checkpoint 不自动绑定到 V2 Operation。"),
        ("workspace/stale_artifacts.json", "legacy_stale_state", "旧 stale artifact 状态不自动覆盖 SQLite Artifact manifest。"),
    ):
        path = context.root / relative
        if not path.exists() or not path.is_file():
            continue
        content = path.read_bytes()
        manifest.append(
            {
                "path": path.relative_to(context.root).as_posix(),
                "sha256": hashlib.sha256(content).hexdigest(),
                "size_bytes": len(content),
                "domain": "orphan",
                "import_pending": True,
            }
        )
        orphan = {"path": path.relative_to(context.root).as_posix(), "kind": kind, "reason": reason}
        if kind == "legacy_pipeline_checkpoint":
            try:
                checkpoint = json.loads(content.decode("utf-8"))
            except Exception:
                checkpoint = None
            if isinstance(checkpoint, dict):
                orphan["state"] = checkpoint
        orphans.append(orphan)
    result = store.migration_dry_run(
        domains,
        orphans=orphans,
        unrecognized=unrecognized,
    )
    result["sources"] = {
        domain: path.relative_to(context.root).as_posix()
        for domain, path in sources.items()
        if path.exists()
    }
    result["source_manifest"] = manifest
    result["source_fingerprint"] = hashlib.sha256(
        json.dumps(manifest, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    ).hexdigest()
    return result


def _migration_snapshot_with_source_state(
    context: WorkspaceContext,
    snapshot: dict[str, Any],
) -> dict[str, Any]:
    """Expose a read-only cutover health state without trusting legacy files."""
    migration = snapshot.get("migration")
    if not isinstance(migration, dict):
        return snapshot
    cutover = migration.get("cutover")
    if not isinstance(cutover, dict) or cutover.get("status") != "active":
        return snapshot
    preview = _v1_migration_dry_run(context)
    expected = str(cutover.get("fingerprint") or "")
    current = str(preview.get("source_fingerprint") or "")
    updated = dict(snapshot)
    updated_migration = dict(migration)
    updated_cutover = dict(cutover)
    updated_cutover["current_fingerprint"] = current
    updated_cutover["source_stale"] = not expected or expected != current
    if updated_cutover["source_stale"]:
        updated_cutover["status"] = "stale"
        updated_migration["status"] = "cutover_stale"
    updated_migration["cutover"] = updated_cutover
    updated["migration"] = updated_migration
    return updated


def _register_legacy_artifact_inventory(context: WorkspaceContext, store: ControlStore) -> int:
    """Hash existing V1 artifacts but never infer that they are reusable stage output."""
    from artifact_manifest import describe_artifact
    from pipeline_registry import workflow_stage_specs

    existing = {item["artifact_key"] for item in store.artifact_states()}
    manifests: list[dict[str, Any]] = []
    for stage in workflow_stage_specs():
        for artifact in stage.produces:
            if artifact.kind == "virtual":
                continue
            current = describe_artifact(context.root, artifact)
            artifact_key = str(current.get("artifact_key") or "")
            if (
                not artifact_key
                or artifact_key in existing
                or current.get("status") != "ready"
                or not current.get("files")
            ):
                continue
            manifests.append(
                {
                    **current,
                    "status": "stale",
                    "producer": stage.command,
                    "stage_id": stage.id,
                    "input_fingerprint": "",
                    "disposition": "legacy_discovered",
                    "legacy_artifact": True,
                    "legacy_readiness": "unverified",
                    "stale_reason": "V1 旧产物仅完成哈希盘点，缺少 V2 成功 StageRun 证据。",
                }
            )
            existing.add(artifact_key)
    if not manifests:
        return 0
    store.upsert_artifact_states(manifests)
    return len(manifests)


def _refresh_migration_report(
    context: WorkspaceContext,
    migration: dict[str, Any],
    *,
    action: dict[str, Any],
) -> None:
    """Keep the file audit projection current after a SQLite migration mutation."""
    path = context.root / "workspace" / "migration_report.json"
    report = _read_json_file(path)
    if not isinstance(report, dict) or str(report.get("workspace_id") or "") != context.workspace_id:
        return
    report["migration"] = migration
    report["last_action"] = action
    report["updated_at"] = datetime.now().astimezone().isoformat(timespec="milliseconds")
    from utils import write_json

    write_json(path, report)


def _request_actor(request: Request, *, source: str) -> dict[str, str]:
    """Bind Command actors on the server; never trust actor fields in JSON payloads."""
    state = getattr(request, "state", None)
    principal = getattr(state, "principal", None) if state is not None else None
    if isinstance(principal, dict):
        principal_id = str(principal.get("id") or "").strip()
        if principal_id:
            actor = {
                "type": str(principal.get("type") or source).strip() or source,
                "id": principal_id[:128],
            }
            principal_role = str(principal.get("role") or "").strip().lower()[:32]
            if principal_role:
                actor["role"] = principal_role
            return actor
    return {"type": source, "id": "anonymous"}


def _request_principal(request: Request) -> dict[str, Any]:
    state = getattr(request, "state", None)
    principal = getattr(state, "principal", None) if state is not None else None
    return dict(principal) if isinstance(principal, dict) else {}


def _v2_gate_can_proceed(context: WorkspaceContext, next_command: str) -> dict[str, Any]:
    """Fail-closed gate evaluation for V2 mutations using the explicit workspace."""
    try:
        from agent.issues import quality_gate_mode

        store = _ensure_v2_issue_import(context)
        issues = store.issue_states()
        open_issues = [
            item for item in issues
            if str(item.get("status") or "") in {"open", "in_progress"}
        ]
        blocks = [
            item for item in open_issues
            if str(item.get("severity") or "") in {"block", "fatal"}
        ]
        mode = quality_gate_mode()
        block_commands = {str(item.get("command") or "") for item in blocks}
        revalidate_allowed = bool(next_command and next_command in block_commands)
        non_overridable = [
            item for item in blocks
            if str(item.get("severity") or "") == "fatal"
            or str(item.get("category") or "") in {"qualification", "disqualification"}
        ]
        can_proceed = (
            not blocks
            or revalidate_allowed
            or (mode == "soft" and not non_overridable)
        )
        gate = {
            "ok": True,
            "can_proceed": can_proceed,
            "mode": mode,
            "block_count": len(blocks),
            "blocks": blocks,
            "next_command": next_command,
            "revalidate_allowed": revalidate_allowed,
            "source": "control.db",
            "message": (
                f"允许重验门禁阶段 `{next_command}`（当前仍有 {len(blocks)} 条 block）"
                if revalidate_allowed
                else "无 open block 问题"
                if not blocks
                else "soft 模式：仅记录可接受风险"
                if can_proceed
                else f"存在 {len(blocks)} 条阻断问题，禁止继续执行。"
            ),
        }
    except ControlPlaneError:
        raise
    except Exception as exc:
        raise ControlPlaneError(
            "STATE_UNAVAILABLE",
            f"门禁状态读取失败，已拒绝执行: {exc}",
            status_code=503,
            retryable=True,
        ) from exc
    if not isinstance(gate, dict):
        raise ControlPlaneError("STATE_UNAVAILABLE", "门禁返回无效状态，已拒绝执行。", status_code=503)
    return gate


def _v2_export_preflight(context: WorkspaceContext) -> dict[str, Any]:
    """Read-only formal-export checks backed by authoritative SQLite Issues."""
    from agent.issues import classify_issue_risk

    store = ControlStore(context)
    store.assert_migration_ready()
    migration_preview = _v1_migration_dry_run(context)
    if migration_preview.get("status") != "ready":
        raise ControlPlaneError(
            "MIGRATION_SCAN_REQUIRED",
            "旧工作区迁移预检发现未登记的状态，请先执行管理员迁移扫描。",
            status_code=409,
            details={"migration": migration_preview.get("counts") or {}},
        )
    cutover = (store.snapshot().get("migration") or {}).get("cutover")
    if isinstance(cutover, dict) and cutover.get("status") == "active":
        expected_fingerprint = str(cutover.get("fingerprint") or "")
        current_fingerprint = str(migration_preview.get("source_fingerprint") or "")
        if not expected_fingerprint or expected_fingerprint != current_fingerprint:
            raise ControlPlaneError(
                "MIGRATION_CUTOVER_STALE",
                "V2 切换后的旧状态源已变化，请重新扫描并完成协调。",
                status_code=409,
                details={"expected_fingerprint": expected_fingerprint, "current_fingerprint": current_fingerprint},
            )
    failed_evaluations = [
        evaluation for evaluation in store.latest_gate_evaluations()
        if str(evaluation.get("verdict") or "") in {"block", "error"}
    ]
    if failed_evaluations:
        error_evaluations = [item for item in failed_evaluations if item.get("verdict") == "error"]
        raise ControlPlaneError(
            "STATE_UNAVAILABLE" if error_evaluations else "GATE_BLOCKED",
            "存在尚未通过或异常的最新 GateEvaluation，已拒绝正式出稿。",
            status_code=503 if error_evaluations else 409,
            details={
                "gate_evaluations": [
                    {
                        "evaluation_id": item.get("evaluation_id"),
                        "command": item.get("command"),
                        "verdict": item.get("verdict"),
                    }
                    for item in failed_evaluations
                ]
            },
        )
    if isinstance(cutover, dict) and cutover.get("status") == "active":
        latest_by_command = {
            str(item.get("command") or ""): item
            for item in store.latest_gate_evaluations()
        }
        missing_evaluations = [
            command
            for command in ("global-review", "compliance-check")
            if str((latest_by_command.get(command) or {}).get("verdict") or "") != "pass"
        ]
        if missing_evaluations:
            raise ControlPlaneError(
                "GATE_BLOCKED",
                "V2 切换后正式出稿必须具有当前通过的质量门禁评估。",
                status_code=409,
                details={"missing_gate_evaluations": missing_evaluations},
            )
        evaluation_inputs = {
            "global-review": context.root / "workspace" / "global_review.json",
            "compliance-check": context.root / "workspace" / "compliance_report.json",
        }
        stale_evaluations: list[str] = []
        for command, source_path in evaluation_inputs.items():
            try:
                evaluated_at = datetime.fromisoformat(str(latest_by_command[command].get("created_at") or ""))
                source_changed_at = datetime.fromtimestamp(source_path.stat().st_mtime, tz=timezone.utc)
            except (OSError, TypeError, ValueError):
                raise ControlPlaneError(
                    "STATE_UNAVAILABLE",
                    "质量门禁评估或其输入时间戳无效，已拒绝正式出稿。",
                    status_code=503,
                )
            if evaluated_at.tzinfo is None or evaluated_at < source_changed_at:
                stale_evaluations.append(command)
        if stale_evaluations:
            raise ControlPlaneError(
                "GATE_BLOCKED",
                "质量门禁评估早于当前报告输入，请重新评估。",
                status_code=409,
                details={"stale_gate_evaluations": stale_evaluations},
            )
    issues = store.issue_states()
    open_issues = [
        item for item in issues
        if str(item.get("status") or "") in {"open", "in_progress"}
    ]
    blocks = [
        item for item in open_issues
        if str(item.get("severity") or "") in {"block", "fatal"}
    ]
    accepted = [item for item in issues if str(item.get("status") or "") == "accepted"]
    checks: list[dict[str, Any]] = []

    global_path = context.root / "workspace" / "global_review.json"
    global_review = _read_json_file(global_path) if global_path.exists() else {}
    if global_path.exists() and (
        not isinstance(global_review, dict)
        or not isinstance(global_review.get("blocking"), bool)
    ):
        raise ControlPlaneError(
            "STATE_UNAVAILABLE",
            "global_review.json 无效或缺少 blocking 布尔值，已拒绝正式出稿。",
            status_code=503,
        )
    global_blocks = [item for item in blocks if str(item.get("stage_id") or "") == "global_review"]
    global_blocking = bool(isinstance(global_review, dict) and global_review.get("blocking"))
    checks.append(
        {
            "id": "global_review",
            "label": "全文审核门禁",
            "ok": global_path.exists() and not global_blocking and not global_blocks,
            "detail": (
                "缺少 global_review.json"
                if not global_path.exists()
                else "通过"
                if not global_blocking and not global_blocks
                else f"阻断 {len(global_blocks)} 项"
            ),
        }
    )

    compliance_path = context.root / "workspace" / "compliance_report.json"
    compliance = _read_json_file(compliance_path) if compliance_path.exists() else {}
    compliance_summary = (
        compliance.get("summary")
        if isinstance(compliance, dict) and isinstance(compliance.get("summary"), dict)
        else {}
    )
    if compliance_path.exists() and (
        not isinstance(compliance, dict)
        or not (
            isinstance(compliance.get("blocking"), bool)
            or isinstance(compliance_summary.get("blocking"), bool)
        )
    ):
        raise ControlPlaneError(
            "STATE_UNAVAILABLE",
            "compliance_report.json 无效或缺少 blocking 布尔值，已拒绝正式出稿。",
            status_code=503,
        )
    compliance_blocking = bool(
        isinstance(compliance, dict)
        and (compliance.get("blocking") or compliance_summary.get("blocking"))
    )
    compliance_blocks = [
        item for item in blocks
        if str(item.get("stage_id") or "") == "compliance_check"
    ]
    checks.append(
        {
            "id": "compliance_check",
            "label": "专项合规门禁",
            "ok": compliance_path.exists() and not compliance_blocking and not compliance_blocks,
            "detail": (
                "缺少 compliance_report.json"
                if not compliance_path.exists()
                else "通过"
                if not compliance_blocking and not compliance_blocks
                else f"阻断 blocking={compliance_blocking}, issues={len(compliance_blocks)}"
            ),
        }
    )

    checks.append(
        {
            "id": "open_blocks",
            "label": "无 open block 问题单",
            "ok": not blocks,
            "detail": "通过" if not blocks else f"仍有 {len(blocks)} 条 block",
        }
    )
    checks.append(
        {
            "id": "accepted_risks",
            "label": "已接受风险披露",
            "ok": True,
            "detail": "无" if not accepted else f"存在 {len(accepted)} 条已接受风险（终稿不得显示全部通过）",
            "count": len(accepted),
        }
    )
    final_md = context.root / "outputs" / "final.md"
    checks.append(
        {
            "id": "final_md",
            "label": "存在 final.md",
            "ok": final_md.exists() and final_md.is_file() and final_md.stat().st_size > 0,
            "detail": str(final_md) if final_md.exists() else "缺失",
        }
    )
    can_export = all(bool(item.get("ok")) for item in checks if item.get("id") != "accepted_risks")
    has_accepted = bool(accepted)
    return {
        "ok": True,
        "can_export": can_export,
        "all_passed": can_export and not has_accepted,
        "has_accepted_risks": has_accepted,
        "accepted_risks": [
            {
                "id": item.get("id"),
                "code": item.get("code"),
                "title": item.get("title"),
                "risk_class": item.get("risk_class") or classify_issue_risk(item),
                "accept_reason": item.get("accept_reason"),
                "accepted_by": item.get("accepted_by"),
                "accepted_at": item.get("accepted_at"),
            }
            for item in accepted[:50]
        ],
        "checks": checks,
        "issues_summary": {
            "open_count": len(open_issues),
            "block_count": len(blocks),
            "warn_count": sum(1 for item in open_issues if str(item.get("severity") or "") == "warn"),
            "source": "control.db",
        },
        "block_issues": [
            {
                "id": item.get("id"),
                "code": item.get("code"),
                "title": item.get("title"),
                "stage_id": item.get("stage_id"),
                "risk_class": item.get("risk_class") or classify_issue_risk(item),
            }
            for item in blocks[:50]
        ],
        "message": (
            f"可以出正式稿，但存在 {len(accepted)} 条已接受风险，不得标注“全部通过”"
            if can_export and accepted
            else "可以出正式稿"
            if can_export
            else "出稿前检查未通过，请先处理阻断项"
        ),
        "source": "control.db",
    }


def _record_v2_stage_artifacts(context: WorkspaceContext, command: str, disposition: str) -> None:
    from artifact_manifest import record_stage_artifacts

    record_stage_artifacts(context, command, disposition=disposition)


def _v2_stage_artifacts_reusable(context: WorkspaceContext, command: str) -> bool:
    from artifact_manifest import stage_artifacts_reusable

    return stage_artifacts_reusable(context, command)


_FORMAL_GATE_RULES_VERSION = "v2-formal-control-2026-07-21"
_FORMAL_GATE_INPUTS = (
    "outputs/final.md",
    "outputs/final.docx",
    "workspace/global_review.json",
    "workspace/compliance_report.json",
    "workspace/format_check_report.json",
)
_FORMAL_GATE_TREES = (
    "inputs/tender",
    "inputs/company",
    "workspace/chapters",
)


def _formal_gate_fingerprint(context: WorkspaceContext) -> tuple[str, str]:
    digest = hashlib.sha256()
    artifact_sha256 = ""
    store = ControlStore(context)
    migration_state = store.migration_state()
    control_domains = {
        "material_states": store.material_states(),
        "issue_states": store.issue_states(),
        "policy_decisions": store.policy_decisions(),
        "artifact_states": store.artifact_states(),
        "latest_gate_evaluations": store.latest_gate_evaluations(),
        "migration": migration_state,
    }
    cutover = migration_state.get("cutover") if isinstance(migration_state, dict) else None
    if isinstance(cutover, dict) and cutover.get("status") == "active":
        control_domains["migration_source_fingerprint"] = (
            _v1_migration_dry_run(context).get("source_fingerprint") or ""
        )
    for domain, value in control_domains.items():
        digest.update(f"control.db:{domain}\0".encode("utf-8"))
        digest.update(
            json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
        )
    for relative in _FORMAL_GATE_INPUTS:
        path = context.root / relative
        digest.update(relative.encode("utf-8"))
        digest.update(b"\0")
        if not path.exists() or not path.is_file():
            digest.update(b"<missing>")
            continue
        content = path.read_bytes()
        file_hash = hashlib.sha256(content).hexdigest()
        digest.update(file_hash.encode("ascii"))
        if relative == "outputs/final.docx":
            artifact_sha256 = file_hash
    for relative_dir in _FORMAL_GATE_TREES:
        directory = context.root / relative_dir
        digest.update(relative_dir.encode("utf-8"))
        digest.update(b"\0")
        if not directory.exists() or not directory.is_dir():
            digest.update(b"<missing-tree>")
            continue
        for path in sorted(item for item in directory.rglob("*") if item.is_file()):
            relative = path.relative_to(context.root).as_posix()
            digest.update(relative.encode("utf-8"))
            digest.update(b"\0")
            digest.update(hashlib.sha256(path.read_bytes()).hexdigest().encode("ascii"))
    digest.update(_FORMAL_GATE_RULES_VERSION.encode("utf-8"))
    return digest.hexdigest(), artifact_sha256


def _assert_formal_materials_verified(context: WorkspaceContext) -> None:
    store = ControlStore(context)
    if store.v1_import_pending("materials"):
        cutover = store.migration_state().get("cutover") or {}
        legacy_materials_path = context.root / "workspace" / "materials_checklist.json"
        if str(cutover.get("status") or "") == "active" and legacy_materials_path.exists():
            raise ControlPlaneError(
                "MIGRATION_SCAN_REQUIRED",
                "材料权威状态尚未迁移，已拒绝签发 GateReceipt。",
                status_code=409,
            )
        # One-version V1 compatibility: read the legacy projection without
        # importing or writing it while a formal gate is being evaluated.
        from materials_checklist import load_materials_checklist

        checklist = load_materials_checklist(context.root)
        items = checklist.get("items") if isinstance(checklist, dict) else []
        items = [dict(item) for item in items if isinstance(item, dict)]
    else:
        items = store.material_states()
    unsafe = [
        str(item.get("item_id") or "")
        for item in items
        if isinstance(item, dict)
        and not _material_fulfillment_verified(item)
        and (
            str(item.get("response_status") or "") == "ready"
            or _protected_material(item)
        )
    ]
    if unsafe:
        raise ControlPlaneError(
            "GATE_BLOCKED",
            "正式稿存在未验证的必交/资格材料或异常 ready 状态，已拒绝签发 GateReceipt。",
            details={"item_ids": unsafe[:50]},
        )


def _assert_formal_artifacts_ready(context: WorkspaceContext) -> None:
    from artifact_manifest import describe_artifact, stage_artifacts_reusable

    store = ControlStore(context)
    states = {item["artifact_key"]: item for item in store.artifact_states()}
    cutover_active = str((store.migration_state().get("cutover") or {}).get("status") or "") == "active"
    blocked: list[dict[str, str]] = []
    for relative in _FORMAL_GATE_INPUTS:
        state = states.get(relative)
        # During the one-version V1 compatibility window, missing manifests are
        # accepted and will be bootstrapped by the V2 Pipeline on first reuse.
        # A workspace that has explicitly cut over to V2 can no longer use this
        # exception: its formal outputs must have V2 manifest evidence.
        if state is None:
            if cutover_active:
                blocked.append({"path": relative, "reason": "manifest_missing_after_cutover"})
            continue
        current = describe_artifact(context.root, RunArtifact(relative))
        if state.get("status") != "ready":
            blocked.append({"path": relative, "reason": str(state.get("status") or "unknown")})
        elif current.get("status") != "ready" or current.get("sha256") != state.get("sha256"):
            blocked.append({"path": relative, "reason": "manifest_mismatch"})
    if states.get("outputs/final.docx") is not None and not stage_artifacts_reusable(context, "build-docx"):
        blocked.append({"path": "outputs/final.docx", "reason": "stale_inputs"})
    if blocked:
        raise ControlPlaneError(
            "GATE_BLOCKED",
            "正式稿依赖的 Artifact 已 stale、缺失或与 SQLite manifest 不一致。",
            details={"artifacts": blocked},
        )


def _validate_formal_gate_receipt(
    context: WorkspaceContext,
    receipt_id: str,
) -> tuple[dict[str, Any], Path]:
    if not str(receipt_id or "").strip():
        raise ControlPlaneError("GATE_RECEIPT_REQUIRED", "正式稿下载必须提供 GateReceipt。", status_code=409)
    receipt = ControlStore(context).gate_receipt(receipt_id)
    if not receipt or receipt.get("verdict") != "pass":
        raise ControlPlaneError("GATE_RECEIPT_INVALID", "GateReceipt 不存在或未通过。", status_code=409)
    if receipt.get("rules_version") != _FORMAL_GATE_RULES_VERSION:
        raise ControlPlaneError("GATE_RECEIPT_STALE", "GateReceipt 规则版本已过期。", status_code=409)
    artifact_path = str(receipt.get("artifact_path") or "")
    if artifact_path != "outputs/final.docx":
        raise ControlPlaneError("GATE_RECEIPT_INVALID", "GateReceipt 关联的正式稿路径无效。", status_code=409)
    artifact = context.root / artifact_path
    fingerprint, artifact_sha256 = _formal_gate_fingerprint(context)
    if (
        fingerprint != receipt.get("gate_input_fingerprint")
        or artifact_sha256 != receipt.get("artifact_sha256")
        or not artifact.exists()
        or not artifact.is_file()
    ):
        raise ControlPlaneError("GATE_RECEIPT_STALE", "正式稿或门禁输入已变化，请重新验收。", status_code=409)
    return receipt, artifact


def _pipeline_status_to_operation(status: str) -> str:
    return {
        "running": "running",
        "recovering": "running",
        "retrying": "running",
        "interrupted": "blocked",
        "pausing": "pausing",
        "paused": "paused",
        "cancelling": "cancelling",
        "cancelled": "cancelled",
        "complete": "succeeded",
        "failed": "failed",
    }.get(status, "")


def _pipeline_snapshot_from_control(
    operations: list[dict[str, Any]],
    checkpoint: dict[str, Any],
) -> dict[str, Any]:
    operation = next(
        (
            item
            for item in operations
            if str(item.get("kind") or "").startswith("pipeline.")
        ),
        None,
    )
    if not operation:
        return {**checkpoint, "source": "v1_checkpoint", "consistent": True} if checkpoint else {}
    status = {
        "queued": "running",
        "running": "running",
        "pausing": "pausing",
        "paused": "paused",
        "cancelling": "cancelling",
        "cancelled": "cancelled",
        "succeeded": "complete",
        "failed": "failed",
        "blocked": "interrupted",
    }.get(str(operation.get("status") or ""), "failed")
    operation_id = str(operation.get("operation_id") or "")
    fencing_token = int(operation.get("fencing_token") or 0)
    checkpoint_matches = bool(
        checkpoint
        and str(checkpoint.get("operation_id") or "") == operation_id
        and int(checkpoint.get("fencing_token") or 0) == fencing_token
    )
    return {
        "run_id": str(checkpoint.get("run_id") or "") if checkpoint_matches else "",
        "operation_id": operation_id,
        "fencing_token": fencing_token,
        "status": status,
        "current_stage": str(checkpoint.get("current_stage") or operation.get("start_command") or "")
        if checkpoint_matches
        else str(operation.get("start_command") or ""),
        "worker_pid": int(checkpoint.get("worker_pid") or 0) if checkpoint_matches else 0,
        "message": str(operation.get("message") or checkpoint.get("message") or ""),
        "error": operation.get("error"),
        "source": "control.db",
        "consistent": checkpoint_matches or not checkpoint,
        "checkpoint_source": "pipeline_control.json" if checkpoint_matches else "ignored_mismatch",
    }


def _sync_pipeline_control_state(root: Path, payload: dict[str, Any]) -> None:
    operation_id = str(payload.get("operation_id") or "").strip()
    operation_status = _pipeline_status_to_operation(str(payload.get("status") or ""))
    if not operation_id or not operation_status:
        return
    context = _workspace_context(root.name)
    if not _same_path(context.root, root):
        return
    error = payload.get("error")
    ControlStore(context).sync_operation(
        operation_id,
        operation_status,
        message=str(payload.get("message") or ""),
        error={"message": str(error)} if error else None,
        fencing_token=int(payload.get("fencing_token") or 0),
    )


def _terminate_workspace_process(context: WorkspaceContext) -> None:
    process = CURRENT_PROCESS if _same_path(CURRENT_RUN_ROOT, context.root) else None
    if process:
        _terminate_process_tree(process)
        return
    control = SUPERVISOR.load(context.root)
    _terminate_pid_tree(int(control.get("worker_pid", 0) or 0))


def _handle_pipeline_start(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    if RUNNING or SUPERVISOR.is_running():
        raise ControlPlaneError("LEASE_CONFLICT", "当前已有任务或流水线正在运行。")
    start_command = str(envelope.payload.get("start_command") or "").strip()
    if envelope.kind == "pipeline.run_stage" and not start_command:
        raise ControlPlaneError("COMMAND_INVALID", "run_stage 缺少 start_command。", status_code=400)
    if start_command and start_command not in auto_run_commands():
        raise ControlPlaneError("COMMAND_INVALID", f"无效起始阶段: {start_command}", status_code=400)
    gate = _v2_gate_can_proceed(context, start_command or "auto_run")
    if not gate.get("can_proceed", False):
        raise ControlPlaneError(
            "GATE_BLOCKED",
            str(gate.get("message") or "质量门禁阻断，禁止启动流水线。"),
            details={"gate": gate},
        )
    operation = ControlStore(context).operation(operation_id) or {}
    fencing_token = int(operation.get("fencing_token") or 0)
    started = SUPERVISOR.start(
        context.workspace_id,
        context.root,
        _run_sync,
        start_command=start_command,
        operation_id=operation_id,
        fencing_token=fencing_token,
        single_command=envelope.kind == "pipeline.run_stage",
        gate_evaluator=lambda _root, command: _v2_gate_can_proceed(context, command),
        artifact_recorder=lambda _root, command, disposition: _record_v2_stage_artifacts(
            context,
            command,
            disposition,
        ),
        artifact_readiness_evaluator=lambda _root, command: _v2_stage_artifacts_reusable(context, command),
    )
    if not started:
        raise ControlPlaneError("LEASE_CONFLICT", "流水线未启动，已有调度线程正在运行。")
    return {
        "accepted": True,
        "operation_status": "running",
        "message": "后端自动流水线已启动。",
    }


def _handle_pipeline_resume(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    payload = dict(envelope.payload)
    if not str(payload.get("start_command") or "").strip():
        control = SUPERVISOR.load(context.root)
        payload["start_command"] = str(control.get("current_stage") or "")
    return _handle_pipeline_start(
        context,
        replace_command_envelope(envelope, kind="pipeline.resume", payload=payload),
        operation_id,
    )


def replace_command_envelope(
    envelope: CommandEnvelope,
    *,
    kind: str | None = None,
    payload: dict[str, Any] | None = None,
) -> CommandEnvelope:
    return CommandEnvelope(
        command_id=envelope.command_id,
        workspace_id=envelope.workspace_id,
        kind=kind or envelope.kind,
        payload=payload if payload is not None else envelope.payload,
        goal_id=envelope.goal_id,
        actor=envelope.actor,
        expected_revision=envelope.expected_revision,
        idempotency_key=envelope.idempotency_key,
        confirmation_id=envelope.confirmation_id,
    )


def _handle_pipeline_pause(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    global PAUSE_REQUESTED
    if _same_path(CURRENT_RUN_ROOT, context.root):
        PAUSE_REQUESTED = True
    if SUPERVISOR.is_running(context.root):
        SUPERVISOR.pause()
        _append_log(f"[暂停] CommandGateway 正在停止: {context.workspace_id}/{CURRENT_TASK}")
        _terminate_workspace_process(context)
        return {"accepted": True, "operation_status": "pausing", "message": "已发送暂停指令。"}
    return {"accepted": True, "operation_status": "paused", "message": "Operation 已处于暂停状态。"}


def _handle_pipeline_cancel(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    global PAUSE_REQUESTED
    if _same_path(CURRENT_RUN_ROOT, context.root):
        PAUSE_REQUESTED = True
    if SUPERVISOR.is_running(context.root):
        SUPERVISOR.cancel()
        _append_log(f"[取消] CommandGateway 正在终止: {context.workspace_id}/{CURRENT_TASK}")
        _terminate_workspace_process(context)
        return {"accepted": True, "operation_status": "cancelling", "message": "已发送取消指令。"}
    return {"accepted": True, "operation_status": "cancelled", "message": "Operation 已取消。"}


def _handle_pipeline_skip(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    stage_id = str(envelope.payload.get("stage_id") or "").strip()
    reason = str(envelope.payload.get("reason") or "").strip()
    if not stage_id or not reason:
        raise ControlPlaneError("COMMAND_INVALID", "跳过阶段必须提供 stage_id 和 reason。", status_code=400)
    # V1 registry has no optional-stage contract. Fail closed until a stage is
    # explicitly declared optional with no required downstream artifacts.
    raise ControlPlaneError(
        "GATE_BLOCKED",
        f"阶段 {stage_id} 尚未声明为 V2 可选阶段，不能跳过。",
        details={"stage_id": stage_id, "operation_id": operation_id},
    )


def _handle_repair_start(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    store = _ensure_v2_repair_import(context)
    operation = store.operation(operation_id) or {}
    fencing_token = int(operation.get("fencing_token") or 0)
    candidates = _minimal_repair_candidates(context.root, context=context)
    if not candidates:
        return {
            "accepted": True,
            "operation_status": "succeeded",
            "message": "当前没有可自动修复的阻断问题。",
        }
    auto_count = sum(1 for issue in candidates if _issue_has_auto_repair(issue))
    job = create_authorized_repair_job(
        context.root,
        operation_id=operation_id,
        issue_fingerprints=[_issue_repair_fingerprint(issue) for issue in candidates],
        total_count=len(candidates),
        auto_count=auto_count,
        manual_count=max(0, len(candidates) - auto_count),
        resume_command=_minimal_repair_resume_command(context.root),
    )
    result = _trigger_repair_job(
        context.root,
        "",
        allow_remint=False,
        control_operation_id=operation_id,
        control_fencing_token=fencing_token,
        # V2 keeps repair and pipeline resume as separate audited Commands.
        resume_pipeline=False,
    )
    if not result.get("ok"):
        code = "LEASE_CONFLICT" if result.get("busy") else "GATE_BLOCKED"
        raise ControlPlaneError(code, str(result.get("message") or "最小修复未能启动。"))
    return {
        "accepted": True,
        "operation_status": "running",
        "message": str(result.get("message") or "已开始最小修复。"),
    }


def _handle_repair_issues(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    from agent.repair import execute_repair_batch

    raw_ids = envelope.payload.get("issue_ids")
    if not isinstance(raw_ids, list):
        raise ControlPlaneError("COMMAND_INVALID", "issue_ids 必须是数组。", status_code=400)
    issue_ids = list(dict.fromkeys(str(item).strip() for item in raw_ids if str(item).strip()))
    if not issue_ids:
        raise ControlPlaneError("COMMAND_INVALID", "issue_ids 不能为空。", status_code=400)
    if len(issue_ids) > 100:
        raise ControlPlaneError("COMMAND_INVALID", "单次最多修复 100 个问题。", status_code=400)
    result = execute_repair_batch(context.root, issue_ids, confirm=True, dry_run=False)
    from artifact_manifest import record_external_chapter_mutation

    record_external_chapter_mutation(context, disposition="issue_repair")
    failed = result.get("failed") if isinstance(result.get("failed"), list) else []
    still_open = result.get("still_open") if isinstance(result.get("still_open"), list) else []
    if not result.get("ok") or failed or still_open:
        raise ControlPlaneError(
            "GATE_BLOCKED",
            str(result.get("message") or "问题修复后仍有阻断项。"),
            details={"failed": failed, "still_open": still_open},
        )
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": str(result.get("message") or f"已修复 {len(issue_ids)} 个问题。"),
    }


def _handle_accept_issue_risk(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    from agent.issues import (
        accept_risk_enabled,
        append_issue_log,
        classify_issue_risk,
        write_open_issues_projection,
        write_risk_register,
    )

    issue_id = str(envelope.payload.get("issue_id") or "").strip()
    reason = str(envelope.payload.get("reason") or "").strip()
    if not issue_id:
        raise ControlPlaneError("COMMAND_INVALID", "缺少 issue_id。", status_code=400)
    if len("".join(reason.split())) < 8:
        raise ControlPlaneError("COMMAND_INVALID", "接受风险原因至少需要 8 个有效字符。", status_code=400)
    if not accept_risk_enabled():
        raise ControlPlaneError("POLICY_DENIED", "未开启接受风险功能。")
    actor = envelope.actor if isinstance(envelope.actor, dict) else {}
    actor_id = str(actor.get("id") or "anonymous").strip()[:128]
    actor_role = str(actor.get("role") or "user").strip().lower()
    store = _ensure_v2_issue_import(context)
    issues = store.issue_states()
    found = next((dict(item) for item in issues if str(item.get("id") or "") == issue_id), None)
    if found is None:
        raise ControlPlaneError("ISSUE_NOT_FOUND", f"未找到问题: {issue_id}", status_code=404)
    risk_class = classify_issue_risk(found)
    if risk_class == "fatal":
        raise ControlPlaneError(
            "POLICY_DENIED",
            "fatal 废标项禁止通过接受风险关闭。",
            details={"policy_code": "fatal_forbidden"},
        )
    if risk_class == "qualification":
        raise ControlPlaneError(
            "POLICY_DENIED",
            "资格材料缺失不可直接接受风险，请补料或标记 deferred。",
            details={"policy_code": "qualification_deferred_only"},
        )
    if risk_class == "critical" and actor_role != "admin":
        raise ControlPlaneError(
            "POLICY_DENIED",
            "critical 合规冲突仅管理员可接受风险。",
            details={"policy_code": "admin_required"},
        )
    if str(found.get("severity") or "") not in {"block", "warn"}:
        raise ControlPlaneError("GATE_BLOCKED", "仅 block/warn 问题支持接受风险。")

    accepted_at = datetime.now().astimezone().isoformat(timespec="seconds")
    evidence = dict(found.get("evidence") or {})
    evidence.setdefault(
        "pre_accept_snapshot",
        {
            "status": found.get("status"),
            "severity": found.get("severity"),
            "detail": str(found.get("detail") or "")[:500],
        },
    )
    found.update(
        {
            "status": "accepted",
            "updated_at": accepted_at,
            "accepted_at": accepted_at,
            "accepted_by": actor_id,
            "accept_reason": reason[:500],
            "evidence": evidence,
            "risk_class": risk_class,
        }
    )
    decision = {
        "risk_class": risk_class,
        "reason": reason[:500],
        "accepted_at": accepted_at,
        "evidence": evidence,
        "confirmation_id": envelope.confirmation_id,
    }
    store.update_issue_state_with_policy(
        found,
        decision_type="accept_risk",
        decision=decision,
        actor={"type": str(actor.get("type") or "authenticated_user"), "id": actor_id, "role": actor_role},
        source="v2_command",
    )
    updated_issues = [found if str(item.get("id") or "") == issue_id else item for item in issues]
    projection_warning = ""
    try:
        append_issue_log(context.root, found)
        write_open_issues_projection(context.root, updated_issues)
        write_risk_register(context.root)
    except Exception as exc:
        projection_warning = f"；V1 兼容投影刷新失败: {exc}"
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": f"风险接受已记录。{projection_warning}",
    }


def _record_quality_gate_evaluation(
    context: WorkspaceContext,
    command: str,
    result: dict[str, Any],
    *,
    verdict: str | None = None,
) -> dict[str, Any]:
    store = _ensure_v2_issue_import(context)
    issues = store.issue_states()
    findings: list[dict[str, Any]] = []
    for issue in issues:
        if not isinstance(issue, dict):
            continue
        target = issue.get("target") if isinstance(issue.get("target"), dict) else {}
        identity = {
            "code": str(issue.get("code") or ""),
            "stage_id": str(issue.get("stage_id") or ""),
            "target": target,
        }
        finding_id = hashlib.sha256(
            json.dumps(identity, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
        ).hexdigest()
        findings.append(
            {
                "finding_id": finding_id,
                "rule_key": identity["code"] or identity["stage_id"] or "unknown",
                "target_key": json.dumps(target, ensure_ascii=False, sort_keys=True, separators=(",", ":")),
                "risk_class": str(issue.get("risk_class") or issue.get("severity") or "warn"),
                "issue_id": str(issue.get("id") or ""),
                "status": str(issue.get("status") or "open"),
                "evidence": {"command": command, "result": result.get("summary") or result.get("message") or ""},
            }
        )
    blocks = [
        item for item in findings
        if item["status"] in {"open", "in_progress"} and item["risk_class"] in {"block", "fatal", "critical"}
    ]
    resolved_verdict = verdict or ("block" if blocks else "pass")
    if resolved_verdict == "error":
        findings.append(
            {
                "finding_id": hashlib.sha256(f"{command}:gate-error".encode("utf-8")).hexdigest(),
                "rule_key": "gate_evaluator",
                "target_key": command,
                "risk_class": "fatal",
                "issue_id": "",
                "status": "open",
                "evidence": {"command": command, "error": str(result.get("error") or result.get("message") or "unknown")},
            }
        )
    fingerprint = hashlib.sha256(
        json.dumps(
            {"command": command, "result": result, "findings": findings},
            ensure_ascii=False, sort_keys=True, separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()
    evaluation = store.record_gate_evaluation(
        command=command,
        verdict=resolved_verdict,
        input_fingerprint=fingerprint,
        findings=findings,
        source="quality.revalidate",
    )
    return evaluation


def _handle_quality_revalidate(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    from agent.repair import revalidate_gate

    command = str(envelope.payload.get("command") or "").strip()
    allowed = {str(step.get("command") or "") for step in WORKFLOW_STEPS}
    if not command or command not in allowed:
        raise ControlPlaneError("COMMAND_INVALID", f"不可重验的门禁命令: {command or '-'}", status_code=400)
    try:
        result = revalidate_gate(context.root, command)
    except Exception as exc:
        try:
            _record_quality_gate_evaluation(
                context,
                command,
                {"ok": False, "error": f"{type(exc).__name__}: {exc}"},
                verdict="error",
            )
        except Exception:
            pass
        raise ControlPlaneError(
            "STATE_UNAVAILABLE",
            f"门禁重验失败，已保持阻断: {exc}",
            status_code=503,
            retryable=True,
        ) from exc
    if not isinstance(result, dict) or not result.get("ok"):
        safe_result = result if isinstance(result, dict) else {"ok": False, "message": "门禁重验返回无效状态。"}
        evaluation = _record_quality_gate_evaluation(
            context,
            command,
            safe_result,
            verdict="error" if safe_result.get("error") else "block",
        )
        raise ControlPlaneError(
            "GATE_BLOCKED",
            str(safe_result.get("message") or "门禁重验返回无效状态。"),
            details={"command": command, "gate_evaluation_id": evaluation["evaluation_id"]},
        )
    evaluation = _record_quality_gate_evaluation(context, command, result)
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": str(result.get("message") or f"门禁 {command} 已重验。"),
        "gate_evaluation": evaluation,
    }


def _handle_goal_resume(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    from agent.goal import load_goal, resume_goal_after_materials

    goal = load_goal(context.root)
    if not isinstance(goal, dict):
        raise ControlPlaneError("GOAL_NOT_FOUND", "当前工作区没有可恢复的 Goal。", status_code=404)
    status = str(goal.get("status") or "")
    if status not in {"blocked_human", "in_progress"}:
        raise ControlPlaneError(
            "GOAL_STATE_INVALID",
            f"Goal 状态 {status or '-'} 不允许恢复。",
            status_code=409,
        )
    if status == "in_progress":
        return {"accepted": True, "operation_status": "succeeded", "message": "Goal 已在进行中。"}
    resumed = resume_goal_after_materials(
        context.root,
        note=str(envelope.payload.get("note") or "v2_goal_resume"),
    )
    if not isinstance(resumed, dict) or str(resumed.get("status") or "") != "in_progress":
        raise ControlPlaneError("GATE_BLOCKED", "Goal 恢复条件未满足。")
    return {"accepted": True, "operation_status": "succeeded", "message": "Goal 已恢复为 in_progress。"}


def _handle_review_update(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    category = str(envelope.payload.get("category") or "").strip()
    payload = envelope.payload.get("payload")
    if category not in {"template_evidence", "score_coverage", "chapter_review", "global_review"}:
        raise ControlPlaneError("COMMAND_INVALID", "人工复核 category 无效。", status_code=400)
    if not isinstance(payload, dict):
        raise ControlPlaneError("COMMAND_INVALID", "人工复核 payload 必须是对象。", status_code=400)
    try:
        result = apply_manual_review_update(context.root, category, payload)
    except ValueError as exc:
        raise ControlPlaneError("COMMAND_INVALID", str(exc), status_code=400) from exc
    item_id = str(result.get("item_id") or payload.get("item_id") or "").strip()
    actor = envelope.actor if isinstance(envelope.actor, dict) else {}
    ControlStore(context).record_policy_decision(
        issue_id=f"manual-review:{category}:{item_id or 'unknown'}",
        decision_type="manual_review",
        decision={
            "category": category,
            "item_id": item_id,
            "status": str(payload.get("status") or ""),
            "operator_instruction": str(payload.get("operator_instruction") or payload.get("note") or "")[:2000],
            "payload": dict(payload),
        },
        actor={
            "type": str(actor.get("type") or "user"),
            "id": str(actor.get("id") or "anonymous")[:128],
        },
    )
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": f"人工复核 {item_id} 已更新。",
    }


def _handle_document_apply_edit(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    if RUNNING:
        raise ControlPlaneError("LEASE_CONFLICT", "当前已有执行任务，不能并发修改终稿。", status_code=409)
    path = context.root / "outputs" / "final.md"
    if not path.exists() or not path.is_file():
        raise ControlPlaneError("ARTIFACT_NOT_FOUND", "final.md 不存在，请先执行 build-md。", status_code=404)
    expected_hash = str(envelope.payload.get("base_sha256") or "").strip()
    actual_hash = hashlib.sha256(path.read_bytes()).hexdigest()
    if not expected_hash or not hmac.compare_digest(expected_hash, actual_hash):
        raise ControlPlaneError("REVISION_CONFLICT", "final.md 已变化，请刷新后重新确认。", status_code=409)

    mode = str(envelope.payload.get("mode") or "").strip()
    instruction = str(envelope.payload.get("instruction") or "").strip()
    result: dict[str, Any]
    try:
        if mode == "line":
            line_number = int(envelope.payload.get("line_number") or 0)
            new_text = str(envelope.payload.get("new_text") or "").rstrip("\n\r")
            if line_number < 1 or not new_text.strip():
                raise ValueError("缺少有效的 line_number 或 new_text。")
            result = _save_final_md_line_edit(context.root, line_number, new_text, instruction, "v2_line_edit")
        elif mode == "block":
            block_id = str(envelope.payload.get("block_id") or "").strip()
            new_text = str(envelope.payload.get("new_text") or "").rstrip("\n\r")
            if not block_id or not new_text.strip():
                raise ValueError("缺少 block_id 或 new_text。")
            result = _replace_final_md_block(context.root, block_id, new_text, instruction, "v2_block_edit")
        elif mode == "overwrite":
            new_md = str(envelope.payload.get("new_md") or "")
            if not new_md.strip():
                raise ValueError("缺少 new_md。")
            result = _overwrite_final_md(context.root, new_md, instruction, "v2_document_overwrite")
        elif mode == "undo":
            relative = str(envelope.payload.get("backup_path") or "").strip()
            backup = (context.root / relative).resolve()
            allowed = (context.root / "workspace" / "manual_line_edits").resolve()
            if not relative or not backup.is_relative_to(allowed) or not backup.is_file():
                raise ValueError("撤销备份不存在或路径无效。")
            rollback = _backup_final_md(context.root, path.read_text(encoding="utf-8"), "undo_operation")
            path.write_text(backup.read_text(encoding="utf-8"), encoding="utf-8")
            result = {"review": {"backup_path": str(rollback.relative_to(context.root)).replace("\\", "/")}}
            _LAST_BACKUP.pop(context.root.resolve(), None)
        else:
            raise ValueError("不支持的文档编辑模式。")
    except (OSError, ValueError) as exc:
        raise ControlPlaneError("COMMAND_INVALID", str(exc), status_code=400) from exc

    exit_code = _run_sync("build-docx", context.workspace_id, context.root)
    if exit_code != 0:
        backup_relative = str(result.get("review", {}).get("backup_path") or result.get("backup_path") or "")
        backup = (context.root / backup_relative).resolve() if backup_relative else None
        if backup and backup.is_file() and backup.is_relative_to(context.root.resolve()):
            path.write_text(backup.read_text(encoding="utf-8"), encoding="utf-8")
        raise ControlPlaneError("COMMAND_DISPATCH_FAILED", "文档编辑后的 Word 重建失败，已恢复 final.md。")

    from artifact_manifest import record_document_edit_artifacts

    record_document_edit_artifacts(context)

    _PENDING_LINE_EDITS.pop(context.root.resolve(), None)
    _PENDING_DOC_EDIT.pop(context.root.resolve(), None)
    return {"accepted": True, "operation_status": "succeeded", "message": "文档修改已保存，Word 已重新生成。"}


def _handle_workspace_set_profile(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    project_type = str(envelope.payload.get("project_type") or "").strip()
    allowed = {str(item.get("project_type") or "") for item in project_profile_choices()}
    if project_type not in allowed:
        raise ControlPlaneError("COMMAND_INVALID", "项目类型无效。", status_code=400)
    previous_type = str(load_project_profile(context.root).get("project_type") or "")
    save_project_profile(context.root, project_type)
    if previous_type != project_type:
        store = ControlStore(context)
        store.mark_artifact_states_stale(
            [item["artifact_key"] for item in store.artifact_states()],
            reason=f"项目类型已从 {previous_type or '未设置'} 切换为 {project_type}",
            source_command="workspace.set_profile",
        )
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": f"项目类型已切换为 {load_project_profile(context.root).get('label', project_type)}。",
    }


def _handle_workspace_run_utility(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    command = str(envelope.payload.get("command") or "").strip()
    allowed = set(COMMANDS) - set(auto_run_commands())
    if command not in allowed:
        raise ControlPlaneError("COMMAND_INVALID", f"不支持的工作区 utility command: {command}", status_code=400)
    if RUNNING or SUPERVISOR.is_running():
        raise ControlPlaneError("LEASE_CONFLICT", "当前已有任务或流水线正在运行。", status_code=409)
    if command not in {"validate", "init"}:
        gate = _v2_gate_can_proceed(context, command)
        if not gate.get("can_proceed", False):
            raise ControlPlaneError(
                "GATE_BLOCKED",
                str(gate.get("message") or "质量门禁阻断。"),
                details={"gate": gate},
            )
    exit_code = _run_sync(command, context.workspace_id, context.root)
    if exit_code != 0:
        raise ControlPlaneError("COMMAND_DISPATCH_FAILED", f"{command} 执行失败，exit_code={exit_code}")
    return {"accepted": True, "operation_status": "succeeded", "message": f"{command} 执行完成。"}


def _handle_workspace_archive(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    if RUNNING or SUPERVISOR.is_running(context.root):
        raise ControlPlaneError("LEASE_CONFLICT", "运行中的工作区不能归档。", status_code=409)
    runs_root = RUNS_DIR.resolve()
    source = context.root.resolve()
    if source.parent != runs_root or source.name != context.workspace_id or not source.is_dir():
        raise ControlPlaneError("WORKSPACE_INVALID", "工作区路径无效，已拒绝归档。", status_code=400)
    trash_root = (runs_root / ".trash").resolve()
    if trash_root.parent != runs_root:
        raise ControlPlaneError("WORKSPACE_INVALID", "归档目录无效。", status_code=400)
    trash_root.mkdir(parents=True, exist_ok=True)
    destination = (trash_root / f"{context.workspace_id}_{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}").resolve()
    if destination.parent != trash_root or destination.exists():
        raise ControlPlaneError("WORKSPACE_INVALID", "归档目标无效。", status_code=400)

    def archive_after_commit() -> None:
        global ACTIVE_RUN_ID, ACTIVE_RUN_ROOT
        close_chat_store(source)
        shutil.move(str(source), str(destination))
        if context.workspace_id == ACTIVE_RUN_ID:
            ACTIVE_RUN_ID = ""
            ACTIVE_RUN_ROOT = None
            ACTIVE_RUN_FILE.parent.mkdir(parents=True, exist_ok=True)
            ACTIVE_RUN_FILE.write_text("", encoding="utf-8")
        _append_log(f"[工作空间] 已归档: {context.workspace_id} -> {destination.name}")

    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": f"工作空间 {context.workspace_id} 已移入可恢复归档区。",
        "_after_commit": archive_after_commit,
    }


def _handle_workspace_clean(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    if RUNNING or SUPERVISOR.is_running(context.root):
        raise ControlPlaneError("LEASE_CONFLICT", "运行中的工作区不能清理。", status_code=409)
    root = context.root.resolve()
    trash_root = (root / ".trash").resolve()
    if trash_root.parent != root:
        raise ControlPlaneError("WORKSPACE_INVALID", "清理归档目录无效。", status_code=400)
    destination = (trash_root / f"clean_{time.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}").resolve()
    if destination.parent != trash_root or destination.exists():
        raise ControlPlaneError("WORKSPACE_INVALID", "清理归档目标无效。", status_code=400)

    def clean_after_commit() -> None:
        trash_workspace = destination / "workspace"
        destination.mkdir(parents=True, exist_ok=False)
        close_chat_store(root)
        outputs = root / "outputs"
        if outputs.exists():
            shutil.move(str(outputs), str(destination / "outputs"))
        workspace = root / "workspace"
        trash_workspace.mkdir(parents=True, exist_ok=True)
        if workspace.exists():
            for child in list(workspace.iterdir()):
                if child.name.startswith("control.db"):
                    continue
                shutil.move(str(child), str(trash_workspace / child.name))
        workspace.mkdir(parents=True, exist_ok=True)
        outputs.mkdir(parents=True, exist_ok=True)
        _append_log(f"[清理] 工作区产物已移入可恢复归档: {destination.name}")

    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": "workspace/ 兼容状态与 outputs/ 已移入可恢复归档区。",
        "_after_commit": clean_after_commit,
    }


def _handle_rewrite_chapters(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    raw_targets = envelope.payload.get("targets")
    targets = [dict(item) for item in raw_targets if isinstance(item, dict)] if isinstance(raw_targets, list) else []
    if not targets:
        raw_ids = envelope.payload.get("chapter_ids")
        if isinstance(raw_ids, list):
            targets = [{"chapter_id": str(item).strip()} for item in raw_ids if str(item).strip()]
    chapter_ids = [str(item.get("chapter_id") or "").strip() for item in targets]
    chapter_ids = list(dict.fromkeys(item for item in chapter_ids if item))
    if not chapter_ids:
        raise ControlPlaneError("COMMAND_INVALID", "定向改稿必须提供 chapter_ids。", status_code=400)
    if len(chapter_ids) > 100:
        raise ControlPlaneError("COMMAND_INVALID", "单次定向改稿最多 100 章。", status_code=400)
    operation = ControlStore(context).operation(operation_id) or {}
    fencing_token = int(operation.get("fencing_token") or 0)
    result = _trigger_rewrite_targets_inline(
        [{"chapter_id": item} for item in chapter_ids],
        root=context.root,
        run_id=context.workspace_id,
        control_operation_id=operation_id,
        control_fencing_token=fencing_token,
    )
    if not result.get("ok"):
        code = "LEASE_CONFLICT" if "正在运行" in str(result.get("message") or "") else "COMMAND_DISPATCH_FAILED"
        raise ControlPlaneError(code, str(result.get("message") or "定向改稿未能启动。"))
    return {
        "accepted": True,
        "operation_status": "running",
        "message": str(result.get("message") or "已开始定向改稿。"),
    }


def _material_fulfillment_verified(item: dict[str, Any]) -> bool:
    lifecycle = str(item.get("lifecycle_status") or "").strip().lower()
    evidence = str(item.get("evidence_status") or "").strip().lower()
    if lifecycle in {"verified", "human_verified"}:
        return True
    if evidence in {"verified", "satisfied"}:
        return True
    # Legacy refill may replace verified with resolved/injected. Preserve that
    # compatibility only when verification provenance is still present; a bare
    # legacy ready flag must never become authoritative by passing through refill.
    return lifecycle in {"resolved", "injected"} and bool(
        item.get("verified_at") or item.get("verification_confidence")
    )


def _protected_material(item: dict[str, Any]) -> bool:
    category = str(item.get("category") or "").strip().lower()
    severity = str(item.get("severity") or "").strip().lower()
    return category in {"qualification", "disqualification", "mandatory_doc"} or severity in {
        "fatal",
        "critical",
        "block",
        "blocker",
    }


def _authoritative_material_state(item: dict[str, Any]) -> dict[str, Any]:
    normalized = dict(item)
    lifecycle = str(normalized.get("lifecycle_status") or "").strip().lower()
    if _material_fulfillment_verified(normalized):
        normalized["response_status"] = "ready"
    elif lifecycle in {"waived", "not_applicable"}:
        normalized["response_status"] = "waived"
    else:
        # V2 never projects submitted/uploaded/rejected/missing evidence as ready.
        normalized["response_status"] = "deferred"
    return normalized


def _material_items(context: WorkspaceContext) -> list[dict[str, Any]]:
    # V2 Commands must not silently convert their V1 projection into authority.
    # migration.scan is the sole import boundary; command handlers may still
    # update the projection after a successful V2 state transition.
    return _ensure_v2_material_import(context).material_states()


def _material_item(context: WorkspaceContext, item_id: str) -> dict[str, Any]:
    value = str(item_id or "").strip()
    if not value:
        raise ControlPlaneError("COMMAND_INVALID", "缺少 item_id。", status_code=400)
    item = next(
        (
            dict(row)
            for row in _material_items(context)
            if isinstance(row, dict) and str(row.get("item_id") or "").strip() == value
        ),
        None,
    )
    if item is None:
        raise ControlPlaneError("COMMAND_INVALID", f"材料项不存在: {value}", status_code=404)
    return item


def _sync_material_state_from_projection(
    context: WorkspaceContext,
    item_id: str,
    *,
    persist: bool = True,
) -> dict[str, Any]:
    from materials_checklist import load_materials_checklist

    checklist = load_materials_checklist(context.root)
    items = checklist.get("items") if isinstance(checklist.get("items"), list) else []
    item = next(
        (dict(row) for row in items if isinstance(row, dict) and str(row.get("item_id") or "") == item_id),
        None,
    )
    if item is None:
        raise ControlPlaneError("STATE_UNAVAILABLE", f"材料 V1 投影缺少条目: {item_id}", status_code=503)
    authoritative = _authoritative_material_state(item)
    if not persist:
        return authoritative
    return ControlStore(context).upsert_material_state(authoritative)


def _workspace_material_path(context: WorkspaceContext, raw_path: Any) -> Path:
    value = str(raw_path or "").strip()
    if not value:
        raise ControlPlaneError("COMMAND_INVALID", "登记上传材料必须提供 uploaded_path。", status_code=400)
    candidate = Path(value)
    if not candidate.is_absolute():
        candidate = context.root / candidate
    resolved = candidate.resolve()
    if not resolved.is_relative_to(context.root.resolve()) or not resolved.is_file():
        raise ControlPlaneError(
            "COMMAND_INVALID",
            "uploaded_path 必须指向当前工作区内已存在的文件。",
            status_code=400,
        )
    return resolved


def _resume_goal_for_verified_material(context: WorkspaceContext, item_id: str, note: str) -> None:
    try:
        from agent.goal import load_goal, resume_goal_after_materials

        goal = load_goal(context.root)
        if goal and str(goal.get("status")) == "blocked_human":
            resume_goal_after_materials(context.root, note=note, item_ids=[item_id])
    except Exception:
        # Goal compatibility state is not authoritative for this material command.
        pass


def _handle_materials_upload(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    from materials_checklist import mark_material_uploaded

    item_id = str(envelope.payload.get("item_id") or "").strip()
    _material_item(context, item_id)
    upload_token = str(envelope.payload.get("upload_token") or "").strip()
    if not upload_token:
        raise ControlPlaneError(
            "UPLOAD_TOKEN_REQUIRED",
            "登记材料必须使用当前工作区签发的一次性 upload_token。",
            status_code=400,
        )
    store = ControlStore(context)
    staged = store.material_upload(upload_token)
    if not staged or str(staged.get("status") or "") != "pending":
        raise ControlPlaneError("UPLOAD_TOKEN_INVALID", "upload_token 不存在或不可用。", status_code=409)
    uploaded_path = _workspace_material_path(context, staged.get("staged_path"))
    actual_sha256 = hashlib.sha256(uploaded_path.read_bytes()).hexdigest()
    if actual_sha256 != str(staged.get("sha256") or ""):
        raise ControlPlaneError("UPLOAD_HASH_MISMATCH", "暂存材料摘要不匹配，已拒绝登记。", status_code=409)
    submission_receipt = store.record_material_submission(
        item_id=item_id,
        upload=staged,
        actor=envelope.actor if isinstance(envelope.actor, dict) else {},
        source="materials.upload",
        consume_upload=True,
    )
    result = mark_material_uploaded(
        context.root,
        item_id,
        uploaded_path=str(uploaded_path),
        note=str(envelope.payload.get("note") or envelope.payload.get("reason") or "").strip(),
        rebuild=True,
        auto_verify=True,
    )
    if not result.get("ok"):
        raise ControlPlaneError("GATE_BLOCKED", str(result.get("message") or "上传材料验证失败。"))
    lifecycle = str(result.get("lifecycle_status") or "uploaded")
    projected = _sync_material_state_from_projection(context, item_id, persist=False)
    projected["lifecycle_status"] = lifecycle
    verification = result.get("verification") if isinstance(result.get("verification"), dict) else {}
    if lifecycle == "verified":
        projected["evidence_status"] = "verified"
        projected["verified_at"] = (
            verification.get("verified_at")
            or result.get("verified_at")
            or datetime.now().isoformat(timespec="seconds")
        )
        projected["verification_confidence"] = verification.get("confidence")
    verification_receipt = store.record_material_verification(
        item_id=item_id,
        verification_type="auto_upload",
        verdict=lifecycle if lifecycle in {"verified", "rejected"} else "uploaded",
        verification={
            "lifecycle_status": lifecycle,
            "upload_sha256": str(staged.get("sha256") or ""),
            "upload_filename": str(staged.get("filename") or ""),
            "confidence": verification.get("confidence"),
            "verification": verification,
        },
        actor=envelope.actor if isinstance(envelope.actor, dict) else {},
        source="materials.upload",
        material_state=_authoritative_material_state(projected),
    )
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": str(result.get("message") or f"材料 {item_id} 已登记，状态={lifecycle}。"),
        "material_submission": submission_receipt,
        "material_verification": verification_receipt,
    }


def _handle_materials_verify(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    from agent.material_verifier import verify_material
    from materials_checklist import update_item_response

    item_id = str(envelope.payload.get("item_id") or "").strip()
    _material_item(context, item_id)
    result = verify_material(
        context.root,
        item_id,
        uploaded_path="",
        note=str(envelope.payload.get("note") or "").strip(),
    )
    if not result.get("ok"):
        raise ControlPlaneError("GATE_BLOCKED", str(result.get("message") or "材料验证失败。"))
    lifecycle = str(result.get("lifecycle_status") or "uploaded").strip().lower()
    if lifecycle not in {"uploaded", "verified", "rejected"}:
        raise ControlPlaneError("STATE_UNAVAILABLE", "材料验证返回了未知生命周期，已拒绝更新。", status_code=503)
    update = update_item_response(
        context.root,
        item_id,
        response_status=lifecycle,
        reason=str(result.get("message") or "材料自动验证"),
        rebuild=True,
    )
    if not update.get("ok"):
        raise ControlPlaneError("COMMAND_DISPATCH_FAILED", str(update.get("message") or "验证状态保存失败。"))
    projected = _sync_material_state_from_projection(context, item_id, persist=False)
    projected["lifecycle_status"] = lifecycle
    projected["verification_confidence"] = result.get("confidence")
    if lifecycle == "verified":
        projected["evidence_status"] = "verified"
        projected["verified_at"] = result.get("verified_at") or datetime.now().isoformat(timespec="seconds")
    store = ControlStore(context)
    verification_receipt = store.record_material_verification(
        item_id=item_id,
        verification_type="automatic",
        verdict=lifecycle if lifecycle in {"verified", "rejected"} else "uploaded",
        verification={
            "lifecycle_status": lifecycle,
            "confidence": result.get("confidence"),
            "verified_at": result.get("verified_at"),
            "message": result.get("message"),
        },
        actor=envelope.actor if isinstance(envelope.actor, dict) else {},
        source="materials.verify",
        material_state=_authoritative_material_state(projected),
    )
    if lifecycle == "verified":
        _resume_goal_for_verified_material(context, item_id, f"material_verified:{item_id}")
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": str(result.get("message") or f"材料 {item_id} 验证完成。"),
        "material_verification": verification_receipt,
    }


def _handle_materials_confirm_verification(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    from agent.material_verifier import human_confirm_verification
    from materials_checklist import update_item_response

    item_id = str(envelope.payload.get("item_id") or "").strip()
    _material_item(context, item_id)
    accept = envelope.payload.get("accept", True)
    if not isinstance(accept, bool):
        raise ControlPlaneError("COMMAND_INVALID", "accept 必须是布尔值。", status_code=400)
    reason = str(envelope.payload.get("reason") or "").strip()
    if not accept and len(reason) < 4:
        raise ControlPlaneError("COMMAND_INVALID", "拒绝材料核验必须填写原因。", status_code=400)
    actor = envelope.actor if isinstance(envelope.actor, dict) else {}
    operator = str(actor.get("id") or actor.get("type") or "unknown-actor").strip()[:80]
    result = human_confirm_verification(
        context.root,
        item_id,
        operator=operator,
        accept=accept,
        reason=reason,
    )
    if not result.get("ok"):
        raise ControlPlaneError("GATE_BLOCKED", str(result.get("message") or "材料核验记录不存在。"))
    lifecycle = "verified" if accept else "rejected"
    update = update_item_response(
        context.root,
        item_id,
        response_status=lifecycle,
        reason=reason or f"人工确认验证通过 by {operator}",
        rebuild=True,
    )
    if not update.get("ok"):
        raise ControlPlaneError("COMMAND_DISPATCH_FAILED", str(update.get("message") or "核验结论保存失败。"))
    projected = _sync_material_state_from_projection(context, item_id, persist=False)
    projected["lifecycle_status"] = lifecycle
    projected["evidence_status"] = "verified" if accept else "rejected"
    if accept:
        projected["verified_at"] = datetime.now().isoformat(timespec="seconds")
    store = ControlStore(context)
    verification_receipt = store.record_material_verification(
        item_id=item_id,
        verification_type="human",
        verdict=lifecycle,
        verification={"accepted": accept, "reason": reason, "operator": operator},
        actor=actor,
        source="materials.confirm_verification",
        material_state=_authoritative_material_state(projected),
    )
    if accept:
        _resume_goal_for_verified_material(context, item_id, f"material_human_verified:{item_id}")
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": f"材料 {item_id} 已人工{'确认通过' if accept else '拒绝'}。",
        "material_verification": verification_receipt,
    }


def _handle_materials_update(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    from materials_checklist import (
        build_materials_checklist,
        update_item_response,
    )

    raw_updates = envelope.payload.get("updates")
    if isinstance(raw_updates, list):
        updates = [dict(item) for item in raw_updates if isinstance(item, dict)]
    else:
        updates = [dict(envelope.payload)]
    if not updates:
        raise ControlPlaneError("COMMAND_INVALID", "材料更新列表不能为空。", status_code=400)
    if len(updates) > 200:
        raise ControlPlaneError("COMMAND_INVALID", "单次最多更新 200 条材料。", status_code=400)

    items = _material_items(context)
    by_id = {str(item.get("item_id") or ""): item for item in items if isinstance(item, dict)}
    normalized: list[dict[str, Any]] = []
    for update in updates:
        item_id = str(update.get("item_id") or "").strip()
        status = str(update.get("response_status") or update.get("status") or "").strip().lower()
        reason = str(update.get("reason") or update.get("note") or "").strip()
        item = by_id.get(item_id)
        if not item:
            raise ControlPlaneError("COMMAND_INVALID", f"材料项不存在: {item_id}", status_code=404)
        if status not in {"deferred", "ready", "waived"}:
            raise ControlPlaneError("COMMAND_INVALID", f"无效材料状态: {status}", status_code=400)
        if status == "ready" and not _material_fulfillment_verified(item):
            raise ControlPlaneError(
                "GATE_BLOCKED",
                f"材料 {item_id} 尚未验证，submitted/uploaded 不能直接标记为 ready。",
                details={"item_id": item_id, "lifecycle_status": item.get("lifecycle_status")},
            )
        if status == "waived":
            if _protected_material(item):
                raise ControlPlaneError(
                    "GATE_BLOCKED",
                    f"材料 {item_id} 属于资格/必交/阻断材料，不允许放弃。",
                    details={"item_id": item_id},
                )
            if len(reason) < 4:
                raise ControlPlaneError("COMMAND_INVALID", "放弃可选材料必须填写原因。", status_code=400)
        normalized.append(
            {
                "item_id": item_id,
                "response_status": status,
                "reason": reason,
                "suggested_attachment": str(update.get("suggested_attachment") or ""),
                "persist_status": (
                    "verified"
                    if status == "ready" and _material_fulfillment_verified(item)
                    else status
                ),
            }
        )

    result: dict[str, Any] = {"ok": True}
    for item in normalized:
        result = update_item_response(
            context.root,
            item["item_id"],
            response_status=item["persist_status"],
            reason=item["reason"],
            suggested_attachment=item["suggested_attachment"],
            rebuild=False,
        )
        if not result.get("ok"):
            raise ControlPlaneError("COMMAND_DISPATCH_FAILED", str(result.get("message") or "材料状态更新失败。"))
    build_materials_checklist(context.root)
    for item in normalized:
        _sync_material_state_from_projection(context, item["item_id"])
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": str(result.get("message") or f"已更新 {len(normalized)} 条材料。"),
    }


def _handle_materials_rebuild(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    from materials_checklist import build_materials_checklist, load_materials_checklist, update_item_response

    authoritative = _material_items(context)
    path = build_materials_checklist(context.root)
    for item in authoritative:
        lifecycle = str(item.get("lifecycle_status") or "").strip().lower()
        response_status = str(item.get("response_status") or "deferred").strip().lower()
        persisted = (
            lifecycle
            if lifecycle
            in {"verified", "uploaded", "rejected", "waived", "requested", "missing", "not_applicable"}
            else response_status
        )
        result = update_item_response(
            context.root,
            str(item.get("item_id") or ""),
            response_status=persisted,
            reason=str(item.get("reason") or ""),
            suggested_attachment=str(item.get("suggested_attachment") or ""),
            rebuild=False,
        )
        if not result.get("ok"):
            raise ControlPlaneError("STATE_UNAVAILABLE", str(result.get("message") or "材料投影更新失败。"), status_code=503)
    if authoritative:
        path = build_materials_checklist(context.root)
    checklist = load_materials_checklist(context.root)
    if not path.exists() or not isinstance(checklist, dict):
        raise ControlPlaneError("STATE_UNAVAILABLE", "材料清单重建后状态不可用。", status_code=503)
    summary = checklist.get("summary") if isinstance(checklist.get("summary"), dict) else {}
    for item in checklist.get("items") or []:
        if isinstance(item, dict):
            ControlStore(context).upsert_material_state(
                _authoritative_material_state(item),
                source="v1_projection",
            )
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": f"材料清单已重建：total={summary.get('total', 0)}。",
    }


def _trigger_material_refill(
    context: WorkspaceContext,
    operation_id: str,
    fencing_token: int,
    *,
    chapter_ids: list[str] | None,
    replan_jobs: bool,
    max_chapters: int,
) -> dict[str, Any]:
    global RUNNING, CURRENT_TASK, CURRENT_RUN_ID, CURRENT_RUN_ROOT, PAUSE_REQUESTED
    if RUNNING or SUPERVISOR.is_running():
        return {"ok": False, "busy": True, "message": "当前已有任务在运行。"}
    RUNNING = True
    CURRENT_TASK = "materials-refill"
    CURRENT_RUN_ID = context.workspace_id
    CURRENT_RUN_ROOT = context.root
    PAUSE_REQUESTED = False

    def worker() -> None:
        global RUNNING, CURRENT_TASK, CURRENT_RUN_ID, CURRENT_RUN_ROOT
        status = "failed"
        message = "材料回填失败。"
        error: dict[str, Any] | None = None
        try:
            from materials_checklist import refill_material_gaps, revalidate_issues_after_materials

            ControlStore(context).sync_operation(
                operation_id,
                "running",
                message="正在按已验证材料回填章节。",
                fencing_token=fencing_token,
            )
            result = refill_material_gaps(
                context.root,
                chapter_ids=chapter_ids,
                replan_jobs=replan_jobs,
                max_chapters=max_chapters,
            )
            from artifact_manifest import record_external_chapter_mutation

            record_external_chapter_mutation(context, disposition="materials_refill")
            try:
                result["revalidate"] = revalidate_issues_after_materials(context.root)
            except Exception as exc:
                result["revalidate_error"] = str(exc)
            failed = result.get("failed") if isinstance(result.get("failed"), list) else []
            status = "succeeded" if result.get("ok") and not failed else "blocked"
            message = str(result.get("message") or "材料回填完成。")
            error = {"failed": failed, "revalidate_error": result.get("revalidate_error")} if status == "blocked" else None
        except Exception as exc:
            error = {"message": str(exc)}
            message = f"材料回填失败: {exc}"
        finally:
            try:
                ControlStore(context).sync_operation(
                    operation_id,
                    status,
                    message=message,
                    error=error,
                    fencing_token=fencing_token,
                )
            finally:
                RUNNING = False
                CURRENT_TASK = ""
                CURRENT_RUN_ID = ""
                CURRENT_RUN_ROOT = None

    threading.Thread(target=worker, daemon=True, name=f"materials-refill-{context.workspace_id}").start()
    return {"ok": True, "message": "材料回填已启动。"}


def _handle_materials_refill(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    from materials_checklist import load_materials_checklist

    items = _material_items(context)
    projection = load_materials_checklist(context.root)
    projection_items = projection.get("items") if isinstance(projection.get("items"), list) else []
    unsafe_ready = [
        str(item.get("item_id") or "")
        for item in [*items, *projection_items]
        if isinstance(item, dict)
        and str(item.get("response_status") or "") == "ready"
        and not _material_fulfillment_verified(item)
    ]
    if unsafe_ready:
        raise ControlPlaneError(
            "GATE_BLOCKED",
            "存在未验证却标记为 ready 的材料，已拒绝回填。",
            details={"item_ids": unsafe_ready[:50]},
        )
    raw_ids = envelope.payload.get("chapter_ids")
    if raw_ids is not None and not isinstance(raw_ids, list):
        raise ControlPlaneError("COMMAND_INVALID", "chapter_ids 必须是数组。", status_code=400)
    chapter_ids = [str(item).strip() for item in raw_ids or [] if str(item).strip()] or None
    try:
        max_chapters = max(1, min(int(envelope.payload.get("max_chapters") or 20), 100))
    except (TypeError, ValueError) as exc:
        raise ControlPlaneError("COMMAND_INVALID", "max_chapters 必须是整数。", status_code=400) from exc
    operation = ControlStore(context).operation(operation_id) or {}
    result = _trigger_material_refill(
        context,
        operation_id,
        int(operation.get("fencing_token") or 0),
        chapter_ids=chapter_ids,
        replan_jobs=bool(envelope.payload.get("replan_jobs", True)),
        max_chapters=max_chapters,
    )
    if not result.get("ok"):
        raise ControlPlaneError("LEASE_CONFLICT", str(result.get("message") or "材料回填未能启动。"))
    return {"accepted": True, "operation_status": "running", "message": result["message"]}


def _handle_gate_revalidate(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    artifact = context.root / "outputs" / "final.docx"
    if not artifact.exists() or not artifact.is_file() or artifact.stat().st_size <= 0:
        raise ControlPlaneError("GATE_BLOCKED", "final.docx 不存在或为空，不能签发正式稿凭据。")
    gate = _v2_gate_can_proceed(context, "build-docx")
    if not gate.get("can_proceed") or gate.get("block_count"):
        raise ControlPlaneError(
            "GATE_BLOCKED",
            str(gate.get("message") or "质量门禁未通过。"),
            details={"blocks": gate.get("blocks") or []},
        )
    try:
        preflight = _v2_export_preflight(context)
        _assert_formal_materials_verified(context)
        _assert_formal_artifacts_ready(context)
    except ControlPlaneError:
        raise
    except Exception as exc:
        raise ControlPlaneError(
            "STATE_UNAVAILABLE",
            f"正式出稿检查读取失败，已拒绝签发: {exc}",
            status_code=503,
            retryable=True,
        ) from exc
    if not isinstance(preflight, dict):
        raise ControlPlaneError("STATE_UNAVAILABLE", "正式出稿检查返回无效状态。", status_code=503)
    if not preflight.get("can_export"):
        raise ControlPlaneError(
            "GATE_BLOCKED",
            str(preflight.get("message") or "正式出稿检查未通过。"),
            details={"checks": preflight.get("checks") or []},
        )
    fingerprint, artifact_sha256 = _formal_gate_fingerprint(context)
    receipt = ControlStore(context).issue_gate_receipt(
        verdict="pass",
        gate_input_fingerprint=fingerprint,
        artifact_path="outputs/final.docx",
        artifact_sha256=artifact_sha256,
        rules_version=_FORMAL_GATE_RULES_VERSION,
        findings=preflight.get("block_issues") or [],
        policy_decisions=ControlStore(context).policy_decisions(),
    )
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": f"正式稿门禁已通过，GateReceipt={receipt.get('receipt_id', '')}",
    }


def _handle_migration_reconcile(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    actor = dict(envelope.actor or {})
    if str(actor.get("role") or "").strip().lower() != "admin":
        raise ControlPlaneError(
            "AUTH_FORBIDDEN",
            "只有管理员可以处理迁移冲突。",
            status_code=403,
        )
    conflict_id = str(envelope.payload.get("conflict_id") or "").strip()
    resolution = str(envelope.payload.get("resolution") or "").strip()
    reason = str(envelope.payload.get("reason") or "").strip()
    if not conflict_id:
        raise ControlPlaneError("COMMAND_INVALID", "迁移协调缺少 conflict_id。", status_code=400)
    store = ControlStore(context)
    conflict = store.resolve_migration_conflict(
        conflict_id,
        resolution=resolution,
        actor=actor,
        reason=reason,
    )
    state = store.migration_state()
    _refresh_migration_report(
        context,
        state,
        action={
            "kind": "migration.reconcile",
            "conflict_id": conflict["conflict_id"],
            "resolution": resolution,
            "actor": actor,
        },
    )
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": f"迁移冲突已处理: {conflict['conflict_id']}",
        "migration": state,
    }


def _handle_migration_scan(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    actor = dict(envelope.actor or {})
    if str(actor.get("role") or "").strip().lower() != "admin":
        raise ControlPlaneError("AUTH_FORBIDDEN", "只有管理员可以扫描旧工作区迁移状态。", status_code=403)
    store = ControlStore(context)
    dry_run = _v1_migration_dry_run(context)
    inventory = dry_run.get("inventory") if isinstance(dry_run.get("inventory"), dict) else {}
    imported = 0
    for item in inventory.get("importable") or []:
        if not isinstance(item, dict):
            continue
        domain = str(item.get("domain") or "")
        legacy = item.get("legacy")
        if domain == "goal":
            imported += store.ensure_goal_state(legacy if isinstance(legacy, dict) else None)
        elif domain == "materials":
            imported += store.ensure_material_states(legacy if isinstance(legacy, list) else [])
        elif domain == "issues":
            imported += store.ensure_issue_states(legacy if isinstance(legacy, list) else [])
        elif domain == "repair_job":
            imported += store.ensure_repair_job_state(legacy if isinstance(legacy, dict) else None)
        elif domain == "agent_activity":
            imported += store.ensure_agent_activity_state(legacy if isinstance(legacy, dict) else None)
    detected = 0
    for item in inventory.get("conflicts") or []:
        if not isinstance(item, dict):
            continue
        store.record_migration_conflict(
            domain=str(item.get("domain") or "unknown"),
            legacy=item.get("legacy"),
            authoritative=item.get("authoritative"),
            reason="管理员迁移扫描发现 V1 与 control.db 状态不一致。",
            exclude_operation_id=operation_id,
        )
        detected += 1
    for category, reason in (("orphans", "旧根目录控制状态未绑定到工作区。"), ("unrecognized", "旧状态文件无法识别。")):
        for item in inventory.get(category) or []:
            store.record_migration_conflict(
                domain="orphan" if category == "orphans" else "unrecognized",
                legacy=item,
                authoritative={},
                reason=reason,
                exclude_operation_id=operation_id,
            )
            detected += 1
    store.record_migration_scan(
        fingerprint=str(dry_run.get("source_fingerprint") or ""),
        manifest=[dict(item) for item in dry_run.get("source_manifest") or [] if isinstance(item, dict)],
        actor=actor,
    )
    legacy_artifact_count = _register_legacy_artifact_inventory(context, store)
    state = store.migration_state()
    from utils import write_json

    write_json(
        context.root / "workspace" / "migration_report.json",
        {
            "version": 1,
            "workspace_id": context.workspace_id,
            "source_fingerprint": dry_run.get("source_fingerprint"),
            "source_manifest": dry_run.get("source_manifest") or [],
            "inventory": inventory,
            "imported_count": imported,
            "legacy_artifact_count": legacy_artifact_count,
            "detected_count": detected,
            "migration": state,
        },
    )
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": (
            f"迁移扫描完成：导入 {imported} 项，登记 {legacy_artifact_count} 个待重建旧 Artifact，"
            f"发现 {detected} 项待协调状态。"
        ),
        "migration": state,
    }


def _handle_migration_cutover(
    context: WorkspaceContext,
    envelope: CommandEnvelope,
    operation_id: str,
) -> dict[str, Any]:
    actor = dict(envelope.actor or {})
    if str(actor.get("role") or "").strip().lower() != "admin":
        raise ControlPlaneError("AUTH_FORBIDDEN", "只有管理员可以切换工作区至 V2 控制面。", status_code=403)
    dry_run = _v1_migration_dry_run(context)
    fingerprint = str(dry_run.get("source_fingerprint") or "")
    store = ControlStore(context)
    cutover = store.activate_migration_cutover(fingerprint=fingerprint, actor=actor)
    state = store.migration_state()
    _refresh_migration_report(
        context,
        state,
        action={"kind": "migration.cutover", "fingerprint": fingerprint, "actor": actor},
    )
    return {
        "accepted": True,
        "operation_status": "succeeded",
        "message": "工作区已完成 V2 控制面切换。",
        "cutover": cutover,
    }


def _command_gateway(context: WorkspaceContext) -> CommandGateway:
    SUPERVISOR.set_status_listener(_sync_pipeline_control_state)
    return CommandGateway(
        context,
        {
            "pipeline.start": _handle_pipeline_start,
            "pipeline.run_stage": _handle_pipeline_start,
            "pipeline.resume": _handle_pipeline_resume,
            "pipeline.pause": _handle_pipeline_pause,
            "pipeline.cancel": _handle_pipeline_cancel,
            "pipeline.skip_stage": _handle_pipeline_skip,
            "repair.start": _handle_repair_start,
            "repair.issues": _handle_repair_issues,
            "issues.accept_risk": _handle_accept_issue_risk,
            "quality.revalidate": _handle_quality_revalidate,
            "goal.resume": _handle_goal_resume,
            "review.update": _handle_review_update,
            "document.apply_edit": _handle_document_apply_edit,
            "workspace.set_profile": _handle_workspace_set_profile,
            "workspace.run_utility": _handle_workspace_run_utility,
            "workspace.archive": _handle_workspace_archive,
            "workspace.clean": _handle_workspace_clean,
            "rewrite.chapters": _handle_rewrite_chapters,
            "materials.upload": _handle_materials_upload,
            "materials.verify": _handle_materials_verify,
            "materials.confirm_verification": _handle_materials_confirm_verification,
            "materials.update": _handle_materials_update,
            "materials.rebuild": _handle_materials_rebuild,
            "materials.refill": _handle_materials_refill,
            "gate.revalidate": _handle_gate_revalidate,
            "migration.scan": _handle_migration_scan,
            "migration.cutover": _handle_migration_cutover,
            "migration.reconcile": _handle_migration_reconcile,
        },
    )


def _command_error_response(exc: ControlPlaneError) -> JSONResponse:
    context_revision = exc.details.get("current_revision") if isinstance(exc.details, dict) else None
    return JSONResponse(
        {
            "ok": False,
            "error": exc.as_dict(),
            "message": exc.message,
            "workspace_revision": context_revision,
        },
        status_code=exc.status_code,
    )


def _artifact_present(root: Path, artifact: Any) -> bool:
    artifact_path = _artifact_path(artifact)
    artifact_kind = _artifact_kind(artifact)
    if artifact_kind == "virtual":
        return True
    if artifact_kind == "glob":
        relative = Path(artifact_path)
        directory = root / relative.parent
        return directory.exists() and any(directory.glob(relative.name))
    target = root / artifact_path
    if target.is_dir():
        return any(target.iterdir())
    return target.exists() and target.is_file() and target.stat().st_size > 0


def _step_by_command(command: str) -> dict[str, Any] | None:
    return next((step for step in WORKFLOW_STEPS if step["command"] == command), None)


def _step_outputs_present(root: Path, command: str) -> bool:
    step = _step_by_command(command)
    if not step:
        return False
    return stage_outputs_ready(root, str(step.get("id", "")))


def _latest_tree_mtime(root: Path) -> float:
    latest = root.stat().st_mtime if root.exists() else 0.0
    for relative in [
        "workspace/run_state.json",
        "workspace/run_state_history.jsonl",
        "workspace/format_check_report.json",
        "workspace/compliance_report.json",
        "outputs/final.md",
        "outputs/final.docx",
    ]:
        path = root / relative
        if path.exists():
            latest = max(latest, path.stat().st_mtime)
    return latest


def _run_progress_summary(run_root: Path) -> dict[str, Any]:
    core_steps = [step for step in WORKFLOW_STEPS if step.get("kind") != "utility"]
    done_count = sum(1 for step in core_steps if stage_outputs_ready(run_root, str(step.get("id", ""))))
    run_state = _read_run_state(run_root)
    failed_command = _command_for_stage(str(run_state.get("stage", "")))
    failed_step = next((step for step in core_steps if step["command"] == failed_command), None)
    failed_outputs_present = _step_outputs_present(run_root, failed_command)
    latest = _latest_tree_mtime(run_root)
    state = str(run_state.get("status", "")).strip()
    is_running_run = RUNNING and _same_path(CURRENT_RUN_ROOT, run_root)
    effective_state = "" if state in {"error", "paused"} and failed_outputs_present else state
    if is_running_run:
        display_state = "running"
        running_step = next((step for step in core_steps if step["command"] == CURRENT_TASK), None)
        display_label = f"运行中 {running_step['label']}" if running_step else "运行中"
    elif failed_step and state in {"error", "paused"} and not failed_outputs_present:
        display_state = "error"
        display_label = f"暂停于 {failed_step['label']}"
    elif done_count >= len(core_steps) and core_steps:
        display_state = "done"
        display_label = "完整流程已完成"
    elif run_state.get("updated_at"):
        display_state = effective_state or "progress"
        display_label = "已开始"
    else:
        display_state = "new"
        display_label = "未开始"
    return {
        "done": done_count,
        "total": len(core_steps),
        "status": display_state,
        "status_label": display_label,
        "stage": run_state.get("stage", ""),
        "message": "" if failed_outputs_present and state in {"error", "paused"} else run_state.get("message", ""),
        "updated_at": run_state.get("updated_at", ""),
        "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(latest)) if latest else "",
    }


async def _create_workspace_request(request: Request, *, select_legacy_active: bool) -> JSONResponse:
    global ACTIVE_RUN_ID, ACTIVE_RUN_ROOT

    if select_legacy_active and RUNNING:
        return JSONResponse(
            {"ok": False, "message": "当前已有任务正在运行，请等待完成。"},
            status_code=409,
        )

    body: dict[str, Any] = {}
    try:
        body = await request.json()
    except Exception:
        body = {}

    run_name = str(body.get("name", "")).strip()
    project_type = str(body.get("project_type", "")).strip()
    expected_pages = int(body.get("expected_pages", 0) or 0)
    if not run_name:
        return JSONResponse(
            {"ok": False, "message": "请先设置工作空间名称。"},
            status_code=400,
        )
    principal = _request_principal(request)
    principal_id = str(principal.get("id") or "").strip()
    if not principal_id:
        return JSONResponse({"ok": False, "message": "缺少服务端认证主体。"}, status_code=401)

    try:
        run_id, run_root = _create_run_workspace(run_name, project_type, expected_pages=expected_pages)
    except Exception as exc:
        return JSONResponse({"ok": False, "message": f"创建运行工作空间失败: {exc}"}, status_code=500)

    ControlStore(WorkspaceContext.resolve(RUNS_DIR, run_id)).grant_workspace_access(principal_id, role="owner")
    if select_legacy_active:
        ACTIVE_RUN_ID = run_id
        ACTIVE_RUN_ROOT = run_root
        ACTIVE_RUN_FILE.write_text(run_id, encoding="utf-8")
    relative_root = run_root.relative_to(ROOT) if run_root.is_relative_to(ROOT) else run_root
    _append_log(f"[运行] 已创建独立工作空间: {relative_root}")
    return JSONResponse({"ok": True, "run": _workspace_payload(run_id, run_root)})


@app.post("/api/v2/workspaces")
async def api_v2_create_workspace(request: Request) -> JSONResponse:
    return await _create_workspace_request(request, select_legacy_active=False)


@app.post("/api/start-run")
async def api_start_run(request: Request) -> JSONResponse:
    return await _create_workspace_request(request, select_legacy_active=True)


def _visible_workspace_runs(request: Request, *, include_legacy_active: bool) -> list[dict[str, Any]]:
    runs: list[dict[str, Any]] = []
    principal = _request_principal(request)
    principal_id = str(principal.get("id") or "").strip() if isinstance(principal, dict) else ""
    is_admin = isinstance(principal, dict) and str(principal.get("role") or "") == "admin"
    if RUNS_DIR.exists():
        run_dirs = [path for path in RUNS_DIR.iterdir() if path.is_dir() and not path.name.startswith(".")]
        for run_root in sorted(run_dirs, key=_latest_tree_mtime, reverse=True):
            if not is_admin:
                try:
                    ControlStore(WorkspaceContext.resolve(RUNS_DIR, run_root.name)).require_workspace_access(
                        principal_id,
                        write=False,
                    )
                except ControlPlaneError:
                    continue
            progress = _run_progress_summary(run_root)
            profile = load_project_profile(run_root)
            runs.append(
                {
                    "id": run_root.name,
                    "root": str(run_root),
                    "relative_root": str(run_root.relative_to(ROOT)) if run_root.is_relative_to(ROOT) else str(run_root),
                    "progress": progress,
                    "project_type": profile.get("project_type", ""),
                    "project_label": profile.get("label", ""),
                    "expected_pages": profile.get("expected_pages", 0),
                }
            )
            if include_legacy_active:
                runs[-1]["active"] = run_root == ACTIVE_RUN_ROOT
    return runs


@app.get("/api/v2/workspaces")
def api_v2_workspaces(request: Request) -> JSONResponse:
    return JSONResponse({"ok": True, "runs": _visible_workspace_runs(request, include_legacy_active=False)})


@app.get("/api/runs")
def api_runs(request: Request) -> JSONResponse:
    _load_active_run_from_disk()
    runs = _visible_workspace_runs(request, include_legacy_active=True)
    return JSONResponse({"ok": True, "active_run_id": ACTIVE_RUN_ID, "runs": runs})


@app.post("/api/select-run")
async def api_select_run(request: Request) -> JSONResponse:
    global ACTIVE_RUN_ID, ACTIVE_RUN_ROOT
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)

    run_id = str(body.get("run_id", "")).strip()
    if not run_id or Path(run_id).name != run_id:
        return JSONResponse({"ok": False, "message": "无效工作空间。"}, status_code=400)

    run_root = (RUNS_DIR / run_id).resolve()
    runs_root = RUNS_DIR.resolve()
    if not run_root.is_relative_to(runs_root) or not run_root.exists() or not run_root.is_dir():
        return JSONResponse({"ok": False, "message": f"工作空间不存在: {run_id}"}, status_code=404)

    try:
        _ensure_workspace_acl(
            WorkspaceContext.resolve(RUNS_DIR, run_id),
            _request_principal(request),
            write=False,
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)

    ACTIVE_RUN_ID = run_id
    ACTIVE_RUN_ROOT = run_root
    RUNS_DIR.mkdir(parents=True, exist_ok=True)
    ACTIVE_RUN_FILE.write_text(run_id, encoding="utf-8")
    _append_log(f"[工作空间] 已切换到: {run_root.relative_to(ROOT)}")
    return JSONResponse({"ok": True, "run": _active_run_payload(), "progress": _run_progress_summary(run_root)})


_MATERIAL_UPLOAD_MAX_BYTES = 50 * 1024 * 1024
_MATERIAL_UPLOAD_EXTENSIONS = {
    ".pdf",
    ".doc",
    ".docx",
    ".xls",
    ".xlsx",
    ".png",
    ".jpg",
    ".jpeg",
    ".txt",
}


def _safe_upload_filename(filename: str) -> str:
    name = Path(str(filename or "")).name.strip()
    if not name or name in {".", ".."}:
        raise ControlPlaneError("UPLOAD_INVALID", "上传文件名无效。", status_code=400)
    extension = Path(name).suffix.lower()
    if extension not in _MATERIAL_UPLOAD_EXTENSIONS:
        raise ControlPlaneError("UPLOAD_TYPE_DENIED", f"不允许的材料文件类型: {extension or '无扩展名'}", status_code=400)
    stem = re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]+", "_", Path(name).stem).strip("._") or "material"
    return f"{stem[:120]}{extension}"


@app.post("/api/v2/workspaces/{workspace_id}/materials/uploads")
async def api_v2_stage_material_upload(
    workspace_id: str,
    file: UploadFile = File(...),
) -> JSONResponse:
    destination: Path | None = None
    try:
        context = _workspace_context(workspace_id)
        filename = _safe_upload_filename(file.filename or "")
        staging_dir = context.root / "workspace" / "material_uploads" / "staging"
        staging_dir.mkdir(parents=True, exist_ok=True)
        destination = (staging_dir / f"{uuid.uuid4().hex}_{filename}").resolve()
        if not destination.is_relative_to(staging_dir.resolve()):
            raise ControlPlaneError("UPLOAD_INVALID", "材料暂存路径越界。", status_code=400)
        digest = hashlib.sha256()
        size_bytes = 0
        with destination.open("xb") as output:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                size_bytes += len(chunk)
                if size_bytes > _MATERIAL_UPLOAD_MAX_BYTES:
                    raise ControlPlaneError(
                        "UPLOAD_TOO_LARGE",
                        "单个材料文件不能超过 50 MB。",
                        status_code=413,
                    )
                digest.update(chunk)
                output.write(chunk)
        if size_bytes <= 0:
            raise ControlPlaneError("UPLOAD_INVALID", "上传文件为空。", status_code=400)
        relative = destination.relative_to(context.root).as_posix()
        staged = ControlStore(context).register_material_upload(
            staged_path=relative,
            filename=filename,
            sha256=digest.hexdigest(),
            size_bytes=size_bytes,
        )
        return JSONResponse(
            {
                "ok": True,
                "upload_token": staged["upload_token"],
                "filename": staged["filename"],
                "sha256": staged["sha256"],
                "size_bytes": staged["size_bytes"],
                "expires_at": staged["expires_at"],
            },
            status_code=201,
        )
    except ControlPlaneError as exc:
        if destination is not None and destination.exists():
            destination.unlink()
        return _command_error_response(exc)
    except Exception as exc:
        if destination is not None and destination.exists():
            destination.unlink()
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"材料暂存失败: {exc}", status_code=503)
        )
    finally:
        await file.close()


@app.post("/api/v2/workspaces/{workspace_id}/commands")
async def api_v2_submit_command(workspace_id: str, request: Request) -> JSONResponse:
    try:
        context = _workspace_context(workspace_id)
        body = await request.json()
        if not isinstance(body, dict):
            raise ControlPlaneError("COMMAND_INVALID", "请求体必须是 JSON 对象。", status_code=400)
        body = dict(body)
        body["actor"] = _request_actor(request, source="v2_api")
        envelope = CommandEnvelope.from_mapping(body, workspace_id=context.workspace_id)
        gateway = _command_gateway(context)
        if envelope.kind in ControlStore.CONFIRMATION_REQUIRED_KINDS:
            if envelope.confirmation_id:
                raise ControlPlaneError(
                    "CONFIRMATION_REQUIRED",
                    "高风险 Command 只能通过已持久化 Action 的确认接口执行。",
                    status_code=400,
                )
            labels = {
                "pipeline.cancel": "确认取消当前任务",
                "pipeline.skip_stage": "确认跳过当前阶段",
                "repair.start": "确认执行最小修复",
                "repair.issues": "确认执行问题最小修复",
                "issues.accept_risk": "确认接受问题风险",
                "rewrite.chapters": "确认执行定向改稿",
                "materials.update": "确认更新材料状态",
                "materials.refill": "确认将已验证材料回填正文",
                "materials.upload": "确认登记并验证上传材料",
                "materials.confirm_verification": "确认材料人工核验结论",
                "review.update": "确认更新人工复核结论",
                "document.apply_edit": "确认修改终稿并重新生成 Word",
                "workspace.set_profile": "确认切换项目类型",
                "workspace.run_utility": "确认执行工作区维护命令",
                "workspace.archive": "确认归档工作区",
                "workspace.clean": "确认清理工作区产物",
                "migration.scan": "确认扫描并登记旧工作区迁移状态",
                "migration.cutover": "确认切换工作区至 V2 控制面",
                "migration.reconcile": "确认处理 V1/V2 迁移冲突",
            }
            label = labels.get(envelope.kind, f"确认执行 {envelope.kind}")
            action = gateway.propose(envelope, label=label, risk="high")
            return JSONResponse(
                {
                    "ok": True,
                    "receipt": {
                        "command_id": envelope.command_id,
                        "operation_id": envelope.payload.get("operation_id"),
                        "status": "requires_confirmation",
                        "workspace_revision": action["expected_revision"],
                        "confirmation_id": action["confirmation_id"],
                        "error": None,
                    },
                    "action": action,
                },
                status_code=202,
            )
        receipt = gateway.submit(envelope)
        return JSONResponse(
            {"ok": receipt.status not in {"rejected"}, "receipt": receipt.as_dict(), "message": receipt.message},
            status_code=202 if receipt.status in {"accepted", "duplicate", "no_op"} else 409,
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError(
                "STATE_UNAVAILABLE",
                f"控制状态不可用，已拒绝执行: {exc}",
                status_code=503,
                retryable=True,
            )
        )


@app.post("/api/v2/workspaces/{workspace_id}/actions/{action_id}/confirm")
@app.post("/api/v2/workspaces/{workspace_id}/confirmations/{action_id}/confirm")
def api_v2_confirm_action(workspace_id: str, action_id: str, request: Request) -> JSONResponse:
    try:
        receipt = _command_gateway(_workspace_context(workspace_id)).confirm(
            action_id,
            actor=_request_actor(request, source="v2_api"),
        )
        return JSONResponse(
            {"ok": receipt.status != "rejected", "receipt": receipt.as_dict(), "message": receipt.message},
            status_code=202 if receipt.status != "rejected" else 409,
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"确认处理失败: {exc}", status_code=503, retryable=True)
        )


@app.post("/api/v2/workspaces/{workspace_id}/actions/{action_id}/decline")
@app.post("/api/v2/workspaces/{workspace_id}/confirmations/{action_id}/decline")
def api_v2_decline_action(workspace_id: str, action_id: str, request: Request) -> JSONResponse:
    try:
        result = _command_gateway(_workspace_context(workspace_id)).decline(
            action_id,
            actor=_request_actor(request, source="v2_api"),
        )
        return JSONResponse({"ok": True, **result})
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"拒绝确认失败: {exc}", status_code=503, retryable=True)
        )


@app.get("/api/v2/workspaces/{workspace_id}/snapshot")
def api_v2_workspace_snapshot(workspace_id: str) -> JSONResponse:
    try:
        context = _workspace_context(workspace_id)
        store = ControlStore(context)
        snapshot = _migration_snapshot_with_source_state(context, store.snapshot())
        snapshot["compatibility_usage"] = store.compatibility_usage()
        snapshot["quality"] = {
            "latest_gate_evaluations": store.latest_gate_evaluations(),
            "source": "control.db",
        }
        migration_preview = _v1_migration_dry_run(context)
        legacy_import_pending = any(
            bool(item.get("import_pending"))
            for item in migration_preview.get("source_manifest") or []
            if isinstance(item, dict)
        )
        # The V1 presentation aggregator initializes several legacy control
        # files. Never invoke it while actual V1 state still awaits an
        # administrator migration scan; the Snapshot endpoint must stay read-only.
        compatibility = (
            {
                "goal": {}, "goal_full": {}, "agent_activity": {}, "repair_job": {},
                "pipeline": {}, "workflow": [], "issues_summary": {}, "outputs": {},
            }
            if legacy_import_pending
            else _status_payload(
                context.root,
                context.workspace_id,
                persist_manual_review_summary=False,
            )
        )
        goal_state = store.goal_state()
        activity_state = store.agent_activity_state()
        repair_state = store.repair_job_state()
        issue_states = store.issue_states()
        material_import_pending = (
            store.v1_import_pending("materials")
            and (context.root / "workspace" / "materials_checklist.json").exists()
        )
        issue_import_pending = (
            store.v1_import_pending("issues")
            and (context.root / "workspace" / "issues" / "open.json").exists()
        )
        material_items = store.material_states()
        material_summary = {
            "exists": bool(material_items),
            "total": len(material_items),
            "ready": sum(1 for item in material_items if item.get("response_status") == "ready"),
            "deferred": sum(1 for item in material_items if item.get("response_status") == "deferred"),
            "waived": sum(1 for item in material_items if item.get("response_status") == "waived"),
            "source": "migration_required" if material_import_pending else "control.db",
        }
        artifact_states = snapshot.get("artifacts") or []
        pipeline_snapshot = _pipeline_snapshot_from_control(
            snapshot.get("operations") or [],
            compatibility.get("pipeline") or SUPERVISOR.load(context.root),
        )
        workflow = []
        for raw_step in compatibility.get("workflow") or []:
            step = dict(raw_step) if isinstance(raw_step, dict) else {}
            command = str(step.get("command") or "")
            manifests = [item for item in artifact_states if str(item.get("producer") or "") == command]
            if manifests and not bool(compatibility.get("running") and compatibility.get("current_task") == command):
                ready = all(str(item.get("status") or "") == "ready" for item in manifests)
                step["done"] = ready
                step["ready"] = True
                step["state"] = "done" if ready else "ready"
                if not ready:
                    stale_count = sum(1 for item in manifests if item.get("status") == "stale")
                    step["message"] = f"SQLite Artifact 已过期（{stale_count or len(manifests)} 项），需重新执行"
                step["artifact_source"] = "control.db"
            workflow.append(step)
        snapshot.update(
            {
                "goal": {
                    **(compatibility.get("goal") or {}),
                    **(goal_state or compatibility.get("goal_full") or {}),
                },
                "activity": activity_state or compatibility.get("agent_activity"),
                "repair_job": repair_state or compatibility.get("repair_job"),
                "manual_review_summary": _v2_manual_review_summary(context),
                "pipeline": pipeline_snapshot,
                "materials": {**material_summary, "items": material_items},
                "findings": {
                    "issues_summary": {
                        **(compatibility.get("issues_summary") or {}),
                        "open_count": sum(
                            1 for item in issue_states
                            if str(item.get("status") or "") in {"open", "in_progress"}
                        ),
                        "block_count": sum(
                            1 for item in issue_states
                            if str(item.get("status") or "") in {"open", "in_progress"}
                            and str(item.get("severity") or "") == "block"
                        ),
                        "warn_count": sum(
                            1 for item in issue_states
                            if str(item.get("status") or "") in {"open", "in_progress"}
                            and str(item.get("severity") or "") == "warn"
                        ),
                        "can_proceed": (
                            str((compatibility.get("issues_summary") or {}).get("mode") or "hard") == "soft"
                            or not any(
                                str(item.get("status") or "") in {"open", "in_progress"}
                                and str(item.get("severity") or "") == "block"
                                for item in issue_states
                            )
                        ),
                        "mode": str((compatibility.get("issues_summary") or {}).get("mode") or "hard"),
                        "top_blocks": [
                            {
                                "id": item.get("id"),
                                "code": item.get("code"),
                                "title": item.get("title"),
                                "stage_id": item.get("stage_id"),
                            }
                            for item in issue_states
                            if str(item.get("status") or "") in {"open", "in_progress"}
                            and str(item.get("severity") or "") == "block"
                        ][:8],
                "source": "migration_required" if issue_import_pending else "control.db",
                    },
                    "issues": issue_states,
                },
                "artifacts": snapshot.get("artifacts") or [],
                "artifact_files": {
                    "inputs": compatibility.get("inputs") or {},
                    "workspace": compatibility.get("workspace") or {},
                    "outputs": compatibility.get("outputs") or {},
                },
                # Presentation-only projections remain file-derived during the
                # one-version adapter window; control fields above are SQLite-first.
                "presentation": {
                    "workflow": workflow,
                    "running": bool(compatibility.get("running")),
                    "current_task": compatibility.get("current_task") or "",
                    "run_state": compatibility.get("run_state") or {},
                    "next_step": compatibility.get("next_step"),
                    "blocked_step": compatibility.get("blocked_step"),
                    "compliance_summary": compatibility.get("compliance_summary") or {},
                    "runtime": compatibility.get("runtime") or {},
                    "product_mode": compatibility.get("product_mode") or "",
                    "product_mode_label": compatibility.get("product_mode_label") or "",
                    "consistent": bool(compatibility.get("consistent", True)),
                    "consistency_warnings": compatibility.get("consistency_warnings") or [],
                },
                "compatibility_source": "v1_projection",
            }
        )
        return JSONResponse({"ok": True, "snapshot": snapshot})
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"快照读取失败: {exc}", status_code=503, retryable=True)
        )


@app.get("/api/v2/workspaces/{workspace_id}/migration/dry-run")
def api_v2_migration_dry_run(workspace_id: str) -> JSONResponse:
    try:
        return JSONResponse({"ok": True, **_v1_migration_dry_run(_workspace_context(workspace_id))})
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError(
                "STATE_UNAVAILABLE",
                f"迁移预检失败: {exc}",
                status_code=503,
                retryable=True,
            )
        )


@app.get("/api/v2/workspaces/{workspace_id}/migration/backups")
def api_v2_migration_backups(workspace_id: str) -> JSONResponse:
    try:
        context = _workspace_context(workspace_id)
        return JSONResponse({"ok": True, "backups": ControlStore(context).migration_backups()})
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"迁移备份读取失败: {exc}", status_code=503, retryable=True)
        )


@app.get("/api/v2/workspaces/{workspace_id}/migration/report")
def api_v2_migration_report(workspace_id: str) -> JSONResponse:
    try:
        context = _workspace_context(workspace_id)
        path = context.root / "workspace" / "migration_report.json"
        if not path.exists() or not path.is_file():
            raise ControlPlaneError("MIGRATION_REPORT_NOT_FOUND", "尚未执行迁移扫描。", status_code=404)
        report = _read_json_file(path)
        if not isinstance(report, dict) or str(report.get("workspace_id") or "") != context.workspace_id:
            raise ControlPlaneError("STATE_UNAVAILABLE", "迁移审计报告无效。", status_code=503)
        return JSONResponse({"ok": True, "report": report})
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"迁移报告读取失败: {exc}", status_code=503, retryable=True)
        )


@app.post("/api/v2/workspaces/{workspace_id}/migration/backups/drill")
async def api_v2_migration_backup_drill(workspace_id: str, request: Request) -> JSONResponse:
    try:
        body = await request.json()
        if not isinstance(body, dict):
            raise ControlPlaneError("COMMAND_INVALID", "请求体必须是 JSON 对象。", status_code=400)
        context = _workspace_context(workspace_id)
        principal = _request_principal(request)
        if str(principal.get("role") or "").lower() != "admin":
            raise ControlPlaneError("AUTH_FORBIDDEN", "只有管理员可以执行恢复演练。", status_code=403)
        result = ControlStore(context).drill_migration_backup(str(body.get("path") or ""))
        return JSONResponse({"ok": True, "backup": result})
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"迁移恢复演练失败: {exc}", status_code=503, retryable=True)
        )


@app.get("/api/v2/workspaces/{workspace_id}/gates/latest")
def api_v2_latest_gate_receipt(workspace_id: str) -> JSONResponse:
    try:
        context = _workspace_context(workspace_id)
        receipt = ControlStore(context).latest_gate_receipt()
        return JSONResponse({"ok": True, "gate_receipt": receipt})
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"GateReceipt 读取失败: {exc}", status_code=503, retryable=True)
        )


@app.get("/api/v2/workspaces/{workspace_id}/exports/final", response_model=None)
def api_v2_download_final(
    workspace_id: str,
    gate_receipt_id: str = Query(..., min_length=1),
) -> FileResponse | JSONResponse:
    try:
        context = _workspace_context(workspace_id)
        _, artifact = _validate_formal_gate_receipt(context, gate_receipt_id)
        return FileResponse(
            str(artifact),
            filename="final.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"正式稿下载校验失败: {exc}", status_code=503)
        )


@app.get("/api/v2/workspaces/{workspace_id}/exports/draft", response_model=None)
def api_v2_download_draft(workspace_id: str) -> FileResponse | JSONResponse:
    try:
        context = _workspace_context(workspace_id)
        artifact = context.root / "outputs" / "final.md"
        if not artifact.exists() or not artifact.is_file() or artifact.stat().st_size <= 0:
            raise ControlPlaneError("ARTIFACT_NOT_FOUND", "final.md 不存在，请先执行 build-md。", status_code=404)
        return FileResponse(str(artifact), filename="final.md", media_type="text/markdown")
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"草稿下载失败: {exc}", status_code=503)
        )


@app.get("/api/v2/workspaces/{workspace_id}/operations/{operation_id}")
def api_v2_operation(workspace_id: str, operation_id: str) -> JSONResponse:
    try:
        context = _workspace_context(workspace_id)
        operation = ControlStore(context).operation(operation_id)
        if not operation:
            raise ControlPlaneError("OPERATION_NOT_FOUND", "Operation 不存在。", status_code=404)
        return JSONResponse({"ok": True, "operation": operation})
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"Operation 读取失败: {exc}", status_code=503, retryable=True)
        )


@app.get("/api/v2/workspaces/{workspace_id}/events", response_model=None)
async def api_v2_workspace_events(
    workspace_id: str,
    request: Request,
    after_seq: int = Query(0, ge=0),
    limit: int = Query(200, ge=1, le=2000),
) -> StreamingResponse | JSONResponse:
    try:
        context = _workspace_context(workspace_id)
        store = ControlStore(context)
        header_cursor = str(request.headers.get("last-event-id") or "").strip()
        if header_cursor:
            try:
                after_seq = max(after_seq, int(header_cursor))
            except ValueError as exc:
                raise ControlPlaneError(
                    "COMMAND_INVALID",
                    "Last-Event-ID 必须是非负整数。",
                    status_code=400,
                ) from exc
    except ControlPlaneError as exc:
        return _command_error_response(exc)

    async def stream():
        cursor = after_seq
        last_keepalive = time.monotonic()
        while True:
            if await request.is_disconnected():
                break
            try:
                events = store.events(cursor, limit=limit)
            except Exception as exc:
                payload = json.dumps(
                    {"code": "STATE_UNAVAILABLE", "message": str(exc), "retryable": True},
                    ensure_ascii=False,
                )
                yield f"event: ControlPlaneError\ndata: {payload}\n\n"
                break
            for event in events:
                cursor = int(event["seq"])
                payload = json.dumps(event, ensure_ascii=False, default=str)
                yield f"id: {cursor}\nevent: WorkspaceEvent\ndata: {payload}\n\n"
            if time.monotonic() - last_keepalive >= 15:
                yield ": keepalive\n\n"
                last_keepalive = time.monotonic()
            await asyncio.sleep(0.5)

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/api/run-command")
async def api_run_command(request: Request) -> JSONResponse:
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)

    command = body.get("command", "").strip()
    if not command:
        return JSONResponse({"ok": False, "message": "缺少 command 字段。"}, status_code=400)
    if command not in COMMANDS:
        return JSONResponse(
            {"ok": False, "message": f"未知命令: {command}，可用: {', '.join(sorted(COMMANDS))}"},
            status_code=400,
        )
    if command not in {"validate", "init-demo"} and ACTIVE_RUN_ROOT is None:
        return JSONResponse(
            {"ok": False, "message": "请先点击“开始生成”，创建本次运行工作空间后再执行流程命令。"},
            status_code=409,
        )
    sent_run_id = str(body.get("run_id", "")).strip()
    if sent_run_id and sent_run_id != ACTIVE_RUN_ID:
        return JSONResponse(
            {"ok": False, "message": "运行工作空间已变化，请刷新页面后重新开始。"},
            status_code=409,
        )

    run_id = sent_run_id or ACTIVE_RUN_ID
    if command in set(auto_run_commands()):
        try:
            context = _workspace_context(run_id)
            store = ControlStore(context)
            envelope = CommandEnvelope.from_mapping(
                {
                    "kind": "pipeline.run_stage",
                    "payload": {"start_command": command},
                    "expected_revision": store.revision(),
                    "idempotency_key": str(body.get("idempotency_key") or f"legacy-run-stage:{uuid.uuid4()}"),
                    "actor": _request_actor(request, source="legacy_api"),
                },
                workspace_id=run_id,
            )
            receipt = _command_gateway(context).submit(envelope)
            return JSONResponse(
                {"ok": receipt.status != "rejected", "message": receipt.message, "receipt": receipt.as_dict()},
                status_code=202 if receipt.status != "rejected" else 409,
            )
        except ControlPlaneError as exc:
            return _command_error_response(exc)

    if run_id:
        try:
            context = _workspace_context(run_id)
            gateway = _command_gateway(context)
            envelope = CommandEnvelope.from_mapping(
                {
                    "kind": "workspace.run_utility",
                    "payload": {"command": command},
                    "expected_revision": gateway.store.revision(),
                    "idempotency_key": str(body.get("idempotency_key") or f"legacy-utility:{uuid.uuid4()}"),
                    "actor": _request_actor(request, source="legacy_api"),
                },
                workspace_id=context.workspace_id,
            )
            action = gateway.propose(envelope, label=f"确认执行维护命令 {command}", risk="high")
            return JSONResponse({"ok": True, "status": "requires_confirmation", "action": action}, status_code=202)
        except ControlPlaneError as exc:
            return _command_error_response(exc)

    # Only root-level validate/init-demo remain outside a workspace control plane.
    if RUNNING or SUPERVISOR.is_running():
        return JSONResponse({"ok": False, "message": "当前已有任务正在运行，请等待完成。"}, status_code=409)
    if command in {"validate", "init-demo"} and not run_id:
        run_root = ROOT
    else:
        try:
            context = _workspace_context(run_id)
            if command not in {"validate", "init", "init-demo", "set-project-profile"}:
                gate = _v2_gate_can_proceed(context, command)
                if not gate.get("can_proceed", False):
                    return JSONResponse(
                        {"ok": False, "message": gate.get("message") or "质量门禁阻断", "gate": gate},
                        status_code=409,
                    )
            run_root = context.root
        except ControlPlaneError as exc:
            return _command_error_response(exc)
    threading.Thread(target=_run_sync, args=(command, run_id, run_root), daemon=True).start()
    return JSONResponse({"ok": True, "message": f"兼容命令已启动: {command}"})


@app.post("/api/start-pipeline")
async def api_start_pipeline(request: Request) -> JSONResponse:
    if ACTIVE_RUN_ROOT is None:
        return JSONResponse({"ok": False, "message": "请先创建并选择工作空间。"}, status_code=409)
    try:
        body = await request.json()
    except Exception:
        body = {}
    sent_run_id = str(body.get("run_id", "")).strip()
    if sent_run_id and sent_run_id != ACTIVE_RUN_ID:
        return JSONResponse({"ok": False, "message": "工作空间已变化，请刷新后重试。"}, status_code=409)
    start_command = str(body.get("start_command", "")).strip()
    if start_command and start_command not in auto_run_commands():
        return JSONResponse({"ok": False, "message": f"无效起始阶段: {start_command}"}, status_code=400)
    try:
        context = _workspace_context(sent_run_id or ACTIVE_RUN_ID)
        store = ControlStore(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "pipeline.start",
                "payload": {"start_command": start_command},
                "expected_revision": store.revision(),
                "idempotency_key": str(body.get("idempotency_key") or f"legacy-start:{uuid.uuid4()}"),
                "actor": _request_actor(request, source="legacy_api"),
            },
            workspace_id=context.workspace_id,
        )
        receipt = _command_gateway(context).submit(envelope)
        return JSONResponse(
            {"ok": receipt.status != "rejected", "message": receipt.message, "receipt": receipt.as_dict()},
            status_code=202 if receipt.status != "rejected" else 409,
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/pause-run")
def api_pause_run(request: Request) -> JSONResponse:
    if not ACTIVE_RUN_ID:
        return JSONResponse({"ok": False, "message": "请先选择工作空间。"}, status_code=409)
    try:
        context = _workspace_context(ACTIVE_RUN_ID)
        store = ControlStore(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "pipeline.pause",
                "payload": {},
                "expected_revision": store.revision(),
                "idempotency_key": f"legacy-pause:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_api"),
            },
            workspace_id=context.workspace_id,
        )
        receipt = _command_gateway(context).submit(envelope)
        return JSONResponse(
            {"ok": receipt.status != "rejected", "message": receipt.message, "receipt": receipt.as_dict()},
            status_code=202 if receipt.status != "rejected" else 409,
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)


# ---------------------------------------------------------------
#  Logs
# ---------------------------------------------------------------

@app.get("/api/v2/workspaces/{workspace_id}/logs")
@app.get("/api/logs")
def api_logs(lines: int = Query(200, ge=1, le=2000), workspace_id: str = "") -> JSONResponse:
    if workspace_id:
        workspace_lines = _workspace_log_lines(_workspace_context(workspace_id).root, LOG_MAX)
        return JSONResponse({"lines": workspace_lines[-lines:], "total": len(workspace_lines)})
    return JSONResponse({"lines": LOG_LINES[-lines:], "total": len(LOG_LINES)})


@app.get("/api/v2/workspaces/{workspace_id}/logs/stream")
@app.get("/api/logs/stream")
async def api_logs_stream(request: Request, workspace_id: str = "") -> StreamingResponse:
    root = _workspace_context(workspace_id).root if workspace_id else None

    async def stream():
        last = 0
        last_event = 0
        while True:
            if await request.is_disconnected():
                break
            available_lines = _workspace_log_lines(root, LOG_MAX) if root else LOG_LINES
            while last < len(available_lines):
                yield f"data: {json.dumps({'type': 'log', 'line': available_lines[last]}, ensure_ascii=False)}\n\n"
                last += 1
            events = load_run_events(root or _active_root())
            while last_event < len(events):
                event = events[last_event]
                event_line = f"[{event.get('stage', '')}] {event.get('event_type', '')} {event.get('message', '')}".strip()
                yield f"data: {json.dumps({'type': 'run_event', 'event': event, 'line': event_line}, ensure_ascii=False)}\n\n"
                last_event += 1
            await asyncio.sleep(0.5)

    return StreamingResponse(stream(), media_type="text/event-stream")


# ---------------------------------------------------------------
#  File upload
# ---------------------------------------------------------------

VALID_CATEGORIES = {"tender", "company", "template"}
_SOURCE_UPLOAD_MAX_BYTES = 100 * 1024 * 1024
_SOURCE_UPLOAD_DENIED_EXTENSIONS = {
    ".bat",
    ".cmd",
    ".com",
    ".dll",
    ".exe",
    ".js",
    ".msi",
    ".ps1",
    ".py",
    ".scr",
    ".vbs",
}


def _safe_source_filename(filename: str) -> str:
    name = Path(str(filename or "")).name.strip()
    if not name or name in {".", ".."}:
        raise ControlPlaneError("UPLOAD_INVALID", "上传文件名无效。", status_code=400)
    extension = Path(name).suffix.lower()
    if extension in _SOURCE_UPLOAD_DENIED_EXTENSIONS:
        raise ControlPlaneError("UPLOAD_TYPE_DENIED", f"不允许上传可执行文件: {extension}", status_code=400)
    stem = re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]+", "_", Path(name).stem).strip("._") or "source"
    return f"{stem[:160]}{extension[:20]}"


@app.post("/api/v2/workspaces/{workspace_id}/sources")
@app.post("/api/upload")
async def api_upload(
    category: str = "tender",
    files: list[UploadFile] = File(...),
    workspace_id: str = "",
) -> JSONResponse:
    if category not in VALID_CATEGORIES:
        return JSONResponse({"ok": False, "message": f"无效 category: {category}"}, status_code=400)

    context = _workspace_context(workspace_id) if workspace_id else None
    active_root = context.root if context is not None else _active_root()
    if active_root == ROOT:
        return JSONResponse({"ok": False, "message": "请先选择或创建工作空间。"}, status_code=400)
    if context is None:
        try:
            context = WorkspaceContext.resolve(RUNS_DIR, active_root.name)
        except ControlPlaneError as exc:
            return _command_error_response(exc)

    dest_dir = active_root / "sources" / category
    dest_dir.mkdir(parents=True, exist_ok=True)

    saved: list[str] = []
    for f in files:
        content = await f.read()
        if not content:
            return JSONResponse({"ok": False, "message": "上传文件为空。"}, status_code=400)
        if len(content) > _SOURCE_UPLOAD_MAX_BYTES:
            return JSONResponse({"ok": False, "message": "单个源文件不能超过 100 MB。"}, status_code=413)
        try:
            filename = _safe_source_filename(f.filename or "")
        except ControlPlaneError as exc:
            return _command_error_response(exc)
        dest = (dest_dir / filename).resolve()
        if not dest.is_relative_to(dest_dir.resolve()):
            return JSONResponse({"ok": False, "message": "上传路径越界。"}, status_code=400)
        if dest.exists():
            dest = dest.with_name(f"{dest.stem}_{uuid.uuid4().hex[:8]}{dest.suffix}")
        dest.write_bytes(content)
        saved.append(dest.name)
        _append_log(f"[上传] {category} → {dest.name}")

    try:
        from artifact_manifest import invalidate_after_source_change

        invalidate_after_source_change(context, category=category)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError(
                "STATE_UNAVAILABLE",
                f"源文件已保存，但 Artifact 失效状态写入失败，已拒绝继续: {exc}",
                status_code=503,
                retryable=True,
            )
        )

    return JSONResponse({"ok": True, "saved": saved, "count": len(saved)})


# ---------------------------------------------------------------
#  Download
# ---------------------------------------------------------------

@app.get("/api/final-md/lines")
def api_final_md_lines() -> JSONResponse:
    root = _active_root()
    path = root / "outputs" / "final.md"
    if not path.exists():
        return JSONResponse({"ok": False, "message": "final.md 不存在，请先执行 build-md。"}, status_code=404)
    text = path.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()
    return JSONResponse(
        {
            "ok": True,
            "path": "outputs/final.md",
            "line_count": len(lines),
            "lines": [{"number": index + 1, "text": line} for index, line in enumerate(lines)],
            "docx_exists": (root / "outputs" / "final.docx").exists(),
        }
    )


def _save_final_md_line_edit(root: Path, line_number: int, new_text: str, instruction: str, source: str) -> dict[str, Any]:
    path = root / "outputs" / "final.md"
    original = path.read_text(encoding="utf-8", errors="ignore")
    lines = original.splitlines()
    if line_number > len(lines):
        raise ValueError(f"行号超出范围：{line_number}")

    old_text = lines[line_number - 1]
    lines[line_number - 1] = new_text
    backup_dir = root / "workspace" / "manual_line_edits"
    backup_dir.mkdir(parents=True, exist_ok=True)
    stamp = time.strftime("%Y%m%d_%H%M%S")
    backup_path = backup_dir / f"final_md_before_line_{line_number}_{stamp}.md"
    backup_path.write_text(original, encoding="utf-8")
    path.write_text("\n".join(lines) + ("\n" if original.endswith("\n") else ""), encoding="utf-8")

    review = {
        "line_number": line_number,
        "old_text": old_text,
        "new_text": new_text,
        "instruction": instruction,
        "source": source,
        "review_status": "accepted_for_docx_rebuild",
        "review_notes": [
            "已完成选中行改写。",
            "已记录修改前备份。",
            "由当前 V2 Operation 执行 build-docx 生成 Word。",
        ],
        "backup_path": str(backup_path.relative_to(root)).replace("\\", "/"),
        "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    review_path = backup_dir / f"line_{line_number}_{stamp}_review.json"
    review_path.write_text(json.dumps(review, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"review": review, "review_path": str(review_path.relative_to(root)).replace("\\", "/")}


def _propose_document_edit(request: Request, root: Path, payload: dict[str, Any], *, label: str) -> dict[str, Any]:
    path = root / "outputs" / "final.md"
    if not path.exists() or not path.is_file():
        raise ControlPlaneError("ARTIFACT_NOT_FOUND", "final.md 不存在，请先执行 build-md。", status_code=404)
    context = _workspace_context(ACTIVE_RUN_ID or root.name)
    gateway = _command_gateway(context)
    command_payload = dict(payload)
    command_payload["base_sha256"] = hashlib.sha256(path.read_bytes()).hexdigest()
    envelope = CommandEnvelope.from_mapping(
        {
            "kind": "document.apply_edit",
            "payload": command_payload,
            "expected_revision": gateway.store.revision(),
            "idempotency_key": f"legacy-document-edit:{uuid.uuid4()}",
            "actor": _request_actor(request, source="legacy_web"),
        },
        workspace_id=context.workspace_id,
    )
    return gateway.propose(envelope, label=label, risk="high")


@app.post("/api/final-md/line-edit")
async def api_final_md_line_edit(request: Request) -> JSONResponse:
    root = _active_root()
    if RUNNING:
        return JSONResponse({"ok": False, "message": "当前已有任务正在运行，请稍后再修改。"}, status_code=409)
    path = root / "outputs" / "final.md"
    if not path.exists():
        return JSONResponse({"ok": False, "message": "final.md 不存在，请先执行 build-md。"}, status_code=404)

    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)

    try:
        line_number = int(body.get("line_number", 0))
    except Exception:
        line_number = 0
    new_text = str(body.get("new_text", "")).rstrip("\n\r")
    instruction = str(body.get("instruction", "")).strip()
    if line_number < 1:
        return JSONResponse({"ok": False, "message": "请选择要修改的行。"}, status_code=400)
    if not new_text.strip():
        return JSONResponse({"ok": False, "message": "请填写修改后的内容。"}, status_code=400)

    try:
        action = _propose_document_edit(
            request,
            root,
            {"mode": "line", "line_number": line_number, "new_text": new_text, "instruction": instruction},
            label=f"确认修改 final.md 第 {line_number} 行",
        )
        return JSONResponse({"ok": True, "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/final-md/line-regenerate")
async def api_final_md_line_regenerate(request: Request) -> JSONResponse:
    root = _active_root()
    if RUNNING:
        return JSONResponse({"ok": False, "message": "当前已有任务正在运行，请稍后再重生成。"}, status_code=409)
    path = root / "outputs" / "final.md"
    if not path.exists():
        return JSONResponse({"ok": False, "message": "final.md 不存在，请先执行 build-md。"}, status_code=404)

    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)

    try:
        line_number = int(body.get("line_number", 0))
    except Exception:
        line_number = 0
    instruction = str(body.get("instruction", "")).strip()
    if line_number < 1:
        return JSONResponse({"ok": False, "message": "请选择要重生成的行。"}, status_code=400)
    if not instruction:
        return JSONResponse({"ok": False, "message": "请填写 AI 生成要求。"}, status_code=400)

    lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    if line_number > len(lines):
        return JSONResponse({"ok": False, "message": f"行号超出范围：{line_number}"}, status_code=400)
    old_text = lines[line_number - 1]
    start = max(0, line_number - 6)
    end = min(len(lines), line_number + 5)
    context = "\n".join(f"{index + 1}: {line}" for index, line in enumerate(lines[start:end], start=start))

    try:
        from llm_client import chat

        generated = chat(
            [
                {"role": "system", "content": "你是标书 Word 定稿改写子 agent。只输出改写后的单行内容，不要解释，不要添加 Markdown 代码块。"},
                {
                    "role": "user",
                    "content": (
                        f"请按用户要求重生成 final.md 第 {line_number} 行。\n"
                        f"用户要求：{instruction}\n\n"
                        f"当前行：{old_text}\n\n"
                        f"上下文：\n{context}"
                    ),
                },
            ],
            temperature=0.3,
        ).strip()
    except Exception as exc:
        return JSONResponse({"ok": False, "message": f"AI 重生成失败: {exc}"}, status_code=500)

    new_text = generated.strip().strip("`").strip()
    if not new_text:
        return JSONResponse({"ok": False, "message": "AI 未生成有效内容。"}, status_code=500)
    _PENDING_LINE_EDITS[root.resolve()] = {
        "line_number": line_number,
        "old_text": old_text,
        "new_text": new_text,
        "instruction": instruction,
        "source": "ai_regenerate",
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    _append_log(f"[AI改写] 已为 final.md 第 {line_number} 行生成预览，等待用户确认后再保存。")
    return JSONResponse(
        {
            "ok": True,
            "message": "AI 已生成预览，请确认后再保存并重建 Word。",
            "line_number": line_number,
            "old_text": old_text,
            "generated_text": new_text,
            "instruction": instruction,
        }
    )


@app.post("/api/final-md/line-confirm")
async def api_final_md_line_confirm(request: Request) -> JSONResponse:
    root = _active_root()
    if RUNNING:
        return JSONResponse({"ok": False, "message": "当前已有任务正在运行，请稍后再确认。"}, status_code=409)
    try:
        body = await request.json()
    except Exception:
        body = {}
    key = root.resolve()
    pending = _PENDING_LINE_EDITS.get(key)
    if not pending:
        return JSONResponse({"ok": False, "message": "没有待确认的修改预览。"}, status_code=404)
    new_text = str(body.get("new_text", pending.get("new_text", ""))).rstrip("\n\r") or pending["new_text"]
    try:
        action = _propose_document_edit(
            request,
            root,
            {
                "mode": "line",
                "line_number": pending["line_number"],
                "new_text": new_text,
                "instruction": pending.get("instruction", ""),
            },
            label=f"确认应用 final.md 第 {pending['line_number']} 行的 AI 改写",
        )
        return JSONResponse({"ok": True, "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/final-md/line-discard")
async def api_final_md_line_discard() -> JSONResponse:
    root = _active_root()
    removed = _PENDING_LINE_EDITS.pop(root.resolve(), None)
    if removed:
        _append_log(f"[AI改写] final.md 第 {removed['line_number']} 行的预览修改已放弃。")
    return JSONResponse({"ok": True, "discarded": bool(removed)})


@app.get("/api/final-md/pending")
def api_final_md_pending() -> JSONResponse:
    root = _active_root()
    pending = _PENDING_LINE_EDITS.get(root.resolve())
    return JSONResponse({"ok": True, "pending": pending})

# ---------------------------------------------------------------
#  Final doc WYSIWYG editor (block-level render / edit / AI rewrite)
# ---------------------------------------------------------------

_PENDING_DOC_EDIT: dict[Path, dict[str, Any]] = {}
_LAST_BACKUP: dict[Path, Path] = {}


def triggerDocRefresh() -> None:
    pass


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^- \s*(.*)$")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)$")
_TBL_LINE_RE = re.compile(r"^\|.*\|\s*$")
_TBL_SEP_RE = re.compile(r"^:?-{3,}:?$")


def _split_md_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _is_md_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _is_md_table_separator(cells: list[str]) -> bool:
    return all(bool(_TBL_SEP_RE.fullmatch(cell.strip())) for cell in cells)


def _parse_final_md_blocks(text: str) -> list[dict[str, Any]]:
    lines = text.splitlines()
    blocks: list[dict[str, Any]] = []
    i = 0
    block_seq = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            level = len(heading.group(1))
            blocks.append({
                "block_id": f"b{block_seq}",
                "type": "heading",
                "level": level,
                "text": heading.group(2).strip(),
                "raw": line,
                "start_line": i + 1,
                "end_line": i + 1,
            })
            block_seq += 1
            i += 1
            continue

        if _is_md_table_line(line) and i + 1 < len(lines) and _is_md_table_line(lines[i + 1]):
            header = _split_md_table_row(line)
            separator = _split_md_table_row(lines[i + 1])
            if _is_md_table_separator(separator):
                rows = [header]
                start = i + 1
                i += 2
                while i < len(lines) and _is_md_table_line(lines[i]):
                    rows.append(_split_md_table_row(lines[i]))
                    i += 1
                blocks.append({
                    "block_id": f"b{block_seq}",
                    "type": "table",
                    "header": header,
                    "rows": rows,
                    "raw": "\n".join(lines[start - 1:i]),
                    "start_line": start,
                    "end_line": i,
                })
                block_seq += 1
                continue

        bullet = _BULLET_RE.match(stripped)
        if bullet:
            blocks.append({
                "block_id": f"b{block_seq}",
                "type": "bullet",
                "text": bullet.group(1).strip(),
                "raw": line,
                "start_line": i + 1,
                "end_line": i + 1,
            })
            block_seq += 1
            i += 1
            continue

        numbered = _NUMBERED_RE.match(stripped)
        if numbered:
            blocks.append({
                "block_id": f"b{block_seq}",
                "type": "numbered",
                "text": numbered.group(1).strip(),
                "raw": line,
                "start_line": i + 1,
                "end_line": i + 1,
            })
            block_seq += 1
            i += 1
            continue

        paragraph_lines = [stripped]
        start = i + 1
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith("#") or _is_md_table_line(nxt):
                break
            if _BULLET_RE.match(nxt) or _NUMBERED_RE.match(nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        blocks.append({
            "block_id": f"b{block_seq}",
            "type": "paragraph",
            "text": " ".join(paragraph_lines),
            "lines": paragraph_lines,
            "raw": "\n".join(paragraph_lines),
            "start_line": start,
            "end_line": i,
        })
        block_seq += 1

    return blocks


def _final_md_path(root: Path) -> Path:
    return root / "outputs" / "final.md"


def _workspace_execution_active(workspace_id: str) -> bool:
    if not workspace_id:
        return RUNNING
    try:
        context = _workspace_context(workspace_id)
        operation = ControlStore(context).snapshot().get("operation")
        status = str(operation.get("status") or "") if isinstance(operation, dict) else ""
        return status in {"queued", "running", "pausing", "cancelling"} or SUPERVISOR.is_running(context.root)
    except Exception:
        return True


def _final_docx_path(root: Path) -> Path:
    return root / "outputs" / "final.docx"


def _backup_final_md(root: Path, original: str, label: str) -> Path:
    backup_dir = root / "workspace" / "manual_line_edits"
    backup_dir.mkdir(parents=True, exist_ok=True)
    stamp = time.strftime("%Y%m%d_%H%M%S")
    backup_path = backup_dir / f"final_md_before_{label}_{stamp}.md"
    backup_path.write_text(original, encoding="utf-8")
    return backup_path


def _block_new_raw(block: dict[str, Any], new_text: str) -> str:
    btype = block.get("type")
    if btype == "heading":
        level = int(block.get("level", 2) or 2)
        return f"{'#' * max(1, min(6, level))} {new_text.strip()}"
    if btype == "bullet":
        return f"- {new_text.strip()}"
    if btype == "numbered":
        return f"1. {new_text.strip()}"
    if btype == "table":
        return new_text.strip("\n")
    return new_text.strip("\n")


def _replace_final_md_block(root: Path, block_id: str, new_text: str, instruction: str, source: str) -> dict[str, Any]:
    path = _final_md_path(root)
    if not path.exists():
        raise FileNotFoundError("final.md 不存在，请先执行 build-md。")
    original = path.read_text(encoding="utf-8", errors="ignore")
    blocks = _parse_final_md_blocks(original)
    target = next((b for b in blocks if b.get("block_id") == block_id), None)
    if target is None:
        raise ValueError(f"找不到 block_id={block_id}")

    lines = original.splitlines()
    new_raw = _block_new_raw(target, new_text)
    new_lines = new_raw.splitlines()
    start_line = int(target.get("start_line", 0))
    end_line = int(target.get("end_line", 0))
    if start_line < 1 or end_line < start_line or end_line > len(lines):
        raise ValueError("行号越界，无法替换。")

    before = lines[: start_line - 1]
    after = lines[end_line:]
    merged = before + new_lines + after
    new_text_full = "\n".join(merged) + ("\n" if original.endswith("\n") else "")

    backup_path = _backup_final_md(root, original, f"block_{block_id}")
    _LAST_BACKUP[root.resolve()] = backup_path
    path.write_text(new_text_full, encoding="utf-8")

    review = {
        "block_id": block_id,
        "block_type": target.get("type", ""),
        "old_text": target.get("raw", ""),
        "new_text": new_raw,
        "instruction": instruction,
        "source": source,
        "review_status": "accepted_for_docx_rebuild",
        "backup_path": str(backup_path.relative_to(root)).replace("\\", "/"),
        "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    backup_dir = root / "workspace" / "manual_line_edits"
    review_path = backup_dir / f"block_{block_id}_{time.strftime('%Y%m%d_%H%M%S')}_review.json"
    review_path.write_text(json.dumps(review, ensure_ascii=False, indent=2), encoding="utf-8")

    return {"review": review, "review_path": str(review_path.relative_to(root)).replace("\\", "/")}


def _overwrite_final_md(root: Path, new_md: str, instruction: str, source: str) -> dict[str, Any]:
    path = _final_md_path(root)
    if not path.exists():
        raise FileNotFoundError("final.md 不存在，请先执行 build-md。")
    original = path.read_text(encoding="utf-8", errors="ignore")
    backup_path = _backup_final_md(root, original, "chat_edit")
    _LAST_BACKUP[root.resolve()] = backup_path
    new_clean = new_md.strip("\n") + "\n"
    path.write_text(new_clean, encoding="utf-8")
    review = {
        "instruction": instruction,
        "source": source,
        "old_md_length": len(original),
        "new_md_length": len(new_clean),
        "backup_path": str(backup_path.relative_to(root)).replace("\\", "/"),
        "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    backup_dir = root / "workspace" / "manual_line_edits"
    review_path = backup_dir / f"chat_edit_{time.strftime('%Y%m%d_%H%M%S')}_review.json"
    review_path.write_text(json.dumps(review, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"review": review, "review_path": str(review_path.relative_to(root)).replace("\\", "/")}


@app.get("/api/v2/workspaces/{workspace_id}/documents/final/render")
@app.get("/api/final-doc/render")
def api_final_doc_render(workspace_id: str = "") -> JSONResponse:
    root = _workspace_context(workspace_id).root if workspace_id else _active_root()
    path = _final_md_path(root)
    if not path.exists():
        return JSONResponse({
            "ok": True,
            "final_md_exists": False,
            "final_docx_exists": _final_docx_path(root).exists(),
            "blocks": [],
            "pending": _PENDING_DOC_EDIT.get(root.resolve()),
        })
    text = path.read_text(encoding="utf-8", errors="ignore")
    blocks = _parse_final_md_blocks(text)
    return JSONResponse({
        "ok": True,
        "final_md_exists": True,
        "final_docx_exists": _final_docx_path(root).exists(),
        "base_sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
        "blocks": blocks,
        "final_md_len": len(text),
        "pending": _PENDING_DOC_EDIT.get(root.resolve()),
    })


@app.post("/api/final-doc/block-edit")
async def api_final_doc_block_edit(request: Request) -> JSONResponse:
    root = _active_root()
    if RUNNING:
        return JSONResponse({"ok": False, "message": "当前已有任务正在运行，请稍后再修改。"}, status_code=409)
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    block_id = str(body.get("block_id", "")).strip()
    new_text = str(body.get("new_text", "")).rstrip("\n\r")
    instruction = str(body.get("instruction", "")).strip()
    if not block_id or not new_text.strip():
        return JSONResponse({"ok": False, "message": "缺少 block_id 或 new_text。"}, status_code=400)
    try:
        action = _propose_document_edit(
            request,
            root,
            {"mode": "block", "block_id": block_id, "new_text": new_text, "instruction": instruction},
            label=f"确认修改终稿块 {block_id}",
        )
        return JSONResponse({"ok": True, "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/v2/workspaces/{workspace_id}/documents/final/selection-rewrite")
@app.post("/api/final-doc/selection-rewrite")
async def api_final_doc_selection_rewrite(request: Request, workspace_id: str = "") -> JSONResponse:
    root = _workspace_context(workspace_id).root if workspace_id else _active_root()
    path = _final_md_path(root)
    if not path.exists():
        return JSONResponse({"ok": False, "message": "final.md 不存在，请先执行 build-md。"}, status_code=404)
    if _workspace_execution_active(workspace_id):
        return JSONResponse({"ok": False, "message": "当前已有任务正在运行，请稍后再改写。"}, status_code=409)
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    block_id = str(body.get("block_id", "")).strip()
    selected_text = str(body.get("selected_text", "")).strip()
    instruction = str(body.get("instruction", "")).strip()
    if not block_id or not selected_text or not instruction:
        return JSONResponse({"ok": False, "message": "请同时提供 block_id、selected_text 和修改意见。"}, status_code=400)

    original = path.read_text(encoding="utf-8", errors="ignore")
    blocks = _parse_final_md_blocks(original)
    target = next((b for b in blocks if b.get("block_id") == block_id), None)
    if target is None:
        return JSONResponse({"ok": False, "message": f"找不到 block_id={block_id}"}, status_code=400)
    full_text = target.get("text") or target.get("raw", "")
    if selected_text not in full_text:
        return JSONResponse({"ok": False, "message": "选区文本未在该块中找到，无法定位。"}, status_code=400)

    try:
        from llm_client import chat

        generated = chat(
            [
                {"role": "system", "content": "你是标书 Word 定稿改写子 agent。按照用户的批注要求，仅改写用户选中的文字片段，输出该块改写后的完整内容（不要仅输出片段）。保留原有结构和措辞风格，直接输出最终文本，不要解释，不要代码块。"},
                {
                    "role": "user",
                    "content": (
                        f"块的完整文本：\n{full_text}\n\n"
                        f"用户选中的片段：\n{selected_text}\n\n"
                        f"用户的批注要求：{instruction}"
                    ),
                },
            ],
            temperature=0.3,
        ).strip().strip("`").strip()
    except Exception as exc:
        return JSONResponse({"ok": False, "message": f"AI 改写失败: {exc}"}, status_code=500)

    if not generated:
        return JSONResponse({"ok": False, "message": "AI 未生成有效内容。"}, status_code=500)

    _PENDING_DOC_EDIT[root.resolve()] = {
        "block_id": block_id,
        "instruction": instruction,
        "selected_text": selected_text,
        "old_text": full_text,
        "new_text": generated,
        "source": "ai_selection_rewrite",
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    return JSONResponse({
        "ok": True,
        "block_id": block_id,
        "old_text": full_text,
        "new_text": generated,
        "selected_text": selected_text,
        "instruction": instruction,
        "message": "AI 已生成预览，预览效果见中间文档区。",
    })


@app.post("/api/final-doc/selection-apply")
async def api_final_doc_selection_apply(request: Request) -> JSONResponse:
    root = _active_root()
    if RUNNING:
        return JSONResponse({"ok": False, "message": "当前已有任务正在运行，请稍后再确认。"}, status_code=409)
    try:
        body = await request.json()
    except Exception:
        body = {}
    key = root.resolve()
    pending = _PENDING_DOC_EDIT.get(key)
    if not pending:
        return JSONResponse({"ok": False, "message": "没有待确认的选区改写预览。"}, status_code=404)
    new_text = str(body.get("new_text", pending.get("new_text", ""))).rstrip("\n\r") or pending["new_text"]
    try:
        action = _propose_document_edit(
            request,
            root,
            {
                "mode": "block",
                "block_id": pending["block_id"],
                "new_text": new_text,
                "instruction": pending.get("instruction", ""),
            },
            label=f"确认应用终稿块 {pending['block_id']} 的 AI 改写",
        )
        return JSONResponse({"ok": True, "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/v2/workspaces/{workspace_id}/documents/final/selection-discard")
@app.post("/api/final-doc/selection-discard")
async def api_final_doc_selection_discard(workspace_id: str = "") -> JSONResponse:
    root = _workspace_context(workspace_id).root if workspace_id else _active_root()
    removed = _PENDING_DOC_EDIT.pop(root.resolve(), None)
    if removed:
        _append_log(f"[WYSIWYG] 块 {removed['block_id']} 的选区改写预览已放弃。")
    return JSONResponse({"ok": True, "discarded": bool(removed)})


@app.post("/api/v2/workspaces/{workspace_id}/documents/final/chat-edit")
@app.post("/api/final-doc/chat-edit")
async def api_final_doc_chat_edit(request: Request, workspace_id: str = "") -> JSONResponse:
    root = _workspace_context(workspace_id).root if workspace_id else _active_root()
    path = _final_md_path(root)
    if not path.exists():
        return JSONResponse({"ok": False, "message": "final.md 不存在，请先执行 build-md。"}, status_code=404)
    if _workspace_execution_active(workspace_id):
        return JSONResponse({"ok": False, "message": "当前已有任务正在运行，请稍后再改写。"}, status_code=409)
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)
    instruction = str(body.get("instruction", "")).strip()
    if not instruction:
        return JSONResponse({"ok": False, "message": "请填写修改意见。"}, status_code=400)

    original = path.read_text(encoding="utf-8", errors="ignore")
    try:
        from llm_client import chat

        generated = chat(
            [
                {"role": "system", "content": "你是标书 Word 全文改写子 agent。按照用户的要求改写整份 Markdown 标书文档。保持 Markdown 结构（# 标题、表格、列表等）不变，直接输出改写后的完整 Markdown，不要解释，不要包裹代码块。"},
                {
                    "role": "user",
                    "content": (
                        f"请按下列要求改写整份标书 final.md：\n{instruction}\n\n"
                        f"原文 Markdown：\n{original}"
                    ),
                },
            ],
            temperature=0.3,
        ).strip()
        if generated.startswith("```"):
            generated = re.sub(r"^```(?:markdown)?\s*\n", "", generated)
            generated = re.sub(r"\n```\s*$", "", generated).strip()
    except Exception as exc:
        return JSONResponse({"ok": False, "message": f"AI 改写失败: {exc}"}, status_code=500)
    if not generated.strip():
        return JSONResponse({"ok": False, "message": "AI 未生成有效内容。"}, status_code=500)

    _PENDING_DOC_EDIT[root.resolve()] = {
        "kind": "chat_edit",
        "instruction": instruction,
        "old_md": original,
        "new_md": generated,
        "source": "ai_chat_edit",
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    return JSONResponse({
        "ok": True,
        "instruction": instruction,
        "new_md": generated,
        "old_md_length": len(original),
        "new_md_length": len(generated),
        "message": "AI 已生成全文改写预览，请确认后再写入并重建 Word。",
    })


@app.post("/api/final-doc/chat-apply")
async def api_final_doc_chat_apply(request: Request) -> JSONResponse:
    root = _active_root()
    if RUNNING:
        return JSONResponse({"ok": False, "message": "当前已有任务正在运行，请稍后再确认。"}, status_code=409)
    try:
        body = await request.json()
    except Exception:
        body = {}
    key = root.resolve()
    pending = _PENDING_DOC_EDIT.get(key)
    instruction = str(body.get("instruction", "")).strip() or (pending.get("instruction", "") if pending else "")
    new_md = str(body.get("new_md", "")).strip()
    if not new_md and pending and pending.get("new_md"):
        new_md = pending["new_md"]
    if not new_md:
        return JSONResponse({"ok": False, "message": "缺少 new_md。"}, status_code=400)
    try:
        action = _propose_document_edit(
            request,
            root,
            {"mode": "overwrite", "new_md": new_md, "instruction": instruction},
            label="确认应用 AI 全文改写",
        )
        return JSONResponse({"ok": True, "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/v2/workspaces/{workspace_id}/documents/final/chat-discard")
@app.post("/api/final-doc/chat-discard")
async def api_final_doc_chat_discard(workspace_id: str = "") -> JSONResponse:
    root = _workspace_context(workspace_id).root if workspace_id else _active_root()
    key = root.resolve()
    pending = _PENDING_DOC_EDIT.pop(key, None) or {}
    if pending.get("kind") == "chat_edit":
        _append_log("[WYSIWYG] 全文改写预览已放弃。")
        return JSONResponse({"ok": True, "discarded": True})
    return JSONResponse({"ok": True, "discarded": False})


@app.get("/api/final-doc/pending")
def api_final_doc_pending() -> JSONResponse:
    root = _active_root()
    pending = _PENDING_DOC_EDIT.get(root.resolve())
    return JSONResponse({"ok": True, "pending": pending})


@app.post("/api/v2/workspaces/{workspace_id}/documents/final/undo")
@app.post("/api/final-doc/undo-rewrite")
def api_final_doc_undo_rewrite(request: Request, workspace_id: str = "") -> JSONResponse:
    root = _workspace_context(workspace_id).root if workspace_id else _active_root()
    key = root.resolve()
    backup_path = _LAST_BACKUP.get(key)
    if not backup_path or not backup_path.exists():
        return JSONResponse({"ok": False, "message": "没有可撤销的改写。"}, status_code=404)
    try:
        action = _propose_document_edit(
            request,
            root,
            {"mode": "undo", "backup_path": str(backup_path.relative_to(root)).replace("\\", "/")},
            label="确认撤销上一次终稿改写",
        )
        return JSONResponse({"ok": True, "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)


def _sse_event(event: str, data: dict[str, Any]) -> str:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


def _sse_error(message: str):
    def _gen():
        yield _sse_event("error", {"message": message})
    return _gen()


@app.post("/api/v2/workspaces/{workspace_id}/documents/final/rewrite-block/stream")
@app.post("/api/final-doc/rewrite-block/stream")
async def api_final_doc_rewrite_block_stream(request: Request, workspace_id: str = "") -> StreamingResponse:
    root = _workspace_context(workspace_id).root if workspace_id else _active_root()
    path = _final_md_path(root)
    if not path.exists():
        return StreamingResponse(
            _sse_error("final.md 不存在，请先执行 build-md。"),
            media_type="text/event-stream",
        )
    if _workspace_execution_active(workspace_id):
        return StreamingResponse(
            _sse_error("当前已有任务正在运行。"),
            media_type="text/event-stream",
        )
    try:
        body = await request.json()
    except Exception:
        return StreamingResponse(
            _sse_error("请求体必须是 JSON。"),
            media_type="text/event-stream",
        )

    line_number = int(body.get("line_number", 0))
    instruction = str(body.get("instruction", "")).strip()
    if line_number <= 0 or not instruction:
        return StreamingResponse(
            _sse_error("请提供 line_number 和修改意见。"),
            media_type="text/event-stream",
        )

    original = path.read_text(encoding="utf-8", errors="ignore")
    blocks = _parse_final_md_blocks(original)
    target = next((b for b in blocks if b.get("start_line") == line_number), None)
    if target is None:
        return StreamingResponse(
            _sse_error(f"未找到第 {line_number} 行对应的内容块。"),
            media_type="text/event-stream",
        )

    full_text = target.get("text") or target.get("raw", "")
    is_table = target.get("type") == "table"
    context_start = max(0, blocks.index(target) - 1)
    context_end = min(len(blocks), blocks.index(target) + 2)
    context_blocks = blocks[context_start:context_end]

    if is_table:
        header = target.get("header", []) or []
        rows = target.get("rows", []) or []
        table_str = "| " + " | ".join(header) + " |"
        for row in rows:
            table_str += "\n| " + " | ".join(row) + " |"
        context_str = "\n\n".join(b.get("raw", "") for b in context_blocks if b != target)
        sys_msg = "你是标书表格改写子 agent。只改写表格中用户要求的单元格内容，严格保留表格结构（行列数不变）。输出完整表格 Markdown，不要解释或代码块。"
        usr_msg = f"上下文：\n{context_str}\n\n当前表格（第 {line_number} 行）：\n{table_str}\n\n修改意见：{instruction}"
    else:
        sys_msg = "你是标书 Word 精确改写子 agent。输出该块的完整改写结果（保持 Markdown 格式），保留原有结构，直接输出最终文本，不要解释或代码块。"
        usr_msg = "\n".join(b.get("raw", "") for b in context_blocks) + f"\n\n需要改写的块（第 {line_number} 行）：\n{full_text}\n\n修改意见：{instruction}"

    async def stream():
        from llm_client import chat_stream_chunks
        import threading

        try:
            yield _sse_event("start", {"line_number": line_number})
            all_chunks = []
            loop = asyncio.get_event_loop()
            q: asyncio.Queue = asyncio.Queue()
            done_flag = object()

            def _run_llm():
                try:
                    for chunk_type, chunk_text in chat_stream_chunks(
                        [{"role": "system", "content": sys_msg}, {"role": "user", "content": usr_msg}],
                        temperature=0.3,
                    ):
                        loop.call_soon_threadsafe(q.put_nowait, (chunk_type, chunk_text))
                except Exception as exc:
                    loop.call_soon_threadsafe(q.put_nowait, ("error", str(exc)))
                finally:
                    loop.call_soon_threadsafe(q.put_nowait, done_flag)

            thread = threading.Thread(target=_run_llm, daemon=True)
            thread.start()
            while True:
                item = await q.get()
                if item is done_flag:
                    break
                chunk_type, chunk_text = item
                if chunk_type == "error":
                    yield _sse_event("error", {"message": chunk_text})
                    return
                if chunk_type == "reasoning":
                    print(f"[SSE] reasoning: {chunk_text[:60]!r}", flush=True)
                    yield _sse_event("reasoning", {"text": chunk_text})
                else:
                    all_chunks.append(chunk_text)
                    print(f"[SSE] content: {chunk_text[:60]!r}", flush=True)
                    yield _sse_event("chunk", {"text": chunk_text})

            generated = "".join(all_chunks).strip().strip("`").strip()
            if generated:
                _PENDING_DOC_EDIT[root.resolve()] = {
                    "kind": "selection_rewrite",
                    "block_id": target["block_id"],
                    "instruction": instruction,
                    "line_number": line_number,
                    "selected_text": full_text,
                    "old_text": full_text,
                    "new_text": generated,
                    "source": "ai_chat_rewrite",
                    "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                }
                _append_log(f"[WYSIWYG] 第 {line_number} 行流式改写完成。")
                triggerDocRefresh()
            yield _sse_event(
                "done",
                {
                    "block_id": target["block_id"],
                    "line_number": line_number,
                    "new_text": generated,
                    "base_sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
                },
            )
        except Exception as exc:
            yield _sse_event("error", {"message": str(exc)})

    return StreamingResponse(stream(), media_type="text/event-stream", headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.get("/api/download/final-md", response_model=None)
def download_final_md() -> FileResponse | JSONResponse:
    path = _active_root() / "outputs" / "final.md"
    if not path.exists():
        return JSONResponse({"ok": False, "message": "final.md 不存在，请先执行 build-md"}, status_code=404)
    return FileResponse(str(path), filename="final.md", media_type="text/markdown")


@app.get("/api/download/final-docx", response_model=None)
def download_final_docx(
    gate_receipt_id: str = Query(""),
) -> FileResponse | JSONResponse:
    try:
        context = _workspace_context(ACTIVE_RUN_ID)
        _, path = _validate_formal_gate_receipt(context, gate_receipt_id)
        return FileResponse(
            str(path),
            filename="final.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Deprecation": "true",
                "Link": f'</api/v2/workspaces/{context.workspace_id}/exports/final>; rel="successor-version"',
            },
        )
    except ControlPlaneError as exc:
        return _command_error_response(exc)
    except Exception as exc:
        return _command_error_response(
            ControlPlaneError("STATE_UNAVAILABLE", f"正式稿下载校验失败: {exc}", status_code=503)
        )


# ---------------------------------------------------------------
#  Global review JSON
# ---------------------------------------------------------------

@app.get("/api/file/global-review")
def api_global_review() -> JSONResponse:
    path = _active_root() / "workspace" / "global_review.json"
    if not path.exists():
        return JSONResponse(
            {"ok": False, "message": "global_review.json 不存在，请先执行 global-review"},
            status_code=404,
        )
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return JSONResponse({"ok": True, "data": data})
    except Exception as exc:
        return JSONResponse({"ok": False, "message": f"读取失败: {exc}"}, status_code=500)


@app.get("/api/file/compliance-report")
def api_file_compliance_report() -> JSONResponse:
    path = _active_root() / "workspace" / "compliance_report.json"
    if not path.exists():
        return JSONResponse(
            {"ok": False, "message": "compliance_report.json 不存在，请先执行 compliance-check"},
            status_code=404,
        )
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return JSONResponse({"ok": True, "data": data})
    except Exception as exc:
        return JSONResponse({"ok": False, "message": f"读取失败: {exc}"}, status_code=500)


# ---------------------------------------------------------------
#  Clean workspace
# ---------------------------------------------------------------

@app.post("/api/delete-run")
async def api_delete_run(request: Request) -> JSONResponse:
    global ACTIVE_RUN_ID, ACTIVE_RUN_ROOT
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False, "message": "请求体必须是 JSON。"}, status_code=400)

    run_id = str(body.get("run_id", "")).strip()
    if not run_id or Path(run_id).name != run_id:
        return JSONResponse({"ok": False, "message": "无效工作空间。"}, status_code=400)

    run_root = (RUNS_DIR / run_id).resolve()
    runs_root = RUNS_DIR.resolve()
    if not run_root.is_relative_to(runs_root) or not run_root.exists() or not run_root.is_dir():
        return JSONResponse({"ok": False, "message": f"工作空间不存在: {run_id}"}, status_code=404)

    try:
        context = WorkspaceContext.resolve(RUNS_DIR, run_id)
        principal = _request_principal(request)
        _ensure_workspace_acl(context, principal, write=True)
        access = ControlStore(context).require_workspace_access(str(principal.get("id") or ""), write=True)
        if access.get("role") != "owner" and str(principal.get("role") or "") != "admin":
            raise ControlPlaneError("WORKSPACE_FORBIDDEN", "只有工作区所有者可以删除工作区。", status_code=403)
    except ControlPlaneError as exc:
        return _command_error_response(exc)

    try:
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "workspace.archive",
                "payload": {},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-workspace-archive:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        action = gateway.propose(envelope, label=f"确认归档工作空间 {run_id}", risk="critical")
        return JSONResponse({"ok": True, "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)


@app.post("/api/clean-workspace")
def api_clean_workspace(request: Request) -> JSONResponse:
    root = _active_root()
    try:
        context = _workspace_context(ACTIVE_RUN_ID or root.name)
        gateway = _command_gateway(context)
        envelope = CommandEnvelope.from_mapping(
            {
                "kind": "workspace.clean",
                "payload": {},
                "expected_revision": gateway.store.revision(),
                "idempotency_key": f"legacy-workspace-clean:{uuid.uuid4()}",
                "actor": _request_actor(request, source="legacy_web"),
            },
            workspace_id=context.workspace_id,
        )
        action = gateway.propose(envelope, label="确认清理工作区产物（可恢复）", risk="critical")
        return JSONResponse({"ok": True, "action": action}, status_code=202)
    except ControlPlaneError as exc:
        return _command_error_response(exc)


# ---------------------------------------------------------------
#  Startup
# ---------------------------------------------------------------


def _reconcile_pipeline_from_control(context: WorkspaceContext) -> bool:
    """Use control.db as restart authority and pipeline_control.json only as a worker checkpoint."""
    store = ControlStore(context)
    operation = store.snapshot().get("operation")
    checkpoint = SUPERVISOR.load(context.root)
    if not isinstance(operation, dict):
        return False
    status = str(operation.get("status") or "")
    operation_id = str(operation.get("operation_id") or "")
    fencing_token = int(operation.get("fencing_token") or 0)
    if status == "pausing":
        store.sync_operation(
            operation_id,
            "paused",
            message="服务重启时完成暂停",
            fencing_token=fencing_token,
        )
        return False
    if status == "cancelling":
        store.sync_operation(
            operation_id,
            "cancelled",
            message="服务重启时完成取消",
            fencing_token=fencing_token,
        )
        return False
    if status not in {"queued", "running"}:
        return False

    checkpoint_operation = str(checkpoint.get("operation_id") or "")
    checkpoint_fencing = int(checkpoint.get("fencing_token") or 0)
    if not checkpoint or checkpoint_operation != operation_id or checkpoint_fencing != fencing_token:
        store.sync_operation(
            operation_id,
            "blocked",
            message="Pipeline checkpoint 与 control.db 不一致，已停止自动恢复。",
            error={
                "code": "STATE_CONFLICT",
                "checkpoint_operation_id": checkpoint_operation,
                "checkpoint_fencing_token": checkpoint_fencing,
            },
            fencing_token=fencing_token,
        )
        return False

    if str(checkpoint.get("status") or "") not in {"running", "recovering", "retrying", "pausing"}:
        SUPERVISOR._save(  # noqa: SLF001 - compatibility checkpoint is normalized from V2 authority here.
            context.root,
            {
                "status": "recovering",
                "operation_id": operation_id,
                "fencing_token": fencing_token,
                "message": "control.db 指示 Operation 仍在运行，准备断点恢复",
            },
        )
    return SUPERVISOR.reconcile(
        context.workspace_id,
        context.root,
        _run_sync,
        gate_evaluator=lambda _root, command: _v2_gate_can_proceed(context, command),
        artifact_recorder=lambda _root, command, disposition: _record_v2_stage_artifacts(
            context,
            command,
            disposition,
        ),
        artifact_readiness_evaluator=lambda _root, command: _v2_stage_artifacts_reusable(context, command),
    )


def _reconcile_inactive_workspace(context: WorkspaceContext) -> dict[str, Any]:
    """Fail closed for V2 work that this single-process worker cannot reclaim."""
    store = ControlStore(context)
    operations = store.snapshot().get("operations") or []
    operation = next(
        (
            item
            for item in operations
            if str(item.get("kind") or "").startswith("pipeline.")
            and str(item.get("status") or "") in {"queued", "running", "pausing", "cancelling"}
        ),
        None,
    )
    if not operation:
        return {"changed": False, "workspace_id": context.workspace_id}
    operation_id = str(operation.get("operation_id") or "")
    fencing_token = int(operation.get("fencing_token") or 0)
    old_status = str(operation.get("status") or "")
    if old_status == "pausing":
        new_status = "paused"
        message = "服务重启时完成暂停；该工作区未被当前 Worker 自动接管。"
    elif old_status == "cancelling":
        new_status = "cancelled"
        message = "服务重启时完成取消；该工作区未被当前 Worker 自动接管。"
    else:
        new_status = "blocked"
        message = "服务重启后该工作区未被当前 Worker 接管，请在工作区内显式继续。"
    store.sync_operation(
        operation_id,
        new_status,
        message=message,
        error={"code": "ORPHANED_AFTER_RESTART", "previous_status": old_status}
        if new_status == "blocked"
        else None,
        fencing_token=fencing_token,
    )
    goal = store.goal_state()
    if goal and str(goal.get("status") or "") in {"planning", "in_progress", "running"}:
        store.upsert_goal_state(
            {
                **goal,
                "status": "blocked_human",
                "blocked_reason": "服务重启后执行 Operation 未被自动接管，请确认后继续。",
                "orphaned_operation_id": operation_id,
            }
        )
    return {
        "changed": True,
        "workspace_id": context.workspace_id,
        "operation_id": operation_id,
        "status": new_status,
    }


def _startup_reconcile() -> None:
    # Install the V2 control-state bridge before any persisted V1 pipeline is
    # reconciled so subsequent status transitions carry the same operation
    # identity and fencing token into control.db.
    SUPERVISOR.set_status_listener(_sync_pipeline_control_state)
    # 启动时把“使用中”的大模型刷进进程环境，避免仅写了 models.json/.env 但进程仍用旧环境变量
    try:
        store = _read_models_store()
        active_id = str(store.get("active_id") or "")
        models = store.get("models") if isinstance(store.get("models"), list) else []
        active = next((m for m in models if isinstance(m, dict) and str(m.get("id")) == active_id), None)
        if active:
            _sync_model_to_env(active)
            _append_log(f"[系统] 已加载使用中大模型: {active.get('name') or active.get('model') or active_id}")
        else:
            _append_log("[系统] 未配置使用中大模型，请在设置页添加并设为使用中")
    except Exception as exc:
        _append_log(f"[警告] 加载大模型配置失败: {exc}")

    try:
        from agent.flags import agent_supervisor_enabled

        mode = "Agent 模式" if agent_supervisor_enabled() else "流水线模式"
        _append_log(f"[系统] 运行模式: {mode} (AGENT_SUPERVISOR_ENABLED)")
    except Exception as exc:
        _append_log(f"[警告] 读取 Agent 开关失败: {exc}")

    _load_active_run_from_disk()
    workspace_roots = [path for path in RUNS_DIR.iterdir() if path.is_dir() and not path.name.startswith(".")]
    if not workspace_roots and ACTIVE_RUN_ROOT is None:
        # Preserve the one-version non-isolated V1 workspace adapter.
        workspace_roots = [ROOT]
    for workspace_root in workspace_roots:
        try:
            reconcile_interrupted_repair(workspace_root)
            from agent.activity import reconcile_interrupted_activity

            act = reconcile_interrupted_activity(workspace_root)
            if str(act.get("status") or "") == "interrupted":
                _append_log(f"[系统] 已清理重启后残留工位: {workspace_root.name}")
            if workspace_root != ROOT and not _same_path(workspace_root, ACTIVE_RUN_ROOT):
                result = _reconcile_inactive_workspace(_workspace_context(workspace_root.name))
                if result.get("changed"):
                    _append_log(f"[系统] 已阻断未接管 Operation: {workspace_root.name}")
        except Exception as exc:
            _append_log(f"[警告] 工作区恢复检查失败 {workspace_root.name}: {exc}")

    # Unified soft-heal + consistency log (replaces ad-hoc goal/pipeline only check)
    if ACTIVE_RUN_ROOT is not None:
        try:
            from agent.runtime_status import soft_heal_inconsistencies

            runtime = soft_heal_inconsistencies(ACTIVE_RUN_ROOT)
            mode = runtime.get("product_mode") or "?"
            warnings = runtime.get("warnings") or []
            heals = runtime.get("heal_actions") or []
            _append_log(
                f"[系统] Runtime 状态: mode={mode} consistent={runtime.get('consistent')} "
                f"warnings={len(warnings)} heals={heals or 'none'}"
            )
            for w in warnings[:5]:
                _append_log(f"[一致性] {w.get('severity')}:{w.get('code')} {w.get('message')}")
        except Exception as exc:
            _append_log(f"[警告] Runtime 一致性检查失败: {exc}")

    if ACTIVE_RUN_ROOT is None:
        return
    try:
        _resumed = _reconcile_pipeline_from_control(_workspace_context(ACTIVE_RUN_ID))
    except Exception as exc:
        _append_log(f"[warn] pipeline resume skipped: {exc}")
        _resumed = False
    if _resumed:
        _append_log(f"[自动恢复] 已接管工作空间流水线: {ACTIVE_RUN_ID}")


@app.on_event("startup")
def reconcile_interrupted_pipeline() -> None:
    """Startup reconcile (Goal/Pipeline consistency, model load, interrupted jobs)."""
    _startup_reconcile()


@app.get("/{full_path:path}")
async def spa_fallback(full_path: str) -> HTMLResponse:
    if USE_VUE and not full_path.startswith("api/") and full_path != "static":
        html = (VUE_DIST_DIR / "index.html").read_text(encoding="utf-8")
        return HTMLResponse(html)
    return HTMLResponse("<h1>Not Found</h1>", status_code=404)


if __name__ == "__main__":
    import socket
    import uvicorn

    def _port_free(host: str, port: int) -> bool:
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
            sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            try:
                sock.bind((host, port))
                return True
            except OSError:
                return False

    host = "127.0.0.1"
    preferred = 7860
    candidates = [preferred, 7861, 7862, 7863, 7870, 8000]
    port = next((p for p in candidates if _port_free(host, p)), None)
    if port is None:
        print("[ERROR] ports busy:", candidates)
        print("Kill process on 7860, e.g.:")
        print("  Get-NetTCPConnection -LocalPort 7860 | Select OwningProcess")
        print("  Stop-Process -Id <PID> -Force")
        raise SystemExit(1)

    if port != preferred:
        print(f"[WARN] port {preferred} busy, using http://{host}:{port}")
    else:
        print(f"[START] bid agent web: http://{host}:{port}")

    _append_log(f"[system] web console starting port={port}")
    uvicorn.run(app, host=host, port=port, log_level="info")
