from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Iterable

from file_loader import load_outline
from template_analyzer import analyze_template
from utils import project_root, read_json, read_nonempty_text, read_text, write_json, write_text

HEADING_NUMBER_PREFIX_RE = re.compile(r"^\s*\d{1,3}(?:[.\uFF0E、]\d{1,3})*[.\uFF0E、]?\s+")
SECTION_HEADING_RE = re.compile(r"^#{1,6}\s+")


def _iter_chapter_paths(root: Path) -> list[Path]:
    jobs_dir = root / "workspace" / "jobs"
    if jobs_dir.exists():
        path_entries: list[tuple[int, str, Path]] = []
        for job_file in sorted(jobs_dir.glob("*.json")):
            order = 0
            try:
                job = read_json(job_file)
                if isinstance(job, dict):
                    order = int(job.get("template_order", 0) or 0)
            except Exception:
                order = 0
            chapter_path = root / "workspace" / "chapters" / f"{job_file.stem}.md"
            if chapter_path.exists():
                path_entries.append((order, job_file.stem, chapter_path))
        if path_entries:
            return [entry[2] for entry in sorted(path_entries, key=lambda entry: (entry[0] or 999999, entry[1]))]

    outline = load_outline(root)
    paths = []
    for chapter in outline.get("chapters", []):
        chapter_id = str(chapter.get("id"))
        chapter_path = root / "workspace" / "chapters" / f"{chapter_id}.md"
        if chapter_path.exists():
            paths.append(chapter_path)
        else:
            print(f"[警告] 章节文件不存在，已跳过: {chapter_path}")
    return paths


def build_markdown(root: Path | None = None) -> Path:
    root = root or project_root()
    chunks: list[str] = []

    for chapter_path in _iter_chapter_paths(root):
        content = read_text(chapter_path).strip()
        if content:
            chunks.append(content)

    output_path = root / "outputs" / "final.md"
    write_text(output_path, "\n\n".join(chunks).strip() + "\n")
    print(f"[完成] 已拼接 Markdown: {output_path}")
    return output_path


def _is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _is_separator_row(cells: Iterable[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _style_exists(document, style_name: str | None) -> bool:
    if not style_name:
        return False
    try:
        document.styles[style_name]
        return True
    except Exception:
        return False


def _resolve_style_name(document, candidates: list[str]) -> str | None:
    for candidate in candidates:
        if _style_exists(document, candidate):
            return candidate
    return None


def _style_has_numbering(style, seen: set[int] | None = None) -> bool:
    if style is None:
        return False
    seen = seen or set()
    marker = id(style)
    if marker in seen:
        return False
    seen.add(marker)

    try:
        p_pr = style.element.pPr
        if p_pr is not None and p_pr.numPr is not None:
            return True
    except Exception:
        pass

    try:
        return _style_has_numbering(style.base_style, seen)
    except Exception:
        return False


def _strip_heading_number_prefix(text: str) -> str:
    stripped = HEADING_NUMBER_PREFIX_RE.sub("", text, count=1).strip()
    return stripped or text


def _extract_template_profile(document) -> dict[str, Any]:
    heading_styles: dict[int, str] = {}
    paragraph_style: str | None = None
    bullet_style: str | None = None
    number_style: str | None = None
    table_style: str | None = None

    for para in document.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        if not style_name:
            continue

        heading_match = re.fullmatch(r"Heading\s+(\d+)", style_name)
        if heading_match:
            level = int(heading_match.group(1))
            heading_styles.setdefault(level, style_name)
            continue

        if paragraph_style is None and text:
            paragraph_style = style_name

        if bullet_style is None and ("Bullet" in style_name or "项目符号" in style_name):
            bullet_style = style_name

        if number_style is None and ("Number" in style_name or "编号" in style_name):
            number_style = style_name

    for table in document.tables:
        try:
            if table.style and table.style.name:
                table_style = table.style.name
                break
        except Exception:
            continue

    return {
        "heading_styles": heading_styles,
        "paragraph_style": paragraph_style,
        "bullet_style": bullet_style,
        "number_style": number_style,
        "table_style": table_style,
    }


def _normalize_template_profile(document, profile: dict[str, Any] | None = None) -> dict[str, Any]:
    profile = profile or {}
    heading_styles = {
        level: style_name
        for level, style_name in (profile.get("heading_styles") or {}).items()
        if isinstance(level, int) and _style_exists(document, style_name)
    }

    normalized = {
        "heading_styles": heading_styles,
        "paragraph_style": profile.get("paragraph_style") if _style_exists(document, profile.get("paragraph_style")) else None,
        "bullet_style": profile.get("bullet_style") if _style_exists(document, profile.get("bullet_style")) else None,
        "number_style": profile.get("number_style") if _style_exists(document, profile.get("number_style")) else None,
        "table_style": profile.get("table_style") if _style_exists(document, profile.get("table_style")) else None,
    }

    for level in range(1, 5):
        if level not in normalized["heading_styles"]:
            fallback = _resolve_style_name(document, [f"Heading {level}"])
            if fallback:
                normalized["heading_styles"][level] = fallback

    if not normalized["paragraph_style"]:
        normalized["paragraph_style"] = _resolve_style_name(document, ["Normal", "正文"])
    if not normalized["bullet_style"]:
        normalized["bullet_style"] = _resolve_style_name(document, ["List Bullet", "Bullet", "列表段落"])
    if not normalized["number_style"]:
        normalized["number_style"] = _resolve_style_name(document, ["List Number", "Number", "列表编号"])
    if not normalized["table_style"]:
        normalized["table_style"] = _resolve_style_name(document, ["Table Grid", "Table Normal"])
    return normalized


def _add_paragraph_with_style(document, text: str, style_name: str | None = None, indent: bool = False) -> None:
    if style_name:
        try:
            p = document.add_paragraph(text, style=style_name)
        except Exception:
            p = document.add_paragraph(text)
    else:
        p = document.add_paragraph(text)
    if indent:
        from docx.shared import Cm

        p.paragraph_format.first_line_indent = Cm(0.74)


def _add_heading_with_style(document, text: str, level: int, profile: dict[str, Any]) -> None:
    style_name = profile.get("heading_styles", {}).get(level)
    if style_name:
        try:
            style = document.styles[style_name]
        except Exception:
            style = None
        heading_text = _strip_heading_number_prefix(text) if _style_has_numbering(style) else text
        _add_paragraph_with_style(document, heading_text, style_name)
        return
    document.add_heading(text, level=min(level, 4))


def _add_table(document, rows: list[list[str]], table_style: str | None = None) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    if table_style:
        try:
            table.style = table_style
        except Exception:
            pass

    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            table.cell(row_index, col_index).text = row[col_index] if col_index < len(row) else ""


def _plain_lines(text: str) -> list[str]:
    cleaned = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    cleaned = cleaned.replace("\u3000", " ")
    return [line.strip() for line in cleaned.splitlines() if line.strip() and line.strip() != "---"]


def _strip_markdown_heading(line: str) -> str:
    return SECTION_HEADING_RE.sub("", line).strip()


def _extract_between(lines: list[str], start_pattern: str, end_patterns: list[str]) -> list[str]:
    start_index = -1
    for index, line in enumerate(lines):
        if re.search(start_pattern, line):
            start_index = index + 1
            break
    if start_index < 0:
        return []

    end_index = len(lines)
    for index in range(start_index, len(lines)):
        text = lines[index]
        if any(re.search(pattern, text) for pattern in end_patterns):
            end_index = index
            break
    return lines[start_index:end_index]


def _extract_field(lines: list[str], label: str) -> str:
    pattern = re.compile(rf"^{re.escape(label)}[：:]\s*(.+)$")
    for line in lines:
        match = pattern.search(_strip_markdown_heading(line))
        if match:
            return match.group(1).strip()
    return ""


def _extract_numbered_items(section_lines: list[str]) -> list[tuple[str, str]]:
    items: list[tuple[str, str]] = []
    current_title = ""
    current_body: list[str] = []

    for raw_line in section_lines:
        line = _strip_markdown_heading(raw_line)
        match = re.match(r"^(\d+)[、.．]\s*(.+)$", line)
        if match:
            if current_title:
                items.append((current_title, " ".join(current_body).strip()))
            current_title = match.group(2).strip()
            current_body = []
            continue
        if current_title:
            current_body.append(line)

    if current_title:
        items.append((current_title, " ".join(current_body).strip()))
    return items


def _supplier_indicator(requirement: str) -> str:
    requirement = requirement.strip(" 。；;")
    if not requirement:
        return "我公司完全响应竞争性磋商文件要求，按采购人要求提供相应服务。"
    return f"我公司完全响应该项要求，将按照竞争性磋商文件要求及项目实施方案落实：{requirement}。"


def _extract_template_context(root: Path) -> dict[str, Any]:
    tender_path = root / "inputs" / "tender.md"
    if not tender_path.exists():
        return {"fields": {}, "work_items": []}

    lines = _plain_lines(read_text(tender_path))
    demand_lines = _extract_between(
        lines,
        r"一、招标需求",
        [r"二、技术参数", r"第七章", r"商务须知"],
    )
    work_lines = _extract_between(
        lines,
        r"（三）工作内容",
        [r"（四）项目依据", r"（五）提交成果", r"第七章"],
    )

    fields = {
        "project_name": _extract_field(demand_lines, "项目名称"),
        "purchaser": _extract_field(demand_lines, "采购单位"),
        "scope": _extract_field(demand_lines, "采购范围"),
        "budget": _extract_field(demand_lines, "采购预算"),
        "deadline": _extract_field(demand_lines, "服务要求"),
        "location": _extract_field(demand_lines, "服务地点") or _extract_field(demand_lines, "采购范围"),
        "payment": _extract_field(demand_lines, "付款方式"),
        "acceptance": _extract_field(demand_lines, "验收要求"),
    }

    rows: list[list[str]] = []
    overview_parts = [
        f"项目名称：{fields['project_name']}" if fields["project_name"] else "",
        fields["scope"],
        fields["acceptance"],
    ]
    overview = "；".join(part for part in overview_parts if part)
    if overview:
        rows.append(["1", "一、招标需求", overview, _supplier_indicator(overview), "完全响应", ""])

    for index, (title, body) in enumerate(_extract_numbered_items(work_lines), start=len(rows) + 1):
        requirement = body or title
        rows.append([str(index), title, requirement, _supplier_indicator(requirement), "完全响应", ""])

    return {"fields": fields, "response_rows": rows}


def _compact_cell_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _table_header_text(table) -> str:
    return "\n".join(_compact_cell_text(cell.text) for row in table.rows[:2] for cell in row.cells)


def _is_requirement_response_table(table) -> bool:
    header_text = _table_header_text(table)
    has_requirement = any(keyword in header_text for keyword in ["文件要求", "采购需求", "服务指标", "技术要求", "规格参数"])
    has_supplier = any(keyword in header_text for keyword in ["供应商提供", "投标响应", "响应内容", "响应指标"])
    has_response = any(keyword in header_text for keyword in ["响应程度", "偏离", "响应情况"])
    return has_requirement and has_supplier and has_response


def _find_header_columns(table) -> dict[str, int]:
    columns: dict[str, int] = {}
    if not table.rows:
        return columns
    header_cells = table.rows[0].cells
    for index, cell in enumerate(header_cells):
        text = _compact_cell_text(cell.text)
        if any(keyword in text for keyword in ["编号", "序号"]):
            columns.setdefault("number", index)
        if any(keyword in text for keyword in ["服务名称", "项目名称", "名称"]):
            columns.setdefault("name", index)
        if any(keyword in text for keyword in ["文件要求", "采购需求", "服务指标", "技术要求", "规格参数"]):
            columns.setdefault("requirement", index)
        if any(keyword in text for keyword in ["供应商提供", "投标响应", "响应内容", "响应指标"]):
            columns.setdefault("supplier", index)
        if any(keyword in text for keyword in ["响应程度", "偏离", "响应情况"]):
            columns.setdefault("response", index)
        if "说明" in text or "备注" in text:
            columns.setdefault("note", index)
    return columns


def _set_cell_text(cell, text: str) -> None:
    cell.text = text


def _fill_requirement_response_table(table, context: dict[str, Any]) -> int:
    rows = context.get("response_rows") or []
    fields = context.get("fields") or {}
    columns = _find_header_columns(table)
    required = {"requirement", "supplier", "response"}
    if not required.issubset(columns):
        return 0

    filled = 0
    data_start = 1
    reserved_tail = 0
    tail_labels = []
    for row_index, row in enumerate(table.rows[1:], start=1):
        row_text = _compact_cell_text("".join(cell.text for cell in row.cells))
        if "服务期限" in row_text or "服务地点" in row_text:
            tail_labels.append((row_index, row_text))
            reserved_tail += 1

    available_rows = max(0, len(table.rows) - data_start - reserved_tail)
    for offset, row_data in enumerate(rows[:available_rows]):
        row = table.rows[data_start + offset]
        values = {
            "number": row_data[0],
            "name": row_data[1],
            "requirement": row_data[2],
            "supplier": row_data[3],
            "response": row_data[4],
            "note": row_data[5],
        }
        for key, value in values.items():
            if key in columns:
                _set_cell_text(row.cells[columns[key]], value)
        filled += 1

    for row_index, row_text in tail_labels:
        row = table.rows[row_index]
        if "服务期限" in row_text and fields.get("deadline"):
            if "requirement" in columns:
                row.cells[columns["requirement"]].text = fields["deadline"]
            if "supplier" in columns:
                row.cells[columns["supplier"]].text = f"我公司承诺{fields['deadline']}。"
            if "response" in columns:
                row.cells[columns["response"]].text = "完全响应"
            filled += 1
        if "服务地点" in row_text and fields.get("location"):
            if "requirement" in columns:
                row.cells[columns["requirement"]].text = fields["location"]
            if "supplier" in columns:
                row.cells[columns["supplier"]].text = f"我公司承诺服务地点满足采购人要求：{fields['location']}。"
            if "response" in columns:
                row.cells[columns["response"]].text = "完全响应"
            filled += 1
    return filled


def _fill_labeled_cells(table, context: dict[str, Any]) -> int:
    fields = context.get("fields") or {}
    label_to_field = {
        "采购项目名称": "project_name",
        "项目名称": "project_name",
        "采购单位": "purchaser",
        "采购人": "purchaser",
        "采购范围": "scope",
        "采购预算": "budget",
        "服务期限": "deadline",
        "服务要求": "deadline",
        "服务地点": "location",
        "付款方式": "payment",
        "验收要求": "acceptance",
    }
    filled = 0
    for row in table.rows:
        cells = row.cells
        for index, cell in enumerate(cells):
            label_text = _compact_cell_text(cell.text).rstrip("：:")
            matched = next((field for label, field in label_to_field.items() if label in label_text), "")
            if not matched or not fields.get(matched):
                continue
            if fields[matched] in cell.text:
                continue
            target_index = index + 1 if index + 1 < len(cells) else index
            if target_index == index:
                cell.text = f"{cell.text.rstrip()} {fields[matched]}"
            elif not cells[target_index].text.strip():
                cells[target_index].text = fields[matched]
            else:
                continue
            filled += 1
    return filled


def _fill_template_structures(document, root: Path) -> dict[str, int]:
    context = _extract_template_context(root)
    stats: dict[str, Any] = {
        "tables": 0,
        "fields": 0,
        "requirement_tables": 0,
        "filled_tables": [],
    }
    for table_index, table in enumerate(document.tables):
        field_count = _fill_labeled_cells(table, context)
        response_count = _fill_requirement_response_table(table, context) if _is_requirement_response_table(table) else 0
        if field_count or response_count:
            stats["tables"] += 1
            stats["fields"] += field_count + response_count
            stats["filled_tables"].append(
                {
                    "table_index": table_index,
                    "labeled_fields": field_count,
                    "response_rows": response_count,
                }
            )
        if response_count:
            stats["requirement_tables"] += 1
    return stats


def _load_or_create_template_schema(root: Path) -> dict[str, Any]:
    schema_path = root / "workspace" / "template_schema.json"
    template_path = root / "inputs" / "template.docx"
    if not schema_path.exists() and template_path.exists():
        analyze_template(root)
    if schema_path.exists():
        try:
            data = read_json(schema_path)
            return data if isinstance(data, dict) else {}
        except Exception:
            return {}
    return {}


def _write_template_fill_report(root: Path, schema: dict[str, Any], fill_stats: dict[str, Any]) -> Path:
    schema_tables = schema.get("tables") if isinstance(schema.get("tables"), list) else []
    unhandled_tables = [
        {
            "table_index": table.get("index"),
            "type": table.get("type"),
            "header_cells": table.get("header_cells", []),
            "empty_cells": table.get("empty_cells", 0),
        }
        for table in schema_tables
        if table.get("type") in {"structured_table", "blank_or_unknown"} and table.get("empty_cells", 0)
    ]
    remaining_placeholders = schema.get("placeholders", [])
    report = {
        "ok": not unhandled_tables and not remaining_placeholders,
        "template_fingerprint": (schema.get("fingerprint") or {}).get("sha256", ""),
        "schema_summary": schema.get("summary", {}),
        "fill_stats": fill_stats,
        "unhandled_tables": unhandled_tables,
        "remaining_template_placeholders": remaining_placeholders,
        "warnings": schema.get("warnings", []),
    }
    output_path = root / "workspace" / "template_fill_report.json"
    write_json(output_path, report)
    return output_path


def _clear_document_body(document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _clear_document_body_after_cover(document) -> bool:
    heading1_seen = 0
    start_element = None
    for para in document.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name == "Heading 1":
            heading1_seen += 1
            if heading1_seen >= 2:
                start_element = para._element
                break

    if start_element is None:
        return False

    body = document._body._element
    remove_mode = False
    for child in list(body):
        if child is start_element:
            remove_mode = True
        if not remove_mode:
            continue
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)
    return True


def _add_markdown_to_document(document, markdown: str, profile: dict[str, Any] | None = None) -> None:
    profile = _normalize_template_profile(document, profile)
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            _add_heading_with_style(document, heading.group(2).strip(), level, profile)
            i += 1
            continue

        if _is_table_line(line) and i + 1 < len(lines) and _is_table_line(lines[i + 1]):
            header = _split_table_row(line)
            separator = _split_table_row(lines[i + 1])
            if _is_separator_row(separator):
                rows = [header]
                i += 2
                while i < len(lines) and _is_table_line(lines[i]):
                    rows.append(_split_table_row(lines[i]))
                    i += 1
                _add_table(document, rows, table_style=profile.get("table_style"))
                continue

        if stripped.startswith("- "):
            bullet_text = stripped[2:].strip()
            _add_paragraph_with_style(document, bullet_text, style_name=profile.get("bullet_style"))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            numberless_text = re.sub(r"^\d+\.\s+", "", stripped)
            _add_paragraph_with_style(
                document,
                numberless_text if profile.get("number_style") else stripped,
                style_name=profile.get("number_style"),
            )
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line or next_line.startswith("#") or _is_table_line(next_line):
                break
            if next_line.startswith("- ") or re.match(r"^\d+\.\s+", next_line):
                break
            paragraph_lines.append(next_line)
            i += 1
        _add_paragraph_with_style(document, " ".join(paragraph_lines), style_name=profile.get("paragraph_style"), indent=True)


def build_docx(root: Path | None = None) -> Path:
    root = root or project_root()
    final_markdown_path = root / "outputs" / "final.md"
    markdown = read_nonempty_text(final_markdown_path, f"最终 Markdown {final_markdown_path}")

    try:
        from docx import Document
        from docx.shared import Cm
    except ImportError as exc:
        raise ImportError("缺少依赖 python-docx，请先执行: pip install -r requirements.txt") from exc

    template_path = root / "inputs" / "template.docx"
    template_profile: dict[str, Any] | None = None
    template_schema: dict[str, Any] = {}
    fill_stats: dict[str, Any] = {"tables": 0, "fields": 0, "requirement_tables": 0, "filled_tables": []}
    if template_path.exists() and template_path.stat().st_size > 0:
        try:
            template_schema = _load_or_create_template_schema(root)
            document = Document(str(template_path))
            template_profile = _extract_template_profile(document)
            fill_stats = _fill_template_structures(document, root)
            if fill_stats["fields"]:
                print(
                    f"[完成] 已按模板结构填充 {fill_stats['tables']} 个表格/"
                    f"{fill_stats['fields']} 项字段"
                )
            if not _clear_document_body_after_cover(document):
                _clear_document_body(document)
        except Exception as exc:
            print(f"[警告] template.docx 无法读取，将新建空白 Word: {exc}")
            document = Document()
    else:
        document = Document()

    _add_markdown_to_document(document, markdown, profile=template_profile)

    output_path = root / "outputs" / "final.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    fill_report_path = _write_template_fill_report(root, template_schema, fill_stats)
    print(f"[完成] 已写入模板填充报告: {fill_report_path}")
    print(f"[完成] 已生成 Word: {output_path}")
    return output_path
