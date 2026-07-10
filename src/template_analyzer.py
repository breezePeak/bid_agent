from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Any

from utils import project_root, stringify, write_json


PLACEHOLDER_RE = re.compile(r"(__{2,}|XXX+|待填写|请填写|请输入|TODO|TBD)", re.IGNORECASE)

FIELD_LABELS = {
    "project_name": ["采购项目名称", "项目名称", "项目"],
    "project_no": ["采购项目编号", "项目编号", "招标编号"],
    "package_no": ["包号", "标包", "包件"],
    "purchaser": ["采购单位", "采购人", "招标人"],
    "bidder_name": ["竞标人名称", "投标人名称", "供应商名称", "公司名称"],
    "deadline": ["服务期限", "服务要求", "履约期限", "交付期限", "工期"],
    "location": ["服务地点", "实施地点", "交付地点", "项目地点"],
    "budget": ["采购预算", "预算金额", "最高限价"],
    "payment": ["付款方式", "支付方式"],
    "acceptance": ["验收要求", "验收标准"],
}

WRITING_TASK_HINTS = [
    (["技术方案", "项目技术"], ["tender.requirements", "company.technical_capability"], "结合技术要求扩写总体技术路线、方法和成果。"),
    (["实施方案", "工作内容"], ["tender.work_content", "company.delivery_experience"], "围绕工作内容拆解实施步骤、职责和交付物。"),
    (["进度", "工期"], ["tender.service_period"], "按服务期限拆解阶段计划和工期控制措施。"),
    (["人员", "机构", "团队"], ["company.team_roles", "company.qualifications"], "填充项目组织、人员职责和专业能力。"),
    (["设备", "投入"], ["company.equipment", "company.resources"], "说明投入设备、工具平台和保障资源。"),
    (["质量", "检查", "验收"], ["tender.acceptance", "company.quality_system"], "扩写质量保证体系、检查流程和验收响应。"),
    (["重难点", "应对"], ["tender.requirements", "company.experience"], "识别项目重难点并给出针对性应对措施。"),
    (["保密", "档案"], ["tender.contract", "company.compliance"], "扩写保密制度、档案管理和数据安全措施。"),
    (["售后", "服务"], ["tender.service_requirements", "company.after_sales"], "扩写售后服务体系、响应时限和保障承诺。"),
    (["承诺", "声明"], ["tender.qualifications", "company.facts"], "根据招标要求和公司事实生成承诺/声明内容。"),
]


def _compact(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _field_key(text: str) -> str:
    compacted = _compact(text).rstrip("：:")
    for key, labels in FIELD_LABELS.items():
        if any(label in compacted for label in labels):
            return key
    return ""


def _writing_hint(title: str) -> dict[str, Any]:
    compacted = _compact(title)
    for keywords, evidence_sources, focus in WRITING_TASK_HINTS:
        if any(keyword in compacted for keyword in keywords):
            return {
                "evidence_sources": evidence_sources,
                "writing_focus": focus,
            }
    return {
        "evidence_sources": ["tender.requirements", "score_points", "company.facts"],
        "writing_focus": "根据模板标题、评分点、招标要求和公司资料补全本节内容。",
    }


def _style_has_numbering(style, seen: set[int] | None = None) -> bool:
    if style is None:
        return False
    seen = seen or set()
    marker = id(style)
    if marker in seen:
        return False
    seen.add(marker)
    try:
        p_pr = style.element.pPr
        if p_pr is not None and p_pr.numPr is not None:
            return True
    except Exception:
        pass
    try:
        return _style_has_numbering(style.base_style, seen)
    except Exception:
        return False


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for chunk in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _semantic_columns(cells: list[str]) -> dict[str, int]:
    columns: dict[str, int] = {}
    for index, raw_text in enumerate(cells):
        text = _compact(raw_text)
        if any(keyword in text for keyword in ["编号", "序号"]):
            columns.setdefault("number", index)
        if any(keyword in text for keyword in ["服务名称", "项目名称", "名称"]):
            columns.setdefault("name", index)
        if any(keyword in text for keyword in ["文件要求", "采购需求", "服务指标", "技术要求", "规格参数"]):
            columns.setdefault("requirement", index)
        if any(keyword in text for keyword in ["供应商提供", "投标响应", "响应内容", "响应指标"]):
            columns.setdefault("supplier", index)
        if any(keyword in text for keyword in ["响应程度", "偏离", "响应情况"]):
            columns.setdefault("response", index)
        if "说明" in text or "备注" in text:
            columns.setdefault("note", index)
    return columns


def _cell_span(cell) -> dict[str, Any]:
    try:
        tc_pr = cell._tc.tcPr
        grid_span = tc_pr.gridSpan.val if tc_pr is not None and tc_pr.gridSpan is not None else 1
        v_merge = tc_pr.vMerge.val if tc_pr is not None and tc_pr.vMerge is not None else None
        return {"grid_span": int(grid_span), "v_merge": v_merge}
    except Exception:
        return {"grid_span": 1, "v_merge": None}


def _table_type(header_cells: list[str], row_labels: list[str]) -> str:
    columns = _semantic_columns(header_cells)
    if {"requirement", "supplier", "response"}.issubset(columns):
        return "requirement_response"
    labels = "".join(_compact(label) for label in row_labels)
    if any(keyword in labels for keyword in ["项目名称", "采购单位", "采购人", "服务期限", "服务地点", "付款方式", "验收要求"]):
        return "labeled_fields"
    if any(_compact(cell) for cell in header_cells):
        return "structured_table"
    return "blank_or_unknown"


def _detect_placeholders(text: str, location: str) -> list[dict[str, str]]:
    return [{"location": location, "text": text.strip()}] if PLACEHOLDER_RE.search(text or "") else []


def _paragraph_fill_slot(text: str, location: str) -> dict[str, Any] | None:
    if not PLACEHOLDER_RE.search(text or ""):
        return None
    before = PLACEHOLDER_RE.split(text, maxsplit=1)[0]
    key = _field_key(before)
    return {
        "id": f"slot_{location}",
        "type": "paragraph_placeholder",
        "location": location,
        "label": before.strip(" ：:") or text.strip(),
        "semantic_key": key or "unknown",
        "required": True,
        "evidence_sources": _field_evidence_sources(key),
        "fill_strategy": "replace_placeholder",
    }


def _field_evidence_sources(key: str) -> list[str]:
    if key in {"project_name", "project_no", "package_no", "purchaser", "deadline", "location", "budget", "payment", "acceptance"}:
        return ["tender.requirements"]
    if key in {"bidder_name"}:
        return ["company.facts"]
    return ["tender.requirements", "company.facts"]


def _table_fill_slots(table_index: int, table_type: str, header_cells: list[str], row_labels: list[str]) -> list[dict[str, Any]]:
    slots: list[dict[str, Any]] = []
    columns = _semantic_columns(header_cells)
    if table_type == "requirement_response":
        source_by_key = {
            "requirement": ["tender.requirements"],
            "supplier": ["tender.requirements", "company.facts"],
            "response": ["tender.requirements"],
        }
        for row_offset, label in enumerate(row_labels, start=1):
            row_label = label or f"第 {row_offset} 行"
            if not _compact(row_label):
                continue
            for key in ["requirement", "supplier", "response"]:
                if key not in columns:
                    continue
                slots.append(
                    {
                        "id": f"table_{table_index}_row_{row_offset}_{key}",
                        "type": "table_cell",
                        "location": f"table[{table_index}].cell[{row_offset},{columns[key]}]",
                        "table_index": table_index,
                        "row_index": row_offset,
                        "column_index": columns[key],
                        "label": row_label,
                        "semantic_key": key,
                        "required": True,
                        "evidence_sources": source_by_key.get(key, ["tender.requirements", "company.facts"]),
                        "fill_strategy": "fill_by_header_semantics",
                    }
                )
        return slots

    if table_type == "labeled_fields":
        for row_offset, label in enumerate(row_labels, start=1):
            key = _field_key(label)
            if not key:
                continue
            slots.append(
                {
                    "id": f"table_{table_index}_row_{row_offset}_{key}",
                    "type": "table_labeled_field",
                    "location": f"table[{table_index}].row[{row_offset}]",
                    "table_index": table_index,
                    "row_index": row_offset,
                    "label": label,
                    "semantic_key": key,
                    "required": True,
                    "evidence_sources": _field_evidence_sources(key),
                    "fill_strategy": "fill_adjacent_blank_cell",
                }
            )
    return slots


def analyze_template(root: Path | None = None) -> Path:
    root = root or project_root()
    template_path = root / "inputs" / "template.docx"
    output_path = root / "workspace" / "template_schema.json"

    if not template_path.exists() or template_path.stat().st_size == 0:
        schema: dict[str, Any] = {
            "exists": False,
            "template_path": str(template_path),
            "headings": [],
            "tables": [],
            "placeholders": [],
            "warnings": ["inputs/template.docx 不存在，无法生成模板 schema"],
        }
        write_json(output_path, schema)
        print(f"[警告] 未发现模板，已写入空模板 schema: {output_path}")
        return output_path

    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("缺少依赖 python-docx，请先执行: pip install -r requirements.txt") from exc

    document = Document(str(template_path))
    headings: list[dict[str, Any]] = []
    placeholders: list[dict[str, str]] = []
    fill_slots: list[dict[str, Any]] = []
    writing_tasks: list[dict[str, Any]] = []
    counters = [0] * 9

    for paragraph_index, para in enumerate(document.paragraphs):
        text = para.text.strip()
        if text:
            placeholders.extend(_detect_placeholders(text, f"paragraph[{paragraph_index}]"))
            slot = _paragraph_fill_slot(text, f"paragraph[{paragraph_index}]")
            if slot:
                fill_slots.append(slot)

        style_name = para.style.name if para.style else ""
        level = None
        heading_match = re.fullmatch(r"Heading\s+(\d+)", style_name)
        if heading_match:
            level = int(heading_match.group(1))
        elif style_name.startswith("标题"):
            match = re.search(r"(\d+)", style_name)
            level = int(match.group(1)) if match else 1
        if level is None or not text:
            continue

        if 1 <= level <= len(counters):
            counters[level - 1] += 1
            for idx in range(level, len(counters)):
                counters[idx] = 0
            number = ".".join(str(value) for value in counters[:level] if value > 0)
            parent_id = ".".join(number.split(".")[:-1])
        else:
            number = ""
            parent_id = ""
        heading = (
            {
                "index": paragraph_index,
                "id": number,
                "title": text,
                "level": level,
                "parent_id": parent_id,
                "style": style_name,
                "style_has_numbering": _style_has_numbering(para.style),
            }
        )
        headings.append(heading)
        hint = _writing_hint(text)
        writing_tasks.append(
            {
                "id": f"write_{number or paragraph_index}",
                "type": "section_expansion",
                "location": f"paragraph[{paragraph_index}]",
                "heading_id": number,
                "parent_id": parent_id,
                "title": text,
                "level": level,
                "required": True,
                **hint,
            }
        )

    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        header_cells = [cell.text.strip() for cell in table.rows[0].cells] if rows else []
        semantic_columns = _semantic_columns(header_cells)
        row_labels = []
        empty_cells = 0
        total_cells = 0
        cells_matrix: list[list[dict[str, Any]]] = []
        for row_index, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            cell_details: list[dict[str, Any]] = []
            if row_index > 0:
                name_col = semantic_columns.get("name")
                preferred_label = cells[name_col] if name_col is not None and name_col < len(cells) else ""
                row_labels.append(preferred_label or next((cell for cell in cells if cell), ""))
            for col_index, cell in enumerate(row.cells):
                text = cells[col_index]
                total_cells += 1
                if not text:
                    empty_cells += 1
                placeholders.extend(_detect_placeholders(text, f"table[{table_index}].cell[{row_index},{col_index}]"))
                cell_details.append(
                    {
                        "row": row_index,
                        "col": col_index,
                        "text": text,
                        "empty": not bool(text),
                        "placeholder": bool(PLACEHOLDER_RE.search(text or "")),
                        **_cell_span(cell),
                    }
                )
            cells_matrix.append(cell_details)
        table_type = _table_type(header_cells, row_labels)
        tables.append(
            {
                "index": table_index,
                "rows": rows,
                "columns": cols,
                "type": table_type,
                "header_cells": header_cells,
                "semantic_columns": semantic_columns,
                "row_labels": row_labels,
                "empty_cells": empty_cells,
                "total_cells": total_cells,
                "cells": cells_matrix,
            }
        )
        fill_slots.extend(_table_fill_slots(table_index, table_type, header_cells, row_labels))

    paragraph_by_element = {paragraph._element: index for index, paragraph in enumerate(document.paragraphs)}
    table_by_element = {table._element: index for index, table in enumerate(document.tables)}
    document_blocks: list[dict[str, Any]] = []
    for child in document._body._element:
        if child.tag.endswith("}p"):
            paragraph_index = paragraph_by_element.get(child)
            if paragraph_index is not None:
                text = document.paragraphs[paragraph_index].text.strip()
                style_name = document.paragraphs[paragraph_index].style.name if document.paragraphs[paragraph_index].style else ""
                document_blocks.append(
                    {
                        "type": "paragraph",
                        "index": paragraph_index,
                        "text": text,
                        "style": style_name,
                    }
                )
        elif child.tag.endswith("}tbl"):
            table_index = table_by_element.get(child)
            if table_index is not None:
                document_blocks.append(
                    {
                        "type": "table",
                        "index": table_index,
                        "table_type": tables[table_index]["type"] if table_index < len(tables) else "",
                    }
                )

    heading_by_paragraph = {int(item["index"]): item for item in headings if isinstance(item.get("index"), int)}
    heading_stack: dict[int, dict[str, Any]] = {}
    location_context: dict[str, dict[str, Any]] = {}

    def _current_section_path() -> list[dict[str, Any]]:
        return [
            {
                "id": stringify(item.get("id")),
                "title": stringify(item.get("title")),
                "level": item.get("level"),
            }
            for _, item in sorted(heading_stack.items())
        ]

    def _context_from_heading(heading: dict[str, Any] | None = None) -> dict[str, Any]:
        active = heading or (heading_stack[max(heading_stack)] if heading_stack else {})
        return {
            "heading_id": stringify(active.get("id")),
            "heading_title": stringify(active.get("title")),
            "heading_level": active.get("level"),
            "section_path": _current_section_path(),
        }

    for block in document_blocks:
        if block["type"] == "paragraph":
            paragraph_index = int(block["index"])
            heading = heading_by_paragraph.get(paragraph_index)
            if heading:
                level = int(heading.get("level", 1) or 1)
                heading_stack = {key: value for key, value in heading_stack.items() if key < level}
                heading_stack[level] = heading
                context = _context_from_heading(heading)
            else:
                context = _context_from_heading()
            block.update(context)
            location_context[f"paragraph[{paragraph_index}]"] = context
        elif block["type"] == "table":
            table_index = int(block["index"])
            context = _context_from_heading()
            block.update(context)
            location_context[f"table[{table_index}]"] = context
            if table_index < len(tables):
                tables[table_index].update(context)

    for slot in fill_slots:
        location = stringify(slot.get("location"))
        context_key = ""
        paragraph_match = re.match(r"paragraph\[\d+\]", location)
        table_match = re.match(r"table\[\d+\]", location)
        if paragraph_match:
            context_key = paragraph_match.group(0)
        elif table_match:
            context_key = table_match.group(0)
        context = location_context.get(context_key)
        if context:
            slot.update({key: value for key, value in context.items() if key not in slot or not slot.get(key)})

    header_footer_text = []
    for section_index, section in enumerate(document.sections):
        header_text = "\n".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
        footer_text = "\n".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
        if header_text or footer_text:
            header_footer_text.append(
                {
                    "section": section_index,
                    "header": header_text,
                    "footer": footer_text,
                }
            )

    schema = {
        "exists": True,
        "template_path": str(template_path),
        "fingerprint": {
            "sha256": _sha256(template_path),
            "size": template_path.stat().st_size,
            "modified": template_path.stat().st_mtime,
        },
        "summary": {
            "paragraph_count": len(document.paragraphs),
            "heading_count": len(headings),
            "table_count": len(tables),
            "placeholder_count": len(placeholders),
            "fill_slot_count": len(fill_slots),
            "writing_task_count": len(writing_tasks),
            "response_table_count": sum(1 for table in tables if table["type"] == "requirement_response"),
            "labeled_field_table_count": sum(1 for table in tables if table["type"] == "labeled_fields"),
            "unknown_table_count": sum(1 for table in tables if table["type"] in {"structured_table", "blank_or_unknown"}),
        },
        "document_blocks": document_blocks,
        "headings": headings,
        "tables": tables,
        "fill_slots": fill_slots,
        "writing_tasks": writing_tasks,
        "placeholders": placeholders,
        "headers_footers": header_footer_text,
        "evidence_plan": {
            "sources": sorted(
                {
                    source
                    for item in [*fill_slots, *writing_tasks]
                    for source in item.get("evidence_sources", [])
                    if source
                }
            ),
            "fill_slot_count": len(fill_slots),
            "writing_task_count": len(writing_tasks),
        },
        "warnings": [
            f"表 {table['index']} 暂未识别为可自动填充结构"
            for table in tables
            if table["type"] in {"structured_table", "blank_or_unknown"} and table["empty_cells"] > 0
        ],
    }
    write_json(output_path, schema)
    print(
        f"[完成] 已解析模板 schema: {output_path} "
        f"(标题 {schema['summary']['heading_count']} 个，表格 {schema['summary']['table_count']} 个)"
    )
    return output_path
