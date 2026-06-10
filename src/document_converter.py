from __future__ import annotations

from pathlib import Path


SUPPORTED_EXTENSIONS = {".md", ".docx", ".pdf"}


def convert_to_markdown(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".md":
        return _read_markdown(file_path)
    elif suffix == ".docx":
        return _convert_docx(file_path)
    elif suffix == ".pdf":
        return _convert_pdf(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}，支持的格式: {SUPPORTED_EXTENSIONS}")


def _read_markdown(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")
    content = path.read_text(encoding="utf-8")
    if not content.strip():
        raise ValueError(f"Markdown 文件内容为空: {path}")
    return content


def _convert_docx(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")
    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要安装 python-docx 才能解析 .docx 文件: pip install python-docx")

    doc = Document(str(path))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            paragraphs.append("")
            continue

        style_name = (para.style.name if para.style else "").lower()
        if "heading 1" in style_name or style_name == "heading1":
            paragraphs.append(f"# {text}")
        elif "heading 2" in style_name or style_name == "heading2":
            paragraphs.append(f"## {text}")
        elif "heading 3" in style_name or style_name == "heading3":
            paragraphs.append(f"### {text}")
        elif "heading" in style_name:
            paragraphs.append(f"#### {text}")
        else:
            paragraphs.append(text)

    for table in doc.tables:
        paragraphs.append(_table_to_markdown(table))

    result = "\n\n".join(paragraphs).strip()
    if not result:
        raise ValueError(f"DOCX 文件内容为空: {path}")
    return result


def _table_to_markdown(table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])

    if not rows:
        return ""

    max_cols = max(len(row) for row in rows)
    for row in rows:
        while len(row) < max_cols:
            row.append("")

    lines: list[str] = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _convert_pdf(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    text = _try_pdf_with_pdfplumber(path)
    if text is not None:
        return text

    text = _try_pdf_with_pymupdf(path)
    if text is not None:
        return text

    raise ImportError(
        "需要安装 pdfplumber 或 PyMuPDF 才能解析 .pdf 文件。\n"
        "  pip install pdfplumber\n"
        "  或\n"
        "  pip install PyMuPDF"
    )


def _try_pdf_with_pdfplumber(path: Path) -> str | None:
    try:
        import pdfplumber
    except ImportError:
        return None

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                parts.append(text)
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    if table:
                        parts.append(_pdf_table_to_markdown(table))

    result = "\n\n".join(parts).strip()
    if not result:
        raise ValueError(f"PDF 文件内容为空或无法提取文本: {path}")
    return result


def _try_pdf_with_pymupdf(path: Path) -> str | None:
    try:
        import fitz
    except ImportError:
        return None

    parts: list[str] = []
    doc = fitz.open(str(path))
    for page in doc:
        text = page.get_text()
        if text:
            parts.append(text)
    doc.close()

    result = "\n\n".join(parts).strip()
    if not result:
        raise ValueError(f"PDF 文件内容为空或无法提取文本: {path}")
    return result


def _pdf_table_to_markdown(table: list[list[str | None]]) -> str:
    if not table:
        return ""

    cleaned: list[list[str]] = []
    for row in table:
        cleaned.append([(cell or "").strip() for cell in row])

    max_cols = max(len(row) for row in cleaned)
    for row in cleaned:
        while len(row) < max_cols:
            row.append("")

    lines: list[str] = []
    lines.append("| " + " | ".join(cleaned[0]) + " |")
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    for row in cleaned[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)
