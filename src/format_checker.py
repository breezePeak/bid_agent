from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from utils import project_root, read_json, read_text, stringify, write_json

HEADING_NUMBER_PREFIX_RE = re.compile(r"^\s*\d{1,3}(?:[.\uFF0E、]\d{1,3})*[.\uFF0E、]?\s+")
LOOSE_HEADING_NUMBER_PREFIX_RE = re.compile(r"^\s*\d{1,3}(?:[.\uFF0E、]\d{1,3})*[.\uFF0E、]?\s*")
DOUBLE_HEADING_NUMBER_RE = re.compile(
    r"^\s*\d{1,3}(?:[.\uFF0E、]\d{1,3})*[.\uFF0E、]?\s+"
    r"\d{1,3}(?:[.\uFF0E、]\d{1,3})*[.\uFF0E、]?\s+"
)
OUTPUT_PLACEHOLDER_RE = re.compile(r"(XXX+|待填写|请填写|请输入|TODO|TBD)", re.IGNORECASE)


def _item(name: str, level: str, message: str, suggestion: str = "") -> dict[str, str]:
    return {
        "name": name,
        "level": level,
        "message": message,
        "suggestion": suggestion,
    }


def _chapter_ids_from_jobs(root: Path) -> list[str]:
    jobs_dir = root / "workspace" / "jobs"
    if not jobs_dir.exists():
        return []
    chapter_ids: list[str] = []
    for job_path in sorted(jobs_dir.glob("*.json")):
        try:
            job = read_json(job_path)
        except Exception:
            chapter_ids.append(job_path.stem)
            continue
        if isinstance(job, dict):
            chapter_ids.append(str(job.get("chapter_id") or job_path.stem))
    return [chapter_id for chapter_id in chapter_ids if chapter_id]


def _style_has_numbering(style, seen: set[int] | None = None) -> bool:
    if style is None:
        return False
    seen = seen or set()
    marker = id(style)
    if marker in seen:
        return False
    seen.add(marker)

    try:
        p_pr = style.element.pPr
        if p_pr is not None and p_pr.numPr is not None:
            return True
    except Exception:
        pass

    try:
        return _style_has_numbering(style.base_style, seen)
    except Exception:
        return False


def _compact_cell_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _semantic_response_columns(table) -> dict[str, int]:
    columns: dict[str, int] = {}
    if not table.rows:
        return columns
    for index, cell in enumerate(table.rows[0].cells):
        text = _compact_cell_text(cell.text)
        if any(keyword in text for keyword in ["文件要求", "采购需求", "服务指标", "技术要求", "规格参数"]):
            columns.setdefault("requirement", index)
        if any(keyword in text for keyword in ["供应商提供", "投标响应", "响应内容", "响应指标"]):
            columns.setdefault("supplier", index)
        if any(keyword in text for keyword in ["响应程度", "偏离", "响应情况"]):
            columns.setdefault("response", index)
    return columns


def _is_semantic_response_table(table) -> bool:
    columns = _semantic_response_columns(table)
    return {"requirement", "supplier", "response"}.issubset(columns)


def _check_response_tables(document, results: list[dict[str, str]]) -> None:
    response_tables = [table for table in document.tables if _is_semantic_response_table(table)]
    if not response_tables:
        return

    empty_issues: list[str] = []
    filled_rows = 0
    for table_index, table in enumerate(response_tables, start=1):
        columns = _semantic_response_columns(table)
        for row_index, row in enumerate(table.rows[1:], start=2):
            cells = [cell.text.strip() for cell in row.cells]
            if max(columns.values()) >= len(cells):
                continue
            row_label = next((cell for cell in cells[:2] if cell), "")
            row_text = "".join(cells)
            if not row_label and not row_text:
                continue
            requirement_ok = bool(cells[columns["requirement"]])
            supplier_ok = bool(cells[columns["supplier"]])
            response_ok = bool(cells[columns["response"]])
            if requirement_ok and supplier_ok and response_ok:
                filled_rows += 1
            else:
                empty_issues.append(f"表{table_index}第{row_index}行({row_label})")

    if empty_issues:
        results.append(
            _item(
                "docx response table",
                "fail",
                "模板响应类表格存在未填写关键列: " + "、".join(empty_issues[:8]),
                "请重新执行 build-docx；生成器会按模板表头语义填充要求、供应商响应和响应程度",
            )
        )
    else:
        results.append(_item("docx response table", "ok", f"模板响应类表格已填写 {filled_rows} 行"))


def _check_template_schema(root: Path, results: list[dict[str, str]]) -> None:
    template_path = root / "inputs" / "template.docx"
    if not template_path.exists() or template_path.stat().st_size == 0:
        results.append(_item("template schema", "warn", "未提供 template.docx，跳过模板 schema 检查"))
        return

    schema_path = root / "workspace" / "template_schema.json"
    if not schema_path.exists():
        results.append(_item("template schema", "fail", "缺少 workspace/template_schema.json", "请重新执行 prepare-inputs 或 analyze-template"))
        return

    try:
        schema = read_json(schema_path)
    except Exception as exc:
        results.append(_item("template schema", "fail", f"template_schema.json 读取失败: {exc}", "请重新执行 analyze-template"))
        return

    summary = schema.get("summary", {}) if isinstance(schema, dict) else {}
    results.append(
        _item(
            "template schema",
            "ok",
            (
                f"模板 schema 已生成: 标题 {summary.get('heading_count', 0)} 个，"
                f"表格 {summary.get('table_count', 0)} 个，"
                f"填充槽位 {summary.get('fill_slot_count', 0)} 个，"
                f"写作任务 {summary.get('writing_task_count', 0)} 个"
            ),
        )
    )

    if summary.get("heading_count", 0) and not summary.get("writing_task_count", 0):
        results.append(_item("template writing tasks", "warn", "模板有标题但未生成写作任务", "请重新执行 analyze-template"))
    fill_slots = schema.get("fill_slots", []) if isinstance(schema, dict) else []
    unknown_slots = [
        slot for slot in fill_slots
        if isinstance(slot, dict) and slot.get("semantic_key") == "unknown"
    ]
    if unknown_slots:
        results.append(
            _item(
                "template unknown fill slots",
                "warn",
                f"模板存在 {len(unknown_slots)} 个未识别语义的填充槽位",
                "请补充字段映射或人工确认这些槽位含义",
            )
        )

    warnings = schema.get("warnings", []) if isinstance(schema, dict) else []
    if warnings:
        results.append(
            _item(
                "template schema warnings",
                "warn",
                "模板存在未识别或需人工确认结构: " + "；".join(str(item) for item in warnings[:5]),
                "如这些区域为必填，请补充模板语义规则或人工确认",
            )
        )


def _check_template_fill_report(root: Path, results: list[dict[str, str]]) -> None:
    template_path = root / "inputs" / "template.docx"
    if not template_path.exists() or template_path.stat().st_size == 0:
        return

    report_path = root / "workspace" / "template_fill_report.json"
    if not report_path.exists():
        results.append(_item("template fill report", "fail", "缺少 workspace/template_fill_report.json", "请重新执行 build-docx"))
        return

    try:
        report = read_json(report_path)
    except Exception as exc:
        results.append(_item("template fill report", "fail", f"template_fill_report.json 读取失败: {exc}", "请重新执行 build-docx"))
        return

    stats = report.get("fill_stats", {}) if isinstance(report, dict) else {}
    results.append(
        _item(
            "template fill report",
            "ok",
            f"模板填充报告已生成: 填充 {stats.get('tables', 0)} 个表格/{stats.get('fields', 0)} 项字段",
        )
    )

    unhandled = report.get("unhandled_tables", []) if isinstance(report, dict) else []
    if unhandled:
        results.append(
            _item(
                "template unhandled tables",
                "warn",
                f"模板中有 {len(unhandled)} 个表格未自动填充或需人工确认",
                "若这些表格不是必填可忽略；否则需补充模板语义规则",
            )
        )


def _check_template_evidence(root: Path, results: list[dict[str, str]]) -> None:
    template_path = root / "inputs" / "template.docx"
    if not template_path.exists() or template_path.stat().st_size == 0:
        return

    evidence_path = root / "workspace" / "template_evidence_map.json"
    quality_path = root / "workspace" / "template_quality_report.json"
    if not evidence_path.exists():
        results.append(_item("template evidence map", "fail", "缺少 workspace/template_evidence_map.json", "请重新执行 build-template-evidence"))
        return
    if not quality_path.exists():
        results.append(_item("template quality report", "fail", "缺少 workspace/template_quality_report.json", "请重新执行 build-template-evidence"))
        return

    try:
        evidence = read_json(evidence_path)
        quality = read_json(quality_path)
    except Exception as exc:
        results.append(_item("template evidence", "fail", f"模板依据映射读取失败: {exc}", "请重新执行 build-template-evidence"))
        return

    summary = evidence.get("summary", {}) if isinstance(evidence, dict) else {}
    results.append(
        _item(
            "template evidence map",
            "ok",
            (
                f"模板依据映射已生成: {summary.get('mapped_count', 0)} 项已映射，"
                f"{summary.get('weak_count', 0)} 项弱映射，"
                f"{summary.get('missing_count', 0)} 项缺依据"
            ),
        )
    )

    if not isinstance(quality, dict):
        results.append(_item("template quality report", "fail", "template_quality_report.json 不是 JSON 对象", "请重新执行 build-template-evidence"))
        return

    fail_count = int(quality.get("fail_count", 0) or 0)
    warn_count = int(quality.get("warn_count", 0) or 0)
    if fail_count:
        messages = [
            stringify(item.get("message"))
            for item in quality.get("results", [])
            if isinstance(item, dict) and item.get("level") == "fail"
        ]
        results.append(
            _item(
                "template quality report",
                "fail",
                "模板分析质量存在失败项: " + "；".join(messages[:5]),
                "请优先修复模板识别、槽位映射或资料缺失问题",
            )
        )
    elif warn_count:
        messages = [
            stringify(item.get("message"))
            for item in quality.get("results", [])
            if isinstance(item, dict) and item.get("level") == "warn"
        ]
        results.append(
            _item(
                "template quality report",
                "warn",
                "模板分析质量存在警告: " + "；".join(messages[:5]),
                "请确认这些模板区域是否需要人工补充",
            )
        )
    else:
        results.append(_item("template quality report", "ok", "模板分析质量报告无警告/失败项"))


def _normalize_contract_heading(value: Any) -> str:
    text = stringify(value)
    text = LOOSE_HEADING_NUMBER_PREFIX_RE.sub("", text)
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"[：:；;，,。、《》<>【】\[\]（）()]", "", text)
    return text.strip()


def _template_required_headings(root: Path) -> list[dict[str, Any]]:
    candidates = [
        root / "workspace" / "template_evidence_map.json",
        root / "workspace" / "template_schema.json",
    ]
    for path in candidates:
        if not path.exists():
            continue
        try:
            data = read_json(path)
        except Exception:
            continue
        if not isinstance(data, dict):
            continue
        contract = data.get("template_contract", {}) if isinstance(data.get("template_contract"), dict) else {}
        headings = contract.get("required_headings") or data.get("headings") or []
        if isinstance(headings, list) and headings:
            return [heading for heading in headings if isinstance(heading, dict)]
    return []


def _check_docx_template_contract(root: Path, document, results: list[dict[str, str]]) -> None:
    required_headings = _template_required_headings(root)
    if not required_headings:
        return

    paragraph_tokens = [
        (index, _normalize_contract_heading(paragraph.text))
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip()
    ]
    cursor = -1
    missing: list[str] = []
    out_of_order: list[str] = []

    for heading in required_headings:
        expected = _normalize_contract_heading(heading.get("title", ""))
        if not expected:
            continue
        label = f"{heading.get('id', '')} {heading.get('title', '')}".strip()
        found_after = next(
            (
                index
                for index, token in paragraph_tokens
                if index > cursor and (expected in token or (len(token) >= 4 and token in expected))
            ),
            None,
        )
        if found_after is not None:
            cursor = found_after
            continue

        found_anywhere = any(
            expected in token or (len(token) >= 4 and token in expected)
            for _, token in paragraph_tokens
        )
        if found_anywhere:
            out_of_order.append(label)
        else:
            missing.append(label)

    if missing or out_of_order:
        messages: list[str] = []
        if missing:
            messages.append("缺少标题: " + "、".join(missing[:8]))
        if out_of_order:
            messages.append("标题顺序异常: " + "、".join(out_of_order[:8]))
        results.append(
            _item(
                "docx template contract",
                "fail",
                "final.docx 未严格保留模板标题结构，" + "；".join(messages),
                "请重新执行 analyze-template、build-template-evidence、build-docx 和 check-format",
            )
        )
    else:
        results.append(
            _item(
                "docx template contract",
                "ok",
                f"final.docx 已按模板保留 {len(required_headings)} 个标题及顺序",
            )
        )


def _check_output_placeholders(document, results: list[dict[str, str]]) -> None:
    hits: list[str] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if OUTPUT_PLACEHOLDER_RE.search(text):
            hits.append(f"段落{paragraph_index}: {text[:40]}")
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                if OUTPUT_PLACEHOLDER_RE.search(text):
                    hits.append(f"表{table_index}({row_index},{col_index}): {text[:40]}")
    if hits:
        results.append(
            _item(
                "docx placeholders",
                "fail",
                "final.docx 仍存在明显占位符: " + "；".join(hits[:8]),
                "请检查模板字段映射或人工补齐后重新生成",
            )
        )
    else:
        results.append(_item("docx placeholders", "ok", "final.docx 未发现明显待填写占位符"))


def _normalize_heading_id(value: str) -> str:
    parts = re.split(r"[.\uFF0E、]", value.strip())
    normalized = []
    for part in parts:
        if not part:
            continue
        normalized.append(str(int(part)) if part.isdigit() else part)
    return ".".join(normalized)


def _heading_has_chapter_id(markdown: str, chapter_id: str) -> bool:
    expected = _normalize_heading_id(chapter_id)
    for heading_id in re.findall(r"(?m)^#{1,6}\s*(\d+(?:[.\uFF0E、]\d+)*)\b", markdown):
        if _normalize_heading_id(heading_id) == expected:
            return True
    return False


def _check_markdown(root: Path, results: list[dict[str, str]]) -> None:
    final_md = root / "outputs" / "final.md"
    if not final_md.exists() or final_md.stat().st_size == 0:
        results.append(_item("final.md", "fail", "final.md 不存在或为空", "请先执行 build-md"))
        return

    markdown = read_text(final_md)
    headings = re.findall(r"(?m)^#{1,6}\s+(.+)$", markdown)
    if not headings:
        results.append(_item("markdown headings", "fail", "final.md 未检测到标题", "请检查章节正文是否以 Markdown 标题开头"))
    else:
        results.append(_item("markdown headings", "ok", f"final.md 检测到 {len(headings)} 个标题"))

    duplicate_number_headings = [heading for heading in headings if DOUBLE_HEADING_NUMBER_RE.search(heading)]
    if duplicate_number_headings:
        sample = "；".join(duplicate_number_headings[:5])
        results.append(
            _item(
                "markdown duplicate heading numbers",
                "fail",
                f"final.md 检测到重复标题编号: {sample}",
                "请去掉标题正文中的重复编号，只保留一个章节编号",
            )
        )
    elif headings:
        results.append(_item("markdown duplicate heading numbers", "ok", "final.md 未发现重复标题编号"))

    expected_ids = _chapter_ids_from_jobs(root)
    if expected_ids:
        missing_ids = [
            chapter_id
            for chapter_id in expected_ids
            if not _heading_has_chapter_id(markdown, chapter_id)
        ]
        if missing_ids:
            results.append(
                _item(
                    "markdown chapter coverage",
                    "fail",
                    f"final.md 缺少章节标题: {', '.join(missing_ids)}",
                    "请检查 workspace/chapters 下是否存在对应章节，并重新执行 build-md",
                )
            )
        else:
            results.append(_item("markdown chapter coverage", "ok", "final.md 已包含所有任务章节标题"))

    if "```" in markdown:
        results.append(_item("markdown code fences", "warn", "final.md 中存在 Markdown 代码块标记", "如非必要，请检查模型是否把正文包进代码块"))
    else:
        results.append(_item("markdown code fences", "ok", "final.md 未发现代码块包裹标记"))


def _check_docx(root: Path, results: list[dict[str, str]]) -> None:
    final_docx = root / "outputs" / "final.docx"
    if not final_docx.exists() or final_docx.stat().st_size == 0:
        results.append(_item("final.docx", "fail", "final.docx 不存在或为空", "请先执行 build-docx"))
        return

    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("缺少依赖 python-docx，请先执行: pip install -r requirements.txt") from exc

    try:
        document = Document(str(final_docx))
    except Exception as exc:
        results.append(_item("docx readable", "fail", f"final.docx 无法打开: {exc}", "请重新执行 build-docx"))
        return

    nonempty_paragraphs = [p for p in document.paragraphs if p.text.strip()]
    if not nonempty_paragraphs:
        results.append(_item("docx content", "fail", "final.docx 没有可见正文段落", "请检查 final.md 并重新生成 Word"))
    else:
        results.append(_item("docx content", "ok", f"final.docx 包含 {len(nonempty_paragraphs)} 个非空段落"))

    heading_paragraphs = [
        p
        for p in document.paragraphs
        if p.text.strip() and p.style and "heading" in p.style.name.lower()
    ]
    if not heading_paragraphs:
        results.append(_item("docx headings", "warn", "final.docx 未检测到 Heading 标题样式", "请检查模板标题样式是否可识别"))
    else:
        results.append(_item("docx headings", "ok", f"final.docx 检测到 {len(heading_paragraphs)} 个标题样式段落"))

    numbered_style_headings = [
        p.text.strip()
        for p in heading_paragraphs
        if _style_has_numbering(p.style) and HEADING_NUMBER_PREFIX_RE.search(p.text.strip())
    ]
    if numbered_style_headings:
        sample = "；".join(numbered_style_headings[:5])
        results.append(
            _item(
                "docx duplicate heading numbers",
                "fail",
                f"final.docx 标题样式自带编号，但标题文本也包含编号: {sample}",
                "请重新执行 build-docx；生成器会在带自动编号的标题样式下去掉标题文本编号",
            )
        )
    elif heading_paragraphs:
        results.append(_item("docx duplicate heading numbers", "ok", "final.docx 未发现自动编号标题叠加正文编号"))

    if document.tables:
        results.append(_item("docx tables", "ok", f"final.docx 包含 {len(document.tables)} 个表格"))
    else:
        results.append(_item("docx tables", "warn", "final.docx 未检测到表格", "如标书应包含表格，请检查章节 Markdown 表格格式"))
    _check_response_tables(document, results)
    _check_output_placeholders(document, results)
    _check_docx_template_contract(root, document, results)


def check_output_format(root: Path | None = None) -> Path:
    root = root or project_root()
    results: list[dict[str, str]] = []
    _check_template_schema(root, results)
    _check_template_evidence(root, results)
    _check_markdown(root, results)
    _check_docx(root, results)
    _check_template_fill_report(root, results)

    fail_count = sum(1 for item in results if item["level"] == "fail")
    warn_count = sum(1 for item in results if item["level"] == "warn")
    ok_count = sum(1 for item in results if item["level"] == "ok")
    report: dict[str, Any] = {
        "ok": fail_count == 0,
        "ok_count": ok_count,
        "warn_count": warn_count,
        "fail_count": fail_count,
        "results": results,
    }

    output_path = root / "workspace" / "format_check_report.json"
    write_json(output_path, report)
    print(f"[完成] 格式检查完成: OK={ok_count}, WARN={warn_count}, FAIL={fail_count} -> {output_path}")
    if fail_count:
        raise RuntimeError(f"格式检查失败: {fail_count} 项 fail，请查看 {output_path}")
    return output_path
