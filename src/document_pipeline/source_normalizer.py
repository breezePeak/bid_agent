from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Any

from control_plane import ControlStore, WorkspaceContext
from docx import Document
from docx.text.paragraph import Paragraph
from utils import write_json

from .contracts import (
    SOURCE_PARSER_VERSION,
    AmendmentRelation,
    InputItem,
    InputRole,
    SourceBlock,
    SourceIndex,
    SourceInputStatus,
    SourceNormalizationCoverage,
    SourceNormalizationCoverageItem,
    SourceAnchor,
)
from .input_manifest import InputManifestService, V3_ROOT
from .source_artifacts import (
    SOURCE_INDEX_PATH,
    derive_by_role_view,
    load_promoted_input_manifest,
    promote_source_artifact,
)


# Re-export for transitional imports.
__all__ = ["SOURCE_INDEX_PATH", "SourceNormalizer"]
NORMALIZABLE_EXTENSIONS = frozenset({".md", ".txt", ".docx", ".pdf"})
_MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_LIST_ITEM = re.compile(r"^\s*(?:[-*+] |\d+[.)] )")


class SourceNormalizer:
    """Recover source order and anchors before any semantic Agent is invoked."""

    def __init__(self, context: WorkspaceContext) -> None:
        self.context = context
        self.root = context.root
        self.inputs = InputManifestService(context)

    def normalize_active_inputs(self) -> dict[str, object]:
        manifest = load_promoted_input_manifest(self.context) or self.inputs.load()
        if load_promoted_input_manifest(self.context) is None and not self.inputs.manifest_path.exists():
            raise ValueError("V3_INPUT_MANIFEST_MISSING: 没有可规范化的输入清单")

        blocks: list[SourceBlock] = []
        coverage_items: list[SourceNormalizationCoverageItem] = []
        input_status: list[SourceInputStatus] = []
        blocked_reasons: list[str] = []
        hard_blocking_roles = {
            InputRole.TENDER,
            InputRole.SCORE,
            InputRole.AMENDMENT,
            InputRole.TEMPLATE,
        }

        for item in manifest.inputs:
            if not item.active:
                continue
            if item.role is InputRole.LEGACY_BID:
                input_status.append(
                    SourceInputStatus(
                        input_id=item.input_id,
                        status="excluded",
                        reason="legacy_bid is reserved for the unreleased rewrite pipeline",
                    )
                )
                continue
            source = self.root / V3_ROOT / "sources" / item.input_id / item.filename
            try:
                source_blocks, coverage = self._blocks_for(item, source)
            except ValueError as exc:
                reason = str(exc)
                input_status.append(SourceInputStatus(input_id=item.input_id, status="blocked", reason=reason))
                if item.role in hard_blocking_roles:
                    blocked_reasons.append(f"{item.role.value}/{item.filename}: {reason}")
                coverage_items.append(
                    SourceNormalizationCoverageItem(
                        element_id=f"{item.input_id}:input",
                        input_id=item.input_id,
                        element_kind="input_file",
                        status="structure_gap",
                        locator=f"input:{item.input_id}",
                        reason=reason,
                    )
                )
                continue
            gap_count = sum(1 for item_cov in coverage if item_cov.status == "structure_gap")
            status = "partial" if gap_count else "processed"
            input_status.append(
                SourceInputStatus(
                    input_id=item.input_id,
                    status=status,
                    block_count=len(source_blocks),
                    reason=f"{gap_count} structure gaps" if gap_count else None,
                )
            )
            blocks.extend(source_blocks)
            coverage_items.extend(coverage)

        active_manifest = ControlStore(self.context).v3_active_artifact("InputManifest")
        manifest_hash = str(active_manifest["artifact_hash"]) if active_manifest is not None else ""
        index = SourceIndex(
            revision=max(1, int(manifest.revision)),
            source_hashes={
                entry.input_id: entry.sha256
                for entry in manifest.inputs
                if entry.active and entry.role is not InputRole.LEGACY_BID
            },
            parser_version=SOURCE_PARSER_VERSION,
            input_manifest_revision=int(manifest.revision),
            input_manifest_artifact_hash=manifest_hash,
            blocks=blocks,
            coverage=SourceNormalizationCoverage(items=coverage_items),
            amendments=self._amendments(manifest.inputs),
            input_status=input_status,
        )

        # Hard-block only when an input is fully blocked (unsupported/corrupt), not for partial OCR gaps.
        if blocked_reasons:
            messages = "; ".join(blocked_reasons)
            raise ValueError(f"V3_SOURCE_NORMALIZATION_BLOCKED: {messages}")

        promote_source_artifact(
            self.context,
            artifact_kind="SourceIndex",
            payload=index.model_dump(mode="json"),
            operation_id=f"source-index:{manifest.revision}:{SOURCE_PARSER_VERSION}:{manifest_hash[:16]}",
            gate_id="G0_SOURCE_STRUCTURE",
            cited_source_ids=sorted(index.source_hashes),
        )
        promoted = ControlStore(self.context).v3_active_artifact("SourceIndex")
        assert promoted is not None
        promoted_index = SourceIndex.model_validate(promoted["payload"])
        # Compatibility projection exposes a derived by_role view (never authoritative).
        projection = promoted_index.model_dump(mode="json")
        projection["by_role"] = derive_by_role_view(promoted_index)
        projection["authority"] = "promoted_artifact_projection"
        write_json(self.root / SOURCE_INDEX_PATH, projection)
        return projection

    @staticmethod
    def _amendments(items: list[InputItem]) -> list[AmendmentRelation]:
        return [
            AmendmentRelation(
                input_id=item.input_id,
                issued_at=str(item.issued_at),
                supersedes_input_ids=list(item.supersedes_input_ids),
                replaces_input_id=item.replaces_input_id,
            )
            for item in sorted(
                (entry for entry in items if entry.active and entry.role is InputRole.AMENDMENT),
                key=lambda entry: (str(entry.issued_at), entry.input_id),
            )
        ]

    def _blocks_for(
        self, item: InputItem, source: Path
    ) -> tuple[list[SourceBlock], list[SourceNormalizationCoverageItem]]:
        suffix = source.suffix.lower()
        if suffix not in NORMALIZABLE_EXTENSIONS:
            supported = ", ".join(sorted(NORMALIZABLE_EXTENSIONS))
            raise ValueError(f"V3_SOURCE_FORMAT_UNSUPPORTED: {suffix or '<无扩展名>'}；支持的格式: {supported}")
        if suffix in {".md", ".txt"}:
            return self._markdown_blocks(item, source)
        if suffix == ".docx":
            return self._docx_blocks(item, source)
        return self._pdf_blocks(item, source)

    def _markdown_blocks(
        self, item: InputItem, source: Path
    ) -> tuple[list[SourceBlock], list[SourceNormalizationCoverageItem]]:
        headings: list[str] = []
        blocks: list[SourceBlock] = []
        coverage: list[SourceNormalizationCoverageItem] = []
        for paragraph_index, line in enumerate(source.read_text(encoding="utf-8").replace("\r\n", "\n").splitlines()):
            content = line.strip()
            locator = f"line:{paragraph_index + 1}"
            if not content:
                coverage.append(
                    SourceNormalizationCoverageItem(
                        element_id=f"{item.input_id}:{locator}",
                        input_id=item.input_id,
                        element_kind="blank_line",
                        status="exempt",
                        locator=locator,
                        reason="empty line",
                    )
                )
                continue
            heading = _MARKDOWN_HEADING.match(content)
            if heading:
                level, title = len(heading.group(1)), heading.group(2).strip()
                headings = headings[: level - 1]
                headings.append(title)
                kind = "heading"
                content = title
            else:
                kind = "list_item" if _LIST_ITEM.match(content) else "paragraph"
            block = self._block(item, kind, len(blocks), content, headings, paragraph_index=paragraph_index)
            blocks.append(block)
            coverage.append(
                SourceNormalizationCoverageItem(
                    element_id=f"{item.input_id}:{locator}",
                    input_id=item.input_id,
                    element_kind=kind,
                    status="normalized",
                    locator=locator,
                    block_id=block.block_id,
                )
            )
        return blocks, coverage

    def _docx_blocks(
        self, item: InputItem, source: Path
    ) -> tuple[list[SourceBlock], list[SourceNormalizationCoverageItem]]:
        document = Document(str(source))
        blocks: list[SourceBlock] = []
        coverage: list[SourceNormalizationCoverageItem] = []
        headings: list[str] = []
        paragraph_index = table_index = 0
        paragraphs = iter(document.paragraphs)
        tables = iter(document.tables)
        for child in document.element.body.iterchildren():
            if child.tag.endswith("}p"):
                paragraph = next(paragraphs)
                text = paragraph.text.strip()
                current_index = paragraph_index
                paragraph_index += 1
                locator = f"paragraph:{current_index + 1}"
                if not text:
                    coverage.append(
                        SourceNormalizationCoverageItem(
                            element_id=f"{item.input_id}:{locator}",
                            input_id=item.input_id,
                            element_kind="blank_paragraph",
                            status="exempt",
                            locator=locator,
                            reason="empty paragraph",
                        )
                    )
                    continue
                level = self._docx_heading_level(paragraph)
                if level is not None:
                    headings = headings[: level - 1]
                    headings.append(text)
                    kind = "heading"
                else:
                    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
                    kind = "list_item" if "list" in style_name else "paragraph"
                block = self._block(item, kind, len(blocks), text, headings, paragraph_index=current_index)
                blocks.append(block)
                coverage.append(
                    SourceNormalizationCoverageItem(
                        element_id=f"{item.input_id}:{locator}",
                        input_id=item.input_id,
                        element_kind=kind,
                        status="normalized",
                        locator=locator,
                        block_id=block.block_id,
                    )
                )
            elif child.tag.endswith("}tbl"):
                table = next(tables)
                current_table = table_index
                table_index += 1
                table_locator = f"table:{current_table + 1}"
                # Table container block preserves document order of the whole table.
                table_block = self._block(
                    item,
                    "table",
                    len(blocks),
                    f"[table {current_table + 1} rows={len(table.rows)}]",
                    headings,
                    table_index=current_table,
                )
                blocks.append(table_block)
                coverage.append(
                    SourceNormalizationCoverageItem(
                        element_id=f"{item.input_id}:{table_locator}",
                        input_id=item.input_id,
                        element_kind="table",
                        status="normalized",
                        locator=table_locator,
                        block_id=table_block.block_id,
                    )
                )
                for row_index, row in enumerate(table.rows):
                    # ``python-docx`` expands a horizontally merged cell once per
                    # occupied grid column.  Those entries wrap the same ``w:tc``
                    # element and would otherwise become duplicate SourceBlocks.
                    #
                    # Keep this registry row-local on purpose.  A vertically
                    # merged cell is also exposed through the same ``w:tc`` in
                    # every covered row, but retaining one occurrence per row
                    # preserves the logical row context used by scoring tables.
                    canonical_cells: dict[int, tuple[str, str | None]] = {}
                    for column_index, cell in enumerate(row.cells):
                        content = cell.text.strip()
                        cell_locator = f"{table_locator}:row:{row_index + 1}:cell:{column_index + 1}"
                        cell_key = id(cell._tc)
                        canonical = canonical_cells.get(cell_key)
                        if canonical is not None:
                            canonical_locator, canonical_block_id = canonical
                            coverage.append(
                                SourceNormalizationCoverageItem(
                                    element_id=f"{item.input_id}:{cell_locator}",
                                    input_id=item.input_id,
                                    element_kind="table_cell",
                                    status="exempt",
                                    locator=cell_locator,
                                    reason=f"merged-cell alias of {canonical_locator}",
                                    block_id=canonical_block_id,
                                )
                            )
                            continue
                        if not content:
                            canonical_cells[cell_key] = (cell_locator, None)
                            coverage.append(
                                SourceNormalizationCoverageItem(
                                    element_id=f"{item.input_id}:{cell_locator}",
                                    input_id=item.input_id,
                                    element_kind="table_cell",
                                    status="exempt",
                                    locator=cell_locator,
                                    reason="empty cell",
                                )
                            )
                            continue
                        block = self._block(
                            item,
                            "table_cell",
                            len(blocks),
                            content,
                            headings,
                            table_index=current_table,
                            row_index=row_index,
                            column_index=column_index,
                        )
                        canonical_cells[cell_key] = (cell_locator, block.block_id)
                        blocks.append(block)
                        coverage.append(
                            SourceNormalizationCoverageItem(
                                element_id=f"{item.input_id}:{cell_locator}",
                                input_id=item.input_id,
                                element_kind="table_cell",
                                status="normalized",
                                locator=cell_locator,
                                block_id=block.block_id,
                            )
                        )
        return blocks, coverage

    @staticmethod
    def _docx_heading_level(paragraph: Paragraph) -> int | None:
        style = paragraph.style
        tokens = f"{style.name} {style.style_id}".lower() if style else ""
        matched = re.search(r"(?:heading|标题)\s*([1-9])", tokens)
        if matched:
            return int(matched.group(1))
        ppr = paragraph._p.pPr
        if ppr is not None and ppr.outlineLvl is not None:
            try:
                return int(ppr.outlineLvl.val) + 1
            except (TypeError, ValueError):
                return None
        return None

    def _pdf_blocks(
        self, item: InputItem, source: Path
    ) -> tuple[list[SourceBlock], list[SourceNormalizationCoverageItem]]:
        try:
            import pdfplumber
        except ImportError as exc:
            raise ValueError("V3_SOURCE_PDF_UNAVAILABLE: 缺少 pdfplumber，无法可靠解析 PDF。") from exc
        blocks: list[SourceBlock] = []
        coverage: list[SourceNormalizationCoverageItem] = []
        pages_with_content = 0
        pages_with_gap = 0
        with pdfplumber.open(str(source)) as pdf:
            if not pdf.pages:
                raise ValueError("V3_SOURCE_PDF_EMPTY: PDF 不包含页面。")
            for page_index, page in enumerate(pdf.pages, start=1):
                page_entries: list[tuple[float, float, SourceBlock, SourceNormalizationCoverageItem]] = []
                words = page.extract_words() or []
                if words:
                    # Group words into approximate lines by top coordinate.
                    lines: dict[int, list[dict[str, Any]]] = {}
                    for word in words:
                        top_key = int(round(float(word.get("top", 0))))
                        lines.setdefault(top_key, []).append(word)
                    for line_index, top_key in enumerate(sorted(lines)):
                        line_words = sorted(lines[top_key], key=lambda item: float(item.get("x0", 0)))
                        content = " ".join(str(word.get("text") or "").strip() for word in line_words).strip()
                        if not content:
                            continue
                        x0 = min(float(word.get("x0", 0)) for word in line_words)
                        x1 = max(float(word.get("x1", 0)) for word in line_words)
                        top = min(float(word.get("top", 0)) for word in line_words)
                        bottom = max(float(word.get("bottom", top)) for word in line_words)
                        bbox = [x0, top, x1, bottom]
                        # Temporary ordinal 0; reassigned after global sort.
                        block = self._block(
                            item,
                            "pdf_text",
                            0,
                            content,
                            [],
                            page=page_index,
                            paragraph_index=line_index,
                            bbox=bbox,
                            reading_order=0,
                        )
                        locator = f"page:{page_index}:text:{line_index + 1}"
                        cov = SourceNormalizationCoverageItem(
                            element_id=f"{item.input_id}:{locator}",
                            input_id=item.input_id,
                            element_kind="pdf_text",
                            status="normalized",
                            locator=locator,
                            block_id=block.block_id,
                        )
                        page_entries.append((top, x0, block, cov))
                else:
                    text = (page.extract_text() or "").strip()
                    if text:
                        for line_index, line in enumerate(text.splitlines()):
                            content = line.strip()
                            if not content:
                                continue
                            block = self._block(
                                item,
                                "pdf_text",
                                0,
                                content,
                                [],
                                page=page_index,
                                paragraph_index=line_index,
                                reading_order=0,
                            )
                            locator = f"page:{page_index}:text:{line_index + 1}"
                            cov = SourceNormalizationCoverageItem(
                                element_id=f"{item.input_id}:{locator}",
                                input_id=item.input_id,
                                element_kind="pdf_text",
                                status="normalized",
                                locator=locator,
                                block_id=block.block_id,
                            )
                            page_entries.append((float(line_index), 0.0, block, cov))

                for table_index, table in enumerate(page.find_tables() if hasattr(page, "find_tables") else []):
                    table_bbox = getattr(table, "bbox", None)
                    extracted = table.extract() if hasattr(table, "extract") else None
                    if not extracted:
                        continue
                    table_top = float(table_bbox[1]) if table_bbox else float(10_000 + table_index)
                    table_x0 = float(table_bbox[0]) if table_bbox else 0.0
                    for row_index, row in enumerate(extracted):
                        for column_index, value in enumerate(row or []):
                            content = str(value or "").strip()
                            if not content:
                                continue
                            bbox = list(table_bbox) if table_bbox and len(table_bbox) == 4 else None
                            block = self._block(
                                item,
                                "pdf_table_cell",
                                0,
                                content,
                                [],
                                page=page_index,
                                table_index=table_index,
                                row_index=row_index,
                                column_index=column_index,
                                bbox=bbox,
                                reading_order=0,
                            )
                            locator = f"page:{page_index}:table:{table_index + 1}:row:{row_index + 1}:cell:{column_index + 1}"
                            cov = SourceNormalizationCoverageItem(
                                element_id=f"{item.input_id}:{locator}",
                                input_id=item.input_id,
                                element_kind="pdf_table_cell",
                                status="normalized",
                                locator=locator,
                                block_id=block.block_id,
                            )
                            page_entries.append((table_top + row_index * 0.01, table_x0 + column_index * 0.01, block, cov))

                if not page_entries:
                    # Blank/low-text page → StructureGap, not whole-document failure.
                    pages_with_gap += 1
                    gap = self._block(
                        item,
                        "ocr_gap",
                        0,
                        f"[STRUCTURE_GAP page={page_index}]",
                        [],
                        page=page_index,
                        reading_order=0,
                    )
                    locator = f"page:{page_index}:ocr_gap"
                    cov = SourceNormalizationCoverageItem(
                        element_id=f"{item.input_id}:{locator}",
                        input_id=item.input_id,
                        element_kind="ocr_gap",
                        status="structure_gap",
                        locator=locator,
                        reason="page has no extractable text/table content",
                        block_id=gap.block_id,
                    )
                    page_entries.append((0.0, 0.0, gap, cov))
                else:
                    pages_with_content += 1

                # Sort text and tables on this page by real position, then emit.
                page_entries.sort(key=lambda entry: (entry[0], entry[1]))
                for _, _, block, cov in page_entries:
                    # Re-stamp ordinal/reading_order for global stream stability.
                    restamped = self._block(
                        item,
                        block.block_kind,
                        len(blocks),
                        block.content,
                        block.heading_path,
                        page=block.page,
                        paragraph_index=block.paragraph_index,
                        table_index=block.table_index,
                        row_index=block.row_index,
                        column_index=block.column_index,
                        bbox=block.bbox,
                        reading_order=len(blocks),
                    )
                    blocks.append(restamped)
                    coverage.append(cov.model_copy(update={"block_id": restamped.block_id}))

        if pages_with_content == 0 and pages_with_gap > 0:
            raise ValueError("V3_SOURCE_OCR_BLOCKED: PDF 全部页面均无可提取文本（疑似扫描件）。")
        return blocks, coverage

    @staticmethod
    def _block(
        item: InputItem,
        kind: str,
        ordinal: int,
        content: str,
        heading_path: list[str],
        **coordinates: Any,
    ) -> SourceBlock:
        location = SourceNormalizer._location(kind, ordinal, coordinates)
        # Identity is content-addressed on file hash + parser + kind + locator (not free agent text).
        identity = f"{item.sha256}:{SOURCE_PARSER_VERSION}:{kind}:{location}:{content}"
        block_id = hashlib.sha256(identity.encode("utf-8")).hexdigest()[:24]
        content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()
        return SourceBlock(
            block_id=block_id,
            input_id=item.input_id,
            input_role=item.role,
            block_kind=kind,  # type: ignore[arg-type]
            ordinal=ordinal,
            content=content,
            heading_path=list(heading_path),
            parser_version=SOURCE_PARSER_VERSION,
            source_anchor=SourceAnchor(
                source_input_id=item.input_id,
                chunk_id=block_id,
                page=coordinates.get("page"),
                location=location,
            ),
            content_hash=content_hash,
            **{key: value for key, value in coordinates.items() if key in {
                "page", "paragraph_index", "table_index", "row_index", "column_index", "bbox", "reading_order",
            }},
        )

    @staticmethod
    def _location(kind: str, ordinal: int, coordinates: dict[str, Any]) -> str:
        page_prefix = (
            f"page:{int(coordinates['page'])}:"
            if coordinates.get("page") is not None
            else ""
        )
        if coordinates.get("table_index") is not None and coordinates.get("row_index") is not None:
            return page_prefix + "table:{table}:row:{row}:cell:{column}".format(
                table=int(coordinates["table_index"]) + 1,
                row=int(coordinates.get("row_index") or 0) + 1,
                column=int(coordinates.get("column_index") or 0) + 1,
            )
        if coordinates.get("table_index") is not None:
            return page_prefix + f"table:{int(coordinates['table_index']) + 1}"
        if coordinates.get("page") is not None:
            return f"page:{coordinates['page']}:block:{ordinal + 1}:{kind}"
        return f"paragraph:{int(coordinates.get('paragraph_index') or 0) + 1}:{kind}"
