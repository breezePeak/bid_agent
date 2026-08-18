"""Word export for the current chapter workbench state."""

from __future__ import annotations

from pathlib import Path

from control_plane import WorkspaceContext

from .chapter_editing import ChapterEditingService


def build_current_word(context: WorkspaceContext) -> Path:
    """Write a downloadable Word draft, retaining headings with no body."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("缺少 Word 导出依赖 python-docx。") from exc

    document = Document()
    composed = ChapterEditingService(context).compose_current_document()
    for chapter in composed["chapters"]:
        title = str(chapter.get("title") or chapter.get("chapter_id") or "未命名章节")
        document.add_heading(title, level=min(int(chapter.get("depth") or 0) + 1, 4))
        for block in chapter.get("blocks") or []:
            content = str(block.get("content") or "").strip()
            for paragraph in content.split("\n\n"):
                if paragraph.strip():
                    document.add_paragraph(paragraph.strip())

    output_dir = context.root / "workspace" / "exports"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "current.docx"
    document.save(str(output_path))
    return output_path
