from __future__ import annotations

import hashlib
import re
from pathlib import Path

from control_plane import WorkspaceContext
from docx import Document
from docx.text.paragraph import Paragraph
from utils import read_json

from .contracts import ContractNode, InputItem, RequirementLedger, TemplateContract, TemplateSlot, TemplateStructureContract
from .input_manifest import V3_ROOT
from .requirement_ledger import load_promoted_requirement_ledger
from .source_artifacts import TEMPLATE_STRUCTURE_PATH, promote_source_artifact, write_template_structure_projection


_PLACEHOLDER = re.compile(r"\{\{([^{}]+)\}\}|【([^】]+)】|\[([^\[\]]+)\]")
_HEADING = re.compile(r"(?:heading|标题)\s*([1-9])", re.IGNORECASE)
_NUMBERING = re.compile(r"^([0-9]+(?:\.[0-9]+)*|[一二三四五六七八九十]+[、.．]|（[0-9]+）|\([0-9]+\))\s*")


class TemplateContractCompiler:
    """Read the original OOXML document; never reconstruct its structure."""

    def __init__(self, context: WorkspaceContext) -> None:
        self.context = context
        self.root = context.root

    def compile(self, template: InputItem) -> TemplateContract:
        ledger = load_promoted_requirement_ledger(self.context)
        structure = self.compile_structure(template)
        blocking_gaps = self._coverage_gaps(structure.nodes, ledger)
        return TemplateContract(
            revision=ledger.revision,
            source_hashes={**ledger.source_hashes, template.input_id: template.sha256},
            template_hash=structure.template_hash,
            structural_fingerprint=structure.structural_fingerprint,
            nodes=structure.nodes,
            slots=structure.slots,
            warnings=[] if structure.slots else ["模板未声明可写 slot；仅允许在后续人工确认的 flow_slot 内填充。"],
            blocking_gaps=blocking_gaps,
        )

    def compile_structure(self, template: InputItem) -> TemplateStructureContract:
        """Freeze template topology before Requirement/Score/Blueprint planning."""
        path = self.root / V3_ROOT / "sources" / template.input_id / template.filename
        if path.suffix.lower() != ".docx":
            raise ValueError("TEMPLATE_INVALID: 严格模板当前只支持 DOCX。")
        try:
            document = Document(str(path))
        except Exception as exc:  # python-docx normalizes malformed/package errors
            raise ValueError(f"TEMPLATE_INVALID: 无法读取活动模板: {exc}") from exc
        nodes, paragraph_node = self._nodes(document)
        if not nodes:
            raise ValueError("TEMPLATE_INVALID: 未能可靠识别模板标题结构")
        structure = TemplateStructureContract(
            revision=template.version,
            source_hashes={template.input_id: template.sha256},
            template_input_id=template.input_id,
            template_hash=template.sha256,
            structural_fingerprint=self._fingerprint(path),
            nodes=nodes,
            slots=self._slots(document, paragraph_node),
        )
        promote_source_artifact(
            self.context,
            artifact_kind="TemplateStructureContract",
            payload=structure.model_dump(mode="json"),
            operation_id=f"template-structure:{template.input_id}:{template.sha256[:16]}:{structure.structural_fingerprint[:16]}",
            gate_id="G0_TEMPLATE_STRUCTURE",
            cited_source_ids=[template.input_id],
        )
        from control_plane import ControlStore

        promoted = ControlStore(self.context).v3_active_artifact("TemplateStructureContract")
        if promoted is None:
            raise ValueError("TEMPLATE_STRUCTURE_PROMOTION_FAILED")
        promoted_structure = TemplateStructureContract.model_validate(promoted["payload"])
        write_template_structure_projection(self.context, promoted_structure)
        return promoted_structure

    @staticmethod
    def _fingerprint(path: Path) -> str:
        document = Document(str(path))
        shape = {
            "paragraphs": [
                {
                    "style": paragraph.style.style_id if paragraph.style else "",
                    # Heading text is part of the immutable template topology.
                    # Body text may contain a declared slot whose replacement is
                    # not recoverable from the rendered document alone.
                    "heading": bool(_HEADING.search(
                        f"{paragraph.style.name} {paragraph.style.style_id}" if paragraph.style else ""
                    )),
                    "text": paragraph.text if _HEADING.search(
                        f"{paragraph.style.name} {paragraph.style.style_id}" if paragraph.style else ""
                    ) else "",
                }
                for paragraph in document.paragraphs
            ],
            "tables": [
                {
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                }
                for table in document.tables
            ],
            "sections": len(document.sections),
        }
        return hashlib.sha256(repr(shape).encode("utf-8")).hexdigest()

    def _nodes(self, document: Document) -> tuple[list[ContractNode], dict[int, str]]:
        nodes: list[ContractNode] = []
        paragraph_node: dict[int, str] = {}
        parents: dict[int, str] = {}
        current_node: str | None = None
        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            level = self._heading_level(paragraph)
            if not text:
                continue
            if level is None:
                if current_node:
                    paragraph_node[index] = current_node
                continue
            node_id = f"p-{index + 1}"
            parent = next((parents[candidate] for candidate in range(level - 1, 0, -1) if candidate in parents), None)
            parents[level] = node_id
            for candidate in list(parents):
                if candidate > level:
                    del parents[candidate]
            numbering_match = _NUMBERING.match(text)
            nodes.append(
                ContractNode(
                    node_id=node_id,
                    parent_node_id=parent,
                    order=len(nodes),
                    level=level,
                    numbering=numbering_match.group(1) if numbering_match else None,
                    writable_target=f"paragraph:{index + 1}",
                    title=text,
                )
            )
            paragraph_node[index] = node_id
            current_node = node_id
        return nodes, paragraph_node

    @staticmethod
    def _heading_level(paragraph: Paragraph) -> int | None:
        style = paragraph.style
        style_tokens = f"{style.name} {style.style_id}" if style else ""
        matched = _HEADING.search(style_tokens)
        if matched:
            return int(matched.group(1))
        ppr = paragraph._p.pPr
        if ppr is not None and ppr.outlineLvl is not None:
            try:
                return int(ppr.outlineLvl.val) + 1
            except (TypeError, ValueError):
                return None
        return None

    def _slots(self, document: Document, paragraph_node: dict[int, str]) -> list[TemplateSlot]:
        slots: list[TemplateSlot] = []
        for index, paragraph in enumerate(document.paragraphs):
            node_id = paragraph_node.get(index)
            if not node_id:
                continue
            for match_index, match in enumerate(_PLACEHOLDER.finditer(paragraph.text)):
                slots.append(
                    TemplateSlot(
                        slot_id=f"text-p-{index + 1}-{match_index + 1}",
                        node_id=node_id,
                        kind="text_slot",
                        anchor=f"paragraph:{index + 1}:placeholder:{match.group(0)}",
                    )
                )
        # Bind each table to the nearest preceding heading in document body order,
        # not the last heading of the entire document.
        nearest_before_table = self._nearest_heading_before_each_table(document, paragraph_node)
        for table_index, table in enumerate(document.tables):
            nearest_node = nearest_before_table.get(table_index) or next(iter(paragraph_node.values()), "p-1")
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    if _PLACEHOLDER.search(cell.text):
                        slots.append(
                            TemplateSlot(
                                slot_id=f"cell-t-{table_index + 1}-{row_index + 1}-{cell_index + 1}",
                                node_id=nearest_node,
                                kind="cell_slot",
                                anchor=f"table:{table_index + 1}:row:{row_index + 1}:cell:{cell_index + 1}",
                            )
                        )
        return slots

    @staticmethod
    def _nearest_heading_before_each_table(document: Document, paragraph_node: dict[int, str]) -> dict[int, str]:
        """Map table_index → nearest upstream heading node_id by body document order."""
        paragraphs = iter(document.paragraphs)
        tables = iter(document.tables)
        paragraph_index = 0
        table_index = 0
        last_heading: str | None = None
        mapping: dict[int, str] = {}
        for child in document.element.body.iterchildren():
            if child.tag.endswith("}p"):
                next(paragraphs)
                node_id = paragraph_node.get(paragraph_index)
                if node_id:
                    last_heading = node_id
                paragraph_index += 1
            elif child.tag.endswith("}tbl"):
                next(tables)
                if last_heading is not None:
                    mapping[table_index] = last_heading
                table_index += 1
        return mapping

    @staticmethod
    def _coverage_gaps(nodes: list[ContractNode], ledger: RequirementLedger) -> list[str]:
        titles = " ".join(node.title.lower() for node in nodes)
        gaps: list[str] = []
        for requirement in ledger.requirements:
            terms = [term for term in re.split(r"[^\w\u4e00-\u9fff]+", requirement.normalized_requirement.lower()) if len(term) >= 2]
            if terms and not any(term in titles for term in terms):
                gaps.append(f"TEMPLATE_COVERAGE_GAP:{requirement.requirement_id}")
        return gaps
