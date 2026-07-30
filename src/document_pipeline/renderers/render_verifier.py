from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document

from control_plane import ControlStore, WorkspaceContext
from utils import read_json, write_json

from ..quality import CONTENT_QUALITY_PATH


RENDER_QUALITY_PATH = Path("workspace/v3/reports/render_quality.json")
RENDER_OUTPUT_PATH = Path("outputs/v3/final.docx")
RENDER_MARKDOWN_PATH = Path("outputs/v3/final.md")


class DeliveryVerifier:
    """Validate the V3-rendered DOCX and record a delivery gate decision."""

    def __init__(
        self,
        context: WorkspaceContext,
        *,
        allow_quality_warnings: bool = False,
        operation_id: str | None = None,
    ) -> None:
        self.context = context
        self.root = context.root
        self.allow_quality_warnings = allow_quality_warnings
        self.operation_id = str(operation_id or "")

    def verify(self) -> dict[str, object]:
        quality = read_json(self.root / CONTENT_QUALITY_PATH)
        quality_verdict = str(quality.get("verdict") or "")
        if (
            quality_verdict != "pass"
            and not self.allow_quality_warnings
        ):
            raise ValueError("DELIVERY_BLOCKED: 内容质量门禁未通过")
        findings = [
            {**item, "stage_id": str(item.get("stage_id") or "verify_document")}
            for item in (quality.get("findings") or [])
            if isinstance(item, dict)
        ]
        store = ControlStore(self.context)
        pipeline_warnings = []
        if self.operation_id:
            for run in store.stage_runs(self.operation_id):
                output = run.get("output")
                value = output if isinstance(output, dict) else {}
                for warning in value.get("warnings") or []:
                    if isinstance(warning, dict):
                        pipeline_warnings.append(
                            {
                                **warning,
                                "stage_id": str(
                                    warning.get("stage_id")
                                    or run.get("stage_command")
                                    or ""
                                ),
                            }
                        )
        warnings: list[dict[str, object]] = []
        seen: set[tuple[str, str, str]] = set()
        for item in [*pipeline_warnings, *findings]:
            key = (
                str(item.get("stage_id") or ""),
                str(item.get("code") or ""),
                str(item.get("message") or ""),
            )
            if key in seen:
                continue
            seen.add(key)
            warnings.append(item)

        output = self.root / RENDER_OUTPUT_PATH
        if not output.is_file() or output.stat().st_size <= 0:
            raise ValueError("DELIVERY_BLOCKED: V3 DOCX 不存在或为空")
        try:
            document = Document(str(output))
        except Exception as exc:
            raise ValueError(f"DELIVERY_BLOCKED: V3 DOCX 无法打开: {exc}") from exc

        visible_content = sum(1 for paragraph in document.paragraphs if paragraph.text.strip())
        if visible_content == 0 and not document.tables:
            raise ValueError("DELIVERY_BLOCKED: V3 DOCX 不包含可见内容")

        markdown = self.root / RENDER_MARKDOWN_PATH
        if not markdown.is_file() or markdown.stat().st_size <= 0:
            raise ValueError("DELIVERY_BLOCKED: V3 Markdown 不存在或为空")
        digest = hashlib.sha256(output.read_bytes()).hexdigest()
        markdown_digest = hashlib.sha256(markdown.read_bytes()).hexdigest()
        report: dict[str, object] = {
            "schema_version": "v3",
            "status": (
                "ready_with_warnings"
                if quality_verdict != "pass" or warnings
                else "ready"
            ),
            "operation_id": self.operation_id,
            "artifact_path": RENDER_OUTPUT_PATH.as_posix(),
            "artifact_sha256": digest,
            "markdown_path": RENDER_MARKDOWN_PATH.as_posix(),
            "markdown_sha256": markdown_digest,
            "visible_paragraphs": visible_content,
            "table_count": len(document.tables),
            "warning_count": len(warnings),
            "warnings": warnings,
        }
        write_json(self.root / RENDER_QUALITY_PATH, report)
        store.record_gate_evaluation(
            command="verify_delivery",
            verdict="pass",
            input_fingerprint=digest,
            findings=[],
            source="v3.render_verifier",
        )
        return report
