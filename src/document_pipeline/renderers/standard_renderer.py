from __future__ import annotations

from pathlib import Path

from docx import Document
from control_plane import WorkspaceContext
from utils import read_json

from ..contracts import DOCUMENT_CONTRACT_ADAPTER, IntegratedDocument, OutlineContract
from ..document_contract import DOCUMENT_CONTRACT_PATH
from ..integrator import INTEGRATED_DOCUMENT_PATH


STANDARD_OUTPUT_PATH = Path("outputs/v3/final.docx")
STANDARD_MARKDOWN_PATH = Path("outputs/v3/final.md")


class StandardRenderer:
    """Build a normal bid document only from the frozen outline and integrated blocks."""

    def __init__(self, context: WorkspaceContext) -> None:
        self.context = context
        self.root = context.root

    def render(self) -> tuple[Path, Path]:
        contract = DOCUMENT_CONTRACT_ADAPTER.validate_python(read_json(self.root / DOCUMENT_CONTRACT_PATH))
        if not isinstance(contract, OutlineContract):
            raise ValueError("STANDARD_RENDER_BLOCKED: 当前文档不是 auto_outline 模式")
        integrated = IntegratedDocument.model_validate(read_json(self.root / INTEGRATED_DOCUMENT_PATH))
        blocks_by_node: dict[str, list[str]] = {}
        for block in integrated.blocks:
            blocks_by_node.setdefault(block.target_node_id, []).append(block.content)
        doc = Document()
        markdown: list[str] = []
        depth = {node.node_id: self._depth(node.node_id, contract.nodes) for node in contract.nodes}
        for node in contract.nodes:
            level = min(depth[node.node_id] + 1, 9)
            doc.add_heading(node.title, level=level)
            markdown.append(f"{'#' * level} {node.title}")
            for content in blocks_by_node.get(node.node_id, []):
                doc.add_paragraph(content)
                markdown.append(content)
        output = self.root / STANDARD_OUTPUT_PATH; output.parent.mkdir(parents=True, exist_ok=True)
        markdown_path = self.root / STANDARD_MARKDOWN_PATH
        doc.save(str(output)); markdown_path.write_text("\n\n".join(markdown) + "\n", encoding="utf-8")
        return output, markdown_path

    @staticmethod
    def _depth(node_id: str, nodes) -> int:
        by_id = {node.node_id: node for node in nodes}; depth = 0; current = by_id[node_id]
        while current.parent_node_id:
            depth += 1; current = by_id[current.parent_node_id]
        return depth
