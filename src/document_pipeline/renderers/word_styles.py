"""Word styles that match the on-page Word preview.

The chapter workbench and document preview render 宋体/黑体, first-line
indent, heading hierarchy, lists and bordered tables in CSS. Draft and
standard DOCX export must apply the same contract so the downloaded file
is not a style-less dump of plain paragraphs.
"""

from __future__ import annotations

import re
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from ..document_preview import parse_markdown_preview


BODY_ASCII_FONT = "Times New Roman"
BODY_EAST_ASIA_FONT = "宋体"
HEADING_ASCII_FONT = "Microsoft YaHei"
HEADING_EAST_ASIA_FONT = "黑体"
TABLE_ASCII_FONT = "Microsoft YaHei"
TABLE_EAST_ASIA_FONT = "微软雅黑"

BODY_SIZE_PT = 12  # 16px preview
BODY_LINE_SPACING = 1.85
FIRST_LINE_INDENT_PT = 24  # 2em at 12pt
PARAGRAPH_SPACE_AFTER_PT = 10

# Preview sizes: h1 28px / h2 23px / h3 19px / h4+ 16px
HEADING_SIZES_PT = {1: 21, 2: 17, 3: 14, 4: 12, 5: 12, 6: 12}
HEADING_SPACE_BEFORE_PT = {1: 0, 2: 18, 3: 14, 4: 12, 5: 10, 6: 10}
HEADING_SPACE_AFTER_PT = {1: 16, 2: 10, 3: 8, 4: 8, 5: 6, 6: 6}

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_LEFT_CM = 2.06
MARGIN_RIGHT_CM = 2.06
MARGIN_TOP_CM = 1.91
MARGIN_BOTTOM_CM = 1.91

TABLE_BORDER_COLOR = "64748B"
TABLE_HEADER_FILL = "F8FAFC"
TABLE_FONT_SIZE_PT = 10.5

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_LIST_ITEM_RE = re.compile(r"^\s*(?:[-*+•]|\d+[.)、．])\s+(.+)$")
_TABLE_SEPARATOR_RE = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$"
)
_HEADING_TYPE_LEVELS = {"heading": 2, "h2": 2, "h3": 3}


def apply_preview_styles(document: Document) -> None:
    """Install page setup and built-in styles used by the on-page preview."""
    _apply_page_setup(document)
    _apply_document_defaults(document)
    _style_normal(document)
    for level in range(1, 7):
        _style_heading(document, level)
    _style_list(document, "List Bullet")
    _style_list(document, "List Number")


def add_styled_heading(document: Document, text: str, level: int) -> Paragraph:
    level = max(1, min(int(level or 1), 6))
    paragraph = document.add_paragraph(style=f"Heading {level}")
    _reset_heading_format(paragraph, level)
    _write_runs(paragraph, text.strip(), HEADING_ASCII_FONT, HEADING_EAST_ASIA_FONT, HEADING_SIZES_PT[level], bold=True)
    return paragraph


def add_styled_paragraph(document: Document, text: str) -> Paragraph:
    paragraph = document.add_paragraph(style="Normal")
    _apply_body_format(paragraph)
    _write_runs(paragraph, text, BODY_ASCII_FONT, BODY_EAST_ASIA_FONT, BODY_SIZE_PT)
    return paragraph


def add_styled_list(document: Document, items: Iterable[str], *, numbered: bool = False) -> list[Paragraph]:
    style_name = "List Number" if numbered else "List Bullet"
    paragraphs: list[Paragraph] = []
    for item in items:
        text = str(item or "").strip()
        if not text:
            continue
        paragraph = document.add_paragraph(style=style_name)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(24)
        paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _write_runs(paragraph, text, BODY_ASCII_FONT, BODY_EAST_ASIA_FONT, BODY_SIZE_PT)
        paragraphs.append(paragraph)
    return paragraphs


def add_styled_table(document: Document, rows: list[list[str]]) -> Table | None:
    normalized = _normalize_table_rows(rows)
    if not normalized:
        return None
    column_count = max(len(row) for row in normalized)
    table = document.add_table(rows=len(normalized), cols=column_count)
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        pass
    table.autofit = False
    content_width_cm = PAGE_WIDTH_CM - MARGIN_LEFT_CM - MARGIN_RIGHT_CM
    _set_table_width(table, content_width_cm)
    col_width = content_width_cm / column_count
    for row_index, row in enumerate(normalized):
        for col_index in range(column_count):
            cell = table.cell(row_index, col_index)
            _set_cell_width(cell, col_width)
            _set_cell_borders(cell)
            _set_cell_margins(cell)
            if row_index == 0:
                _shade_cell(cell, TABLE_HEADER_FILL)
            _write_cell_text(
                cell,
                row[col_index] if col_index < len(row) else "",
                header=row_index == 0,
            )
    return table


def add_markdown_content(document: Document, markdown: str) -> None:
    """Render markdown with the same block types as the on-page Word preview."""
    text = str(markdown or "").strip()
    if not text:
        return
    _render_preview_blocks(document, _preview_blocks_from_text(text))


def add_content_block(document: Document, block: dict[str, Any], *, chapter_level: int = 1) -> None:
    content = str(block.get("content") or "").strip()
    if not content:
        return
    block_type = str(block.get("type") or "paragraph").strip().lower()
    if block_type in _HEADING_TYPE_LEVELS:
        heading = _HEADING_RE.match(content)
        title = heading.group(2).strip() if heading else content
        level = len(heading.group(1)) if heading else _HEADING_TYPE_LEVELS[block_type]
        add_styled_heading(document, title, min(max(level, chapter_level + 1), 6))
        return
    if block_type == "list":
        items, numbered = _list_items(content)
        add_styled_list(document, items, numbered=numbered)
        return
    if block_type == "table":
        rows = _table_rows_from_text(content)
        if rows:
            add_styled_table(document, rows)
            return
    add_markdown_content(document, content)


def write_composed_document(composed: dict[str, Any], output_path) -> Any:
    """Write a downloadable draft DOCX for the current chapter tree."""
    document = Document()
    apply_preview_styles(document)
    for chapter in composed.get("chapters") or []:
        if not isinstance(chapter, dict):
            continue
        title = str(chapter.get("title") or chapter.get("chapter_id") or "未命名章节").strip()
        try:
            depth = int(chapter.get("depth") or 0)
        except (TypeError, ValueError):
            depth = 0
        level = min(max(depth, 0) + 1, 4)
        add_styled_heading(document, title, level)
        for block in chapter.get("blocks") or []:
            if isinstance(block, dict):
                add_content_block(document, block, chapter_level=level)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def _preview_blocks_from_text(text: str) -> list[dict[str, Any]]:
    _, blocks = parse_markdown_preview(text)
    if (
        blocks
        and blocks[0].get("type") == "heading"
        and blocks[0].get("id") == "document-start"
        and blocks[0].get("text") == "标书全文"
    ):
        blocks = blocks[1:]
    numbered = _looks_numbered(text)
    for block in blocks:
        if block.get("type") == "list":
            block["numbered"] = numbered
    if any(block.get("type") == "table" for block in blocks):
        return blocks
    # Preview parser requires a markdown separator. Consecutive pipe rows
    # from writer blocks should still become a table in the exported file.
    rows = _table_rows_from_text(text)
    if rows and all("|" in line for line in text.splitlines() if line.strip()):
        return [{"id": "block-1", "type": "table", "section_id": "", "rows": rows}]
    return blocks


def _render_preview_blocks(document: Document, blocks: list[dict[str, Any]]) -> None:
    for block in blocks:
        block_type = str(block.get("type") or "")
        if block_type == "heading":
            add_styled_heading(document, str(block.get("text") or ""), int(block.get("level") or 1))
        elif block_type == "list":
            items = [str(item) for item in (block.get("items") or []) if str(item).strip()]
            add_styled_list(document, items, numbered=bool(block.get("numbered")))
        elif block_type == "table":
            add_styled_table(document, list(block.get("rows") or []))
        elif str(block.get("text") or "").strip():
            add_styled_paragraph(document, str(block.get("text") or ""))


def _list_items(content: str) -> tuple[list[str], bool]:
    items: list[str] = []
    numbered = False
    for line in str(content or "").splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        match = _LIST_ITEM_RE.match(stripped)
        if match:
            numbered = numbered or bool(re.match(r"^\d+", stripped))
            items.append(match.group(1).strip())
        else:
            items.append(stripped)
    return items, numbered


def _looks_numbered(content: str) -> bool:
    return bool(re.search(r"^\s*\d+[.)、．]\s+", content, re.M))


def _table_rows_from_text(content: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in str(content or "").splitlines():
        if "|" not in line or _TABLE_SEPARATOR_RE.match(line):
            continue
        rows.append(_table_cells(line))
    return rows


def _table_cells(line: str) -> list[str]:
    text = line.strip().strip("|")
    return [cell.strip().replace(r"\|", "|") for cell in text.split("|")]


def _normalize_table_rows(rows: list[list[str]]) -> list[list[str]]:
    cleaned: list[list[str]] = []
    for row in rows:
        cells = [str(cell or "").strip() for cell in (row or [])]
        if cells:
            cleaned.append(cells)
    return cleaned


def _apply_page_setup(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin = Cm(MARGIN_LEFT_CM)
    section.right_margin = Cm(MARGIN_RIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)


def _apply_document_defaults(document: Document) -> None:
    styles_el = document.styles.element
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    _set_rfonts(rpr, BODY_ASCII_FONT, BODY_EAST_ASIA_FONT)
    _set_sz(rpr, BODY_SIZE_PT)


def _style_normal(document: Document) -> None:
    style = document.styles["Normal"]
    _set_style_fonts(style, BODY_ASCII_FONT, BODY_EAST_ASIA_FONT, BODY_SIZE_PT, bold=False)
    fmt = style.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = BODY_LINE_SPACING
    fmt.first_line_indent = Pt(FIRST_LINE_INDENT_PT)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(PARAGRAPH_SPACE_AFTER_PT)
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _style_heading(document: Document, level: int) -> None:
    try:
        style = document.styles[f"Heading {level}"]
    except KeyError:
        return
    _set_style_fonts(
        style,
        HEADING_ASCII_FONT,
        HEADING_EAST_ASIA_FONT,
        HEADING_SIZES_PT[level],
        bold=True,
    )
    fmt = style.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.space_before = Pt(HEADING_SPACE_BEFORE_PT[level])
    fmt.space_after = Pt(HEADING_SPACE_AFTER_PT[level])
    fmt.line_spacing = 1.45
    fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    fmt.outline_level = level - 1


def _style_list(document: Document, style_name: str) -> None:
    try:
        style = document.styles[style_name]
    except KeyError:
        return
    _set_style_fonts(style, BODY_ASCII_FONT, BODY_EAST_ASIA_FONT, BODY_SIZE_PT, bold=False)
    fmt = style.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.line_spacing = BODY_LINE_SPACING
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _set_style_fonts(style, ascii_font: str, east_asia_font: str, size_pt: float, *, bold: bool) -> None:
    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    rpr = style.element.get_or_add_rPr()
    _set_rfonts(rpr, ascii_font, east_asia_font)
    _set_sz(rpr, size_pt)


def _apply_body_format(paragraph: Paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(FIRST_LINE_INDENT_PT)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = BODY_LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(PARAGRAPH_SPACE_AFTER_PT)
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _reset_heading_format(paragraph: Paragraph, level: int) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.space_before = Pt(HEADING_SPACE_BEFORE_PT[level])
    fmt.space_after = Pt(HEADING_SPACE_AFTER_PT[level])
    fmt.line_spacing = 1.45
    fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT


def _write_runs(
    paragraph: Paragraph,
    text: str,
    ascii_font: str,
    east_asia_font: str,
    size_pt: float,
    *,
    bold: bool = False,
) -> None:
    parts = str(text or "").split("\n")
    for index, part in enumerate(parts):
        if index:
            if paragraph.runs:
                paragraph.runs[-1].add_break()
            else:
                paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        _apply_run_font(run, ascii_font, east_asia_font, size_pt, bold=bold)


def _apply_run_font(run: Run, ascii_font: str, east_asia_font: str, size_pt: float, *, bold: bool = False) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    rpr = run._element.get_or_add_rPr()
    _set_rfonts(rpr, ascii_font, east_asia_font)
    _set_sz(rpr, size_pt)


def _set_rfonts(rpr, ascii_font: str, east_asia_font: str) -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)
    rfonts.set(qn("w:cs"), ascii_font)


def _set_sz(rpr, size_pt: float) -> None:
    half_points = str(int(round(size_pt * 2)))
    for tag in ("w:sz", "w:szCs"):
        node = rpr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            rpr.append(node)
        node.set(qn("w:val"), half_points)


def _write_cell_text(cell, text: str, *, header: bool) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.55
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _write_runs(
        paragraph,
        text,
        TABLE_ASCII_FONT,
        TABLE_EAST_ASIA_FONT,
        TABLE_FONT_SIZE_PT,
        bold=header,
    )


def _set_table_width(table: Table, width_cm: float) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_cm * 567)))


def _set_cell_width(cell, width_cm: float) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))


def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "8")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), TABLE_BORDER_COLOR)
        borders.append(node)
    tc_pr.append(borders)


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    margins = OxmlElement("w:tcMar")
    for edge, value in (("top", "80"), ("left", "120"), ("bottom", "80"), ("right", "120")):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)
