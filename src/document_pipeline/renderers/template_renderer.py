from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from control_plane import WorkspaceContext
from utils import read_json, write_json

from ..contracts import DOCUMENT_CONTRACT_ADAPTER, IntegratedDocument, TemplateContract
from ..document_contract import DOCUMENT_CONTRACT_PATH
from ..input_manifest import InputManifestService, V3_ROOT
from ..integrator import INTEGRATED_DOCUMENT_PATH
from ..template_contract import TemplateContractCompiler


TEMPLATE_OUTPUT_PATH = Path("outputs/v3/final.docx")
TEMPLATE_RENDER_REPORT_PATH = V3_ROOT / "reports" / "template_render.json"


class StrictTemplateRenderer:
    """Fill only declared slots in a copied template; never rebuild a document."""

    def __init__(self, context: WorkspaceContext) -> None:
        self.context = context
        self.root = context.root

    def render(self) -> Path:
        contract = DOCUMENT_CONTRACT_ADAPTER.validate_python(read_json(self.root / DOCUMENT_CONTRACT_PATH))
        if not isinstance(contract, TemplateContract):
            raise ValueError("TEMPLATE_RENDER_BLOCKED: 当前文档不是 template_strict 模式")
        if contract.blocking_gaps:
            raise ValueError("TEMPLATE_RENDER_BLOCKED: 模板覆盖存在缺口")
        document = IntegratedDocument.model_validate(read_json(self.root / INTEGRATED_DOCUMENT_PATH))
        template = next(item for item in InputManifestService(self.context).load().inputs if item.active and item.role.value == "template")
        source = self.root / V3_ROOT / "sources" / template.input_id / template.filename
        output = self.root / TEMPLATE_OUTPUT_PATH
        output.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, output)
        docx = Document(str(output))
        slot_values: dict[str, list[str]] = {}
        slot_ids = {slot.slot_id for slot in contract.slots}
        for block in document.blocks:
            if block.target_node_id not in slot_ids:
                raise ValueError(f"TEMPLATE_RENDER_BLOCKED: ContentBlock 未指向合法 slot: {block.target_node_id}")
            slot_values.setdefault(block.target_node_id, []).append(block.content)
        filled: list[str] = []
        for slot in contract.slots:
            values = "\n".join(slot_values.get(slot.slot_id, []))
            if not values:
                continue
            if slot.kind == "text_slot":
                self._fill_text_slot(docx, slot.anchor, values)
            elif slot.kind == "cell_slot":
                self._fill_cell_slot(docx, slot.anchor, values)
            else:
                raise ValueError(f"TEMPLATE_RENDER_BLOCKED: 尚未实现的 slot 类型 {slot.kind}")
            filled.append(slot.slot_id)
        docx.save(str(output))
        current_fingerprint = TemplateContractCompiler._fingerprint(output)
        if current_fingerprint != contract.structural_fingerprint:
            raise ValueError("TEMPLATE_RENDER_BLOCKED: 模板结构指纹发生未授权变化")
        write_json(self.root / TEMPLATE_RENDER_REPORT_PATH, {"schema_version": "v3", "filled_slots": filled, "structural_fingerprint": current_fingerprint, "ok": True})
        return output

    @staticmethod
    def _fill_text_slot(document: Document, anchor: str, value: str) -> None:
        pieces = anchor.split(":")
        index = int(pieces[1]) - 1
        placeholder = anchor.split(":placeholder:", 1)[1]
        paragraph = document.paragraphs[index]
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                return
        if placeholder in paragraph.text:
            raise ValueError("TEMPLATE_RENDER_BLOCKED: 占位符跨 run，无法安全替换")
        raise ValueError("TEMPLATE_RENDER_BLOCKED: 未找到 text_slot 锚点")

    @staticmethod
    def _fill_cell_slot(document: Document, anchor: str, value: str) -> None:
        _, table_index, _, row_index, _, cell_index = anchor.split(":")
        cell = document.tables[int(table_index) - 1].rows[int(row_index) - 1].cells[int(cell_index) - 1]
        if not cell.paragraphs:
            raise ValueError("TEMPLATE_RENDER_BLOCKED: cell_slot 不存在段落")
        cell.paragraphs[0].text = value
